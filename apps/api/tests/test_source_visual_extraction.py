from __future__ import annotations

import base64
import io
from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile

import pytest
from docx import Document
from docx.shared import Inches
from PIL import Image, ImageDraw
from reportlab.lib.utils import ImageReader
from reportlab.pdfgen import canvas

from app.models import (
    SourceChapter,
    SourceChunk,
    SourceIngestionRecord,
    SourceStructure,
    SourceVisualAsset,
)
from app.services import source_visual_extraction as extraction_module
from app.services import workspace_state
from app.services.source_archive import SafeSourceArchive, SourceArchiveError
from app.services.source_structure_indexer import SourceStructureIndexer
from app.services.source_structure_store import SourceStructureStore
from app.services.source_visual_extraction import (
    CURRENT_SOURCE_VISUAL_INDEX_VERSION,
    SourceVisualExtractor,
)
from app.services.source_visual_extraction_markup import (
    extract_markup_visuals,
    extract_standalone_image,
)
from app.services.source_visual_extraction_office import extract_office_visuals
from app.services.source_visual_extraction_pdf import extract_pdf_visuals
from app.services.source_visual_extraction_types import RawSourceVisual, SourceVisualAdapterResult
from app.services.source_visual_libreoffice import (
    LibreOfficeRenderError,
    _validate_ooxml_relationships,
)
from app.services.source_visual_storage import (
    SourceVisualStorageError,
    persist_source_visual_asset,
    resolve_source_visual_storage_key,
    source_visual_staging,
)
from app.services.source_xml import SourceXmlError, parse_untrusted_xml


_PNG_1X1 = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNk+A8A"
    "AQUBAScY42YAAAAASUVORK5CYII="
)


def _record(path: Path, *, mime_type: str) -> SourceIngestionRecord:
    return SourceIngestionRecord(
        id=f"source_{path.stem}",
        owner_user_id="user_visuals",
        package_id="package_visuals",
        title=path.name,
        source_type="local_file",
        file_name=path.name,
        mime_type=mime_type,
        size_bytes=path.stat().st_size,
        status="ready",
        metadata={"local_source_path": str(path)},
    )


def _png(width: int = 120, height: int = 80, *, label: str = "figure") -> bytes:
    image = Image.new("RGB", (width, height), "white")
    drawing = ImageDraw.Draw(image)
    drawing.rectangle((4, 4, width - 5, height - 5), outline="navy", width=3)
    drawing.line((8, height - 12, width - 8, 12), fill="red", width=4)
    drawing.text((10, 10), label, fill="black")
    output = io.BytesIO()
    image.save(output, format="PNG")
    return output.getvalue()


def test_pdf_extractor_rejects_full_page_scan_layer_but_keeps_real_figure(
    tmp_path: Path,
) -> None:
    path = tmp_path / "scan-with-figure.pdf"
    background = ImageReader(io.BytesIO(_png(600, 800, label="scanned page")))
    figure = ImageReader(io.BytesIO(_png(220, 140, label="real figure")))
    pdf = canvas.Canvas(str(path), pagesize=(600, 800))
    pdf.drawImage(background, 0, 0, width=600, height=800)
    pdf.drawImage(figure, 170, 330, width=220, height=140)
    pdf.drawString(170, 310, "Figure 1: retained teaching visual")
    pdf.save()

    result = extract_pdf_visuals(path)

    image_visuals = [visual for visual in result.visuals if visual.kind == "image"]
    assert result.status == "ready"
    assert len(image_visuals) == 1
    assert image_visuals[0].bbox[0] > 0.20
    assert image_visuals[0].bbox[2] < 0.80


def test_materialization_preserves_cross_page_and_verified_anchor_fields(
    monkeypatch,
    tmp_path: Path,
) -> None:
    monkeypatch.setattr(workspace_state, "UPLOAD_DIR", tmp_path / "uploads")
    source_path = tmp_path / "source.pdf"
    source_path.write_bytes(b"placeholder")
    record = _record(source_path, mime_type="application/pdf")
    structure = SourceStructure(
        id="structure_visuals",
        owner_user_id=record.owner_user_id,
        package_id=record.package_id,
        source_ingestion_id=record.id,
        status="ready",
    )
    wrong_prefix_chapter = SourceChapter(
        id="chapter_1",
        owner_user_id=record.owner_user_id,
        package_id=record.package_id,
        source_ingestion_id=record.id,
        title="One",
        source_locator="docx:paragraph:1",
        body_start_offset=0,
        body_end_offset=100,
        anchor_status="verified",
    )
    exact_chapter = SourceChapter(
        id="chapter_10",
        owner_user_id=record.owner_user_id,
        package_id=record.package_id,
        source_ingestion_id=record.id,
        title="Ten",
        source_locator="docx:paragraph:10",
        body_start_offset=100,
        body_end_offset=200,
        order_index=1,
        anchor_status="verified",
    )
    chunk = SourceChunk(
        id="chunk_10",
        owner_user_id=record.owner_user_id,
        package_id=record.package_id,
        source_ingestion_id=record.id,
        chapter_id=exact_chapter.id,
        text="Figure discussion",
        start_offset=100,
        end_offset=120,
    )
    raw = RawSourceVisual(
        kind="image",
        source_locator="docx:paragraph:10:drawing:0:image:0",
        native_order=0,
        content=_png(),
        mime_type="image/png",
        page_no=4,
        paragraph_index=10,
        bbox=[0.1, 0.2, 0.8, 0.7],
        caption="Figure",
        metadata={"page_end": 5},
    )

    visuals, warnings = SourceVisualExtractor()._materialize(
        record=record,
        structure=structure,
        raw_visuals=[raw],
        chapters=[wrong_prefix_chapter, exact_chapter],
        chunks=[chunk],
    )

    assert warnings == []
    assert len(visuals) == 1
    assert visuals[0].chapter_id == exact_chapter.id
    assert visuals[0].anchor_status == "verified"
    assert visuals[0].before_chunk_id == chunk.id
    assert visuals[0].after_chunk_id == chunk.id
    assert visuals[0].page_start == 4
    assert visuals[0].page_end == 5
    assert visuals[0].paragraph_index == 10
    assert visuals[0].position_hash


def test_pdf_visual_uses_chunk_locator_interval_when_page_fields_are_unset(
    monkeypatch,
    tmp_path: Path,
) -> None:
    monkeypatch.setattr(workspace_state, "UPLOAD_DIR", tmp_path / "uploads")
    source_path = tmp_path / "source.pdf"
    source_path.write_bytes(b"placeholder")
    record = _record(source_path, mime_type="application/pdf")
    structure = SourceStructure(
        id="structure_pdf_locator_anchor",
        owner_user_id=record.owner_user_id,
        package_id=record.package_id,
        source_ingestion_id=record.id,
        status="ready",
    )
    chunks = [
        SourceChunk(
            id="chunk_page_134",
            owner_user_id=record.owner_user_id,
            package_id=record.package_id,
            source_ingestion_id=record.id,
            order_index=0,
            source_locator="page:134",
            text="Discussion spanning pages 134 and 135",
            start_offset=0,
            end_offset=100,
        ),
        SourceChunk(
            id="chunk_page_136",
            owner_user_id=record.owner_user_id,
            package_id=record.package_id,
            source_ingestion_id=record.id,
            order_index=1,
            source_locator="page:136",
            text="Following discussion",
            start_offset=100,
            end_offset=180,
        ),
    ]
    raw = RawSourceVisual(
        kind="image",
        source_locator="pdf:page:135:image:1:occurrence:0",
        native_order=0,
        content=_png(),
        mime_type="image/png",
        page_no=135,
        bbox=[0.2, 0.2, 0.8, 0.7],
        caption="Grounded figure",
    )

    visuals, warnings = SourceVisualExtractor()._materialize(
        record=record,
        structure=structure,
        raw_visuals=[raw],
        chapters=[],
        chunks=chunks,
    )

    assert warnings == []
    assert visuals[0].anchor_status == "verified"
    assert visuals[0].before_chunk_id == "chunk_page_134"
    assert visuals[0].after_chunk_id == "chunk_page_136"


def test_markup_and_standalone_adapters_cover_html_markdown_csv_and_image(
    tmp_path: Path,
) -> None:
    data_uri = "data:image/png;base64," + base64.b64encode(_PNG_1X1).decode("ascii")
    html_path = tmp_path / "page.html"
    html_path.write_text(
        f'<h1>Heading</h1><p>Before <img src="{data_uri}" alt="Diagram"></p>'
        "<table><tr><th>A</th><th>B</th></tr>"
        '<tr><td colspan="2">merged</td></tr></table>',
        encoding="utf-8",
    )
    markdown_path = tmp_path / "notes.md"
    markdown_path.write_text(
        f"# Notes\n\n![Diagram]({data_uri})\n\n| A | B |\n| --- | --- |\n| 1 | 2 |\n",
        encoding="utf-8",
    )
    csv_path = tmp_path / "table.csv"
    csv_path.write_text("name,value\nalpha,1\n", encoding="utf-8")
    image_path = tmp_path / "figure.png"
    image_path.write_bytes(_PNG_1X1)

    html = extract_markup_visuals(html_path, _record(html_path, mime_type="text/html"))
    markdown = extract_markup_visuals(
        markdown_path,
        _record(markdown_path, mime_type="text/markdown"),
    )
    csv_result = extract_markup_visuals(csv_path, _record(csv_path, mime_type="text/csv"))
    standalone = extract_standalone_image(
        image_path,
        _record(image_path, mime_type="image/png"),
    )

    assert html.status == "partial"
    assert [visual.kind for visual in html.visuals] == ["image", "table"]
    assert html.visuals[0].paragraph_index == 1
    assert html.visuals[1].metadata["force_unverified"] is True
    assert [visual.kind for visual in markdown.visuals] == ["image", "table"]
    assert markdown.visuals[0].text_offset < markdown.visuals[1].text_offset
    assert csv_result.visuals[0].table_data == [["name", "value"], ["alpha", "1"]]
    assert standalone.visuals[0].metadata["standalone_image"] is True


def test_office_adapters_cover_docx_pptx_and_xlsx_visuals(tmp_path: Path) -> None:
    image_path = tmp_path / "figure.png"
    image_path.write_bytes(_png())

    docx_path = tmp_path / "document.docx"
    document = Document()
    document.add_heading("1 Visuals", level=1)
    document.add_paragraph("Image context")
    document.add_picture(str(image_path), width=Inches(1.5))
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "A"
    table.cell(0, 1).text = "B"
    table.cell(1, 0).text = "1"
    table.cell(1, 1).text = "2"
    document.save(docx_path)

    pptx_path = tmp_path / "slides.pptx"
    _write_pptx_with_image(pptx_path, _png())
    xlsx_path = tmp_path / "workbook.xlsx"
    _write_xlsx_with_table(xlsx_path)

    docx = extract_office_visuals(docx_path)
    pptx = extract_office_visuals(pptx_path)
    xlsx = extract_office_visuals(xlsx_path)
    extensionless_docx_path = tmp_path / "document-blob"
    extensionless_docx_path.write_bytes(docx_path.read_bytes())
    extensionless_docx = SourceVisualExtractor()._adapter_result(
        path=extensionless_docx_path,
        record=_record(
            extensionless_docx_path,
            mime_type=(
                "application/vnd.openxmlformats-officedocument."
                "wordprocessingml.document"
            ),
        ),
    )

    assert {visual.kind for visual in docx.visuals} == {"image", "table"}
    assert {visual.kind for visual in extensionless_docx.visuals} == {"image", "table"}
    assert next(visual for visual in docx.visuals if visual.kind == "image").paragraph_index == 2
    assert len(pptx.visuals) == 1
    assert pptx.visuals[0].slide_no == 1
    assert pptx.visuals[0].bbox == [0.1, 0.2, 0.4, 0.6]
    assert len(xlsx.visuals) == 1
    assert xlsx.visuals[0].kind == "table"
    assert xlsx.visuals[0].sheet_name == "Metrics"
    assert xlsx.visuals[0].table_data == [["Metric", "Value"], ["Retention", "0.82"]]


def test_legacy_structure_is_lazily_upgraded_with_verified_visuals(tmp_path: Path) -> None:
    csv_path = tmp_path / "metrics.csv"
    csv_path.write_text("metric,value\nretention,0.82\n", encoding="utf-8")
    record = _record(csv_path, mime_type="text/csv")
    store = SourceStructureStore(tmp_path / "openclass.sqlite3")
    legacy = SourceStructure(
        id="structure_legacy",
        owner_user_id=record.owner_user_id,
        package_id=record.package_id,
        source_ingestion_id=record.id,
        status="linear_only",
        visual_index_status="pending",
        visual_index_version=0,
    )
    chunk = SourceChunk(
        id="chunk_legacy",
        owner_user_id=record.owner_user_id,
        package_id=record.package_id,
        source_ingestion_id=record.id,
        text="metric,value\nretention,0.82",
        start_offset=0,
        end_offset=27,
    )
    store.save_structure_bundle(structure=legacy, chapters=[], chunks=[chunk])

    upgraded = SourceStructureIndexer(store=store).ensure_structure(record)
    view = store.get_structure_view(source=record)

    assert upgraded is not None
    assert upgraded.visual_index_version == CURRENT_SOURCE_VISUAL_INDEX_VERSION
    assert upgraded.visual_index_status == "ready"
    assert upgraded.visual_count == 1
    assert view.visuals[0].anchor_status == "verified"
    assert view.visuals[0].before_chunk_id == view.chunks[0].id
    assert view.visuals[0].structure_id == upgraded.id
    assert view.visuals[0].structure_version == CURRENT_SOURCE_VISUAL_INDEX_VERSION


def test_visual_storage_archive_xml_and_office_render_boundaries(
    monkeypatch,
    tmp_path: Path,
) -> None:
    monkeypatch.setattr(workspace_state, "UPLOAD_DIR", tmp_path / "uploads")
    first_key, first_hash = persist_source_visual_asset(_PNG_1X1, mime_type="image/png")
    second_key, second_hash = persist_source_visual_asset(_PNG_1X1, mime_type="image/png")
    assert (first_key, first_hash) == (second_key, second_hash)
    assert resolve_source_visual_storage_key(first_key).read_bytes() == _PNG_1X1
    with pytest.raises(SourceVisualStorageError):
        resolve_source_visual_storage_key("../escape.png", must_exist=False)

    unsafe_archive = tmp_path / "unsafe.epub"
    with ZipFile(unsafe_archive, "w", ZIP_DEFLATED) as archive:
        archive.writestr("../escape.xhtml", "unsafe")
    with pytest.raises(SourceArchiveError):
        SafeSourceArchive(unsafe_archive)

    with pytest.raises(SourceXmlError):
        parse_untrusted_xml('<!DOCTYPE root [<!ENTITY x SYSTEM "file:///etc/passwd">]><root>&x;</root>')

    external_office = tmp_path / "external.docx"
    with ZipFile(external_office, "w", ZIP_DEFLATED) as archive:
        archive.writestr(
            "word/_rels/document.xml.rels",
            '<Relationships xmlns="rels"><Relationship Id="rId1" '
            'TargetMode="External" Target="https://example.com/image.png"/></Relationships>',
        )
    with pytest.raises(LibreOfficeRenderError, match="external relationship"):
        _validate_ooxml_relationships(external_office)


def test_store_reads_visual_bytes_with_scope_path_and_hash_validation(
    monkeypatch,
    tmp_path: Path,
) -> None:
    monkeypatch.setattr(workspace_state, "UPLOAD_DIR", tmp_path / "uploads")
    store = SourceStructureStore(tmp_path / "openclass.sqlite3")
    content = _png()
    storage_key, content_hash = persist_source_visual_asset(
        content,
        mime_type="image/png",
    )
    stored_path = resolve_source_visual_storage_key(storage_key)
    stored_asset = SourceVisualAsset(
        id="visual_stored",
        owner_user_id="owner_1",
        package_id="package_1",
        source_ingestion_id="source_stored",
        structure_id="structure_stored",
        structure_version=CURRENT_SOURCE_VISUAL_INDEX_VERSION,
        storage_key=storage_key,
        asset_path=str(stored_path),
        mime_type="image/png",
        content_hash=content_hash,
        anchor_status="verified",
    )
    store.save_structure_bundle(
        structure=SourceStructure(
            id="structure_stored",
            owner_user_id="owner_1",
            package_id="package_1",
            source_ingestion_id="source_stored",
            status="linear_only",
        ),
        chapters=[],
        chunks=[],
        visuals=[stored_asset],
    )

    result = store.read_visual_bytes(
        owner_user_id="owner_1",
        package_id="package_1",
        source_id="source_stored",
        visual_id=stored_asset.id,
    )
    assert result is not None
    assert result[0].id == stored_asset.id
    assert result[1] == content
    assert store.get_visual(
        owner_user_id="owner_1",
        package_id="package_1",
        source_id="source_stored",
        visual_id=stored_asset.id,
    ) is not None
    assert (
        store.get_visual(
            owner_user_id="owner_1",
            package_id="other_package",
            source_id="source_stored",
            visual_id=stored_asset.id,
        )
        is None
    )
    assert (
        store.get_visual(
            owner_user_id="owner_1",
            package_id="package_1",
            source_id="other_source",
            visual_id=stored_asset.id,
        )
        is None
    )
    assert (
        store.read_visual_bytes(
            owner_user_id="other_owner",
            package_id="package_1",
            source_id="source_stored",
            visual_id=stored_asset.id,
        )
        is None
    )

    legacy_directory = (
        tmp_path
        / "external-source"
        / ".openclass-source-visuals"
        / "source-visuals"
        / "source_legacy"
    )
    legacy_directory.mkdir(parents=True)
    legacy_path = legacy_directory / f"{content_hash}.png"
    legacy_path.write_bytes(content)
    legacy_asset = SourceVisualAsset(
        id="visual_legacy",
        owner_user_id="owner_1",
        package_id="package_1",
        source_ingestion_id="source_legacy",
        structure_id="structure_legacy_path",
        structure_version=1,
        asset_path=str(legacy_path),
        mime_type="image/png",
        content_hash=content_hash,
        anchor_status="verified",
    )
    store.save_structure_bundle(
        structure=SourceStructure(
            id="structure_legacy_path",
            owner_user_id="owner_1",
            package_id="package_1",
            source_ingestion_id="source_legacy",
            status="linear_only",
        ),
        chapters=[],
        chunks=[],
        visuals=[legacy_asset],
    )
    legacy_result = store.read_visual_bytes(
        owner_user_id="owner_1",
        package_id="package_1",
        source_id="source_legacy",
        visual_id=legacy_asset.id,
    )
    assert legacy_result is not None
    assert legacy_result[1] == content

    linked_directory = legacy_directory.parent / "source_linked"
    linked_directory.mkdir()
    linked_path = linked_directory / "linked.png"
    linked_path.symlink_to(legacy_path)
    linked_asset = SourceVisualAsset(
        id="visual_linked",
        owner_user_id="owner_1",
        package_id="package_1",
        source_ingestion_id="source_linked",
        asset_path=str(linked_path),
        mime_type="image/png",
        content_hash=content_hash,
    )
    store.save_structure_bundle(
        structure=SourceStructure(
            owner_user_id="owner_1",
            package_id="package_1",
            source_ingestion_id="source_linked",
            status="linear_only",
        ),
        chapters=[],
        chunks=[],
        visuals=[linked_asset],
    )
    assert (
        store.read_visual_bytes(
            owner_user_id="owner_1",
            package_id="package_1",
            source_id="source_linked",
            visual_id=linked_asset.id,
        )
        is None
    )

    stored_path.write_bytes(b"tampered")
    assert (
        store.read_visual_bytes(
            owner_user_id="owner_1",
            package_id="package_1",
            source_id="source_stored",
            visual_id=stored_asset.id,
        )
        is None
    )


def test_office_render_mapping_accepts_raster_output_and_rejects_unanchored_objects(
    monkeypatch,
    tmp_path: Path,
) -> None:
    class _Renderer:
        available = True

        def render_pdf(self, _source_path: Path, *, output_dir: Path) -> Path:
            output = output_dir / "rendered.pdf"
            output.write_bytes(b"rendered")
            return output

    rendered_candidate = RawSourceVisual(
        kind="image",
        source_locator="pdf:page:1:image:0",
        native_order=0,
        content=_png(),
        mime_type="image/png",
        page_no=1,
        bbox=[0.1, 0.2, 0.5, 0.6],
        confidence=0.9,
    )
    monkeypatch.setattr(
        extraction_module,
        "extract_pdf_visuals",
        lambda _path: SourceVisualAdapterResult(visuals=[rendered_candidate]),
    )
    extractor = SourceVisualExtractor(office_renderer=_Renderer())
    source_path = tmp_path / "slides.pptx"
    source_path.write_bytes(b"placeholder")
    anchored = RawSourceVisual(
        kind="chart",
        source_locator="pptx:slide:1:native-chart:0",
        native_order=0,
        page_no=1,
        slide_no=1,
        bbox=[0.1, 0.2, 0.5, 0.6],
    )
    unanchored = RawSourceVisual(
        kind="diagram",
        source_locator="docx:paragraph:2:native-diagram:0",
        native_order=1,
    )

    result = extractor._render_native_office_visuals(
        source_path=source_path,
        anchors=[anchored, unanchored],
    )

    assert result.status == "partial"
    assert len(result.visuals) == 1
    assert result.visuals[0].kind == "chart"
    assert result.visuals[0].source_locator == anchored.source_locator
    assert result.visuals[0].metadata["office_anchor_mapping_verified"] is True
    assert any("1 of 2" in warning for warning in result.warnings)


def _write_pptx_with_image(path: Path, image: bytes) -> None:
    with ZipFile(path, "w", ZIP_DEFLATED) as archive:
        archive.writestr(
            "ppt/presentation.xml",
            '<p:presentation xmlns:p="p"><p:sldSz cx="1000" cy="1000"/></p:presentation>',
        )
        archive.writestr(
            "ppt/slides/slide1.xml",
            """<p:sld xmlns:p="p" xmlns:a="a" xmlns:r="r">
            <p:cSld><p:spTree><p:pic><p:nvPicPr><p:cNvPr id="1" name="Figure"/>
            </p:nvPicPr><p:blipFill><a:blip r:embed="rId1"/></p:blipFill>
            <p:spPr><a:xfrm><a:off x="100" y="200"/><a:ext cx="300" cy="400"/>
            </a:xfrm><a:prstGeom prst="rect"/><a:ln><a:noFill/></a:ln></p:spPr>
            </p:pic></p:spTree></p:cSld></p:sld>""",
        )
        archive.writestr(
            "ppt/slides/_rels/slide1.xml.rels",
            '<Relationships xmlns="rels"><Relationship Id="rId1" '
            'Target="../media/image1.png"/></Relationships>',
        )
        archive.writestr("ppt/media/image1.png", image)


def _write_xlsx_with_table(path: Path) -> None:
    with ZipFile(path, "w", ZIP_DEFLATED) as archive:
        archive.writestr(
            "xl/workbook.xml",
            '<workbook xmlns="x" xmlns:r="r"><sheets><sheet name="Metrics" r:id="rId1"/>'
            "</sheets></workbook>",
        )
        archive.writestr(
            "xl/_rels/workbook.xml.rels",
            '<Relationships xmlns="rels"><Relationship Id="rId1" '
            'Target="worksheets/sheet1.xml"/></Relationships>',
        )
        archive.writestr(
            "xl/sharedStrings.xml",
            '<sst xmlns="x"><si><t>Metric</t></si><si><t>Value</t></si>'
            "<si><t>Retention</t></si></sst>",
        )
        archive.writestr(
            "xl/worksheets/sheet1.xml",
            '<worksheet xmlns="x" xmlns:r="r"><sheetData>'
            '<row r="1"><c r="A1" t="s"><v>0</v></c><c r="B1" t="s"><v>1</v></c></row>'
            '<row r="2"><c r="A2" t="s"><v>2</v></c><c r="B2"><v>0.82</v></c></row>'
            '</sheetData><tableParts><tablePart r:id="rId2"/></tableParts></worksheet>',
        )
        archive.writestr(
            "xl/worksheets/_rels/sheet1.xml.rels",
            '<Relationships xmlns="rels"><Relationship Id="rId2" '
            'Target="../tables/table1.xml"/></Relationships>',
        )
        archive.writestr(
            "xl/tables/table1.xml",
            '<table xmlns="x" name="MetricsTable" displayName="MetricsTable" ref="A1:B2"/>',
        )


def test_visual_scope_uses_exclusive_page_end(tmp_path: Path) -> None:
    store = SourceStructureStore(tmp_path / "page-range.sqlite3")
    structure = SourceStructure(
        id="structure_page_range",
        owner_user_id="owner_page_range",
        package_id="package_page_range",
        source_ingestion_id="source_page_range",
        status="ready",
    )
    visuals = [
        SourceVisualAsset(
            id=f"visual_page_{page}",
            owner_user_id=structure.owner_user_id,
            package_id=structure.package_id,
            source_ingestion_id=structure.source_ingestion_id,
            structure_id=structure.id,
            kind="table",
            page_start=page,
            page_end=page,
            anchor_status="verified",
            table_data=[["page", str(page)]],
            order_index=page,
        )
        for page in (4, 5)
    ]
    store.save_structure_bundle(
        structure=structure,
        chapters=[],
        chunks=[],
        visuals=visuals,
    )

    selected = store.visual_evidence_for_scope(
        owner_user_id=structure.owner_user_id,
        package_id=structure.package_id,
        source_ingestion_id=structure.source_ingestion_id,
        page_start=4,
        page_end=5,
    )

    assert [item.visual_id for item in selected] == ["visual_page_4"]


def test_visual_cleanup_does_not_delete_an_inflight_staged_blob(
    monkeypatch: pytest.MonkeyPatch,
    tmp_path: Path,
) -> None:
    monkeypatch.setattr(workspace_state, "UPLOAD_DIR", tmp_path / "uploads")
    store = SourceStructureStore(tmp_path / "cleanup.sqlite3")

    with source_visual_staging():
        storage_key, _content_hash = persist_source_visual_asset(
            _PNG_1X1,
            mime_type="image/png",
        )
        path = resolve_source_visual_storage_key(storage_key)
        store.cleanup_unreferenced_visual_assets([storage_key])
        assert path.is_file()

    store.cleanup_unreferenced_visual_assets([storage_key])
    assert not path.exists()
