from __future__ import annotations

import zipfile

import pytest
from docx import Document as DocxDocument
from reportlab.pdfgen import canvas

from app.models import CoursePackage, ResourceLibraryItem
from app.services.resource_copyright_audit import (
    ResourcePublicationBlockedError,
    ResourceCopyrightReviewer,
    SearchResult,
    assert_package_publication_allowed,
    audit_resource_public_distribution,
    copyright_metadata_query,
)
from app.services.resource_library import build_resource_item


class FakeProvider:
    name = "fake"

    def __init__(self, results: list[SearchResult] | None = None, *, fail: bool = False) -> None:
        self.results = results or []
        self.fail = fail
        self.queries: list[str] = []

    def search(self, query: str, *, limit: int = 8) -> list[SearchResult]:
        self.queries.append(query)
        if self.fail:
            raise RuntimeError("search unavailable")
        return self.results


def test_metadata_query_does_not_include_resource_body_text() -> None:
    resource = ResourceLibraryItem(
        name="original-notes.md",
        mime_type="text/markdown",
        resource_type="document",
        size_bytes=12,
        text_content="private body text that must not be sent to search",
    )

    query = copyright_metadata_query(resource)

    assert "original notes" in query
    assert "private body" not in query


def test_markdown_probe_packet_extracts_rights_metadata_without_body(tmp_path) -> None:
    private_body = "PRIVATE BODY SENTENCE THAT MUST NOT BE SENT"
    resource_path = tmp_path / "packet-notes.md"
    resource_path.write_text(
        "\n".join(
            [
                "# Packet Notes",
                "Copyright 2026 Generic Publisher",
                "ISBN 9781234567897",
                "License: Creative Commons CC BY 4.0",
                "",
                "## Working Section",
                private_body,
            ]
        ),
        encoding="utf-8",
    )

    resource = build_resource_item(resource_path, "packet-notes.md")
    packet_text = _packet_text(resource)

    assert "Packet Notes" in resource.copyright_probe.title_candidates
    assert "9781234567897" in resource.copyright_probe.isbn_candidates
    assert resource.copyright_probe.license_candidates
    assert private_body not in packet_text


def test_docx_probe_packet_extracts_core_properties_and_rights(tmp_path) -> None:
    resource_path = tmp_path / "rights.docx"
    document = DocxDocument()
    document.core_properties.title = "Rights Packet"
    document.core_properties.author = "Example Author"
    document.add_heading("Rights Packet", level=1)
    document.add_paragraph("Copyright 2026 Example Publisher")
    document.add_paragraph("ISBN 9781234567897")
    document.add_heading("Body", level=1)
    document.add_paragraph("PRIVATE DOCX BODY SENTENCE THAT MUST NOT BE SENT")
    document.save(resource_path)

    resource = build_resource_item(resource_path, "rights.docx")
    packet_text = _packet_text(resource)

    assert "Rights Packet" in resource.copyright_probe.title_candidates
    assert "Example Author" in resource.copyright_probe.author_candidates
    assert "9781234567897" in resource.copyright_probe.isbn_candidates
    assert "PRIVATE DOCX BODY" not in packet_text


def test_pdf_probe_packet_reads_metadata_and_selected_pages_without_body(tmp_path) -> None:
    resource_path = tmp_path / "frontmatter.pdf"
    pdf = canvas.Canvas(str(resource_path))
    pdf.setTitle("PDF Rights Packet")
    pdf.setAuthor("PDF Author")
    pdf.drawString(72, 760, "PDF Rights Packet")
    pdf.drawString(72, 736, "Copyright 2026 PDF Publisher")
    pdf.drawString(72, 712, "ISBN 9781234567897")
    pdf.showPage()
    for _ in range(5):
        pdf.drawString(72, 760, "introductory material")
        pdf.showPage()
    pdf.drawString(72, 760, "PRIVATE PDF BODY SENTENCE THAT MUST NOT BE SENT")
    pdf.showPage()
    pdf.drawString(72, 760, "closing page")
    pdf.save()

    resource = build_resource_item(resource_path, "frontmatter.pdf")
    packet_text = _packet_text(resource)

    assert "PDF Rights Packet" in resource.copyright_probe.title_candidates
    assert "PDF Author" in resource.copyright_probe.author_candidates
    assert "9781234567897" in resource.copyright_probe.isbn_candidates
    assert "PRIVATE PDF BODY" not in packet_text


def test_epub_probe_packet_extracts_opf_metadata_and_copyright_item(tmp_path) -> None:
    resource_path = tmp_path / "packet.epub"
    _write_minimal_epub(resource_path)

    resource = build_resource_item(resource_path, "packet.epub")
    packet_text = _packet_text(resource)

    assert "EPUB Rights Packet" in resource.copyright_probe.title_candidates
    assert "EPUB Publisher" in resource.copyright_probe.publisher_candidates
    assert "9781234567897" in resource.copyright_probe.isbn_candidates
    assert resource.copyright_probe.rights_candidates
    assert "PRIVATE EPUB BODY" not in packet_text


def test_search_provider_receives_only_probe_metadata(tmp_path) -> None:
    resource_path = tmp_path / "private-notes.md"
    private_body = "PRIVATE SEARCH BODY THAT MUST NOT BE SENT"
    resource_path.write_text(
        "# Private Notes\n\nCopyright 2026 Generic Publisher\nISBN 9781234567897\n\n## Body\n"
        + private_body,
        encoding="utf-8",
    )
    resource = build_resource_item(resource_path, "private-notes.md")
    provider = FakeProvider(
        [
            SearchResult(
                title="Catalog listing",
                url="https://example.com/catalog",
                snippet="Publisher listing with ISBN and ebook editions.",
            )
        ]
    )

    audit = audit_resource_public_distribution(resource, search_provider=provider)

    assert audit.public_distribution == "blocked"
    assert provider.queries
    assert "9781234567897" in provider.queries[0]
    assert private_body not in provider.queries[0]
    assert resource.body_blocks


def test_unconfigured_ai_reviewer_does_not_block_deterministic_audit() -> None:
    provider = FakeProvider(
        [
            SearchResult(
                title="Open material",
                url="https://example.org/open",
                snippet="Released under Creative Commons CC BY 4.0.",
            )
        ]
    )

    audit = audit_resource_public_distribution(
        ResourceLibraryItem(name="open-material.md", mime_type="text/markdown", resource_type="document", size_bytes=10),
        search_provider=provider,
        reviewer=ResourceCopyrightReviewer(),
    )

    assert audit.status == "clear"
    assert audit.public_distribution == "allowed"


def test_audit_blocks_commercial_publication_match() -> None:
    provider = FakeProvider(
        [
            SearchResult(
                title="Catalog listing",
                url="https://example.com/catalog",
                snippet="Publisher listing with ISBN, hardcover, paperback, and ebook editions.",
            )
        ]
    )

    audit = audit_resource_public_distribution(
        ResourceLibraryItem(name="catalog-title.pdf", mime_type="application/pdf", resource_type="document", size_bytes=10),
        search_provider=provider,
    )

    assert audit.status == "public_blocked"
    assert audit.public_distribution == "blocked"
    assert "commercial_publication_match" in audit.signals


def test_audit_allows_open_license_evidence() -> None:
    provider = FakeProvider(
        [
            SearchResult(
                title="Open course material",
                url="https://example.org/open",
                snippet="Released under Creative Commons CC BY 4.0.",
            )
        ]
    )

    audit = audit_resource_public_distribution(
        ResourceLibraryItem(name="open-material.md", mime_type="text/markdown", resource_type="document", size_bytes=10),
        search_provider=provider,
    )

    assert audit.status == "clear"
    assert audit.public_distribution == "allowed"
    assert "open_license_evidence" in audit.signals


def test_audit_marks_provider_errors_without_blocking_private_upload() -> None:
    audit = audit_resource_public_distribution(
        ResourceLibraryItem(name="notes.md", mime_type="text/markdown", resource_type="document", size_bytes=10),
        search_provider=FakeProvider(fail=True),
    )

    assert audit.status == "error"
    assert audit.public_distribution == "pending"
    assert "external_search_failed" in audit.signals


def test_publication_gate_requires_allowed_resources() -> None:
    blocked = ResourceLibraryItem(name="blocked.pdf", mime_type="application/pdf", resource_type="document", size_bytes=1)
    blocked.copyright_audit.public_distribution = "blocked"
    allowed = ResourceLibraryItem(name="allowed.md", mime_type="text/markdown", resource_type="document", size_bytes=1)
    allowed.copyright_audit.public_distribution = "allowed"

    with pytest.raises(ResourcePublicationBlockedError) as exc_info:
        assert_package_publication_allowed(CoursePackage(title="Package", summary="", lessons=[], resources=[allowed, blocked]))

    assert exc_info.value.resource_names == ["blocked.pdf"]


def _packet_text(resource: ResourceLibraryItem) -> str:
    return "\n".join(section.text_excerpt for section in resource.copyright_probe.probe_sections)


def _write_minimal_epub(path) -> None:
    with zipfile.ZipFile(path, "w") as archive:
        archive.writestr("mimetype", "application/epub+zip")
        archive.writestr(
            "META-INF/container.xml",
            """<?xml version="1.0"?>
            <container version="1.0" xmlns="urn:oasis:names:tc:opendocument:xmlns:container">
              <rootfiles>
                <rootfile full-path="OEBPS/content.opf" media-type="application/oebps-package+xml"/>
              </rootfiles>
            </container>""",
        )
        archive.writestr(
            "OEBPS/content.opf",
            """<?xml version="1.0"?>
            <package version="3.0" xmlns="http://www.idpf.org/2007/opf"
              xmlns:dc="http://purl.org/dc/elements/1.1/">
              <metadata>
                <dc:title>EPUB Rights Packet</dc:title>
                <dc:creator>EPUB Author</dc:creator>
                <dc:publisher>EPUB Publisher</dc:publisher>
                <dc:identifier>ISBN 9781234567897</dc:identifier>
                <dc:rights>Copyright 2026 EPUB Publisher</dc:rights>
              </metadata>
              <manifest>
                <item id="copyright" href="copyright.xhtml" media-type="application/xhtml+xml"/>
                <item id="body" href="body.xhtml" media-type="application/xhtml+xml"/>
              </manifest>
              <spine>
                <itemref idref="copyright"/>
                <itemref idref="body"/>
              </spine>
            </package>""",
        )
        archive.writestr(
            "OEBPS/copyright.xhtml",
            """<html xmlns="http://www.w3.org/1999/xhtml">
              <body><h1>Copyright</h1><p>All rights reserved. ISBN 9781234567897.</p></body>
            </html>""",
        )
        archive.writestr(
            "OEBPS/body.xhtml",
            """<html xmlns="http://www.w3.org/1999/xhtml">
              <body><h1>Body</h1><p>PRIVATE EPUB BODY SENTENCE THAT MUST NOT BE SENT</p></body>
            </html>""",
        )
