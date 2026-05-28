import json

import pytest
from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor
from fastapi import HTTPException
from fastapi.testclient import TestClient
from pydantic import BaseModel

import app.main as main_module
from app.models import (
    AIModelSelection,
    BoardTeachingProgress,
    ChatRequest,
    CreateBranchRequest,
    DocumentSaveRequest,
    InteractionRuleDraft,
    InteractionSession,
    InteractionTurnDecision,
    LearningRequirementChecklistItem,
    LearningRequirementKeyFact,
    RealtimeTranscriptLogRequest,
    SelectionRef,
    UserView,
)
from app.routers.auth import current_user
from app.routers import documents as documents_router
from app.routers import realtime as realtime_router
from app.services.ai_logging import ai_log_context, ai_usage_logger, current_ai_log_context
from app.services import chat_service, workspace_state
from app.services.course_runtime import refresh_lesson_runtime
from app.services.course_store import SqliteCourseStore, build_initial_workspace_state
from app.services.lesson_factory import build_requirements, create_empty_lesson
from app.services.openai_course_ai import (
    BoardDocumentEditResult,
    ChatbotReply,
    ComplexProblemSolution,
    GeneratedResourceCatalog,
    LearningRequirementUpdate,
    OpenAICourseAI,
    bind_text_model_selection,
    openai_course_ai,
)
from app.services.rich_document import build_document
from app.services.resource_library import build_resource_item


TEST_USER = UserView(
    id="user_test",
    email="test@example.com",
    role="user",
    created_at="2026-01-01T00:00:00+00:00",
)


def _read_log_entries(path):
    if not path.exists():
        return []
    return [json.loads(line) for line in path.read_text(encoding="utf-8").splitlines() if line.strip()]


def _seed_test_user_workspace(store: SqliteCourseStore):
    workspace = build_initial_workspace_state()
    workspace.packages[0].title = "测试课程工作台"
    lesson = create_empty_lesson("测试页面")
    package = workspace.packages[0]
    package.lessons.append(lesson)
    package.open_lesson_ids.append(lesson.id)
    package.workspace_tab_order.append(lesson.id)
    package.active_lesson_id = lesson.id
    store.save_for_user(TEST_USER.id, workspace)
    return workspace


def _fake_chatbot_message(**kwargs) -> str:
    return "AI生成：这是一段测试讲解。"


def _fake_requirement_update(**kwargs) -> LearningRequirementUpdate:
    return LearningRequirementUpdate(
        progress=100,
        summary="用户已经说明当前学习目标，可以进入后续板书阶段。",
        key_facts=[
            LearningRequirementKeyFact(
                label="学习请求",
                value="用户提出了当前要解决的学习问题。",
                evidence="来自用户输入。",
                category="other",
            )
        ],
        checklist=[
            LearningRequirementChecklistItem(
                title="用户已经说明当前学习目标",
                is_clear=True,
                evidence="用户提出了当前要解决的学习问题。",
            ),
            LearningRequirementChecklistItem(
                title="后续板书可以围绕该目标组织",
                is_clear=True,
                evidence="对话已经给出可继续展开的学习方向。",
            ),
        ],
        missing_items=[],
        next_question="",
        ready_for_board=True,
    )


@pytest.fixture
def isolated_ai_log(monkeypatch: pytest.MonkeyPatch, tmp_path):
    log_path = tmp_path / "logs" / "ai-usage.jsonl"
    log_path.parent.mkdir(parents=True, exist_ok=True)
    monkeypatch.setattr(ai_usage_logger, "path", log_path)
    return log_path


def test_openai_parse_logs_prompt_and_output(isolated_ai_log) -> None:
    class _Output(BaseModel):
        title: str

    class _LessonOutput:
        id = "resp_123"
        output_text = '{"title":"工作台标题"}'
        usage = {"total_tokens": 42}

        def __init__(self) -> None:
            self.output_parsed = _Output(title="工作台标题")

    class _FakeResponses:
        def __init__(self) -> None:
            self.payload = None

        def parse(self, **kwargs):
            self.payload = kwargs
            return _LessonOutput()

    class _FakeClient:
        def __init__(self) -> None:
            self.responses = _FakeResponses()

    ai = OpenAICourseAI()
    ai.client = _FakeClient()
    ai.config.default_model = "gpt-5.3"
    ai.config.lesson_model = "gpt-5.3"
    ai.config.compat_api = "responses"

    with ai_log_context(trace_id="trace_unit", route="unit_test"), bind_text_model_selection(
        AIModelSelection(provider="openai", model="gpt-5.3")
    ):
        generated = ai._parse(
            "lesson",
            system_prompt="Return a structured title.",
            user_prompt='{"request":"unit"}',
            schema=_Output,
        )

    assert generated is not None
    entries = _read_log_entries(isolated_ai_log)
    assert len(entries) == 1
    entry = entries[0]
    assert entry["event_type"] == "openai_text_call"
    assert entry["context"]["trace_id"] == "trace_unit"
    assert entry["payload"]["model"] == "gpt-5.3"
    assert isinstance(entry["payload"]["duration_ms"], int)
    assert entry["payload"]["duration_ms"] >= 0
    assert entry["payload"]["user_prompt"]
    assert entry["payload"]["parsed_output"]["title"] == "工作台标题"


def test_openai_parse_retries_model_not_found_with_fallback(isolated_ai_log) -> None:
    class _Output(BaseModel):
        title: str

    class _Response:
        id = "resp_retry"
        output_text = '{"title":"勾股定理"}'
        usage = {"total_tokens": 21}

        def __init__(self) -> None:
            self.output_parsed = _Output(title="勾股定理")

    class _FakeResponses:
        def __init__(self) -> None:
            self.calls: list[str] = []

        def parse(self, **kwargs):
            model = kwargs["model"]
            self.calls.append(model)
            if model == "gpt-5.3":
                raise Exception(
                    "Error code: 400 - {'error': {'message': \"The requested model 'gpt-5.3' does not exist.\", "
                    "'type': 'invalid_request_error', 'param': 'model', 'code': 'model_not_found'}}"
                )
            return _Response()

    class _FakeClient:
        def __init__(self) -> None:
            self.responses = _FakeResponses()

    ai = OpenAICourseAI()
    ai.client = _FakeClient()
    ai.config.default_model = "gpt-5.3"
    ai.config.pm_model = "gpt-5.3"
    ai.config.fallback_model = "gpt-5.4"
    ai.config.compat_api = "responses"

    with ai_log_context(trace_id="trace_retry", route="unit_test"), bind_text_model_selection(
        AIModelSelection(provider="openai", model="gpt-5.3")
    ):
        result = ai._parse("pm", "system", "user", _Output)

    assert result is not None
    assert result.title == "勾股定理"
    assert ai.client.responses.calls == ["gpt-5.3", "gpt-5.4"]

    entries = _read_log_entries(isolated_ai_log)
    assert [entry["event_type"] for entry in entries] == ["openai_text_call_retry", "openai_text_call"]
    assert entries[0]["payload"]["model"] == "gpt-5.3"
    assert isinstance(entries[0]["payload"]["duration_ms"], int)
    assert entries[0]["payload"]["retry_model"] == "gpt-5.4"
    assert entries[1]["payload"]["model"] == "gpt-5.4"
    assert isinstance(entries[1]["payload"]["duration_ms"], int)
    assert entries[1]["payload"]["fallback_from_model"] == "gpt-5.3"


def test_chatbot_reply_prompt_uses_chatbot_identity(monkeypatch: pytest.MonkeyPatch) -> None:
    captured: dict[str, str] = {}
    ai = OpenAICourseAI()

    def _fake_parse(role, *, system_prompt, user_prompt, schema):
        captured["role"] = role
        captured["system_prompt"] = system_prompt
        captured["user_prompt"] = user_prompt
        return ChatbotReply(chatbot_message="你好，我是 OpenClass 的 Chatbot。")

    monkeypatch.setattr(ai, "_parse", _fake_parse)

    reply = ai.generate_chatbot_reply(
        lesson_title="测试页",
        learning_goal="先澄清学习需求。",
        board_summary="测试页",
        resource_summary="暂无已上传资料摘要",
        conversation_summary="",
        user_message="你好",
    )

    assert reply is not None
    assert captured["role"] == "chatbot"
    assert captured["system_prompt"].startswith("你是 OpenClass 的 Chatbot，")
    assert "AI Chatbot" not in captured["system_prompt"]


def test_board_document_generation_prompt_requests_substantial_default_length(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    captured: dict[str, str] = {}
    ai = OpenAICourseAI()

    def _fake_parse(role, *, system_prompt, user_prompt, schema):
        captured["role"] = role
        captured["system_prompt"] = system_prompt
        captured["user_prompt"] = user_prompt
        return BoardDocumentEditResult(
            operation="replace_document",
            title="通用主题板书",
            content_text="# 通用主题板书\n## 第一节\n正文",
            summary="生成了完整板书。",
            section_titles=["第一节"],
        )

    monkeypatch.setattr(ai, "_parse", _fake_parse)

    result = ai.generate_board_document_edit(
        intent="generate_from_requirements",
        lesson_title="测试页",
        learning_requirement_context={"summary": "用户已经给出学习目标。"},
        current_document_title="空白板书",
        current_document_text="",
        resource_summary="暂无已上传资料摘要",
        conversation_summary="",
        user_instruction="开始生成板书",
    )

    assert result is not None
    assert captured["role"] == "board"
    assert "完整文档篇幅生成" in captured["system_prompt"]
    assert "多个相互衔接的 H2 小节" in captured["system_prompt"]
    assert "足以支撑一节课直接教学" in captured["system_prompt"]
    assert "较完整篇幅展开" in captured["user_prompt"]


def test_openai_parse_falls_back_to_google_on_provider_auth_error(isolated_ai_log) -> None:
    class _Output(BaseModel):
        title: str

    class _GoogleResponse:
        id = "google_123"
        output_text = '{"title":"勾股定理"}'
        usage = {"totalTokenCount": 12}
        output_parsed = _Output(title="勾股定理")

    class _FakeOpenAIResponses:
        def parse(self, **kwargs):
            raise Exception(
                "Error code: 401 - {'error': {'message': 'Incorrect API key provided', "
                "'type': 'invalid_request_error', 'code': 'invalid_api_key'}}"
            )

    class _FakeOpenAIClient:
        def __init__(self) -> None:
            self.responses = _FakeOpenAIResponses()

    class _FakeGoogleClient:
        def __init__(self) -> None:
            self.payload = None

        def parse(self, **kwargs):
            self.payload = kwargs
            return _GoogleResponse()

    ai = OpenAICourseAI()
    ai.client = _FakeOpenAIClient()
    ai.google_client = _FakeGoogleClient()
    ai.google_config.default_model = "gemini-good"
    ai.config.compat_api = "responses"

    with bind_text_model_selection(AIModelSelection(provider="openai", model="gpt-bad")):
        result = ai._parse("pm", "system", "user", _Output)

    assert result is not None
    assert result.title == "勾股定理"
    assert ai.google_client.payload["model"] == "gemini-good"

    entries = _read_log_entries(isolated_ai_log)
    assert [entry["event_type"] for entry in entries] == ["openai_text_call_provider_retry", "google_text_call"]
    assert entries[0]["payload"]["retry_provider"] == "google"
    assert entries[0]["payload"]["retry_model"] == "gemini-good"
    assert entries[1]["payload"]["fallback_from_provider"] == "openai"
    assert entries[1]["payload"]["fallback_from_model"] == "gpt-bad"


def test_unavailable_provider_falls_back_to_google(isolated_ai_log) -> None:
    class _Output(BaseModel):
        title: str

    class _GoogleResponse:
        id = "google_456"
        output_text = '{"title":"函数"}'
        usage = {"totalTokenCount": 10}
        output_parsed = _Output(title="函数")

    class _FakeGoogleClient:
        def __init__(self) -> None:
            self.payload = None

        def parse(self, **kwargs):
            self.payload = kwargs
            return _GoogleResponse()

    ai = OpenAICourseAI()
    ai.client = None
    ai.google_client = _FakeGoogleClient()
    ai.google_config.default_model = "gemini-ready"

    with bind_text_model_selection(AIModelSelection(provider="openai", model="gpt-missing")):
        result = ai._parse("pm", "system", "user", _Output)

    assert result is not None
    assert result.title == "函数"
    assert ai.google_client.payload["model"] == "gemini-ready"

    entries = _read_log_entries(isolated_ai_log)
    assert [entry["event_type"] for entry in entries] == [
        "openai_text_call_skipped",
        "openai_text_call_provider_retry",
        "google_text_call",
    ]
    assert entries[0]["payload"]["reason"] == "client_disabled"
    assert entries[1]["payload"]["retry_provider"] == "google"
    assert entries[1]["payload"]["retry_model"] == "gemini-ready"
    assert entries[2]["payload"]["fallback_from_provider"] == "openai"
    assert entries[2]["payload"]["fallback_from_model"] == "gpt-missing"


def test_openai_compat_chat_completions_mode_parses_json(isolated_ai_log) -> None:
    class _Output(BaseModel):
        title: str

    class _Message:
        content = '{"title":"勾股定理"}'

    class _Choice:
        message = _Message()

    class _Response:
        id = "chatcmpl_123"
        choices = [_Choice()]
        usage = {"total_tokens": 12}

    class _FakeChatCompletions:
        def __init__(self) -> None:
            self.payload = None

        def create(self, **kwargs):
            self.payload = kwargs
            return _Response()

    class _FakeChat:
        def __init__(self) -> None:
            self.completions = _FakeChatCompletions()

    class _FakeClient:
        def __init__(self) -> None:
            self.chat = _FakeChat()

    ai = OpenAICourseAI()
    ai.client = _FakeClient()
    ai.config.compat_api = "chat_completions"
    ai.config.default_model = "gpt-5.4"
    ai.config.pm_model = "gpt-5.4"

    with ai_log_context(trace_id="trace_chat_compat", route="unit_test"), bind_text_model_selection(
        AIModelSelection(provider="openai", model="gpt-5.4")
    ):
        result = ai._parse("pm", "system", "user", _Output)

    assert result is not None
    assert result.title == "勾股定理"
    assert ai.client.chat.completions.payload["model"] == "gpt-5.4"
    assert ai.client.chat.completions.payload["response_format"]["type"] == "json_schema"
    entries = _read_log_entries(isolated_ai_log)
    assert entries[0]["event_type"] == "openai_text_call"
    assert entries[0]["payload"]["output_text"] == '{"title":"勾股定理"}'


def test_chat_completions_accepts_jsonish_object_response(isolated_ai_log) -> None:
    class _Output(BaseModel):
        title: str

    class _Message:
        content = '{title:"勾股定理"}'

    class _Choice:
        message = _Message()

    class _Response:
        id = "chatcmpl_jsonish"
        choices = [_Choice()]
        usage = {"total_tokens": 12}

    class _FakeChatCompletions:
        def __init__(self) -> None:
            self.calls = 0

        def create(self, **kwargs):
            self.calls += 1
            return _Response()

    class _FakeChat:
        def __init__(self) -> None:
            self.completions = _FakeChatCompletions()

    class _FakeClient:
        def __init__(self) -> None:
            self.chat = _FakeChat()

    ai = OpenAICourseAI()
    ai.client = _FakeClient()
    ai.config.compat_api = "chat_completions"
    ai.config.default_model = "gpt-5.4"
    ai.config.pm_model = "gpt-5.4"

    with bind_text_model_selection(AIModelSelection(provider="openai", model="gpt-5.4")):
        result = ai._parse("pm", "system", "user", _Output)

    assert result is not None
    assert result.title == "勾股定理"
    assert ai.client.chat.completions.calls == 1


def test_chat_completions_repairs_unparseable_structured_response(isolated_ai_log) -> None:
    class _Output(BaseModel):
        title: str

    class _Message:
        def __init__(self, content: str) -> None:
            self.content = content

    class _Choice:
        def __init__(self, content: str) -> None:
            self.message = _Message(content)

    class _Response:
        def __init__(self, content: str, response_id: str) -> None:
            self.id = response_id
            self.choices = [_Choice(content)]
            self.usage = {"total_tokens": 12}

    class _FakeChatCompletions:
        def __init__(self) -> None:
            self.payloads = []

        def create(self, **kwargs):
            self.payloads.append(kwargs)
            if len(self.payloads) == 1:
                return _Response("勾股定理", "chatcmpl_bad")
            return _Response('{"title":"勾股定理"}', "chatcmpl_repaired")

    class _FakeChat:
        def __init__(self) -> None:
            self.completions = _FakeChatCompletions()

    class _FakeClient:
        def __init__(self) -> None:
            self.chat = _FakeChat()

    ai = OpenAICourseAI()
    ai.client = _FakeClient()
    ai.config.compat_api = "chat_completions"
    ai.config.default_model = "gpt-5.4"
    ai.config.pm_model = "gpt-5.4"

    with bind_text_model_selection(AIModelSelection(provider="openai", model="gpt-5.4")):
        result = ai._parse("pm", "system", "user", _Output)

    assert result is not None
    assert result.title == "勾股定理"
    assert len(ai.client.chat.completions.payloads) == 2
    repair_messages = ai.client.chat.completions.payloads[1]["messages"]
    assert repair_messages[-2]["role"] == "assistant"
    assert repair_messages[-2]["content"] == "勾股定理"

    entries = _read_log_entries(isolated_ai_log)
    assert entries[0]["event_type"] == "openai_text_call"
    assert entries[0]["payload"]["response_id"] == "chatcmpl_repaired"
    assert entries[0]["payload"]["output_text"] == '{"title":"勾股定理"}'


def test_openai_defaults_to_official_openai_and_gpt_image_2(monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.delenv("OPENAI_API_KEY", raising=False)
    monkeypatch.delenv("OPENAI_BASE_URL", raising=False)
    monkeypatch.delenv("OPENAI_COMPAT_API", raising=False)
    monkeypatch.delenv("OPENAI_IMAGE_MODEL", raising=False)

    ai = OpenAICourseAI()

    assert ai.config.base_url == "https://api.openai.com/v1"
    assert ai.config.compat_api == "chat_completions"
    assert ai.config.image_model == "gpt-image-2"


@pytest.mark.parametrize(
    ("provider", "client_attr", "model"),
    [
        ("deepseek", "deepseek_client", "deepseek-v4-pro"),
        ("kimi", "kimi_client", "kimi-k2.6"),
        ("minimax", "minimax_client", "MiniMax-M2.7-highspeed"),
        ("openai_compatible", "openai_compatible_client", "router-model"),
    ],
)
def test_openai_compatible_style_providers_route_to_selected_client(
    provider: str, client_attr: str, model: str, isolated_ai_log
) -> None:
    class _Output(BaseModel):
        title: str

    class _Message:
        content = '{"title":"统一接口"}'

    class _Choice:
        message = _Message()

    class _Response:
        id = f"{provider}_123"
        choices = [_Choice()]
        usage = {"total_tokens": 12}

    class _FakeChatCompletions:
        def __init__(self) -> None:
            self.payload = None

        def create(self, **kwargs):
            self.payload = kwargs
            return _Response()

    class _FakeChat:
        def __init__(self) -> None:
            self.completions = _FakeChatCompletions()

    class _FakeClient:
        def __init__(self) -> None:
            self.chat = _FakeChat()

    ai = OpenAICourseAI()
    fake_client = _FakeClient()
    setattr(ai, client_attr, fake_client)

    with bind_text_model_selection(AIModelSelection(provider=provider, model=model)):  # type: ignore[arg-type]
        result = ai._parse("pm", "system", "user", _Output)

    assert result is not None
    assert result.title == "统一接口"
    assert fake_client.chat.completions.payload["model"] == model
    entries = _read_log_entries(isolated_ai_log)
    assert entries[0]["event_type"] == f"{provider}_text_call"
    assert entries[0]["payload"]["provider"] == provider
    assert entries[0]["payload"]["model"] == model


def test_anthropic_compatible_provider_routes_to_selected_client(isolated_ai_log) -> None:
    class _Output(BaseModel):
        title: str

    class _FakeAnthropicCompatibleClient:
        def __init__(self) -> None:
            self.payload = None

        def parse(self, **kwargs):
            self.payload = kwargs

            class _Response:
                id = "anthropic_compatible_123"
                output_text = '{"title":"统一接口"}'
                usage = {"input_tokens": 3, "output_tokens": 5}
                output_parsed = _Output(title="统一接口")

            return _Response()

    ai = OpenAICourseAI()
    fake_client = _FakeAnthropicCompatibleClient()
    ai.anthropic_compatible_client = fake_client

    with bind_text_model_selection(
        AIModelSelection(provider="anthropic_compatible", model="claude-router")
    ):
        result = ai._parse("pm", "system", "user", _Output)

    assert result is not None
    assert result.title == "统一接口"
    assert fake_client.payload["model"] == "claude-router"
    entries = _read_log_entries(isolated_ai_log)
    assert entries[0]["event_type"] == "anthropic_compatible_text_call"
    assert entries[0]["payload"]["provider"] == "anthropic_compatible"
    assert entries[0]["payload"]["model"] == "claude-router"


def test_catalog_role_uses_dedicated_openai_model_even_with_text_selection(isolated_ai_log) -> None:
    class _FakeResponses:
        def __init__(self) -> None:
            self.payload = None

        def parse(self, **kwargs):
            self.payload = kwargs

            class _Response:
                id = "catalog_123"
                usage = {"input_tokens": 5, "output_tokens": 8}

                def __init__(self) -> None:
                    payload = {
                        "chapters": [
                            {
                                "title": "入口",
                                "summary": "资料的起点。",
                                "keywords": ["入口"],
                                "level": 1,
                            }
                        ]
                    }
                    self.output_text = json.dumps(payload, ensure_ascii=False)
                    self.output_parsed = GeneratedResourceCatalog.model_validate(payload)

            return _Response()

    class _FakeClient:
        def __init__(self) -> None:
            self.responses = _FakeResponses()

    ai = OpenAICourseAI()
    ai.client = _FakeClient()
    ai.config.default_model = "gpt-5.5"
    ai.config.catalog_model = "gpt-5.4-mini"
    ai.config.compat_api = "responses"

    with bind_text_model_selection(AIModelSelection(provider="openai", model="gpt-5.5")):
        result = ai.generate_resource_outline(
            resource_name="material.txt",
            extracted_text="入口说明。" * 80,
        )

    assert result is not None
    assert result.chapters[0].title == "入口"
    assert ai.client.responses.payload["model"] == "gpt-5.4-mini"


def test_chat_route_returns_chatbot_reply(monkeypatch: pytest.MonkeyPatch, isolated_ai_log, tmp_path) -> None:
    store = SqliteCourseStore(tmp_path / "openclass.sqlite3", legacy_json_path=None)
    monkeypatch.setattr(workspace_state, "STORE", store)
    captured_context: dict[str, object] = {}

    def _fake_chatbot_reply(**kwargs):
        captured_context.update(current_ai_log_context())
        return ChatbotReply(chatbot_message="AI生成：这是一段测试讲解。")

    monkeypatch.setattr(openai_course_ai, "generate_chatbot_reply", _fake_chatbot_reply)
    monkeypatch.setattr(openai_course_ai, "generate_learning_requirement_update", _fake_requirement_update)

    lesson_id = _seed_test_user_workspace(store).packages[0].lessons[0].id
    response = chat_service.process_chat_on_lesson(
        lesson_id,
        ChatRequest(message="请解释一下当前主题的核心问题"),
        user_id=TEST_USER.id,
    )

    assert response.chatbot_message == "AI生成：这是一段测试讲解。"
    lesson = response.course_package.lessons[0]
    commit = lesson.history_graph.commits[-1]
    assert commit.metadata["kind"] == "chat_flow"
    assert commit.metadata["user_message"] == "请解释一下当前主题的核心问题"
    assert commit.metadata["assistant_message_source"] == "chatbot"
    assert commit.metadata["learning_clarification"]["ready_for_board"] is True
    assert commit.metadata["learning_clarification"]["key_facts"][0]["label"] == "学习请求"
    assert response.learning_clarification.progress == 100
    assert response.learning_clarification.key_facts[0].value == "用户提出了当前要解决的学习问题。"
    assert captured_context["route"] == "/api/lessons/{lesson_id}/chat"
    assert captured_context["lesson_id"] == lesson_id
    assert captured_context["user_id"] == TEST_USER.id
    assert str(captured_context["trace_id"]).startswith("chat_")
    assert response.learning_clarification.checklist[0].is_clear is True
    assert _read_log_entries(isolated_ai_log) == []


def test_chat_route_binds_requested_text_model_selection(
    monkeypatch: pytest.MonkeyPatch, isolated_ai_log, tmp_path
) -> None:
    store = SqliteCourseStore(tmp_path / "openclass.sqlite3", legacy_json_path=None)
    monkeypatch.setattr(workspace_state, "STORE", store)
    captured_models: dict[str, tuple[str, str]] = {}

    def _fake_chatbot_reply(**kwargs):
        captured_models["chatbot"] = openai_course_ai._model_for("chatbot")
        return ChatbotReply(chatbot_message="AI生成：这是一段测试讲解。")

    def _fake_requirement_update_with_model(**kwargs):
        captured_models["pm"] = openai_course_ai._model_for("pm")
        return _fake_requirement_update(**kwargs)

    monkeypatch.setattr(openai_course_ai, "generate_chatbot_reply", _fake_chatbot_reply)
    monkeypatch.setattr(openai_course_ai, "generate_learning_requirement_update", _fake_requirement_update_with_model)

    lesson_id = _seed_test_user_workspace(store).packages[0].lessons[0].id
    response = chat_service.process_chat_on_lesson(
        lesson_id,
        ChatRequest(
            message="请解释一下当前主题的核心问题",
            text_model=AIModelSelection(provider="deepseek", model="deepseek-v4-pro"),
        ),
        user_id=TEST_USER.id,
    )

    assert response.chatbot_message == "AI生成：这是一段测试讲解。"
    assert captured_models["chatbot"] == ("deepseek", "deepseek-v4-pro")
    assert captured_models["pm"] == ("deepseek", "deepseek-v4-pro")
    assert _read_log_entries(isolated_ai_log) == []


def test_chatbot_runtime_empty_reply_does_not_show_canned_template(
    monkeypatch: pytest.MonkeyPatch, isolated_ai_log, tmp_path
) -> None:
    store = SqliteCourseStore(tmp_path / "openclass.sqlite3", legacy_json_path=None)
    monkeypatch.setattr(workspace_state, "STORE", store)
    monkeypatch.setattr(openai_course_ai, "generate_chatbot_reply", lambda **kwargs: None)

    captured: dict[str, str] = {}

    def _fake_requirement_update(**kwargs):
        captured["chatbot_message"] = kwargs["chatbot_message"]
        return LearningRequirementUpdate(
            progress=45,
            summary="用户已经提出学习请求，但还需要确认下一步重点。",
            key_facts=[
                LearningRequirementKeyFact(
                    category="learning",
                    label="学习请求",
                    value="用户提出了一个学习问题。",
                    evidence="来自用户输入。",
                )
            ],
            checklist=[
                LearningRequirementChecklistItem(
                    title="用户已经给出学习请求",
                    is_clear=True,
                    evidence="来自用户输入。",
                ),
                LearningRequirementChecklistItem(
                    title="下一步重点仍需确认",
                    is_clear=False,
                    evidence="用户还没有说明希望先看例子、步骤还是应用。",
                ),
            ],
            missing_items=["下一步讲解重点"],
            next_question="你希望我先从概念、步骤，还是一个具体例子开始？",
            ready_for_board=False,
        )

    monkeypatch.setattr(openai_course_ai, "generate_learning_requirement_update", _fake_requirement_update)

    lesson_id = _seed_test_user_workspace(store).packages[0].lessons[0].id
    response = chat_service.process_chat_on_lesson(
        lesson_id,
        ChatRequest(message="我想学一个新内容"),
        user_id=TEST_USER.id,
    )

    assert captured["chatbot_message"] == ""
    assert response.chatbot_message == ""
    assert "如果你希望我继续" not in response.chatbot_message
    assert "再具体一点" not in response.chatbot_message
    commit = response.course_package.lessons[0].history_graph.commits[-1]
    assert commit.metadata["assistant_message"] == response.chatbot_message
    assert commit.metadata["assistant_message_source"] == "chatbot_empty"
    assert _read_log_entries(isolated_ai_log) == []


def test_requirement_manager_keeps_low_substance_chat_unclear(
    monkeypatch: pytest.MonkeyPatch, isolated_ai_log, tmp_path
) -> None:
    store = SqliteCourseStore(tmp_path / "openclass.sqlite3", legacy_json_path=None)
    monkeypatch.setattr(workspace_state, "STORE", store)
    monkeypatch.setattr(
        openai_course_ai,
        "generate_chatbot_reply",
        lambda **kwargs: ChatbotReply(chatbot_message="你好，我们可以先明确学习目标。"),
    )
    monkeypatch.setattr(openai_course_ai, "generate_learning_requirement_update", lambda **kwargs: None)

    lesson_id = _seed_test_user_workspace(store).packages[0].lessons[0].id
    response = chat_service.process_chat_on_lesson(
        lesson_id,
        ChatRequest(message="你好"),
        user_id=TEST_USER.id,
    )

    assert response.learning_clarification.progress < 50
    assert response.learning_clarification.ready_for_board is False
    assert response.learning_clarification.next_question
    assert response.learning_clarification.key_facts == []
    assert response.learning_clarification.checklist
    assert response.learning_clarification.checklist[0].is_clear is False
    assert response.learning_clarification.checklist[0].title == "用户具体想学什么内容或解决什么问题"
    assert response.learning_clarification.missing_items == ["具体学习内容", "当前水平", "学习目的或使用场景"]
    assert _read_log_entries(isolated_ai_log) == []


def test_requirement_manager_records_learning_content_from_first_intent(
    monkeypatch: pytest.MonkeyPatch, isolated_ai_log, tmp_path
) -> None:
    store = SqliteCourseStore(tmp_path / "openclass.sqlite3", legacy_json_path=None)
    monkeypatch.setattr(workspace_state, "STORE", store)
    monkeypatch.setattr(
        openai_course_ai,
        "generate_chatbot_reply",
        lambda **kwargs: ChatbotReply(chatbot_message="好的，我先确认你的学习目标。"),
    )
    monkeypatch.setattr(
        openai_course_ai,
        "generate_learning_requirement_update",
        lambda **kwargs: LearningRequirementUpdate(
            progress=10,
            summary="用户想学习一个主题，但还需要澄清水平和用途。",
            key_facts=[],
            checklist=[
                LearningRequirementChecklistItem(
                    title="用户具体想学什么内容或解决什么问题",
                    is_clear=False,
                    evidence="模型没有结构化提取。",
                )
            ],
            missing_items=["当前水平", "学习目的或使用场景"],
            next_question="你目前是什么水平？",
            ready_for_board=False,
        ),
    )

    lesson_id = _seed_test_user_workspace(store).packages[0].lessons[0].id
    response = chat_service.process_chat_on_lesson(
        lesson_id,
        ChatRequest(message="你好，我想学一门新的内容"),
        user_id=TEST_USER.id,
    )

    assert response.learning_clarification.key_facts[0].label == "学习内容"
    assert response.learning_clarification.key_facts[0].value == "一门新的内容"
    commit = response.course_package.lessons[0].history_graph.commits[-1]
    assert commit.metadata["learning_clarification"]["key_facts"][0]["label"] == "学习内容"
    assert commit.metadata["learning_clarification"]["key_facts"][0]["value"] == "一门新的内容"
    assert _read_log_entries(isolated_ai_log) == []


def test_requirement_theme_uses_learning_content_not_lesson_title(
    monkeypatch: pytest.MonkeyPatch, isolated_ai_log, tmp_path
) -> None:
    store = SqliteCourseStore(tmp_path / "openclass.sqlite3", legacy_json_path=None)
    monkeypatch.setattr(workspace_state, "STORE", store)
    monkeypatch.setattr(
        openai_course_ai,
        "generate_chatbot_reply",
        lambda **kwargs: ChatbotReply(chatbot_message="好的，我先按你的目标开始。"),
    )
    monkeypatch.setattr(
        openai_course_ai,
        "generate_learning_requirement_update",
        lambda **kwargs: LearningRequirementUpdate(
            progress=60,
            summary="高中生想学习反应平衡，用于预习。",
            key_facts=[
                LearningRequirementKeyFact(
                    label="学习内容",
                    value="反应平衡",
                    evidence="用户说想学习反应平衡。",
                    category="learning",
                )
            ],
            checklist=[
                LearningRequirementChecklistItem(
                    title="学习主题",
                    is_clear=True,
                    evidence="用户说明了要学习的内容。",
                )
            ],
            missing_items=["当前水平"],
            next_question="你目前是什么水平？",
            ready_for_board=False,
        ),
    )

    workspace = _seed_test_user_workspace(store)
    lesson = workspace.packages[0].lessons[0]
    lesson.title = "化学测试"
    lesson.board_document.title = "化学测试"
    store.save_for_user(TEST_USER.id, workspace)

    response = chat_service.process_chat_on_lesson(
        lesson.id,
        ChatRequest(message="我是高中生，我想学习化学中的反应平衡"),
        user_id=TEST_USER.id,
    )

    assert response.learning_requirement_sheet.theme == "反应平衡"
    assert response.course_package.lessons[0].learning_requirements.theme == "反应平衡"
    assert _read_log_entries(isolated_ai_log) == []


def test_requirement_manager_does_not_extract_learning_content_from_student_identity(
    monkeypatch: pytest.MonkeyPatch, isolated_ai_log, tmp_path
) -> None:
    store = SqliteCourseStore(tmp_path / "openclass.sqlite3", legacy_json_path=None)
    monkeypatch.setattr(workspace_state, "STORE", store)
    monkeypatch.setattr(
        openai_course_ai,
        "generate_chatbot_reply",
        lambda **kwargs: ChatbotReply(chatbot_message="好的，我先确认你的学习目标。"),
    )
    monkeypatch.setattr(
        openai_course_ai,
        "generate_learning_requirement_update",
        lambda **kwargs: LearningRequirementUpdate(
            progress=35,
            summary="用户说明了身份和学习主题。",
            key_facts=[],
            checklist=[
                LearningRequirementChecklistItem(
                    title="具体学习内容",
                    is_clear=True,
                    evidence="用户说了想学什么。",
                )
            ],
            missing_items=["学习目的或使用场景"],
            next_question="你希望用于什么场景？",
            ready_for_board=False,
        ),
    )

    lesson_id = _seed_test_user_workspace(store).packages[0].lessons[0].id
    response = chat_service.process_chat_on_lesson(
        lesson_id,
        ChatRequest(message="我是大学生，我想学一门新的内容"),
        user_id=TEST_USER.id,
    )

    assert response.learning_clarification.key_facts[0].label == "学习内容"
    assert response.learning_clarification.key_facts[0].value == "一门新的内容"
    assert _read_log_entries(isolated_ai_log) == []


def test_requirement_manager_keeps_structured_learning_fact_over_identity_phrase(
    monkeypatch: pytest.MonkeyPatch, isolated_ai_log, tmp_path
) -> None:
    store = SqliteCourseStore(tmp_path / "openclass.sqlite3", legacy_json_path=None)
    monkeypatch.setattr(workspace_state, "STORE", store)
    def _unexpected_chatbot_reply(**kwargs):
        raise AssertionError("document artifact request should write the blank document directly")

    monkeypatch.setattr(openai_course_ai, "generate_chatbot_reply", _unexpected_chatbot_reply)
    monkeypatch.setattr(
        openai_course_ai,
        "generate_learning_requirement_update",
        lambda **kwargs: LearningRequirementUpdate(
            progress=100,
            summary="用户需要一篇面向真实场景的对话课文。",
            key_facts=[
                LearningRequirementKeyFact(
                    category="learning",
                    label="学习主题",
                    value="真实场景对话课文",
                    evidence="用户要求生成一篇情景对话课文。",
                )
            ],
            checklist=[
                LearningRequirementChecklistItem(
                    title="学习主题",
                    is_clear=True,
                    evidence="用户说明了要生成的内容。",
                )
            ],
            missing_items=[],
            next_question="",
            ready_for_board=True,
        ),
    )

    def _fake_board_edit(**kwargs):
        return BoardDocumentEditResult(
            operation="replace_document",
            title="真实场景对话课文",
            content_text="# 真实场景对话课文\n## 第一节\n这是一段面向真实场景的右侧文档内容。",
            summary="已生成右侧文档。",
            section_titles=["第一节"],
        )

    monkeypatch.setattr(openai_course_ai, "generate_board_document_edit", _fake_board_edit)

    lesson_id = _seed_test_user_workspace(store).packages[0].lessons[0].id
    response = chat_service.process_chat_on_lesson(
        lesson_id,
        ChatRequest(message="我是一名学习者，请给我生成一篇真实场景对话课文"),
        user_id=TEST_USER.id,
    )

    assert response.learning_clarification.key_facts[0].label == "学习内容"
    assert response.learning_clarification.key_facts[0].value == "真实场景对话课文"
    assert response.learning_requirement_sheet.theme == "真实场景对话课文"
    assert response.learning_requirement_sheet.theme != "者"
    assert response.board_decision.action == "edit_board"
    assert "真实场景对话课文" in response.course_package.lessons[0].board_document.content_text
    assert _read_log_entries(isolated_ai_log) == []


def test_requirement_manager_direct_teaching_start_does_not_keep_clarifying(
    monkeypatch: pytest.MonkeyPatch, isolated_ai_log, tmp_path
) -> None:
    store = SqliteCourseStore(tmp_path / "openclass.sqlite3", legacy_json_path=None)
    monkeypatch.setattr(workspace_state, "STORE", store)
    monkeypatch.setattr(
        openai_course_ai,
        "generate_chatbot_reply",
        lambda **kwargs: ChatbotReply(chatbot_message="好的，我们从第一小节开始讲。"),
    )
    monkeypatch.setattr(
        openai_course_ai,
        "generate_learning_requirement_update",
        lambda **kwargs: LearningRequirementUpdate(
            progress=40,
            summary="用户已有学习主题，要求按零基础直接开始讲解。",
            key_facts=[
                LearningRequirementKeyFact(
                    label="学习内容",
                    value="一个通用主题",
                    evidence="来自前文。",
                    category="learning",
                ),
                LearningRequirementKeyFact(
                    label="当前水平",
                    value="0基础",
                    evidence="用户说按0基础。",
                    category="level",
                ),
            ],
            checklist=[
                LearningRequirementChecklistItem(
                    title="学习主题",
                    is_clear=True,
                    evidence="已有学习内容。",
                ),
                LearningRequirementChecklistItem(
                    title="当前水平",
                    is_clear=True,
                    evidence="用户要求按0基础讲。",
                ),
            ],
            missing_items=["具体子知识点偏好"],
            next_question="你想先讲哪个子知识点？",
            ready_for_board=False,
        ),
    )

    def _unexpected_board_edit(**kwargs):
        raise AssertionError("existing board content must not be overwritten from a direct teaching request")

    monkeypatch.setattr(openai_course_ai, "generate_board_document_edit", _unexpected_board_edit)

    workspace = _seed_test_user_workspace(store)
    lesson = workspace.packages[0].lessons[0]
    lesson.board_document = build_document(title="已有板书", content_text="已有内容")
    lesson.history_graph.commits[-1].snapshot = lesson.board_document
    store.save_for_user(TEST_USER.id, workspace)
    response = chat_service.process_chat_on_lesson(
        lesson.id,
        ChatRequest(message="你就当我是0基础，直接为我开始讲解"),
        user_id=TEST_USER.id,
    )

    assert response.chatbot_message == "好的，我们从第一小节开始讲。"
    assert response.learning_clarification.label == "可以开始讲解"
    assert response.learning_clarification.can_start is True
    assert response.learning_clarification.forced_start is True
    assert response.learning_clarification.ready_for_board is False
    assert response.learning_clarification.progress == 90
    assert response.learning_clarification.missing_items == []
    assert response.learning_clarification.next_question == ""
    assert response.board_decision.action == "no_change"
    assert response.course_package.lessons[0].board_document.content_text == "已有内容"
    assert _read_log_entries(isolated_ai_log) == []


def test_direct_teaching_on_blank_board_does_not_auto_generate_board(
    monkeypatch: pytest.MonkeyPatch, isolated_ai_log, tmp_path
) -> None:
    store = SqliteCourseStore(tmp_path / "openclass.sqlite3", legacy_json_path=None)
    monkeypatch.setattr(workspace_state, "STORE", store)
    monkeypatch.setattr(
        openai_course_ai,
        "generate_chatbot_reply",
        lambda **kwargs: ChatbotReply(chatbot_message="AI生成：第一小节讲解。"),
    )
    monkeypatch.setattr(
        openai_course_ai,
        "generate_learning_requirement_update",
        lambda **kwargs: LearningRequirementUpdate(
            progress=40,
            summary="用户要求直接开始讲解一个通用主题。",
            key_facts=[
                LearningRequirementKeyFact(
                    label="学习内容",
                    value="一个通用主题",
                    evidence="用户说要直接讲解。",
                    category="learning",
                )
            ],
            checklist=[
                LearningRequirementChecklistItem(
                    title="学习主题",
                    is_clear=True,
                    evidence="用户说明了要学的内容。",
                )
            ],
            missing_items=["具体子知识点偏好"],
            next_question="你想先讲哪个子知识点？",
            ready_for_board=False,
        ),
    )

    def _unexpected_board_edit(**kwargs):
        raise AssertionError("direct teaching requests must not auto-generate a board document")

    monkeypatch.setattr(openai_course_ai, "generate_board_document_edit", _unexpected_board_edit)

    workspace = _seed_test_user_workspace(store)
    lesson = workspace.packages[0].lessons[0]
    lesson.board_document = build_document(title="空白板书")
    lesson.history_graph.commits[-1].snapshot = lesson.board_document
    store.save_for_user(TEST_USER.id, workspace)

    response = chat_service.process_chat_on_lesson(
        lesson.id,
        ChatRequest(message="直接为我讲解一个通用主题"),
        user_id=TEST_USER.id,
    )

    assert response.chatbot_message == "AI生成：第一小节讲解。"
    assert response.board_decision.action == "no_change"
    assert response.requirement_cleared is False
    assert response.learning_clarification.forced_start is True
    assert response.learning_clarification.ready_for_board is False
    updated_lesson = response.course_package.lessons[0]
    assert updated_lesson.board_document.content_text == ""
    commit = updated_lesson.history_graph.commits[-1]
    assert commit.metadata["kind"] == "chat_flow"
    assert commit.metadata["requirement_cleared"] is False
    assert "auto_board_generation" not in commit.metadata
    saved_lesson = store.load_for_user(TEST_USER.id).packages[0].lessons[0]
    assert saved_lesson.learning_requirements is not None
    assert saved_lesson.board_teaching_guide is None
    assert saved_lesson.board_teaching_progress is None
    assert _read_log_entries(isolated_ai_log) == []


def test_plain_zero_basis_request_updates_requirements_without_board_generation(
    monkeypatch: pytest.MonkeyPatch, isolated_ai_log, tmp_path
) -> None:
    store = SqliteCourseStore(tmp_path / "openclass.sqlite3", legacy_json_path=None)
    monkeypatch.setattr(workspace_state, "STORE", store)
    monkeypatch.setattr(
        openai_course_ai,
        "generate_chatbot_reply",
        lambda **kwargs: ChatbotReply(chatbot_message="你想先从哪个部分开始？"),
    )
    monkeypatch.setattr(
        openai_course_ai,
        "generate_learning_requirement_update",
        lambda **kwargs: LearningRequirementUpdate(
            progress=20,
            summary="用户想学一个通用主题，当前是零基础，还需要澄清学习目的。",
            key_facts=[
                LearningRequirementKeyFact(
                    label="学习内容",
                    value="一个通用主题",
                    evidence="用户说明想学的内容。",
                    category="learning",
                ),
                LearningRequirementKeyFact(
                    label="当前水平",
                    value="零基础",
                    evidence="用户说自己是零基础。",
                    category="level",
                ),
            ],
            checklist=[
                LearningRequirementChecklistItem(
                    title="具体学习范围",
                    is_clear=False,
                    evidence="还没有说明想先学哪一部分。",
                ),
                LearningRequirementChecklistItem(
                    title="当前水平",
                    is_clear=True,
                    evidence="用户说自己是零基础。",
                ),
                LearningRequirementChecklistItem(
                    title="学习目的",
                    is_clear=False,
                    evidence="还没有说明学习目的。",
                ),
            ],
            missing_items=["学习目的"],
            next_question="你学这个主要是为了什么？",
            ready_for_board=False,
        ),
    )

    def _unexpected_board_edit(**kwargs):
        raise AssertionError("plain zero-basis requirement collection must not generate board content")

    monkeypatch.setattr(openai_course_ai, "generate_board_document_edit", _unexpected_board_edit)

    workspace = _seed_test_user_workspace(store)
    lesson = workspace.packages[0].lessons[0]
    lesson.board_document = build_document(title="空白板书")
    lesson.history_graph.commits[-1].snapshot = lesson.board_document
    store.save_for_user(TEST_USER.id, workspace)

    response = chat_service.process_chat_on_lesson(
        lesson.id,
        ChatRequest(message="我想学一个通用主题，我是零基础"),
        user_id=TEST_USER.id,
    )

    assert response.chatbot_message == "你想先从哪个部分开始？"
    assert response.board_decision.action == "no_change"
    assert response.requirement_cleared is False
    assert response.learning_clarification.progress == 20
    assert response.learning_clarification.ready_for_board is False
    assert response.learning_clarification.forced_start is False
    assert response.learning_requirement_sheet.theme == "一个通用主题"
    assert response.learning_requirement_sheet.level == "零基础"
    updated_lesson = response.course_package.lessons[0]
    assert updated_lesson.board_document.content_text == ""
    saved_lesson = store.load_for_user(TEST_USER.id).packages[0].lessons[0]
    assert saved_lesson.learning_requirements is not None
    assert saved_lesson.learning_requirements.theme == "一个通用主题"
    assert saved_lesson.learning_requirements.level == "零基础"
    commit = updated_lesson.history_graph.commits[-1]
    assert commit.metadata["requirement_cleared"] is False
    assert commit.metadata["active_requirement_sheet_after"] is not None
    assert _read_log_entries(isolated_ai_log) == []


def test_explicit_board_generation_request_generates_and_clears_requirements(
    monkeypatch: pytest.MonkeyPatch, isolated_ai_log, tmp_path
) -> None:
    store = SqliteCourseStore(tmp_path / "openclass.sqlite3", legacy_json_path=None)
    monkeypatch.setattr(workspace_state, "STORE", store)
    captured: dict[str, object] = {}
    monkeypatch.setattr(
        openai_course_ai,
        "generate_learning_requirement_update",
        lambda **kwargs: LearningRequirementUpdate(
            progress=40,
            summary="用户要求生成一份通用主题的入门板书。",
            key_facts=[
                LearningRequirementKeyFact(
                    label="学习内容",
                    value="一个通用主题",
                    evidence="用户说明要生成的内容。",
                    category="learning",
                )
            ],
            checklist=[
                LearningRequirementChecklistItem(
                    title="学习主题",
                    is_clear=True,
                    evidence="用户说明了要学的内容。",
                )
            ],
            missing_items=["学习目的"],
            next_question="你希望面向什么场景？",
            ready_for_board=False,
        ),
    )

    def _fake_board_edit(**kwargs):
        captured["intent"] = kwargs.get("intent")
        captured["user_instruction"] = kwargs.get("user_instruction")
        return BoardDocumentEditResult(
            operation="replace_document",
            title="通用主题板书",
            content_text=(
                "# 通用主题板书\n"
                "## 第一节\n"
                "这是一段足够长的第一节讲解正文，用来形成稳定的小节计划。\n"
                "## 第二节\n"
                "这是一段足够长的第二节讲解正文，用来形成稳定的小节计划。"
            ),
            summary="生成了通用主题板书。",
            section_titles=["第一节", "第二节"],
        )

    monkeypatch.setattr(openai_course_ai, "generate_board_document_edit", _fake_board_edit)

    workspace = _seed_test_user_workspace(store)
    lesson = workspace.packages[0].lessons[0]
    lesson.board_document = build_document(title="空白板书")
    lesson.history_graph.commits[-1].snapshot = lesson.board_document
    store.save_for_user(TEST_USER.id, workspace)

    response = chat_service.process_chat_on_lesson(
        lesson.id,
        ChatRequest(message="请生成一份入门板书"),
        user_id=TEST_USER.id,
    )

    assert captured == {
        "intent": "generate_from_requirements",
        "user_instruction": "请生成一份入门板书",
    }
    assert response.board_decision.action == "edit_board"
    assert response.requirement_cleared is True
    assert response.active_requirement_sheet is None
    updated_lesson = response.course_package.lessons[0]
    assert "第一节" in updated_lesson.board_document.content_text
    assert "第二节" in updated_lesson.board_document.content_text
    commit = updated_lesson.history_graph.commits[-1]
    assert commit.metadata["kind"] == "board_document_generation"
    assert commit.metadata["board_generation_action"] == "explicit_board_request"
    assert commit.metadata["requirement_cleared"] is True
    saved_lesson = store.load_for_user(TEST_USER.id).packages[0].lessons[0]
    assert saved_lesson.learning_requirements is None
    assert saved_lesson.board_teaching_guide is not None
    assert [plan.heading for plan in saved_lesson.board_teaching_guide.section_plans] == ["第一节", "第二节"]
    assert saved_lesson.board_teaching_progress is None
    assert _read_log_entries(isolated_ai_log) == []


def test_low_information_chat_does_not_auto_generate_blank_board(
    monkeypatch: pytest.MonkeyPatch, isolated_ai_log, tmp_path
) -> None:
    store = SqliteCourseStore(tmp_path / "openclass.sqlite3", legacy_json_path=None)
    monkeypatch.setattr(workspace_state, "STORE", store)
    monkeypatch.setattr(
        openai_course_ai,
        "generate_chatbot_reply",
        lambda **kwargs: ChatbotReply(chatbot_message="你好，我们可以先明确学习目标。"),
    )
    monkeypatch.setattr(
        openai_course_ai,
        "generate_learning_requirement_update",
        lambda **kwargs: LearningRequirementUpdate(
            progress=15,
            summary="用户还没有透露足够具体的学习需求。",
            key_facts=[],
            checklist=[
                LearningRequirementChecklistItem(
                    title="用户具体想学什么内容或解决什么问题",
                    is_clear=False,
                    evidence="最近对话还没有说明学习主题。",
                )
            ],
            missing_items=["具体学习内容"],
            next_question="你想围绕哪个主题、资料或具体问题开始学习？",
            ready_for_board=False,
        ),
    )

    def _unexpected_board_edit(**kwargs):
        raise AssertionError("low-information chat must not trigger board generation")

    monkeypatch.setattr(openai_course_ai, "generate_board_document_edit", _unexpected_board_edit)

    lesson_id = _seed_test_user_workspace(store).packages[0].lessons[0].id
    response = chat_service.process_chat_on_lesson(
        lesson_id,
        ChatRequest(message="你好"),
        user_id=TEST_USER.id,
    )

    updated_lesson = response.course_package.lessons[0]
    assert updated_lesson.board_document.content_text == ""
    assert response.board_decision.action == "no_change"
    assert "右侧板书" not in response.chatbot_message
    assert _read_log_entries(isolated_ai_log) == []


def test_requirement_manager_progress_does_not_drop_for_same_topic(
    monkeypatch: pytest.MonkeyPatch, isolated_ai_log, tmp_path
) -> None:
    store = SqliteCourseStore(tmp_path / "openclass.sqlite3", legacy_json_path=None)
    monkeypatch.setattr(workspace_state, "STORE", store)
    monkeypatch.setattr(
        openai_course_ai,
        "generate_chatbot_reply",
        lambda **kwargs: ChatbotReply(chatbot_message="好的，我会继续帮你推进。"),
    )
    updates = iter(
        [
            LearningRequirementUpdate(
                progress=50,
                summary="用户说明了学习主题。",
                key_facts=[
                    LearningRequirementKeyFact(
                        label="学习内容",
                        value="一个通用主题",
                        evidence="用户说明了学习主题。",
                        category="learning",
                    )
                ],
                checklist=[
                    LearningRequirementChecklistItem(title="学习主题", is_clear=True, evidence="来自用户。")
                ],
                missing_items=["学习场景"],
                next_question="你是为什么学？",
                ready_for_board=False,
            ),
            LearningRequirementUpdate(
                progress=30,
                summary="用户继续围绕同一主题补充了学习场景。",
                key_facts=[
                    LearningRequirementKeyFact(
                        label="学习内容",
                        value="一个通用主题",
                        evidence="来自前文。",
                        category="learning",
                    ),
                    LearningRequirementKeyFact(
                        label="面向场景",
                        value="预习",
                        evidence="来自用户。",
                        category="scenario",
                    ),
                ],
                checklist=[
                    LearningRequirementChecklistItem(title="学习主题", is_clear=True, evidence="来自前文。"),
                    LearningRequirementChecklistItem(title="学习场景", is_clear=True, evidence="来自用户。"),
                ],
                missing_items=["当前水平"],
                next_question="你目前是什么水平？",
                ready_for_board=False,
            ),
        ]
    )
    monkeypatch.setattr(openai_course_ai, "generate_learning_requirement_update", lambda **kwargs: next(updates))

    lesson_id = _seed_test_user_workspace(store).packages[0].lessons[0].id
    first = chat_service.process_chat_on_lesson(
        lesson_id,
        ChatRequest(message="我想学习一个通用主题"),
        user_id=TEST_USER.id,
    )
    second = chat_service.process_chat_on_lesson(
        lesson_id,
        ChatRequest(
            message="我是为了预习",
            conversation=[
                {"role": "user", "content": "我想学习一个通用主题"},
                {"role": "assistant", "content": first.chatbot_message},
            ],
        ),
        user_id=TEST_USER.id,
    )

    assert first.learning_clarification.progress == 50
    assert second.learning_clarification.progress == 50
    assert second.learning_clarification.key_facts[0].value == "一个通用主题"
    assert second.learning_clarification.key_facts[1].value == "预习"
    assert _read_log_entries(isolated_ai_log) == []


def test_requirement_manager_immediate_board_request_sets_progress_to_complete(
    monkeypatch: pytest.MonkeyPatch, isolated_ai_log, tmp_path
) -> None:
    store = SqliteCourseStore(tmp_path / "openclass.sqlite3", legacy_json_path=None)
    monkeypatch.setattr(workspace_state, "STORE", store)
    monkeypatch.setattr(
        openai_course_ai,
        "generate_chatbot_reply",
        lambda **kwargs: ChatbotReply(chatbot_message="好的，我会按当前信息准备板书。"),
    )
    monkeypatch.setattr(
        openai_course_ai,
        "generate_learning_requirement_update",
        lambda **kwargs: LearningRequirementUpdate(
            progress=35,
            summary="用户还没有补齐全部学习背景。",
            key_facts=[
                LearningRequirementKeyFact(
                    label="当前请求",
                    value="用户希望进入板书生成。",
                    evidence="来自用户输入。",
                    category="other",
                )
            ],
            checklist=[
                LearningRequirementChecklistItem(
                    title="用户具体想学什么或解决什么问题",
                    is_clear=False,
                    evidence="还需要从上下文继续确认。",
                )
            ],
            missing_items=["当前水平"],
            next_question="你目前对这个内容了解多少？",
            ready_for_board=False,
        ),
    )

    def _unexpected_board_edit(**kwargs):
        raise AssertionError("existing board content must not be overwritten from a status update")

    monkeypatch.setattr(openai_course_ai, "generate_board_document_edit", _unexpected_board_edit)

    workspace = _seed_test_user_workspace(store)
    lesson = workspace.packages[0].lessons[0]
    lesson.board_document = build_document(title="已有板书", content_text="已有内容")
    lesson.history_graph.commits[-1].snapshot = lesson.board_document
    store.save_for_user(TEST_USER.id, workspace)
    response = chat_service.process_chat_on_lesson(
        lesson.id,
        ChatRequest(message="不用再问了，直接生成版书"),
        user_id=TEST_USER.id,
    )

    assert response.learning_clarification.progress == 100
    assert response.learning_clarification.ready_for_board is True
    assert response.learning_clarification.can_start is True
    assert response.learning_clarification.forced_start is True
    assert response.learning_clarification.missing_items == []
    assert response.learning_clarification.next_question == ""
    assert response.board_decision.action == "no_change"
    assert response.course_package.lessons[0].board_document.content_text == "已有内容"
    commit = response.course_package.lessons[0].history_graph.commits[-1]
    assert commit.metadata["learning_clarification"]["progress"] == 100
    assert commit.metadata["learning_clarification"]["forced_start"] is True
    assert _read_log_entries(isolated_ai_log) == []


def test_requirement_manager_explicit_board_generation_sets_progress_to_complete(
    monkeypatch: pytest.MonkeyPatch, isolated_ai_log, tmp_path
) -> None:
    store = SqliteCourseStore(tmp_path / "openclass.sqlite3", legacy_json_path=None)
    monkeypatch.setattr(workspace_state, "STORE", store)
    monkeypatch.setattr(
        openai_course_ai,
        "generate_chatbot_reply",
        lambda **kwargs: ChatbotReply(chatbot_message="好的，我会开始讲解并准备板书。"),
    )
    monkeypatch.setattr(
        openai_course_ai,
        "generate_learning_requirement_update",
        lambda **kwargs: LearningRequirementUpdate(
            progress=45,
            summary="用户已有学习主题，希望进入讲解和板书生成。",
            key_facts=[
                LearningRequirementKeyFact(
                    label="学习内容",
                    value="一门新的内容",
                    evidence="来自前文。",
                    category="learning",
                )
            ],
            checklist=[
                LearningRequirementChecklistItem(
                    title="具体学习内容",
                    is_clear=True,
                    evidence="已有学习内容。",
                )
            ],
            missing_items=["当前水平"],
            next_question="你目前是什么水平？",
            ready_for_board=False,
        ),
    )

    def _unexpected_board_edit(**kwargs):
        raise AssertionError("existing board content must not be overwritten from a status update")

    monkeypatch.setattr(openai_course_ai, "generate_board_document_edit", _unexpected_board_edit)

    workspace = _seed_test_user_workspace(store)
    lesson = workspace.packages[0].lessons[0]
    lesson.board_document = build_document(title="已有板书", content_text="已有内容")
    lesson.history_graph.commits[-1].snapshot = lesson.board_document
    store.save_for_user(TEST_USER.id, workspace)
    response = chat_service.process_chat_on_lesson(
        lesson.id,
        ChatRequest(message="先为我讲解，生成板书"),
        user_id=TEST_USER.id,
    )

    assert response.learning_clarification.progress == 100
    assert response.learning_clarification.ready_for_board is True
    assert response.learning_clarification.forced_start is True
    assert response.learning_clarification.next_question == ""
    assert response.board_decision.action == "no_change"
    assert response.course_package.lessons[0].board_document.content_text == "已有内容"
    assert _read_log_entries(isolated_ai_log) == []


def test_requirement_manager_start_generation_request_writes_blank_document_from_context(
    monkeypatch: pytest.MonkeyPatch, isolated_ai_log, tmp_path
) -> None:
    store = SqliteCourseStore(tmp_path / "openclass.sqlite3", legacy_json_path=None)
    monkeypatch.setattr(workspace_state, "STORE", store)
    captured: dict[str, object] = {}

    def _unexpected_chatbot_reply(**kwargs):
        raise AssertionError("generation control should write the blank document instead of handoff chat")

    monkeypatch.setattr(openai_course_ai, "generate_chatbot_reply", _unexpected_chatbot_reply)
    monkeypatch.setattr(
        openai_course_ai,
        "generate_learning_requirement_update",
        lambda **kwargs: LearningRequirementUpdate(
            progress=60,
            summary="用户已有学习内容、水平和输出需求，但还没有指定全部细节。",
            key_facts=[
                LearningRequirementKeyFact(
                    label="学习内容",
                    value="一门新的内容",
                    evidence="来自前文。",
                    category="learning",
                ),
                LearningRequirementKeyFact(
                    label="当前水平",
                    value="中级",
                    evidence="来自前文。",
                    category="level",
                ),
                LearningRequirementKeyFact(
                    label="学习内容需求",
                    value="生成一份练习材料",
                    evidence="来自前文。",
                    category="output",
                ),
            ],
            checklist=[
                LearningRequirementChecklistItem(
                    title="具体学习内容",
                    is_clear=True,
                    evidence="已有学习内容。",
                )
            ],
            missing_items=["具体场景"],
            next_question="你希望面向什么场景？",
            ready_for_board=False,
        ),
    )

    def _fake_board_edit(**kwargs):
        captured["intent"] = kwargs.get("intent")
        captured["user_instruction"] = kwargs.get("user_instruction")
        return BoardDocumentEditResult(
            operation="replace_document",
            title="生成后的文档",
            content_text="# 生成后的文档\n## 第一节\n这是一段根据已有学习需求生成的右侧文档内容。",
            summary="已生成右侧文档。",
            section_titles=["第一节"],
        )

    monkeypatch.setattr(openai_course_ai, "generate_board_document_edit", _fake_board_edit)

    lesson_id = _seed_test_user_workspace(store).packages[0].lessons[0].id
    response = chat_service.process_chat_on_lesson(
        lesson_id,
        ChatRequest(message="看你发挥，开始生成"),
        user_id=TEST_USER.id,
    )

    assert captured == {
        "intent": "generate_from_requirements",
        "user_instruction": "看你发挥，开始生成",
    }
    assert response.learning_clarification.progress == 100
    assert response.learning_clarification.ready_for_board is True
    assert response.learning_clarification.can_start is True
    assert response.learning_clarification.forced_start is True
    assert response.learning_clarification.missing_items == []
    assert response.learning_clarification.next_question == ""
    assert response.board_decision.action == "edit_board"
    assert "第一节" in response.course_package.lessons[0].board_document.content_text
    assert response.requirement_cleared is True
    assert _read_log_entries(isolated_ai_log) == []


def test_generation_control_request_writes_blank_document_from_existing_context(
    monkeypatch: pytest.MonkeyPatch, isolated_ai_log, tmp_path
) -> None:
    store = SqliteCourseStore(tmp_path / "openclass.sqlite3", legacy_json_path=None)
    monkeypatch.setattr(workspace_state, "STORE", store)
    captured: dict[str, object] = {}

    def _unexpected_chatbot_reply(**kwargs):
        raise AssertionError("direct generation should not ask Chatbot to produce board-like content")

    monkeypatch.setattr(openai_course_ai, "generate_chatbot_reply", _unexpected_chatbot_reply)
    monkeypatch.setattr(
        openai_course_ai,
        "generate_learning_requirement_update",
        lambda **kwargs: LearningRequirementUpdate(
            progress=65,
            summary="用户允许系统决定未指定细节，并希望基于当前需求进入生成。",
            key_facts=[
                LearningRequirementKeyFact(
                    label="当前水平",
                    value="中级",
                    evidence="来自前文。",
                    category="level",
                ),
                LearningRequirementKeyFact(
                    label="学习需求",
                    value="生成一份练习材料",
                    evidence="来自前文。",
                    category="output",
                ),
            ],
            checklist=[
                LearningRequirementChecklistItem(
                    title="用户允许系统决定未指定细节",
                    is_clear=True,
                    evidence="用户说可以看系统发挥。",
                )
            ],
            missing_items=["具体场景"],
            next_question="你希望面向什么场景？",
            ready_for_board=False,
        ),
    )

    def _fake_board_edit(**kwargs):
        captured["intent"] = kwargs.get("intent")
        captured["user_instruction"] = kwargs.get("user_instruction")
        return BoardDocumentEditResult(
            operation="replace_document",
            title="练习材料",
            content_text="# 练习材料\n## 第一节\n这是一段根据当前需求生成的练习材料。",
            summary="已生成练习材料。",
            section_titles=["第一节"],
        )

    monkeypatch.setattr(openai_course_ai, "generate_board_document_edit", _fake_board_edit)

    workspace = _seed_test_user_workspace(store)
    lesson = workspace.packages[0].lessons[0]
    lesson.board_document = build_document(title="空白板书")
    lesson.history_graph.commits[-1].snapshot = lesson.board_document
    lesson.history_graph.commits[-1].metadata["learning_clarification"] = {
        "progress": 40,
        "label": "继续澄清",
        "reason": "用户还没有指定全部细节。",
        "missing_items": ["具体场景"],
        "can_start": False,
        "forced_start": False,
        "summary": "用户需要一篇情景对话。",
        "key_facts": [
            {"label": "学习内容", "value": "一个通用主题", "evidence": "来自用户。"},
            {"label": "当前水平", "value": "中级", "evidence": "来自用户。"},
        ],
        "checklist": [],
        "next_question": "你希望面向什么场景？",
        "ready_for_board": False,
    }
    store.save_for_user(TEST_USER.id, workspace)

    response = chat_service.process_chat_on_lesson(
        lesson.id,
        ChatRequest(message="都行，看你发挥，直接生成"),
        user_id=TEST_USER.id,
    )

    assert captured == {
        "intent": "generate_from_requirements",
        "user_instruction": "都行，看你发挥，直接生成",
    }
    assert response.board_decision.action == "edit_board"
    assert response.learning_clarification.ready_for_board is True
    assert response.learning_clarification.forced_start is True
    assert response.learning_clarification.missing_items == []
    assert response.learning_clarification.next_question == ""
    assert response.learning_requirement_sheet.level == "中级"
    updated_lesson = response.course_package.lessons[0]
    assert "练习材料" in updated_lesson.board_document.content_text
    commit = updated_lesson.history_graph.commits[-1]
    assert commit.metadata["kind"] == "board_document_generation"
    assert commit.metadata["assistant_message_source"] == "board_document_editor_ai"
    assert commit.metadata["requirement_cleared"] is True
    assert _read_log_entries(isolated_ai_log) == []


def test_document_artifact_request_writes_blank_document_without_chatbot_generation(
    monkeypatch: pytest.MonkeyPatch, isolated_ai_log, tmp_path
) -> None:
    store = SqliteCourseStore(tmp_path / "openclass.sqlite3", legacy_json_path=None)
    monkeypatch.setattr(workspace_state, "STORE", store)
    captured: dict[str, object] = {}

    def _unexpected_chatbot_reply(**kwargs):
        raise AssertionError("document artifact generation should be written by the document editor")

    monkeypatch.setattr(openai_course_ai, "generate_chatbot_reply", _unexpected_chatbot_reply)
    monkeypatch.setattr(
        openai_course_ai,
        "generate_learning_requirement_update",
        lambda **kwargs: LearningRequirementUpdate(
            progress=100,
            summary="用户希望生成一篇面向真实任务的情景对话课文。",
            key_facts=[
                LearningRequirementKeyFact(
                    category="learning",
                    label="学习内容",
                    value="面向真实任务的情景对话课文",
                    evidence="用户要求生成一篇情景对话课文。",
                ),
                LearningRequirementKeyFact(
                    category="output",
                    label="输出需求",
                    value="生成一篇课文",
                    evidence="用户要求生成一篇课文。",
                ),
            ],
            checklist=[
                LearningRequirementChecklistItem(
                    title="用户已说明要生成的内容形态",
                    is_clear=True,
                    evidence="用户要求生成一篇情景对话课文。",
                )
            ],
            missing_items=[],
            next_question="",
            ready_for_board=True,
        ),
    )

    def _fake_board_edit(**kwargs):
        captured["intent"] = kwargs.get("intent")
        captured["user_instruction"] = kwargs.get("user_instruction")
        return BoardDocumentEditResult(
            operation="replace_document",
            title="任务材料",
            content_text="# 任务材料\n## 第一节\n这是一段面向真实任务的右侧文档内容。",
            summary="已生成任务材料。",
            section_titles=["第一节"],
        )

    monkeypatch.setattr(openai_course_ai, "generate_board_document_edit", _fake_board_edit)

    workspace = _seed_test_user_workspace(store)
    lesson = workspace.packages[0].lessons[0]
    lesson.board_document = build_document(title="空白板书")
    lesson.history_graph.commits[-1].snapshot = lesson.board_document
    store.save_for_user(TEST_USER.id, workspace)

    response = chat_service.process_chat_on_lesson(
        lesson.id,
        ChatRequest(message="请给我生成一篇用于真实任务的情景对话课文"),
        user_id=TEST_USER.id,
    )

    assert captured == {
        "intent": "generate_from_requirements",
        "user_instruction": "请给我生成一篇用于真实任务的情景对话课文",
    }
    assert response.chatbot_message == "已生成任务材料。"
    assert response.board_decision.action == "edit_board"
    assert "任务材料" in response.course_package.lessons[0].board_document.content_text
    commit = response.course_package.lessons[0].history_graph.commits[-1]
    assert commit.metadata["kind"] == "board_document_generation"
    assert commit.metadata["assistant_message_source"] == "board_document_editor_ai"
    assert commit.metadata["requirement_cleared"] is True
    assert _read_log_entries(isolated_ai_log) == []


def test_accepting_defaults_without_explicit_generation_stays_in_chatbot(
    monkeypatch: pytest.MonkeyPatch, isolated_ai_log, tmp_path
) -> None:
    store = SqliteCourseStore(tmp_path / "openclass.sqlite3", legacy_json_path=None)
    monkeypatch.setattr(workspace_state, "STORE", store)
    monkeypatch.setattr(
        openai_course_ai,
        "generate_chatbot_reply",
        lambda **kwargs: ChatbotReply(chatbot_message="好的，我会继续帮你把需求讲清楚。"),
    )
    monkeypatch.setattr(
        openai_course_ai,
        "generate_learning_requirement_update",
        lambda **kwargs: LearningRequirementUpdate(
            progress=70,
            summary="用户接受系统建议，但没有明确要求生成板书。",
            key_facts=[
                LearningRequirementKeyFact(
                    category="learning",
                    label="学习内容",
                    value="一个通用主题",
                    evidence="来自前文。",
                )
            ],
            checklist=[
                LearningRequirementChecklistItem(
                    title="用户已确认学习内容",
                    is_clear=True,
                    evidence="来自前文。",
                )
            ],
            missing_items=["输出要求"],
            next_question="你希望最后得到什么形式的内容？",
            ready_for_board=False,
        ),
    )

    lesson_id = _seed_test_user_workspace(store).packages[0].lessons[0].id
    response = chat_service.process_chat_on_lesson(
        lesson_id,
        ChatRequest(message="都行，看你发挥"),
        user_id=TEST_USER.id,
    )

    assert response.chatbot_message == "好的，我会继续帮你把需求讲清楚。"
    assert response.learning_clarification.forced_start is False
    assert response.board_decision.action == "no_change"
    commit = response.course_package.lessons[0].history_graph.commits[-1]
    assert commit.metadata["assistant_message_source"] == "chatbot"
    assert _read_log_entries(isolated_ai_log) == []


def test_requirement_manager_prefers_structured_key_fact_categories(
    monkeypatch: pytest.MonkeyPatch, isolated_ai_log, tmp_path
) -> None:
    store = SqliteCourseStore(tmp_path / "openclass.sqlite3", legacy_json_path=None)
    monkeypatch.setattr(workspace_state, "STORE", store)
    monkeypatch.setattr(
        openai_course_ai,
        "generate_chatbot_reply",
        lambda **kwargs: ChatbotReply(chatbot_message="好的，我先确认这些需求。"),
    )
    monkeypatch.setattr(
        openai_course_ai,
        "generate_learning_requirement_update",
        lambda **kwargs: LearningRequirementUpdate(
            progress=100,
            summary="用户已经提供学习内容、水平、能力指标、场景和输出需求。",
            key_facts=[
                LearningRequirementKeyFact(category="learning", label="目标", value="一个通用主题", evidence="来自用户。"),
                LearningRequirementKeyFact(category="level", label="能力", value="进阶", evidence="来自用户。"),
                LearningRequirementKeyFact(category="vocabulary", label="数量", value="3500", evidence="来自用户。"),
                LearningRequirementKeyFact(category="scenario", label="用途", value="真实任务", evidence="来自用户。"),
                LearningRequirementKeyFact(category="output", label="交付", value="生成一套练习", evidence="来自用户。"),
            ],
            checklist=[
                LearningRequirementChecklistItem(
                    title="用户已经提供关键需求",
                    is_clear=True,
                    evidence="来自用户。",
                )
            ],
            missing_items=[],
            next_question="",
            ready_for_board=True,
        ),
    )

    lesson_id = _seed_test_user_workspace(store).packages[0].lessons[0].id
    response = chat_service.process_chat_on_lesson(
        lesson_id,
        ChatRequest(message="我已经说明需求"),
        user_id=TEST_USER.id,
    )

    facts = response.learning_clarification.key_facts
    assert [fact.category for fact in facts[:5]] == ["learning", "level", "vocabulary", "scenario", "output"]
    assert [fact.label for fact in facts[:5]] == ["学习内容", "当前水平", "词汇量", "面向场景", "输出需求"]
    assert response.learning_requirement_sheet.level == "进阶"
    assert response.learning_requirement_sheet.output_preference == "生成一套练习"
    assert _read_log_entries(isolated_ai_log) == []


def test_requirement_manager_filters_internal_key_fact_labels(
    monkeypatch: pytest.MonkeyPatch, isolated_ai_log, tmp_path
) -> None:
    store = SqliteCourseStore(tmp_path / "openclass.sqlite3", legacy_json_path=None)
    monkeypatch.setattr(workspace_state, "STORE", store)
    monkeypatch.setattr(
        openai_course_ai,
        "generate_chatbot_reply",
        lambda **kwargs: ChatbotReply(chatbot_message="好的，我先帮你确认学习需求。"),
    )
    monkeypatch.setattr(
        openai_course_ai,
        "generate_learning_requirement_update",
        lambda **kwargs: LearningRequirementUpdate(
            progress=70,
            summary="用户表达了一个学习目标，还需要继续澄清。",
            key_facts=[
                LearningRequirementKeyFact(
                    label="preferred_output",
                    value="想要讲义、练习、复习、对话还是项目",
                    evidence="这只是可选输出方向，不是用户已透露的信息。",
                ),
                LearningRequirementKeyFact(
                    label="学习目标",
                    value="用户想理解当前问题。",
                    evidence="来自用户输入。",
                    category="learning",
                ),
            ],
            checklist=[
                LearningRequirementChecklistItem(
                    title="用户具体想学什么内容或解决什么问题",
                    is_clear=True,
                    evidence="用户想理解当前问题。",
                )
            ],
            missing_items=["当前水平"],
            next_question="你目前对这个内容了解多少？",
            ready_for_board=False,
        ),
    )

    lesson_id = _seed_test_user_workspace(store).packages[0].lessons[0].id
    response = chat_service.process_chat_on_lesson(
        lesson_id,
        ChatRequest(message="我想理解当前问题"),
        user_id=TEST_USER.id,
    )

    assert [item.label for item in response.learning_clarification.key_facts] == ["学习内容"]
    assert response.learning_clarification.key_facts[0].value == "当前问题"
    assert _read_log_entries(isolated_ai_log) == []


def test_document_save_route_keeps_autosave_metadata(monkeypatch: pytest.MonkeyPatch, tmp_path) -> None:
    store = SqliteCourseStore(tmp_path / "openclass.sqlite3", legacy_json_path=None)
    monkeypatch.setattr(workspace_state, "STORE", store)

    workspace = _seed_test_user_workspace(store)
    lesson = workspace.packages[0].lessons[0]
    document = lesson.board_document.model_copy(deep=True)
    document.content_html = "<p>自动保存后的内容</p>"
    document.content_text = "自动保存后的内容"

    package = documents_router.save_document(
        lesson.id,
        DocumentSaveRequest(
            document=document,
            label="Auto Save",
            message="Auto-saved Word-like rich document changes from the editor",
            metadata={
                "kind": "auto_document_save",
                "autosave": True,
                "autosave_reason": "pagehide",
                "source": "word_board_editor",
            },
        ),
        user=TEST_USER,
    )

    updated_lesson = next(current for current in package.lessons if current.id == lesson.id)
    commit = updated_lesson.history_graph.commits[-1]
    assert commit.snapshot.content_text == "自动保存后的内容"
    assert commit.metadata["kind"] == "auto_document_save"
    assert commit.metadata["autosave"] is True
    assert commit.metadata["autosave_reason"] == "pagehide"


def test_stale_autosave_does_not_overwrite_newer_board_commit(
    monkeypatch: pytest.MonkeyPatch, tmp_path
) -> None:
    store = SqliteCourseStore(tmp_path / "openclass.sqlite3", legacy_json_path=None)
    monkeypatch.setattr(workspace_state, "STORE", store)

    workspace = _seed_test_user_workspace(store)
    lesson = workspace.packages[0].lessons[0]
    base_commit_id = lesson.history_graph.commits[-1].id
    stale_document = lesson.board_document.model_copy(deep=True)
    stale_document.content_html = "<p>旧草稿自动保存</p>"
    stale_document.content_text = "旧草稿自动保存"

    lesson.board_document = build_document(
        title=lesson.board_document.title,
        content_text="AI 编辑后的新版书内容",
        document_id=lesson.board_document.id,
        page_settings=lesson.board_document.page_settings,
    )
    workspace_state.commit_document_snapshot(
        lesson,
        label="Board target edit",
        message="Edited selected board content",
        metadata={"kind": "board_document_edit"},
    )
    store.save_for_user(TEST_USER.id, workspace)
    newer_head_id = lesson.history_graph.commits[-1].id

    package = documents_router.save_document(
        lesson.id,
        DocumentSaveRequest(
            document=stale_document,
            label="Auto Save",
            message="Auto-saved Word-like rich document changes from the editor",
            metadata={
                "kind": "auto_document_save",
                "autosave": True,
                "autosave_reason": "debounce",
                "source": "word_board_editor",
            },
            base_commit_id=base_commit_id,
        ),
        user=TEST_USER,
    )

    updated_lesson = next(current for current in package.lessons if current.id == lesson.id)
    assert updated_lesson.board_document.content_text == "AI 编辑后的新版书内容"
    assert updated_lesson.history_graph.commits[-1].id == newer_head_id
    assert len(updated_lesson.history_graph.commits) == 2


def test_stale_manual_document_save_is_rejected(monkeypatch: pytest.MonkeyPatch, tmp_path) -> None:
    store = SqliteCourseStore(tmp_path / "openclass.sqlite3", legacy_json_path=None)
    monkeypatch.setattr(workspace_state, "STORE", store)

    workspace = _seed_test_user_workspace(store)
    lesson = workspace.packages[0].lessons[0]
    base_commit_id = lesson.history_graph.commits[-1].id
    stale_document = lesson.board_document.model_copy(deep=True)
    stale_document.content_text = "手动保存的旧草稿"

    lesson.board_document = build_document(
        title=lesson.board_document.title,
        content_text="服务端已有更新",
        document_id=lesson.board_document.id,
        page_settings=lesson.board_document.page_settings,
    )
    workspace_state.commit_document_snapshot(
        lesson,
        label="Board target edit",
        message="Edited selected board content",
        metadata={"kind": "board_document_edit"},
    )
    store.save_for_user(TEST_USER.id, workspace)

    with pytest.raises(HTTPException) as exc_info:
        documents_router.save_document(
            lesson.id,
            DocumentSaveRequest(
                document=stale_document,
                label="Manual document edit",
                message="Saved Word-like rich document changes from the editor",
                metadata={"kind": "manual_document_save"},
                base_commit_id=base_commit_id,
            ),
            user=TEST_USER,
        )

    assert exc_info.value.status_code == 409


def test_document_save_beacon_accepts_plain_text_json(monkeypatch: pytest.MonkeyPatch, tmp_path) -> None:
    store = SqliteCourseStore(tmp_path / "openclass.sqlite3", legacy_json_path=None)
    monkeypatch.setattr(workspace_state, "STORE", store)

    workspace = _seed_test_user_workspace(store)
    lesson = workspace.packages[0].lessons[0]
    document = lesson.board_document.model_copy(deep=True)
    document.content_html = "<p>关闭页面前保存</p>"
    document.content_text = "关闭页面前保存"
    save_request = DocumentSaveRequest(
        document=document,
        label="Auto Save",
        message="Auto-saved Word-like rich document changes from the editor",
        metadata={
            "kind": "auto_document_save",
            "autosave": True,
            "autosave_reason": "pagehide",
        },
    )

    main_module.app.dependency_overrides[current_user] = lambda: TEST_USER
    try:
        response = TestClient(main_module.app).post(
            f"/api/lessons/{lesson.id}/document/save-beacon",
            content=save_request.model_dump_json(),
            headers={"content-type": "text/plain;charset=UTF-8"},
        )
    finally:
        main_module.app.dependency_overrides.pop(current_user, None)

    assert response.status_code == 200
    updated_lesson = next(current for current in response.json()["lessons"] if current["id"] == lesson.id)
    commit = updated_lesson["history_graph"]["commits"][-1]
    assert commit["snapshot"]["content_text"] == "关闭页面前保存"
    assert commit["metadata"]["autosave"] is True
    assert commit["metadata"]["autosave_reason"] == "pagehide"


def test_board_generation_from_ready_sheet_writes_empty_document_without_chat_board_content(
    monkeypatch: pytest.MonkeyPatch, isolated_ai_log, tmp_path
) -> None:
    store = SqliteCourseStore(tmp_path / "openclass.sqlite3", legacy_json_path=None)
    monkeypatch.setattr(workspace_state, "STORE", store)
    captured: dict[str, object] = {}

    def _fake_board_edit(**kwargs):
        captured["intent"] = kwargs.get("intent")
        captured["learning_requirement_context"] = kwargs.get("learning_requirement_context")
        return BoardDocumentEditResult(
            operation="replace_document",
            title="生成后的板书",
            content_text="# 生成后的板书\n## 第一节\n**讲解重点:** 这是一段足够长的第一节正文，用来避免被识别成额外章节并保持测试稳定。\n## 第二节\n这是一段足够长的第二节正文，用来避免被识别成额外章节并保持测试稳定。",
            content_html="<p>不应采用模型 HTML</p>",
            summary="生成了完整板书。",
            section_titles=["第一节", "第二节"],
        )

    monkeypatch.setattr(openai_course_ai, "generate_board_document_edit", _fake_board_edit)

    def _unexpected_chatbot_reply(**kwargs):
        raise AssertionError("board generation must not ask the chatbot to generate board-like content")

    monkeypatch.setattr(openai_course_ai, "generate_chatbot_reply", _unexpected_chatbot_reply)
    monkeypatch.setattr(openai_course_ai, "generate_learning_requirement_update", _fake_requirement_update)

    workspace = _seed_test_user_workspace(store)
    lesson = workspace.packages[0].lessons[0]
    lesson.board_document = build_document(title="空白板书")
    lesson.history_graph.commits[-1].snapshot = lesson.board_document
    lesson.history_graph.commits[-1].metadata["learning_clarification"] = {
        "progress": 100,
        "label": "需求已清晰",
        "reason": "用户已说明学习需求。",
        "missing_items": [],
        "can_start": True,
        "forced_start": False,
        "summary": "用户已说明学习需求。",
        "key_facts": [{"label": "学习内容", "value": "一个通用主题", "evidence": "来自用户。"}],
        "checklist": [{"title": "学习内容", "is_clear": True, "evidence": "来自用户。"}],
        "next_question": "",
        "ready_for_board": True,
    }
    store.save_for_user(TEST_USER.id, workspace)

    response = chat_service.process_chat_on_lesson(
        lesson.id,
        ChatRequest(message="开始生成板书", board_generation_action="start"),
        user_id=TEST_USER.id,
    )

    assert captured["intent"] == "generate_from_requirements"
    context = captured["learning_requirement_context"]
    assert isinstance(context, dict)
    assert context["key_facts"][0]["label"] == "学习内容"
    updated_lesson = response.course_package.lessons[0]
    assert "第一节" in updated_lesson.board_document.content_text
    assert "不应采用模型 HTML" not in updated_lesson.board_document.content_html
    assert "<strong>讲解重点:</strong>" in updated_lesson.board_document.content_html
    saved_lesson = store.load_for_user(TEST_USER.id).packages[0].lessons[0]
    assert saved_lesson.board_teaching_guide is not None
    assert [plan.heading for plan in saved_lesson.board_teaching_guide.section_plans] == ["第一节", "第二节"]
    assert "检查问题" in saved_lesson.board_teaching_guide.lecture_handout
    assert response.board_decision.action == "edit_board"
    assert response.active_requirement_sheet is None
    assert response.requirement_cleared is True
    assert response.teaching_progress is None
    assert response.chatbot_message == "生成了完整板书。"
    commit = updated_lesson.history_graph.commits[-1]
    assert commit.metadata["kind"] == "board_document_generation"
    assert commit.metadata["requirement_cleared"] is True
    assert commit.metadata["task_requirement_sheet"]["action_type"] == "generate_board"
    assert commit.metadata["active_requirement_sheet_after"] is None
    assert "teaching_progress" not in commit.metadata
    assert _read_log_entries(isolated_ai_log) == []


def test_board_generation_does_not_overwrite_existing_document(
    monkeypatch: pytest.MonkeyPatch, isolated_ai_log, tmp_path
) -> None:
    store = SqliteCourseStore(tmp_path / "openclass.sqlite3", legacy_json_path=None)
    monkeypatch.setattr(workspace_state, "STORE", store)

    def _unexpected_board_edit(**kwargs):
        raise AssertionError("non-empty documents must not be overwritten by board generation")

    monkeypatch.setattr(openai_course_ai, "generate_board_document_edit", _unexpected_board_edit)
    monkeypatch.setattr(openai_course_ai, "generate_learning_requirement_update", _fake_requirement_update)

    workspace = _seed_test_user_workspace(store)
    lesson = workspace.packages[0].lessons[0]
    lesson.board_document = build_document(title="已有板书", content_text="已有内容")
    lesson.history_graph.commits[-1].snapshot = lesson.board_document
    store.save_for_user(TEST_USER.id, workspace)

    response = chat_service.process_chat_on_lesson(
        lesson.id,
        ChatRequest(message="开始生成板书", board_generation_action="start"),
        user_id=TEST_USER.id,
    )

    updated_lesson = response.course_package.lessons[0]
    assert updated_lesson.board_document.content_text == "已有内容"
    assert response.board_decision.action == "no_change"
    assert response.chatbot_message == ""
    assert _read_log_entries(isolated_ai_log) == []


def test_document_ai_edit_updates_selected_board_content(
    monkeypatch: pytest.MonkeyPatch, isolated_ai_log, tmp_path
) -> None:
    store = SqliteCourseStore(tmp_path / "openclass.sqlite3", legacy_json_path=None)
    monkeypatch.setattr(workspace_state, "STORE", store)
    captured: dict[str, str | None] = {}

    def _fake_board_edit(**kwargs):
        captured["selection_excerpt"] = kwargs.get("selection_excerpt")
        captured["intent"] = kwargs.get("intent")
        return BoardDocumentEditResult(
            operation="replace_selection",
            content_text="改写后的内容",
            summary="改写了选中内容。",
        )

    def _unexpected_chatbot_reply(**kwargs):
        raise AssertionError("direct document edits should use the board document editor AI")

    monkeypatch.setattr(openai_course_ai, "generate_chatbot_reply", _unexpected_chatbot_reply)
    monkeypatch.setattr(openai_course_ai, "generate_board_document_edit", _fake_board_edit)
    monkeypatch.setattr(openai_course_ai, "generate_learning_requirement_update", _fake_requirement_update)

    workspace = _seed_test_user_workspace(store)
    lesson = workspace.packages[0].lessons[0]
    lesson.board_document = build_document(title="已有板书", content_html="<p>原文</p>")
    lesson.history_graph.commits[-1].snapshot = lesson.board_document
    store.save_for_user(TEST_USER.id, workspace)
    response = chat_service.document_ai_edit_request(
        lesson.id,
        "改写选中内容",
        "原文",
        [],
        user_id=TEST_USER.id,
    )

    assert response.chatbot_message == "改写了选中内容。"
    assert response.board_decision.action == "edit_board"
    assert response.active_requirement_sheet is None
    assert response.requirement_cleared is True
    assert response.resolved_focus is not None
    updated_lesson = response.course_package.lessons[0]
    assert updated_lesson.learning_requirements is None
    assert updated_lesson.board_document.content_text == "改写后的内容"
    commit = updated_lesson.history_graph.commits[-1]
    assert commit.metadata["kind"] == "board_document_edit"
    assert commit.metadata["board_edit_operation"] == "replace_selection"
    assert commit.metadata["requirement_cleared"] is True
    assert commit.metadata["task_requirement_sheet"]["action_type"] == "rewrite_target"
    assert commit.metadata["active_requirement_sheet_after"] is None
    assert captured == {"selection_excerpt": "原文", "intent": "edit_existing_document"}
    assert _read_log_entries(isolated_ai_log) == []


def test_document_ai_edit_uses_rich_markdown_context_and_keeps_html_marks(
    monkeypatch: pytest.MonkeyPatch, isolated_ai_log, tmp_path
) -> None:
    store = SqliteCourseStore(tmp_path / "openclass.sqlite3", legacy_json_path=None)
    monkeypatch.setattr(workspace_state, "STORE", store)
    captured: dict[str, str | None] = {}

    def _fake_board_edit(**kwargs):
        captured["current_document_text"] = kwargs.get("current_document_text")
        captured["selection_excerpt"] = kwargs.get("selection_excerpt")
        return BoardDocumentEditResult(
            operation="replace_selection",
            content_text="Speaker A: Simpler target.",
            content_html="<p><strong>Speaker A:</strong> Simpler target.</p>",
            summary="改写了选中内容。",
        )

    monkeypatch.setattr(openai_course_ai, "generate_board_document_edit", _fake_board_edit)
    monkeypatch.setattr(openai_course_ai, "generate_learning_requirement_update", _fake_requirement_update)

    workspace = _seed_test_user_workspace(store)
    lesson = workspace.packages[0].lessons[0]
    lesson.board_document = build_document(
        title="已有板书",
        content_html=(
            "<h2>Dialogue</h2>"
            "<p><strong>Speaker A:</strong> Original target.</p>"
            "<ul><li><strong>Goal:</strong> Keep structure</li></ul>"
        ),
    )
    lesson.history_graph.commits[-1].snapshot = lesson.board_document
    store.save_for_user(TEST_USER.id, workspace)

    response = chat_service.document_ai_edit_request(
        lesson.id,
        "改写选中内容",
        "Speaker A: Original target.",
        [],
        user_id=TEST_USER.id,
    )

    updated_document = response.course_package.lessons[0].board_document
    assert captured["selection_excerpt"] == "Speaker A: Original target."
    assert "## Dialogue" in (captured["current_document_text"] or "")
    assert "**Speaker A:** Original target." in (captured["current_document_text"] or "")
    assert "<h2>Dialogue</h2>" in updated_document.content_html
    assert "<strong>Speaker A:</strong> Simpler target." in updated_document.content_html
    assert "<strong>Goal:</strong> Keep structure" in updated_document.content_html
    assert _read_log_entries(isolated_ai_log) == []


def test_selected_simplify_request_in_chat_routes_to_board_editor(
    monkeypatch: pytest.MonkeyPatch, isolated_ai_log, tmp_path
) -> None:
    store = SqliteCourseStore(tmp_path / "openclass.sqlite3", legacy_json_path=None)
    monkeypatch.setattr(workspace_state, "STORE", store)
    captured: dict[str, str | None] = {}

    def _fake_board_edit(**kwargs):
        captured["selection_excerpt"] = kwargs.get("selection_excerpt")
        captured["intent"] = kwargs.get("intent")
        captured["user_instruction"] = kwargs.get("user_instruction")
        return BoardDocumentEditResult(
            operation="replace_selection",
            content_text="更简单的说法",
            summary="把选中内容改得更简单。",
            chatbot_message="AI生成：选中内容已经简化。",
        )

    def _unexpected_chatbot_reply(**kwargs):
        raise AssertionError("selected simplify requests must use the board document editor AI")

    monkeypatch.setattr(openai_course_ai, "generate_chatbot_reply", _unexpected_chatbot_reply)
    monkeypatch.setattr(openai_course_ai, "generate_board_document_edit", _fake_board_edit)
    monkeypatch.setattr(
        openai_course_ai,
        "generate_learning_requirement_update",
        lambda **kwargs: LearningRequirementUpdate(
            progress=100,
            summary="用户要求简化选中的板书内容。",
            key_facts=[
                LearningRequirementKeyFact(
                    label="局部编辑",
                    value="简化选中内容",
                    evidence="用户说把这里改的简单点。",
                    category="other",
                )
            ],
            checklist=[
                LearningRequirementChecklistItem(
                    title="目标位置来自选区",
                    is_clear=True,
                    evidence="请求携带选中内容。",
                )
            ],
            missing_items=[],
            next_question="",
            ready_for_board=True,
            action_type="simplify_target",
            action_instruction="把选中内容改得更简单。",
            target_hint="选中内容",
        ),
    )

    workspace = _seed_test_user_workspace(store)
    lesson = workspace.packages[0].lessons[0]
    lesson.board_document = build_document(
        title="已有板书",
        content_text="# 主线\n## 第一节\n原文复杂句子\n## 第二节\n其他内容",
    )
    lesson.history_graph.commits[-1].snapshot = lesson.board_document
    store.save_for_user(TEST_USER.id, workspace)

    response = chat_service.process_chat_on_lesson(
        lesson.id,
        ChatRequest(
            message="把这里改的简单点",
            selection=SelectionRef(kind="board", excerpt="原文复杂句子", lesson_id=lesson.id),
        ),
        user_id=TEST_USER.id,
    )

    assert response.chatbot_message == "AI生成：选中内容已经简化。"
    assert response.board_decision.action == "edit_board"
    assert response.resolved_focus is not None
    assert response.requirement_cleared is True
    updated_lesson = response.course_package.lessons[0]
    assert "更简单的说法" in updated_lesson.board_document.content_text
    assert "原文复杂句子" not in updated_lesson.board_document.content_text
    commit = updated_lesson.history_graph.commits[-1]
    assert commit.metadata["kind"] == "board_document_edit"
    assert commit.metadata["task_requirement_sheet"]["action_type"] == "simplify_target"
    assert commit.metadata["assistant_message_source"] == "board_document_editor_ai"
    assert captured == {
        "selection_excerpt": "原文复杂句子",
        "intent": "edit_existing_document",
        "user_instruction": "把这里改的简单点",
    }
    assert _read_log_entries(isolated_ai_log) == []


def test_numbered_blank_edit_routes_to_board_editor(
    monkeypatch: pytest.MonkeyPatch, isolated_ai_log, tmp_path
) -> None:
    store = SqliteCourseStore(tmp_path / "openclass.sqlite3", legacy_json_path=None)
    monkeypatch.setattr(workspace_state, "STORE", store)
    captured: dict[str, str | None] = {}

    def _fake_board_edit(**kwargs):
        captured["selection_excerpt"] = kwargs.get("selection_excerpt")
        captured["intent"] = kwargs.get("intent")
        captured["user_instruction"] = kwargs.get("user_instruction")
        return BoardDocumentEditResult(
            operation="replace_selection",
            content_text="You should list all your tasks and then (3)______ which responsibilities deserve immediate attention.",
            summary="已调整编号空格。",
            chatbot_message="AI生成：第3个空已经调整。",
        )

    def _unexpected_chatbot_reply(**kwargs):
        raise AssertionError("resolved numbered edits should use the board document editor AI")

    monkeypatch.setattr(openai_course_ai, "generate_chatbot_reply", _unexpected_chatbot_reply)
    monkeypatch.setattr(openai_course_ai, "generate_board_document_edit", _fake_board_edit)
    monkeypatch.setattr(openai_course_ai, "generate_learning_requirement_update", _fake_requirement_update)

    workspace = _seed_test_user_workspace(store)
    lesson = workspace.packages[0].lessons[0]
    lesson.board_document = build_document(
        title="已有板书",
        content_text=(
            "# 练习\n"
            "Many learners make a plan first. You should list all your tasks and then (3)______ "
            "which ones are most important. Focus on those first.\n"
            "3. A. decide  B. discuss  C. discover  D. differ"
        ),
    )
    lesson.history_graph.commits[-1].snapshot = lesson.board_document
    store.save_for_user(TEST_USER.id, workspace)

    response = chat_service.process_chat_on_lesson(
        lesson.id,
        ChatRequest(message="把第三个空改得更难一些"),
        user_id=TEST_USER.id,
    )

    assert response.chatbot_message == "AI生成：第3个空已经调整。"
    assert response.board_decision.action == "edit_board"
    assert response.resolved_focus is not None
    assert "(3)______ which ones are most important" in response.resolved_focus.excerpt
    assert captured == {
        "selection_excerpt": "You should list all your tasks and then (3)______ which ones are most important.",
        "intent": "edit_existing_document",
        "user_instruction": "把第三个空改得更难一些",
    }
    updated_lesson = response.course_package.lessons[0]
    assert "which responsibilities deserve immediate attention" in updated_lesson.board_document.content_text
    commit = updated_lesson.history_graph.commits[-1]
    assert commit.metadata["kind"] == "board_document_edit"
    assert commit.metadata["task_requirement_sheet"]["action_type"] == "rewrite_target"
    assert commit.metadata["requirement_cleared"] is True
    assert _read_log_entries(isolated_ai_log) == []


def test_ambiguous_numbered_edit_asks_for_focus_confirmation(
    monkeypatch: pytest.MonkeyPatch, isolated_ai_log, tmp_path
) -> None:
    store = SqliteCourseStore(tmp_path / "openclass.sqlite3", legacy_json_path=None)
    monkeypatch.setattr(workspace_state, "STORE", store)

    def _fake_chatbot_reply(**kwargs):
        return ChatbotReply(chatbot_message="请确认要修改哪一个编号位置。")

    def _unexpected_board_edit(**kwargs):
        raise AssertionError("ambiguous numbered edits must not modify the board")

    monkeypatch.setattr(openai_course_ai, "generate_chatbot_reply", _fake_chatbot_reply)
    monkeypatch.setattr(openai_course_ai, "generate_board_document_edit", _unexpected_board_edit)
    monkeypatch.setattr(openai_course_ai, "generate_learning_requirement_update", _fake_requirement_update)

    workspace = _seed_test_user_workspace(store)
    lesson = workspace.packages[0].lessons[0]
    lesson.board_document = build_document(
        title="已有板书",
        content_text="# 第一组\n3）第一组目标\n# 第二组\n3）第二组目标",
    )
    lesson.history_graph.commits[-1].snapshot = lesson.board_document
    store.save_for_user(TEST_USER.id, workspace)

    response = chat_service.process_chat_on_lesson(
        lesson.id,
        ChatRequest(message="修改第3题"),
        user_id=TEST_USER.id,
    )

    assert response.chatbot_message == "请确认要修改哪一个编号位置。"
    assert response.board_decision.action == "await_focus_choice"
    assert len(response.focus_candidates) == 2
    assert response.course_package.lessons[0].board_document.content_text == lesson.board_document.content_text
    commit = response.course_package.lessons[0].history_graph.commits[-1]
    assert commit.metadata["kind"] == "chat_flow"
    assert commit.metadata["focus_candidates"]
    assert commit.metadata["requirement_cleared"] is False
    assert _read_log_entries(isolated_ai_log) == []


def test_numbered_target_explanation_uses_structured_focus(
    monkeypatch: pytest.MonkeyPatch, isolated_ai_log, tmp_path
) -> None:
    store = SqliteCourseStore(tmp_path / "openclass.sqlite3", legacy_json_path=None)
    monkeypatch.setattr(workspace_state, "STORE", store)
    captured: dict[str, str | None] = {}

    def _fake_chatbot_reply(**kwargs):
        captured["selection_excerpt"] = kwargs.get("selection_excerpt")
        return ChatbotReply(chatbot_message="AI生成：这是第3题讲解。")

    def _unexpected_board_edit(**kwargs):
        raise AssertionError("numbered explanations must not edit the board")

    monkeypatch.setattr(openai_course_ai, "generate_chatbot_reply", _fake_chatbot_reply)
    monkeypatch.setattr(openai_course_ai, "generate_board_document_edit", _unexpected_board_edit)
    monkeypatch.setattr(openai_course_ai, "generate_learning_requirement_update", _fake_requirement_update)

    workspace = _seed_test_user_workspace(store)
    lesson = workspace.packages[0].lessons[0]
    lesson.board_document = build_document(
        title="已有板书",
        content_text="# 练习\n1）第一题内容\n2）第二题内容\n3）第三题内容",
    )
    lesson.history_graph.commits[-1].snapshot = lesson.board_document
    store.save_for_user(TEST_USER.id, workspace)

    response = chat_service.process_chat_on_lesson(
        lesson.id,
        ChatRequest(message="讲解第3题"),
        user_id=TEST_USER.id,
    )

    assert response.chatbot_message == "AI生成：这是第3题讲解。"
    assert response.board_decision.action == "no_change"
    assert response.resolved_focus is not None
    assert response.resolved_focus.excerpt == "3）第三题内容"
    assert "3）第三题内容" in (captured["selection_excerpt"] or "")
    commit = response.course_package.lessons[0].history_graph.commits[-1]
    assert commit.metadata["kind"] == "chat_flow"
    assert commit.metadata["task_requirement_sheet"]["action_type"] == "explain_target"
    assert commit.metadata["requirement_cleared"] is True
    assert _read_log_entries(isolated_ai_log) == []


def test_targeted_explanation_uses_resolved_board_focus_and_clears_task_sheet(
    monkeypatch: pytest.MonkeyPatch, isolated_ai_log, tmp_path
) -> None:
    store = SqliteCourseStore(tmp_path / "openclass.sqlite3", legacy_json_path=None)
    monkeypatch.setattr(workspace_state, "STORE", store)
    captured: dict[str, str | None] = {}

    def _fake_chatbot_reply(**kwargs):
        captured["selection_excerpt"] = kwargs.get("selection_excerpt")
        return ChatbotReply(chatbot_message="AI生成：针对目标文段的讲解。")

    def _unexpected_board_edit(**kwargs):
        raise AssertionError("targeted explanation must not edit the board")

    monkeypatch.setattr(openai_course_ai, "generate_chatbot_reply", _fake_chatbot_reply)
    monkeypatch.setattr(openai_course_ai, "generate_board_document_edit", _unexpected_board_edit)
    monkeypatch.setattr(openai_course_ai, "generate_learning_requirement_update", _fake_requirement_update)

    workspace = _seed_test_user_workspace(store)
    lesson = workspace.packages[0].lessons[0]
    lesson.board_document = build_document(
        title="已有板书",
        content_text="# 主线\n## 形成机制\n这里说明影响因素和形成过程。\n## 示例\n这里给出一个例子。",
    )
    lesson.history_graph.commits[-1].snapshot = lesson.board_document
    store.save_for_user(TEST_USER.id, workspace)

    response = chat_service.process_chat_on_lesson(
        lesson.id,
        ChatRequest(message="帮我讲一下为什么会这样"),
        user_id=TEST_USER.id,
    )

    assert response.chatbot_message == "AI生成：针对目标文段的讲解。"
    assert response.board_decision.action == "no_change"
    assert response.resolved_focus is not None
    assert response.requirement_cleared is True
    assert response.course_package.lessons[0].learning_requirements is None
    assert "目标文段" in (captured["selection_excerpt"] or "")
    commit = response.course_package.lessons[0].history_graph.commits[-1]
    assert commit.metadata["kind"] == "chat_flow"
    assert commit.metadata["requirement_cleared"] is True
    assert commit.metadata["task_requirement_sheet"]["action_type"] == "explain_target"
    assert commit.metadata["active_requirement_sheet_after"] is None
    assert _read_log_entries(isolated_ai_log) == []


def test_complex_request_prompts_for_strong_reasoning_before_solving(
    monkeypatch: pytest.MonkeyPatch, isolated_ai_log, tmp_path
) -> None:
    store = SqliteCourseStore(tmp_path / "openclass.sqlite3", legacy_json_path=None)
    monkeypatch.setattr(workspace_state, "STORE", store)
    monkeypatch.setattr(openai_course_ai, "client", object())
    captured: dict[str, str] = {}

    def _fake_chatbot_reply(**kwargs):
        captured["user_message"] = kwargs["user_message"]
        return ChatbotReply(chatbot_message="AI生成：建议先确认深度推理。")

    def _unexpected_solver(**kwargs):
        raise AssertionError("strong reasoning must wait for explicit confirmation")

    monkeypatch.setattr(openai_course_ai, "generate_chatbot_reply", _fake_chatbot_reply)
    monkeypatch.setattr(openai_course_ai, "generate_learning_requirement_update", _fake_requirement_update)
    monkeypatch.setattr(openai_course_ai, "solve_complex_problem", _unexpected_solver)

    workspace = _seed_test_user_workspace(store)
    lesson = workspace.packages[0].lessons[0]
    store.save_for_user(TEST_USER.id, workspace)

    response = chat_service.process_chat_on_lesson(
        lesson.id,
        ChatRequest(message="这是一道复杂难题，请完整推导并严谨证明"),
        user_id=TEST_USER.id,
    )

    assert response.chatbot_message == "AI生成：建议先确认深度推理。"
    assert response.strong_reasoning_prompt is not None
    assert response.strong_reasoning_prompt.confirm_label == "确认推理"
    assert "请不要解题" in captured["user_message"]
    commit = response.course_package.lessons[0].history_graph.commits[-1]
    assert commit.label == "Strong reasoning prompt"
    assert commit.metadata["strong_reasoning_prompt"]["confirm_label"] == "确认推理"
    assert commit.metadata["strong_reasoning_action"] is None
    assert "strong_reasoning_tool" not in commit.metadata
    assert _read_log_entries(isolated_ai_log) == []


def test_confirmed_strong_reasoning_uses_solver_then_chatbot(
    monkeypatch: pytest.MonkeyPatch, isolated_ai_log, tmp_path
) -> None:
    store = SqliteCourseStore(tmp_path / "openclass.sqlite3", legacy_json_path=None)
    monkeypatch.setattr(workspace_state, "STORE", store)
    monkeypatch.setattr(openai_course_ai, "client", object())
    captured: dict[str, str] = {}

    def _fake_solver(**kwargs):
        captured["target_excerpt"] = kwargs["target_excerpt"]
        return ComplexProblemSolution(
            summary="强推理摘要",
            answer="强推理可转述答案",
            confidence="high",
            limits="",
            model="gpt-5.5",
            reasoning_effort="high",
        )

    def _fake_chatbot_reply(**kwargs):
        captured["chatbot_user_message"] = kwargs["user_message"]
        return ChatbotReply(chatbot_message="AI生成：这是深度推理后的讲解。")

    monkeypatch.setattr(openai_course_ai, "solve_complex_problem", _fake_solver)
    monkeypatch.setattr(openai_course_ai, "generate_chatbot_reply", _fake_chatbot_reply)
    monkeypatch.setattr(openai_course_ai, "generate_learning_requirement_update", _fake_requirement_update)

    workspace = _seed_test_user_workspace(store)
    lesson = workspace.packages[0].lessons[0]
    lesson.board_document = build_document(
        title="已有板书",
        content_text="# 题目\n这是需要多步推导的题目原文。",
    )
    lesson.history_graph.commits[-1].snapshot = lesson.board_document
    store.save_for_user(TEST_USER.id, workspace)

    response = chat_service.process_chat_on_lesson(
        lesson.id,
        ChatRequest(
            message="请完整推导这道复杂难题",
            selection=SelectionRef(kind="board", excerpt="这是需要多步推导的题目原文。", lesson_id=lesson.id),
            strong_reasoning_action="confirm",
        ),
        user_id=TEST_USER.id,
    )

    assert response.chatbot_message == "AI生成：这是深度推理后的讲解。"
    assert response.strong_reasoning_prompt is None
    assert "题目原文" in captured["target_excerpt"]
    assert "强推理可转述答案" in captured["chatbot_user_message"]
    commit = response.course_package.lessons[0].history_graph.commits[-1]
    assert commit.metadata["strong_reasoning_action"] == "confirm"
    assert commit.metadata["strong_reasoning_tool"]["model"] == "gpt-5.5"
    assert commit.metadata["strong_reasoning_tool"]["confidence"] == "high"
    assert _read_log_entries(isolated_ai_log) == []


def test_skipped_strong_reasoning_answers_without_solver(
    monkeypatch: pytest.MonkeyPatch, isolated_ai_log, tmp_path
) -> None:
    store = SqliteCourseStore(tmp_path / "openclass.sqlite3", legacy_json_path=None)
    monkeypatch.setattr(workspace_state, "STORE", store)
    monkeypatch.setattr(openai_course_ai, "client", object())

    def _unexpected_solver(**kwargs):
        raise AssertionError("skip must not call the strong reasoning solver")

    monkeypatch.setattr(openai_course_ai, "solve_complex_problem", _unexpected_solver)
    monkeypatch.setattr(
        openai_course_ai,
        "generate_chatbot_reply",
        lambda **kwargs: ChatbotReply(chatbot_message="AI生成：普通讲解。"),
    )
    monkeypatch.setattr(openai_course_ai, "generate_learning_requirement_update", _fake_requirement_update)

    workspace = _seed_test_user_workspace(store)
    lesson = workspace.packages[0].lessons[0]
    store.save_for_user(TEST_USER.id, workspace)

    response = chat_service.process_chat_on_lesson(
        lesson.id,
        ChatRequest(message="这是一道复杂难题，请完整推导", strong_reasoning_action="skip"),
        user_id=TEST_USER.id,
    )

    assert response.chatbot_message == "AI生成：普通讲解。"
    assert response.strong_reasoning_prompt is None
    commit = response.course_package.lessons[0].history_graph.commits[-1]
    assert commit.metadata["strong_reasoning_action"] == "skip"
    assert "strong_reasoning_tool" not in commit.metadata
    assert _read_log_entries(isolated_ai_log) == []


def test_numbered_target_explanation_skips_requirement_update(
    monkeypatch: pytest.MonkeyPatch, isolated_ai_log, tmp_path
) -> None:
    store = SqliteCourseStore(tmp_path / "openclass.sqlite3", legacy_json_path=None)
    monkeypatch.setattr(workspace_state, "STORE", store)
    captured: dict[str, str | None] = {}

    def _fake_chatbot_reply(**kwargs):
        captured["selection_excerpt"] = kwargs.get("selection_excerpt")
        return ChatbotReply(chatbot_message="AI生成：这是第四节讲解。")

    def _unexpected_requirement_update(**kwargs):
        raise AssertionError("numbered target explanation should not run the requirement updater")

    monkeypatch.setattr(openai_course_ai, "generate_chatbot_reply", _fake_chatbot_reply)
    monkeypatch.setattr(openai_course_ai, "generate_learning_requirement_update", _unexpected_requirement_update)

    workspace = _seed_test_user_workspace(store)
    lesson = workspace.packages[0].lessons[0]
    lesson.board_document = build_document(
        title="已有板书",
        content_text=(
            "# 主线\n"
            "## 1. 起点\n第一节正文。\n"
            "## 2. 推进\n第二节正文。\n"
            "## 3. 例子\n第三节正文。\n"
            "## 4. 检查问题\n第四节正文。"
        ),
    )
    lesson.history_graph.commits[-1].snapshot = lesson.board_document
    store.save_for_user(TEST_USER.id, workspace)

    response = chat_service.process_chat_on_lesson(
        lesson.id,
        ChatRequest(message="为我讲解第4节"),
        user_id=TEST_USER.id,
    )

    assert response.chatbot_message == "AI生成：这是第四节讲解。"
    assert response.resolved_focus is not None
    assert response.resolved_focus.excerpt == "4. 检查问题"
    assert "4. 检查问题" in (captured["selection_excerpt"] or "")
    commit = response.course_package.lessons[0].history_graph.commits[-1]
    assert commit.metadata["kind"] == "chat_flow"
    assert commit.metadata["task_requirement_sheet"]["action_type"] == "explain_target"
    assert commit.metadata["requirement_cleared"] is True
    assert _read_log_entries(isolated_ai_log) == []


def test_append_section_request_writes_to_existing_board_without_requirement_update(
    monkeypatch: pytest.MonkeyPatch, isolated_ai_log, tmp_path
) -> None:
    store = SqliteCourseStore(tmp_path / "openclass.sqlite3", legacy_json_path=None)
    monkeypatch.setattr(workspace_state, "STORE", store)
    captured: dict[str, str | None] = {}

    def _fake_board_edit(**kwargs):
        captured["intent"] = kwargs.get("intent")
        captured["selection_excerpt"] = kwargs.get("selection_excerpt")
        captured["user_instruction"] = kwargs.get("user_instruction")
        return BoardDocumentEditResult(
            operation="replace_document",
            content_text="## 新增章节\n这是追加到末尾的新内容。",
            summary="已在右侧板书末尾续写。",
            chatbot_message="AI生成：已续写到右侧板书。",
            section_titles=["新增章节"],
        )

    def _unexpected_requirement_update(**kwargs):
        raise AssertionError("append requests should not run the requirement updater before writing")

    def _unexpected_chatbot_reply(**kwargs):
        raise AssertionError("append requests should use the board document editor AI")

    monkeypatch.setattr(openai_course_ai, "generate_board_document_edit", _fake_board_edit)
    monkeypatch.setattr(openai_course_ai, "generate_learning_requirement_update", _unexpected_requirement_update)
    monkeypatch.setattr(openai_course_ai, "generate_chatbot_reply", _unexpected_chatbot_reply)

    workspace = _seed_test_user_workspace(store)
    lesson = workspace.packages[0].lessons[0]
    lesson.board_document = build_document(title="已有板书", content_text="# 已有板书\n## 第一节\n已有内容。")
    lesson.history_graph.commits[-1].snapshot = lesson.board_document
    store.save_for_user(TEST_USER.id, workspace)

    response = chat_service.process_chat_on_lesson(
        lesson.id,
        ChatRequest(message="在右侧续写板书"),
        user_id=TEST_USER.id,
    )

    updated_lesson = response.course_package.lessons[0]
    assert response.chatbot_message == "AI生成：已续写到右侧板书。"
    assert response.board_decision.action == "edit_board"
    assert response.requirement_cleared is True
    assert updated_lesson.learning_requirements is None
    assert "已有内容" in updated_lesson.board_document.content_text
    assert "这是追加到末尾的新内容" in updated_lesson.board_document.content_text
    commit = updated_lesson.history_graph.commits[-1]
    assert commit.metadata["kind"] == "board_document_edit"
    assert commit.metadata["board_edit_operation"] == "append_section"
    assert commit.metadata["task_requirement_sheet"]["action_type"] == "append_section"
    assert commit.metadata["active_requirement_sheet_after"] is None
    assert captured == {
        "intent": "edit_existing_document",
        "selection_excerpt": None,
        "user_instruction": "在右侧续写板书",
    }
    assert _read_log_entries(isolated_ai_log) == []


def test_followup_write_executes_existing_append_requirement(
    monkeypatch: pytest.MonkeyPatch, isolated_ai_log, tmp_path
) -> None:
    store = SqliteCourseStore(tmp_path / "openclass.sqlite3", legacy_json_path=None)
    monkeypatch.setattr(workspace_state, "STORE", store)

    def _fake_board_edit(**kwargs):
        return BoardDocumentEditResult(
            operation="append_section",
            content_text="## 继续内容\n按前文需求追加的内容。",
            summary="已按前文需求续写。",
            chatbot_message="AI生成：已按前文需求续写。",
            section_titles=["继续内容"],
        )

    def _unexpected_requirement_update(**kwargs):
        raise AssertionError("follow-up execution should reuse the active task sheet")

    monkeypatch.setattr(openai_course_ai, "generate_board_document_edit", _fake_board_edit)
    monkeypatch.setattr(openai_course_ai, "generate_learning_requirement_update", _unexpected_requirement_update)

    workspace = _seed_test_user_workspace(store)
    lesson = workspace.packages[0].lessons[0]
    lesson.board_document = build_document(title="已有板书", content_text="# 已有板书\n## 第一节\n已有内容。")
    requirements = build_requirements(lesson.title)
    requirements.action_type = "expand_target"
    requirements.action_instruction = "在右侧续写板书"
    requirements.learning_goal = "用户希望基于当前文档继续写后续章节。"
    lesson.learning_requirements = requirements
    lesson.history_graph.commits[-1].snapshot = lesson.board_document
    store.save_for_user(TEST_USER.id, workspace)

    response = chat_service.process_chat_on_lesson(
        lesson.id,
        ChatRequest(message="写啊"),
        user_id=TEST_USER.id,
    )

    updated_lesson = response.course_package.lessons[0]
    assert response.chatbot_message == "AI生成：已按前文需求续写。"
    assert "按前文需求追加的内容" in updated_lesson.board_document.content_text
    commit = updated_lesson.history_graph.commits[-1]
    assert commit.metadata["board_edit_operation"] == "append_section"
    assert commit.metadata["task_requirement_sheet"]["action_type"] == "append_section"
    assert _read_log_entries(isolated_ai_log) == []


def test_rule_based_interaction_start_creates_session_and_clears_task_sheet(
    monkeypatch: pytest.MonkeyPatch, isolated_ai_log, tmp_path
) -> None:
    store = SqliteCourseStore(tmp_path / "openclass.sqlite3", legacy_json_path=None)
    monkeypatch.setattr(workspace_state, "STORE", store)
    captured_contexts: list[dict | None] = []

    def _fake_chatbot_reply(**kwargs):
        captured_contexts.append(kwargs.get("interaction_context"))
        return ChatbotReply(chatbot_message="AI生成：已按你的规则开始互动。")

    monkeypatch.setattr(openai_course_ai, "generate_chatbot_reply", _fake_chatbot_reply)
    monkeypatch.setattr(
        openai_course_ai,
        "generate_learning_requirement_update",
        lambda **kwargs: LearningRequirementUpdate(
            progress=100,
            summary="用户要求按自定义规则和原文互动。",
            key_facts=[],
            checklist=[
                LearningRequirementChecklistItem(
                    title="用户给出互动规则",
                    is_clear=True,
                    evidence="来自用户输入。",
                )
            ],
            missing_items=[],
            next_question="",
            ready_for_board=True,
            interaction_rule_draft=InteractionRuleDraft(
                should_start=True,
                rule_text="按用户指定的规则参考原文逐轮互动。",
                interaction_goal="围绕选中原文进行规则互动。",
                target_hint="选中内容",
                expected_user_behavior="用户每轮按规则给出输入。",
                assistant_behavior="Chatbot 每轮参考规则和原文回应。",
                reference_instruction="优先参考选中原文。",
            ),
        ),
    )

    workspace = _seed_test_user_workspace(store)
    lesson = workspace.packages[0].lessons[0]
    lesson.board_document = build_document(
        title="已有板书",
        content_text="# 原文\n## 第一段\n目标原文内容\n## 第二段\n其他内容",
    )
    lesson.history_graph.commits[-1].snapshot = lesson.board_document
    store.save_for_user(TEST_USER.id, workspace)

    response = chat_service.process_chat_on_lesson(
        lesson.id,
        ChatRequest(
            message="我们按这个规则和选中内容互动",
            selection=SelectionRef(kind="board", excerpt="目标原文内容", lesson_id=lesson.id),
        ),
        user_id=TEST_USER.id,
    )

    assert response.requirement_cleared is True
    assert response.active_requirement_sheet is None
    assert response.active_interaction_session is not None
    assert response.active_interaction_session.rule_text == "按用户指定的规则参考原文逐轮互动。"
    assert response.active_interaction_session.target_focus is not None
    assert "目标原文内容" in response.active_interaction_session.reference_context
    assert captured_contexts[-1] is not None
    assert captured_contexts[-1]["rule_text"] == "按用户指定的规则参考原文逐轮互动。"
    updated_lesson = response.course_package.lessons[0]
    assert updated_lesson.learning_requirements is None
    assert updated_lesson.active_interaction_session is not None
    commit = updated_lesson.history_graph.commits[-1]
    assert commit.metadata["kind"] == "interaction_flow"
    assert commit.metadata["assistant_message"] == "AI生成：已按你的规则开始互动。"
    assert commit.metadata["assistant_message_source"] == "chatbot_interaction"
    assert commit.metadata["requirement_cleared"] is True
    assert commit.metadata["active_requirement_sheet_after"] is None
    assert commit.metadata["active_interaction_session_after"]["status"] == "active"
    assert _read_log_entries(isolated_ai_log) == []


def test_rule_based_interaction_start_uses_whole_document_for_broad_reference(
    monkeypatch: pytest.MonkeyPatch, isolated_ai_log, tmp_path
) -> None:
    store = SqliteCourseStore(tmp_path / "openclass.sqlite3", legacy_json_path=None)
    monkeypatch.setattr(workspace_state, "STORE", store)
    captured_contexts: list[dict | None] = []

    def _fake_chatbot_reply(**kwargs):
        captured_contexts.append(kwargs.get("interaction_context"))
        if kwargs.get("interaction_context"):
            return ChatbotReply(chatbot_message="AI生成：按角色规则回应下一句。")
        return ChatbotReply(chatbot_message="AI生成：普通聊天先给出下一句。")

    monkeypatch.setattr(openai_course_ai, "generate_chatbot_reply", _fake_chatbot_reply)
    monkeypatch.setattr(
        openai_course_ai,
        "generate_learning_requirement_update",
        lambda **kwargs: LearningRequirementUpdate(
            progress=100,
            summary="用户要求基于已有对话逐句轮流练习。",
            key_facts=[],
            checklist=[
                LearningRequirementChecklistItem(
                    title="用户给出互动规则",
                    is_clear=True,
                    evidence="来自用户输入。",
                )
            ],
            missing_items=[],
            next_question="",
            ready_for_board=True,
            interaction_rule_draft=InteractionRuleDraft(
                should_start=True,
                rule_text="按用户指定的角色分工逐句轮流练习。",
                interaction_goal="通过一人一句的方式练习已有对话。",
                target_hint="已有对话全文",
                expected_user_behavior="用户输入自己角色的下一句。",
                assistant_behavior="Chatbot 输入另一个角色的下一句。",
                reference_instruction="依据当前已有对话推进。",
            ),
        ),
    )

    workspace = _seed_test_user_workspace(store)
    lesson = workspace.packages[0].lessons[0]
    lesson.board_document = build_document(
        title="已有板书",
        content_text=(
            "# 情景对话\n"
            "A: Bonjour, je vous appelle au sujet de l'annonce.\n"
            "B: Bonjour. Oui, elle est encore disponible.\n"
        ),
    )
    lesson.history_graph.commits[-1].snapshot = lesson.board_document
    store.save_for_user(TEST_USER.id, workspace)

    response = chat_service.process_chat_on_lesson(
        lesson.id,
        ChatRequest(message="我一句你一句轮流练习这篇对话"),
        user_id=TEST_USER.id,
    )

    assert response.chatbot_message == "AI生成：普通聊天先给出下一句。"
    assert response.focus_candidates == []
    assert response.requirement_cleared is True
    assert response.active_interaction_session is not None
    assert response.active_interaction_session.target_focus is None
    assert "Bonjour. Oui, elle est encore disponible." in response.active_interaction_session.reference_context
    assert captured_contexts == [None]
    updated_lesson = response.course_package.lessons[0]
    commit = updated_lesson.history_graph.commits[-1]
    assert commit.label == "Interaction session start"
    assert commit.metadata["assistant_message"] == "AI生成：普通聊天先给出下一句。"
    assert commit.metadata["assistant_message_source"] == "chatbot"
    assert commit.metadata["focus_candidates"] == []
    assert commit.metadata["active_interaction_session_after"]["target_focus"] is None
    assert _read_log_entries(isolated_ai_log) == []


def test_active_rule_interaction_continues_with_rule_context(
    monkeypatch: pytest.MonkeyPatch, isolated_ai_log, tmp_path
) -> None:
    store = SqliteCourseStore(tmp_path / "openclass.sqlite3", legacy_json_path=None)
    monkeypatch.setattr(workspace_state, "STORE", store)
    captured: dict[str, object | None] = {}

    monkeypatch.setattr(
        openai_course_ai,
        "generate_interaction_turn_decision",
        lambda **kwargs: InteractionTurnDecision(
            route="continue_rule",
            reason="用户输入仍在当前互动规则内。",
            progress_note="已经完成第一轮互动。",
            user_intent="继续互动",
        ),
    )

    def _fake_chatbot_reply(**kwargs):
        captured["interaction_context"] = kwargs.get("interaction_context")
        return ChatbotReply(chatbot_message="AI生成：按规则继续。")

    monkeypatch.setattr(openai_course_ai, "generate_chatbot_reply", _fake_chatbot_reply)

    workspace = _seed_test_user_workspace(store)
    lesson = workspace.packages[0].lessons[0]
    lesson.active_interaction_session = InteractionSession(
        status="active",
        rule_text="按用户指定规则逐轮互动。",
        interaction_goal="完成一段规则互动。",
        reference_context="目标原文内容",
        expected_user_behavior="用户按规则输入。",
        assistant_behavior="Chatbot 按规则回应。",
        turn_count=1,
    )
    store.save_for_user(TEST_USER.id, workspace)

    response = chat_service.process_chat_on_lesson(
        lesson.id,
        ChatRequest(message="这是我的下一轮输入"),
        user_id=TEST_USER.id,
    )

    assert response.chatbot_message == "AI生成：按规则继续。"
    assert response.interaction_decision is not None
    assert response.interaction_decision.route == "continue_rule"
    assert response.active_interaction_session is not None
    assert response.active_interaction_session.status == "active"
    assert response.active_interaction_session.turn_count == 2
    assert response.active_interaction_session.progress_note == "已经完成第一轮互动。"
    context = captured["interaction_context"]
    assert isinstance(context, dict)
    assert context["rule_text"] == "按用户指定规则逐轮互动。"
    assert context["turn_decision"]["route"] == "continue_rule"
    commit = response.course_package.lessons[0].history_graph.commits[-1]
    assert commit.metadata["kind"] == "interaction_flow"
    assert commit.metadata["interaction_decision"]["route"] == "continue_rule"
    assert commit.metadata["active_interaction_session_after"]["turn_count"] == 2
    assert _read_log_entries(isolated_ai_log) == []


def test_rule_violation_keeps_session_active(
    monkeypatch: pytest.MonkeyPatch, isolated_ai_log, tmp_path
) -> None:
    store = SqliteCourseStore(tmp_path / "openclass.sqlite3", legacy_json_path=None)
    monkeypatch.setattr(workspace_state, "STORE", store)
    monkeypatch.setattr(
        openai_course_ai,
        "generate_interaction_turn_decision",
        lambda **kwargs: InteractionTurnDecision(
            route="rule_violation",
            reason="用户输入不符合当前互动规则。",
            progress_note="等待用户修正本轮输入。",
            user_intent="规则内输入偏离",
        ),
    )
    monkeypatch.setattr(
        openai_course_ai,
        "generate_chatbot_reply",
        lambda **kwargs: ChatbotReply(chatbot_message="AI生成：这是规则内纠错。"),
    )

    workspace = _seed_test_user_workspace(store)
    lesson = workspace.packages[0].lessons[0]
    lesson.active_interaction_session = InteractionSession(
        status="active",
        rule_text="按用户指定规则逐轮互动。",
        interaction_goal="完成一段规则互动。",
        reference_context="目标原文内容",
        expected_user_behavior="用户按规则输入。",
        assistant_behavior="Chatbot 按规则回应。",
    )
    store.save_for_user(TEST_USER.id, workspace)

    response = chat_service.process_chat_on_lesson(
        lesson.id,
        ChatRequest(message="偏离规则的输入"),
        user_id=TEST_USER.id,
    )

    assert response.interaction_decision is not None
    assert response.interaction_decision.route == "rule_violation"
    assert response.active_interaction_session is not None
    assert response.active_interaction_session.status == "active"
    assert response.active_interaction_session.turn_count == 1
    assert response.course_package.lessons[0].active_interaction_session is not None
    assert _read_log_entries(isolated_ai_log) == []


def test_side_learning_request_pauses_interaction_session(
    monkeypatch: pytest.MonkeyPatch, isolated_ai_log, tmp_path
) -> None:
    store = SqliteCourseStore(tmp_path / "openclass.sqlite3", legacy_json_path=None)
    monkeypatch.setattr(workspace_state, "STORE", store)
    monkeypatch.setattr(
        openai_course_ai,
        "generate_interaction_turn_decision",
        lambda **kwargs: InteractionTurnDecision(
            route="side_learning_request",
            reason="用户临时询问原文内容。",
            progress_note="互动暂停在当前轮。",
            user_intent="临时讲解",
        ),
    )
    monkeypatch.setattr(
        openai_course_ai,
        "generate_chatbot_reply",
        lambda **kwargs: ChatbotReply(chatbot_message="AI生成：这是临时讲解。"),
    )

    workspace = _seed_test_user_workspace(store)
    lesson = workspace.packages[0].lessons[0]
    lesson.active_interaction_session = InteractionSession(
        status="active",
        rule_text="按用户指定规则逐轮互动。",
        interaction_goal="完成一段规则互动。",
        reference_context="目标原文内容",
        expected_user_behavior="用户按规则输入。",
        assistant_behavior="Chatbot 按规则回应。",
    )
    store.save_for_user(TEST_USER.id, workspace)

    response = chat_service.process_chat_on_lesson(
        lesson.id,
        ChatRequest(message="先解释一下这里的一个词"),
        user_id=TEST_USER.id,
    )

    assert response.interaction_decision is not None
    assert response.interaction_decision.route == "side_learning_request"
    assert response.active_interaction_session is not None
    assert response.active_interaction_session.status == "paused"
    assert response.active_interaction_session.pause_reason == "用户临时询问原文内容。"
    assert response.course_package.lessons[0].active_interaction_session is not None
    assert _read_log_entries(isolated_ai_log) == []


def test_paused_interaction_can_resume_and_exit(
    monkeypatch: pytest.MonkeyPatch, isolated_ai_log, tmp_path
) -> None:
    store = SqliteCourseStore(tmp_path / "openclass.sqlite3", legacy_json_path=None)
    monkeypatch.setattr(workspace_state, "STORE", store)
    decisions = iter(
        [
            InteractionTurnDecision(
                route="resume_rule",
                reason="用户要恢复互动。",
                progress_note="恢复到暂停前进度。",
                user_intent="恢复互动",
            ),
            InteractionTurnDecision(
                route="exit_rule",
                reason="用户结束互动。",
                progress_note="互动结束。",
                user_intent="结束互动",
            ),
        ]
    )
    monkeypatch.setattr(openai_course_ai, "generate_interaction_turn_decision", lambda **kwargs: next(decisions))
    monkeypatch.setattr(
        openai_course_ai,
        "generate_chatbot_reply",
        lambda **kwargs: ChatbotReply(chatbot_message="AI生成：互动状态已处理。"),
    )

    workspace = _seed_test_user_workspace(store)
    lesson = workspace.packages[0].lessons[0]
    lesson.active_interaction_session = InteractionSession(
        status="paused",
        rule_text="按用户指定规则逐轮互动。",
        interaction_goal="完成一段规则互动。",
        reference_context="目标原文内容",
        expected_user_behavior="用户按规则输入。",
        assistant_behavior="Chatbot 按规则回应。",
        pause_reason="临时讲解",
        turn_count=2,
    )
    store.save_for_user(TEST_USER.id, workspace)

    resumed = chat_service.process_chat_on_lesson(
        lesson.id,
        ChatRequest(message="继续刚才的互动"),
        user_id=TEST_USER.id,
    )
    assert resumed.interaction_decision is not None
    assert resumed.interaction_decision.route == "resume_rule"
    assert resumed.active_interaction_session is not None
    assert resumed.active_interaction_session.status == "active"
    assert resumed.active_interaction_session.turn_count == 3

    exited = chat_service.process_chat_on_lesson(
        lesson.id,
        ChatRequest(message="结束这个互动"),
        user_id=TEST_USER.id,
    )
    assert exited.interaction_decision is not None
    assert exited.interaction_decision.route == "exit_rule"
    assert exited.active_interaction_session is None
    assert exited.course_package.lessons[0].active_interaction_session is None
    commit = exited.course_package.lessons[0].history_graph.commits[-1]
    assert commit.metadata["active_interaction_session_after"] is None
    assert _read_log_entries(isolated_ai_log) == []


def test_chatbot_uses_confirmed_uploaded_resource_context_without_editing_board(
    monkeypatch: pytest.MonkeyPatch, isolated_ai_log, tmp_path
) -> None:
    store = SqliteCourseStore(tmp_path / "openclass.sqlite3", legacy_json_path=None)
    monkeypatch.setattr(workspace_state, "STORE", store)
    captured: dict[str, str | None] = {}

    def _fake_chatbot_reply(**kwargs):
        captured["selection_excerpt"] = kwargs.get("selection_excerpt")
        captured["resource_summary"] = kwargs.get("resource_summary")
        return ChatbotReply(chatbot_message="AI生成：参考上传资料讲解第一章。")

    def _unexpected_board_edit(**kwargs):
        raise AssertionError("resource-backed explanation must not edit the board")

    monkeypatch.setattr(openai_course_ai, "generate_chatbot_reply", _fake_chatbot_reply)
    monkeypatch.setattr(openai_course_ai, "generate_board_document_edit", _unexpected_board_edit)
    monkeypatch.setattr(openai_course_ai, "generate_learning_requirement_update", _fake_requirement_update)

    workspace = _seed_test_user_workspace(store)
    package = workspace.packages[0]
    lesson = package.lessons[0]
    lesson.board_document = build_document(title="已有板书", content_text="已有板书内容")
    lesson.history_graph.commits[-1].snapshot = lesson.board_document
    resource_path = tmp_path / "resource.md"
    resource_path.write_text(
        "# 第一章\n这是上传资料第一章正文，说明核心概念和学习主线。",
        encoding="utf-8",
    )
    resource = build_resource_item(resource_path, "resource.md")
    resource.scope_lesson_id = lesson.id
    package.resources.append(resource)
    store.save_for_user(TEST_USER.id, workspace)

    response = chat_service.process_chat_on_lesson(
        lesson.id,
        ChatRequest(message="根据上传资料讲一下第一章"),
        user_id=TEST_USER.id,
    )

    assert response.chatbot_message == "AI生成：参考上传资料讲解第一章。"
    assert response.board_decision.action == "no_change"
    assert response.selected_reference is not None
    assert response.selected_reference.chapter_title == "第一章"
    assert response.resource_matches
    assert "参考资料" in (captured["selection_excerpt"] or "")
    assert "上传资料第一章正文" in (captured["selection_excerpt"] or "")
    assert response.course_package.lessons[0].board_document.content_text == "已有板书内容"
    commit = response.course_package.lessons[0].history_graph.commits[-1]
    assert commit.metadata["kind"] == "chat_flow"
    assert commit.metadata["selected_reference"]["chapter_title"] == "第一章"
    assert _read_log_entries(isolated_ai_log) == []


def test_chat_imports_uploaded_resource_text_into_empty_board_without_pm_update(
    monkeypatch: pytest.MonkeyPatch, isolated_ai_log, tmp_path
) -> None:
    store = SqliteCourseStore(tmp_path / "openclass.sqlite3", legacy_json_path=None)
    monkeypatch.setattr(workspace_state, "STORE", store)
    captured: dict[str, str | None] = {}

    def _fake_chatbot_reply(**kwargs):
        captured["user_message"] = kwargs.get("user_message")
        captured["selection_excerpt"] = kwargs.get("selection_excerpt")
        return ChatbotReply(chatbot_message="AI生成：已把上传文件原文放进文档框。")

    def _unexpected_requirement_update(**kwargs):
        raise AssertionError("resource document import must not update learning requirements")

    def _unexpected_board_edit(**kwargs):
        raise AssertionError("resource document import must not use AI board generation")

    monkeypatch.setattr(openai_course_ai, "generate_chatbot_reply", _fake_chatbot_reply)
    monkeypatch.setattr(openai_course_ai, "generate_learning_requirement_update", _unexpected_requirement_update)
    monkeypatch.setattr(openai_course_ai, "generate_board_document_edit", _unexpected_board_edit)

    workspace = _seed_test_user_workspace(store)
    package = workspace.packages[0]
    lesson = package.lessons[0]
    resource_path = tmp_path / "resource.md"
    resource_path.write_text(
        "# 第一章\n上传资料第一章全文。\n\n## 第二章\n上传资料第二章全文。",
        encoding="utf-8",
    )
    resource = build_resource_item(resource_path, "resource.md")
    resource.scope_lesson_id = lesson.id
    package.resources.append(resource)
    store.save_for_user(TEST_USER.id, workspace)

    response = chat_service.process_chat_on_lesson(
        lesson.id,
        ChatRequest(message="将我上传的文件全部显示在文档框中"),
        user_id=TEST_USER.id,
    )

    imported_text = response.course_package.lessons[0].board_document.content_text
    assert response.chatbot_message == "AI生成：已把上传文件原文放进文档框。"
    assert response.board_decision.action == "edit_board"
    assert response.learning_clarification.ready_for_board is True
    assert "上传资料第一章全文" in imported_text
    assert "上传资料第二章全文" in imported_text
    assert "资料导入黑板" in (captured["user_message"] or "")
    assert "上传资料第一章全文" in (captured["selection_excerpt"] or "")
    commit = response.course_package.lessons[0].history_graph.commits[-1]
    assert commit.metadata["kind"] == "board_document_import"
    assert commit.metadata["import_status"] == "imported"
    assert commit.metadata["resource_name"] == "resource.md"
    assert commit.metadata["resource_import_scope"] == "full_resource"
    assert commit.metadata["board_edit_operation"] == "replace_document"
    assert _read_log_entries(isolated_ai_log) == []


def test_chat_imports_uploaded_docx_with_rich_styles_into_empty_board(
    monkeypatch: pytest.MonkeyPatch, isolated_ai_log, tmp_path
) -> None:
    store = SqliteCourseStore(tmp_path / "openclass.sqlite3", legacy_json_path=None)
    monkeypatch.setattr(workspace_state, "STORE", store)

    def _fake_chatbot_reply(**kwargs):
        return ChatbotReply(chatbot_message="AI生成：已保留格式导入。")

    def _unexpected_requirement_update(**kwargs):
        raise AssertionError("resource document import must not update learning requirements")

    monkeypatch.setattr(openai_course_ai, "generate_chatbot_reply", _fake_chatbot_reply)
    monkeypatch.setattr(openai_course_ai, "generate_learning_requirement_update", _unexpected_requirement_update)

    workspace = _seed_test_user_workspace(store)
    package = workspace.packages[0]
    lesson = package.lessons[0]
    resource_path = tmp_path / "styled.docx"
    docx = DocxDocument()
    docx.add_heading("彩色标题", level=1)
    paragraph = docx.add_paragraph()
    run = paragraph.add_run("加粗红色大字")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(194, 65, 12)
    docx.save(resource_path)
    resource = build_resource_item(resource_path, "styled.docx")
    resource.scope_lesson_id = lesson.id
    package.resources.append(resource)
    store.save_for_user(TEST_USER.id, workspace)

    response = chat_service.process_chat_on_lesson(
        lesson.id,
        ChatRequest(message="将我上传的文件全部显示在文档框中"),
        user_id=TEST_USER.id,
    )

    document = response.course_package.lessons[0].board_document
    assert response.board_decision.action == "edit_board"
    assert "<h1>彩色标题</h1>" in document.content_html
    assert "<strong>加粗红色大字</strong>" in document.content_html
    assert "font-size: 18pt" in document.content_html
    assert "color: #C2410C" in document.content_html
    heading = document.content_json["content"][0]
    styled_paragraph = document.content_json["content"][1]
    styled_text = styled_paragraph["content"][0]
    assert heading["type"] == "heading"
    assert heading["attrs"]["level"] == 1
    assert any(mark["type"] == "bold" for mark in styled_text["marks"])
    text_style = next(mark for mark in styled_text["marks"] if mark["type"] == "textStyle")
    assert text_style["attrs"]["fontSize"] == "18pt"
    assert text_style["attrs"]["color"] == "#C2410C"
    commit = response.course_package.lessons[0].history_graph.commits[-1]
    assert commit.metadata["kind"] == "board_document_import"
    assert commit.metadata["resource_import_scope"] == "full_resource"
    assert _read_log_entries(isolated_ai_log) == []


def test_chat_resource_visibility_question_does_not_import_uploaded_resource(
    monkeypatch: pytest.MonkeyPatch, isolated_ai_log, tmp_path
) -> None:
    store = SqliteCourseStore(tmp_path / "openclass.sqlite3", legacy_json_path=None)
    monkeypatch.setattr(workspace_state, "STORE", store)

    def _fake_chatbot_reply(**kwargs):
        return ChatbotReply(chatbot_message="AI生成：我可以读到你上传的资料。")

    def _unexpected_board_edit(**kwargs):
        raise AssertionError("visibility check must not edit the board")

    monkeypatch.setattr(openai_course_ai, "generate_chatbot_reply", _fake_chatbot_reply)
    monkeypatch.setattr(openai_course_ai, "generate_learning_requirement_update", _fake_requirement_update)
    monkeypatch.setattr(openai_course_ai, "generate_board_document_edit", _unexpected_board_edit)

    workspace = _seed_test_user_workspace(store)
    package = workspace.packages[0]
    lesson = package.lessons[0]
    resource_path = tmp_path / "resource.md"
    resource_path.write_text("# 第一章\n上传资料正文。", encoding="utf-8")
    resource = build_resource_item(resource_path, "resource.md")
    resource.scope_lesson_id = lesson.id
    package.resources.append(resource)
    store.save_for_user(TEST_USER.id, workspace)

    response = chat_service.process_chat_on_lesson(
        lesson.id,
        ChatRequest(message="你能看见我上传的文件吗？"),
        user_id=TEST_USER.id,
    )

    assert response.chatbot_message == "AI生成：我可以读到你上传的资料。"
    assert response.board_decision.action == "no_change"
    assert response.course_package.lessons[0].board_document.content_text == ""
    commit = response.course_package.lessons[0].history_graph.commits[-1]
    assert commit.metadata["kind"] == "chat_flow"
    assert _read_log_entries(isolated_ai_log) == []


def test_chat_asks_write_mode_before_importing_resource_into_non_empty_board(
    monkeypatch: pytest.MonkeyPatch, isolated_ai_log, tmp_path
) -> None:
    store = SqliteCourseStore(tmp_path / "openclass.sqlite3", legacy_json_path=None)
    monkeypatch.setattr(workspace_state, "STORE", store)
    captured: dict[str, str | None] = {}

    def _fake_chatbot_reply(**kwargs):
        captured["user_message"] = kwargs.get("user_message")
        return ChatbotReply(chatbot_message="AI生成：当前文档已有内容，请说明追加还是替换。")

    def _unexpected_requirement_update(**kwargs):
        raise AssertionError("resource document import must not update learning requirements")

    def _unexpected_board_edit(**kwargs):
        raise AssertionError("ambiguous import into non-empty board must not use AI board generation")

    monkeypatch.setattr(openai_course_ai, "generate_chatbot_reply", _fake_chatbot_reply)
    monkeypatch.setattr(openai_course_ai, "generate_learning_requirement_update", _unexpected_requirement_update)
    monkeypatch.setattr(openai_course_ai, "generate_board_document_edit", _unexpected_board_edit)

    workspace = _seed_test_user_workspace(store)
    package = workspace.packages[0]
    lesson = package.lessons[0]
    lesson.board_document = build_document(title="已有板书", content_text="已有黑板内容")
    lesson.history_graph.commits[-1].snapshot = lesson.board_document
    resource_path = tmp_path / "resource.md"
    resource_path.write_text("# 第一章\n上传资料正文。", encoding="utf-8")
    resource = build_resource_item(resource_path, "resource.md")
    resource.scope_lesson_id = lesson.id
    package.resources.append(resource)
    store.save_for_user(TEST_USER.id, workspace)

    response = chat_service.process_chat_on_lesson(
        lesson.id,
        ChatRequest(message="将上传文件显示在文档框中"),
        user_id=TEST_USER.id,
    )

    assert response.chatbot_message == "AI生成：当前文档已有内容，请说明追加还是替换。"
    assert response.board_decision.action == "no_change"
    assert response.course_package.lessons[0].board_document.content_text == "已有黑板内容"
    assert "需要用户说明是追加" in (captured["user_message"] or "")
    commit = response.course_package.lessons[0].history_graph.commits[-1]
    assert commit.metadata["kind"] == "board_document_import"
    assert commit.metadata["import_status"] == "await_write_mode"
    assert _read_log_entries(isolated_ai_log) == []


def test_learning_request_generates_board_after_resource_reference_confirmation(
    monkeypatch: pytest.MonkeyPatch, isolated_ai_log, tmp_path
) -> None:
    store = SqliteCourseStore(tmp_path / "openclass.sqlite3", legacy_json_path=None)
    monkeypatch.setattr(workspace_state, "STORE", store)
    captured: dict[str, str | None] = {}

    def _fake_board_edit(**kwargs):
        captured["resource_summary"] = kwargs.get("resource_summary")
        captured["intent"] = kwargs.get("intent")
        return BoardDocumentEditResult(
            operation="replace_document",
            title="公式板书",
            content_text="# 公式板书\n## 核心概念\n根据资料生成的板书内容。",
            summary="已参考上传资料生成公式板书。",
            chatbot_message="AI生成：已参考上传资料生成公式板书。",
            section_titles=["核心概念"],
        )

    monkeypatch.setattr(
        openai_course_ai,
        "generate_chatbot_reply",
        lambda **kwargs: ChatbotReply(chatbot_message="确认参考资料。"),
    )
    monkeypatch.setattr(openai_course_ai, "generate_board_document_edit", _fake_board_edit)
    monkeypatch.setattr(openai_course_ai, "generate_learning_requirement_update", _fake_requirement_update)

    workspace = _seed_test_user_workspace(store)
    package = workspace.packages[0]
    lesson = package.lessons[0]
    resource_path = tmp_path / "resource.md"
    resource_path.write_text(
        "# 定积分\n这一节先说明面积问题。\n\n牛顿莱布尼茨公式连接原函数与定积分，可以辅助生成板书建议。",
        encoding="utf-8",
    )
    resource = build_resource_item(resource_path, "resource.md")
    resource.scope_lesson_id = lesson.id
    package.resources.append(resource)
    store.save_for_user(TEST_USER.id, workspace)

    first = chat_service.process_chat_on_lesson(
        lesson.id,
        ChatRequest(message="我要学牛顿莱布尼茨公式"),
        user_id=TEST_USER.id,
    )

    assert first.board_decision.action == "await_reference_choice"
    assert first.reference_prompt is not None
    assert first.reference_prompt.segment_id is not None
    assert first.resource_matches
    assert first.resource_matches[0].segment_id == first.reference_prompt.segment_id
    assert "牛顿莱布尼茨公式" in first.resource_matches[0].excerpt
    assert first.course_package.lessons[0].board_document.content_text == ""

    confirmed = chat_service.process_chat_on_lesson(
        lesson.id,
        ChatRequest(
            message="我要学牛顿莱布尼茨公式",
            resource_reference_action="confirm",
            resource_reference_resource_id=first.reference_prompt.resource_id,
            resource_reference_chapter_id=first.reference_prompt.chapter_id,
            resource_reference_segment_id=first.reference_prompt.segment_id,
        ),
        user_id=TEST_USER.id,
    )

    assert confirmed.chatbot_message == "AI生成：已参考上传资料生成公式板书。"
    assert confirmed.board_decision.action == "edit_board"
    assert confirmed.selected_reference is not None
    assert confirmed.selected_reference.segment_id == first.reference_prompt.segment_id
    assert confirmed.requirement_cleared is True
    assert "牛顿莱布尼茨公式" in (captured["resource_summary"] or "")
    assert "根据资料生成的板书内容" in confirmed.course_package.lessons[0].board_document.content_text
    commit = confirmed.course_package.lessons[0].history_graph.commits[-1]
    assert commit.metadata["kind"] == "board_document_generation"
    assert commit.metadata["resource_backed_generation"] is True
    assert commit.metadata["selected_reference"]["chapter_title"] == "定积分"
    assert commit.metadata["selected_reference"]["segment_id"] == first.reference_prompt.segment_id
    assert commit.metadata["active_requirement_sheet_after"] is None
    assert captured["intent"] == "generate_from_requirements"
    assert _read_log_entries(isolated_ai_log) == []


def test_board_teaching_continue_advances_to_next_section(
    monkeypatch: pytest.MonkeyPatch, isolated_ai_log, tmp_path
) -> None:
    store = SqliteCourseStore(tmp_path / "openclass.sqlite3", legacy_json_path=None)
    monkeypatch.setattr(workspace_state, "STORE", store)
    monkeypatch.setattr(
        openai_course_ai,
        "generate_chatbot_reply",
        lambda **kwargs: ChatbotReply(chatbot_message="AI生成：第二节讲解。"),
    )
    monkeypatch.setattr(openai_course_ai, "generate_learning_requirement_update", _fake_requirement_update)

    workspace = _seed_test_user_workspace(store)
    lesson = workspace.packages[0].lessons[0]
    document = build_document(
        title="分节板书",
        content_text="# 分节板书\n## 第一节\n这是一段足够长的第一节正文，用来避免被识别成额外章节并保持测试稳定。\n## 第二节\n这是一段足够长的第二节正文，用来避免被识别成额外章节并保持测试稳定。",
    )
    refresh_lesson_runtime(lesson, document=document, requirements=lesson.learning_requirements)
    lesson.board_teaching_progress = BoardTeachingProgress(
        board_document_id=document.id,
        board_snapshot_hash="hash-1",
        current_section_index=0,
        completed_section_indexes=[0],
        waiting_for_continue=True,
    )
    lesson.history_graph.commits[-1].snapshot = document
    store.save_for_user(TEST_USER.id, workspace)

    response = chat_service.process_chat_on_lesson(
        lesson.id,
        ChatRequest(message="继续下一节", teaching_action="continue"),
        user_id=TEST_USER.id,
    )

    assert response.teaching_progress is not None
    assert response.teaching_progress.section_index == 1
    assert response.teaching_progress.has_next_section is False
    assert "AI生成：第二节讲解。" in response.chatbot_message
    commit = response.course_package.lessons[0].history_graph.commits[-1]
    assert commit.metadata["teaching_progress"]["section_index"] == 1
    assert _read_log_entries(isolated_ai_log) == []


def test_chat_http_endpoint_returns_chatbot_reply(
    monkeypatch: pytest.MonkeyPatch, isolated_ai_log, tmp_path
) -> None:
    store = SqliteCourseStore(tmp_path / "openclass.sqlite3", legacy_json_path=None)
    monkeypatch.setattr(workspace_state, "STORE", store)
    monkeypatch.setattr(
        openai_course_ai,
        "generate_chatbot_reply",
        lambda **kwargs: ChatbotReply(chatbot_message="AI生成：HTTP 路由回复。"),
    )
    monkeypatch.setattr(openai_course_ai, "generate_learning_requirement_update", _fake_requirement_update)

    workspace = _seed_test_user_workspace(store)
    lesson_id = workspace.packages[0].lessons[0].id

    main_module.app.dependency_overrides[current_user] = lambda: TEST_USER
    try:
        response = TestClient(main_module.app).post(
            f"/api/lessons/{lesson_id}/chat",
            json={"message": "请解释一下核心公式"},
        )
    finally:
        main_module.app.dependency_overrides.pop(current_user, None)

    assert response.status_code == 200
    assert response.json()["chatbot_message"] == "AI生成：HTTP 路由回复。"
    assert _read_log_entries(isolated_ai_log) == []


def test_realtime_transcript_route_reports_removed_workflow(
    monkeypatch: pytest.MonkeyPatch, isolated_ai_log, tmp_path
) -> None:
    monkeypatch.delenv("OPENCLASS_REALTIME_ENABLED", raising=False)
    store = SqliteCourseStore(tmp_path / "openclass.sqlite3", legacy_json_path=None)
    monkeypatch.setattr(workspace_state, "STORE", store)
    lesson = _seed_test_user_workspace(store).packages[0].lessons[0]

    with pytest.raises(HTTPException) as exc_info:
        realtime_router.log_realtime_event(
            lesson.id,
            RealtimeTranscriptLogRequest(
                client_session_id="realtime_session_1",
                lesson_title="测试课",
                role="assistant",
                transport_event_type="response.audio_transcript.done",
                transcript="测试转写",
            ),
            user=TEST_USER,
        )

    assert exc_info.value.status_code == 410
    assert _read_log_entries(isolated_ai_log) == []
