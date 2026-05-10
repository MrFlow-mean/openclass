import pytest
from docx import Document as DocxDocument
from docx.enum.section import WD_ORIENT
from pypdf import PdfWriter
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from zipfile import ZipFile

from app.models import BoardTeachingGuide, BoardTeachingProgress, BoardTeachingSelectedItem, ChatRequest, ConversationTurn, PatchOperation, SelectionRef
from app.services.chart_generation import extract_chart_data_fragments
from app.services.ai_workflow import _board_snapshot_hash, _is_append_document_request, classify_scope, course_workflow, match_resources
from app.services.course_runtime import effective_requirements
from app.services.course_store import build_initial_course_package
from app.services.document_ops import apply_patch
from app.services.history import create_branch, restore_commit
from app.services.lesson_factory import build_requirements, create_empty_lesson, create_lesson
from app.services.openai_course_ai import DocumentEditOutput, openai_course_ai
from app.services.resource_library import _keywords_from_text, build_resource_item, extract_reference_context
from app.services.rich_document import build_document, export_docx, import_docx, replace_selection_in_document


@pytest.fixture(autouse=True)
def disable_openai_for_tests(monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.setattr(openai_course_ai, "client", None)


def test_apply_patch_is_a_document_snapshot_compatibility_shim() -> None:
    document = build_document(title="Test", content_text="first\nsecond")
    next_document, diff = apply_patch(
        document,
        [PatchOperation(op="update_block_content", content="ignored by rich document mode")],
    )

    assert next_document.content_text == document.content_text
    assert diff == []


def test_branch_and_restore_keep_history() -> None:
    lesson = create_lesson("勾股定理")
    first_commit_id = lesson.history_graph.commits[0].id

    create_branch(lesson, "alt-proof", first_commit_id)
    assert lesson.history_graph.current_branch == "alt-proof"
    assert lesson.history_graph.branches["alt-proof"].base_commit_id == first_commit_id

    restore_commit(lesson, first_commit_id, "Restore origin")
    assert lesson.history_graph.branches["alt-proof"].head_commit_id == lesson.history_graph.commits[-1].id
    restore_metadata = lesson.history_graph.commits[-1].metadata
    assert restore_metadata["kind"] == "restore_snapshot"
    assert restore_metadata["restored_commit_id"] == first_commit_id


def test_scope_escalation_detects_out_of_domain_question() -> None:
    lesson = create_lesson("微积分入门")
    assert classify_scope("什么是袋代数中的环和层？什么是光滑？", lesson) == "scope_escalation"
    assert classify_scope("为这些内容出几道习题", lesson) == "in_scope"


def test_create_empty_lesson_starts_with_blank_rich_document() -> None:
    lesson = create_empty_lesson("抽象代数导论")

    assert lesson.board_document.title == "抽象代数导论"
    assert lesson.board_document.content_text == ""
    assert lesson.history_graph.commits[0].snapshot.content_text == ""


def test_workflow_asks_for_clarification_when_request_is_too_vague() -> None:
    package = build_initial_course_package()
    lesson = package.lessons[0]

    result = course_workflow.invoke(
        {
            "lesson": lesson,
            "course_package": package,
            "request": ChatRequest(message="这里没懂"),
        }
    )

    assert result["needs_clarification"] is True
    assert result["board_decision"].action == "clarify_request"
    assert result.get("document_updated") is False
    assert result["teacher_message"].strip()
    assert "什么水平" not in result["teacher_message"]


def test_workflow_asks_for_topic_keyword_on_greeting_without_level_refrain() -> None:
    package = build_initial_course_package()
    lesson = package.lessons[0]

    result = course_workflow.invoke(
        {
            "lesson": lesson,
            "course_package": package,
            "request": ChatRequest(message="你好"),
        }
    )

    assert result["learning_clarification"].progress == 0
    assert result["needs_clarification"] is True
    assert result["board_decision"].action == "clarify_request"
    assert result.get("document_updated") is False
    assert "想学的主题" in result["learning_clarification"].missing_items
    assert "什么水平" not in result["teacher_message"]
    assert "给我一个关键词" not in result["teacher_message"]
    assert "从那里开讲" not in result["teacher_message"]


def test_workflow_probes_level_and_goal_on_first_subject_only_learning_goal() -> None:
    package = build_initial_course_package()
    lesson = package.lessons[0]

    result = course_workflow.invoke(
        {
            "lesson": lesson,
            "course_package": package,
            "request": ChatRequest(message="我想学生物学"),
        }
    )

    assert result["learning_clarification"].progress == 35
    assert result["needs_clarification"] is False
    assert result["board_decision"].action == "no_change"
    assert result["document_updated"] is False
    assert result["board_edit_prompt"] is not None
    assert result["board_edit_prompt"].topic == "生物学"
    assert result["board_teaching_guide"] is not None
    assert result["board_teaching_guide"].lecture_handout
    assert result["learning_requirement_sheet"].learning_need_checklist
    assert result["teacher_message"].strip()
    assert "生物学" in result["teacher_message"]
    assert "起点" in result["teacher_message"] or "具体问题" in result["teacher_message"]
    assert "比如高中" not in result["teacher_message"]


def test_workflow_first_broad_math_goal_only_asks_for_level_and_specific_topic() -> None:
    package = build_initial_course_package()
    lesson = create_empty_lesson("测试12")
    package.lessons.append(lesson)

    result = course_workflow.invoke(
        {
            "lesson": lesson,
            "course_package": package,
            "request": ChatRequest(message="我想学数学"),
        }
    )

    assert result["learning_requirement_sheet"].theme == "数学"
    assert result["learning_clarification"].progress == 35
    assert result["needs_clarification"] is False
    assert result["board_decision"].action == "no_change"
    assert "数学" in result["teacher_message"]
    assert "起点" in result["teacher_message"] or "具体问题" in result["teacher_message"]
    assert "什么是数学" not in result["teacher_message"]
    assert "我们直接抓这次最该讲的重点" not in result["teacher_message"]
    assert "教师模型" not in result["teacher_message"]


def test_workflow_probes_learning_purpose_after_greeting_then_broad_math_goal() -> None:
    package = build_initial_course_package()
    lesson = create_empty_lesson("测试12")
    package.lessons.append(lesson)

    result = course_workflow.invoke(
        {
            "lesson": lesson,
            "course_package": package,
            "request": ChatRequest(
                message="我想学数学",
                conversation=[
                    ConversationTurn(role="user", content="你好"),
                    ConversationTurn(
                        role="assistant",
                        content="请直接发这次要学的内容、材料片段或卡住的题目。",
                    ),
                ],
            ),
        }
    )

    assert result["learning_requirement_sheet"].theme == "数学"
    assert result["needs_clarification"] is True
    assert result["board_decision"].action == "clarify_request"
    assert "数学" in result["teacher_message"]
    assert "起点" in result["teacher_message"] or "具体问题" in result["teacher_message"]
    assert "给我一个关键词" not in result["teacher_message"]
    assert "从那里开讲" not in result["teacher_message"]
    assert "理想" not in result["teacher_message"]
    assert "素理想" not in result["teacher_message"]
    assert "教师模型" not in result["teacher_message"]


def test_workflow_updates_topic_and_starts_after_high_school_concept_answer() -> None:
    package = build_initial_course_package()
    lesson = create_empty_lesson("测试13")
    package.lessons.append(lesson)

    result = course_workflow.invoke(
        {
            "lesson": lesson,
            "course_package": package,
            "request": ChatRequest(
                message="我是高中生，我想学什么事库仑力",
                conversation=[
                    ConversationTurn(role="user", content="你好"),
                    ConversationTurn(role="assistant", content="你具体想学什么内容？可以直接说一个主题、章节，或者把卡住的题目发给我。"),
                    ConversationTurn(role="user", content="我想学数学"),
                    ConversationTurn(
                        role="assistant",
                        content="你当前是什么水平、几年级？另外你具体想学数学里的什么内容，比如函数、几何、代数、微积分、概率统计，还是某类题？",
                    ),
                    ConversationTurn(role="user", content="其实我想学物理"),
                    ConversationTurn(role="assistant", content="你当前是什么水平或背景？这次具体想学什么内容，想达到什么目标？"),
                ],
            ),
        }
    )

    assert result["learning_requirement_sheet"].theme == "库仑力"
    assert result["learning_requirement_sheet"].level == "高中生"
    assert result["needs_clarification"] is False
    assert result["board_decision"].action == "no_change"
    assert "库仑力" in result["teacher_message"]
    assert "具体想学" not in result["teacher_message"]
    assert "你当前是什么水平或背景？这次具体想学什么内容，想达到什么目标？" not in result["teacher_message"]


def test_workflow_probes_background_for_first_advanced_subject_goal() -> None:
    package = build_initial_course_package()
    lesson = create_empty_lesson("测试9")
    package.lessons.append(lesson)

    result = course_workflow.invoke(
        {
            "lesson": lesson,
            "course_package": package,
            "request": ChatRequest(message="我要学代数几何的环的内容"),
        }
    )

    assert result["learning_requirement_sheet"].theme == "代数几何的环"
    assert result["learning_clarification"].progress == 35
    assert set(result["learning_clarification"].missing_items) == {"当前水平或背景", "学习目的或应用场景"}
    assert result["needs_clarification"] is False
    assert result["board_decision"].action == "no_change"
    assert "代数几何的环" in result["teacher_message"]
    assert "起点" in result["teacher_message"] or "具体问题" in result["teacher_message"]
    assert "测试9" not in result["teacher_message"]


def test_workflow_marks_detailed_learning_goal_as_fully_clarified() -> None:
    package = build_initial_course_package()
    lesson = create_empty_lesson("板书测试")
    package.lessons.append(lesson)

    result = course_workflow.invoke(
        {
            "lesson": lesson,
            "course_package": package,
            "request": ChatRequest(
                message=(
                    "我是一名法语的学习者，我的法语水平是B2，词汇是3500左右。"
                    "我要去法国旅游了，你能不能给我生成一篇在法国咖啡厅点餐的一篇情景对话课文，"
                    "要用上关于过去将来时的语法"
                ),
            ),
        }
    )

    assert result["learning_clarification"].progress == 100
    assert result["needs_clarification"] is False
    assert result["board_decision"].action == "edit_board"
    assert result["document_updated"] is False
    assert "没有写入板书" in result["teacher_message"]
    assert "适用水平：" not in result["teacher_message"]
    assert "学习目标：" not in result["teacher_message"]


@pytest.mark.parametrize(
    ("title", "message", "expected_level", "expected_goal"),
    [
        (
            "操作系统学习",
            "我是计算机大二，正在学操作系统，准备操作系统期末考试，想把虚拟内存、页表和 TLB 讲清楚。",
            "大二",
            "操作系统期末考试",
        ),
        (
            "统计方法学习",
            "我是社会学研一，量化方法刚入门，论文阅读里总看到 p-value 和显著性检验，想知道怎么解释结论。",
            "研一",
            "论文阅读",
        ),
        (
            "法国大革命",
            "我是高二学生，准备历史考试，想学法国大革命的原因、过程和影响。",
            "高二",
            "历史考试",
        ),
        (
            "红楼梦导读",
            "我是中文系大一，古代小说课要做课程展示，想理清《红楼梦》人物关系和主题。",
            "大一",
            "课程展示",
        ),
        (
            "有机化学学习",
            "我是化学专业大一，实验课前要补 SN1 和 SN2 反应机理，想知道怎么判断底物和条件。",
            "大一",
            "实验",
        ),
        (
            "合同法学习",
            "我是法学本科二年级，合同法期末考试要考要约和承诺，想会判断案例里的合同是否成立。",
            "本科二年级",
            "期末考试",
        ),
    ],
)
def test_workflow_recognizes_varied_student_personas_as_clear_learning_needs(
    title: str,
    message: str,
    expected_level: str,
    expected_goal: str,
) -> None:
    package = build_initial_course_package()
    lesson = create_empty_lesson(title)
    package.lessons.append(lesson)

    result = course_workflow.invoke(
        {
            "lesson": lesson,
            "course_package": package,
            "request": ChatRequest(message=message),
        }
    )

    status = result["learning_clarification"]
    requirements = result["learning_requirement_sheet"]

    assert status.progress >= 90
    assert "当前水平或背景" not in status.missing_items
    assert "学习目的或应用场景" not in status.missing_items
    assert result["needs_clarification"] is False
    assert requirements.level == expected_level
    assert expected_goal in requirements.success_criteria


@pytest.mark.parametrize(
    "message",
    [
        "我是医学生，背诵能力强，但机制理解不足。我想学习免疫系统的细胞机制，用于理解感染、疫苗和自身免疫病。",
        "我是土木工程学生，学过材料力学。我想学习结构力学中的梁、桁架、弯矩图和稳定性分析。",
        "我是建筑学学生，设计能力强，但结构知识弱。我想学习建筑结构受力原理，理解柱、梁、拱、壳体结构。",
        "我是软件工程学生，会写代码，但系统设计经验少。我想学习后端架构、数据库设计、缓存、消息队列和微服务。",
        "我是社会学学生，对社会现象感兴趣。我想学习社会分层、现代性、韦伯、涂尔干和马克思的社会理论。",
        "我是音乐学院学生，会演奏但乐理薄弱。我想学习和声学、调式、转调和曲式分析。",
        "我是公共卫生学生，医学基础一般。我想学习流行病学中的发病率、患病率、队列研究和病例对照研究。",
        "我是商科学生，对创业感兴趣。我想学习商业模式、市场定位、用户需求分析和 MVP 产品验证。",
    ],
)
def test_workflow_recognizes_natural_background_phrases_from_student_personas(message: str) -> None:
    package = build_initial_course_package()
    lesson = create_empty_lesson("多画像需求识别")
    package.lessons.append(lesson)

    result = course_workflow.invoke(
        {
            "lesson": lesson,
            "course_package": package,
            "request": ChatRequest(message=message),
        }
    )

    status = result["learning_clarification"]
    assert status.progress >= 90
    assert "当前水平或背景" not in status.missing_items
    assert "学习目的或应用场景" not in status.missing_items
    assert result["needs_clarification"] is False


def test_workflow_does_not_misread_environment_science_as_ring_algebra() -> None:
    package = build_initial_course_package()
    lesson = create_empty_lesson("环境科学")
    package.lessons.append(lesson)

    result = course_workflow.invoke(
        {
            "lesson": lesson,
            "course_package": package,
            "request": ChatRequest(message="我想学环境科学"),
        }
    )

    assert "环境科学" in result["teacher_message"]
    assert "起点" in result["teacher_message"] or "具体问题" in result["teacher_message"]
    assert "理想或素理想" not in result["teacher_message"]
    assert "抽象代数" not in result["teacher_message"]


@pytest.mark.parametrize(
    "message",
    [
        "请生成一份虚拟内存讲义，覆盖地址空间、页表、TLB、缺页异常、页面置换，生成后先只讲第一小节。",
        "请生成一份法国大革命讲义，覆盖三级会议、攻占巴士底狱、雅各宾派、拿破仑，生成后先只讲第一小节。",
        "请生成一份合同法要约与承诺专题讲义，覆盖要约、承诺、撤回、撤销和案例判断，生成后先只讲第一小节。",
        "请生成一份 p-value 与显著性检验讲义，覆盖零假设、备择假设、一类错误、置信区间，生成后先只讲第一小节。",
        "请生成一份集合、映射、群、环、域系统讲义，覆盖集合、映射、群、环、域，为大学数学打基础。生成后先只讲第一小节。",
        "请生成一份项目复盘讲义，覆盖目标、假设、指标、风险和下一步行动。生成后先只讲第一小节。",
    ],
)
def test_workflow_does_not_persist_local_template_when_generation_model_unavailable(message: str) -> None:
    package = build_initial_course_package()
    lesson = create_empty_lesson("跨学科生成测试")
    package.lessons.append(lesson)

    result = course_workflow.invoke(
        {
            "lesson": lesson,
            "course_package": package,
            "request": ChatRequest(message=message),
        }
    )

    doc_text = result["teacher_document"].content_text
    guide = result["board_teaching_guide"]

    assert result["board_decision"].action == "edit_board"
    assert result["document_updated"] is False
    assert guide is not None
    assert doc_text == ""
    assert "没有写入板书" in result["teacher_message"]
    assert "问题入口" not in result["teacher_message"]
    assert "核心概念" not in result["teacher_message"]


def test_workflow_can_start_when_user_forces_teaching_before_goal_is_clear() -> None:
    package = build_initial_course_package()
    lesson = package.lessons[0]

    result = course_workflow.invoke(
        {
            "lesson": lesson,
            "course_package": package,
            "request": ChatRequest(message="我想学项目管理，直接开始教学"),
        }
    )

    assert result["learning_clarification"].progress < 60
    assert result["learning_clarification"].forced_start is True
    assert result["needs_clarification"] is False


def test_workflow_extracts_level_and_goal_from_user_message() -> None:
    package = build_initial_course_package()
    lesson = create_empty_lesson("项目沟通")
    package.lessons.append(lesson)

    result = course_workflow.invoke(
        {
            "lesson": lesson,
            "course_package": package,
            "request": ChatRequest(message="教我项目沟通 B2 我想用于跨部门协作"),
        }
    )

    assert result["learning_requirement_sheet"].level == "B2"
    assert "跨部门协作" in result["learning_requirement_sheet"].success_criteria
    assert result["needs_clarification"] is False


def test_workflow_treats_integrated_math_learning_goal_as_purpose() -> None:
    package = build_initial_course_package()
    lesson = create_empty_lesson("环论学习")
    package.lessons.append(lesson)

    result = course_workflow.invoke(
        {
            "lesson": lesson,
            "course_package": package,
            "request": ChatRequest(
                message=(
                    "我是数学系大三，本科抽象代数学过群、环的基本定义，也知道一点理想和商环，"
                    "但代数几何还没系统学。我想把环作为连接抽象代数、交换代数和代数几何的主线学扎实。"
                )
            ),
        }
    )

    assert result["learning_clarification"].progress >= 90
    assert "学习目的或应用场景" not in result["learning_clarification"].missing_items
    assert "连接抽象代数" in result["learning_requirement_sheet"].success_criteria
    assert result["needs_clarification"] is False


def test_workflow_preserves_blank_board_when_generation_model_unavailable() -> None:
    package = build_initial_course_package()
    lesson = create_empty_lesson("环论学习")
    package.lessons.append(lesson)

    result = course_workflow.invoke(
        {
            "lesson": lesson,
            "course_package": package,
            "request": ChatRequest(
                message=(
                    "我是数学系大三，本科抽象代数学过群、环的基本定义，也知道一点理想和商环，"
                    "但代数几何还没系统学。我想把环作为连接抽象代数、交换代数和代数几何的主线学扎实。"
                    "请生成一份系统的 Word 式板书讲义，至少 21 个小节，覆盖 Spec、Zariski 拓扑、"
                    "Hilbert 零点定理、局部化、Noether 环和仿射概形。生成后先只讲第一个小节。"
                )
            ),
        }
    )

    doc_text = result["teacher_document"].content_text
    guide = result["board_teaching_guide"]
    progress = result["teaching_progress"]

    assert result["board_decision"].action == "edit_board"
    assert result["document_updated"] is False
    assert guide is not None
    assert progress.section_count == 0
    assert doc_text == ""
    assert "Hilbert 零点定理" not in doc_text
    assert "Zariski 拓扑" not in doc_text
    assert "仿射概形" not in doc_text
    assert "没有写入板书" in result["teacher_message"]
    assert "问题入口" not in result["teacher_message"]
    assert "核心概念" not in result["teacher_message"]


def test_topicless_board_generation_keeps_prior_learning_topic_when_model_unavailable() -> None:
    package = build_initial_course_package()
    lesson = create_empty_lesson("测试2")
    lesson.learning_requirements = build_requirements("线性代数")
    lesson.learning_requirements.level = "大一"
    lesson.learning_requirements.known_background = "用户自述背景：大一"
    package.lessons.append(lesson)

    result = course_workflow.invoke(
        {
            "lesson": lesson,
            "course_package": package,
            "request": ChatRequest(
                message="开始生成板书",
                conversation=[
                    ConversationTurn(role="user", content="我想学线性代数"),
                    ConversationTurn(role="assistant", content="我们可以从线性代数开始。"),
                    ConversationTurn(role="user", content="我是大一新生"),
                ],
            ),
        }
    )

    assert result["board_decision"].action == "edit_board"
    assert result["document_updated"] is False
    assert result["learning_requirement_sheet"].theme == "线性代数"
    assert result["learning_requirement_sheet"].level == "大一"
    assert "开始生成板书" not in result["learning_requirement_sheet"].learning_need_checklist
    assert result["teacher_document"].content_text == ""
    assert "没有写入板书" in result["teacher_message"]


def test_workflow_generates_board_for_blank_lesson_when_user_requests_direct_open_lecture(tmp_path) -> None:
    package = build_initial_course_package()
    lesson = create_empty_lesson("cs测试2")
    package.lessons.append(lesson)
    resource_path = tmp_path / "csapp-mini.md"
    resource_path.write_text(
        "\n".join(
            [
                "# Chapter 1",
                "## Section 1.1",
                "Intro 1.1",
                "## Section 1.2",
                "Intro 1.2",
                "# Chapter 2",
                "## Section 2.1",
                "Intro 2.1",
                "## Section 2.2",
                "Intro 2.2",
                "# Chapter 3",
                "## Section 3.1",
                "Intro 3.1",
                "## Section 3.2",
                "Intro 3.2",
                "# Chapter 4",
                "## Section 4.1",
                "Intro 4.1",
                "## Section 4.2",
                "Intro 4.2",
                "# Chapter 5",
                "## Expressing Program Performance",
                "This section explains how to express program performance clearly.",
                "## Program Example",
                "This section walks through a concrete program example.",
            ]
        ),
        encoding="utf-8",
    )
    resource = build_resource_item(resource_path, "CSAPP mini.md")
    package.resources.append(resource)
    chapter = next(chapter for chapter in resource.outline if chapter.title == "Program Example")

    result = course_workflow.invoke(
        {
            "lesson": lesson,
            "course_package": package,
            "request": ChatRequest(
                message="为我讲解教材中的第5章第2节的内容，直接开讲",
                resource_reference_action="confirm",
                resource_reference_resource_id=resource.id,
                resource_reference_chapter_id=chapter.id,
            ),
        }
    )

    assert result["needs_clarification"] is False
    assert result["board_decision"].action == "edit_board"
    assert result["document_updated"] is True
    assert result["selected_reference"] is not None
    assert result["selected_reference"].chapter_title == "Program Example"
    assert result["board_teaching_guide"] is not None
    assert "Program Example" in result["board_teaching_guide"].lecture_handout
    assert "Program Example" in result["teacher_document"].content_text
    assert "什么水平" not in result["teacher_message"]


def test_workflow_teaches_generated_board_one_h2_section_at_a_time(monkeypatch: pytest.MonkeyPatch) -> None:
    package = build_initial_course_package()
    lesson = create_empty_lesson("分节讲义")
    package.lessons.append(lesson)
    content_html = """
    <h1>分节讲义</h1>
    <h2>一、学习动机</h2><p>先说明为什么要学这个主题。</p>
    <h2>二、判断标准</h2><p>解释最重要的定义和边界。</p>
    <h2>三、迁移任务</h2><p>用一个最小例子走完整流程。</p>
    <h2>四、检查练习</h2><p>让学生自己判断是否掌握。</p>
    """.strip()
    monkeypatch.setattr(
        openai_course_ai,
        "generate_document_edit",
        lambda **kwargs: DocumentEditOutput(
            rationale="生成四节板书",
            replacement_html=content_html,
        ),
    )
    monkeypatch.setattr(
        openai_course_ai,
        "generate_teacher_message",
        lambda **kwargs: pytest.fail("section teaching should not call the generic teacher message"),
    )

    result = course_workflow.invoke(
        {
            "lesson": lesson,
            "course_package": package,
            "request": ChatRequest(message="请生成一版四个小节的板书"),
        }
    )

    guide = result["board_teaching_guide"]
    progress = result["board_teaching_progress"]
    assert guide is not None
    assert [plan.heading for plan in guide.section_plans] == ["一、学习动机", "二、判断标准", "三、迁移任务", "四、检查练习"]
    assert progress.current_section_index == 0
    assert result["teaching_progress"].has_next_section is True
    assert "第 1 小节" in result["teacher_message"]
    assert "一、学习动机" in result["teacher_message"]
    assert "二、判断标准" not in result["teacher_message"]

    lesson.board_document = result["teacher_document"]
    lesson.board_teaching_guide = guide
    lesson.board_teaching_progress = progress
    followup = course_workflow.invoke(
        {
            "lesson": lesson,
            "course_package": package,
            "request": ChatRequest(message="继续下一节", teaching_action="continue"),
        }
    )

    assert followup["board_teaching_progress"].current_section_index == 1
    assert followup["teaching_progress"].current_section_title == "二、判断标准"
    assert "第 2 小节" in followup["teacher_message"]
    assert "二、判断标准" in followup["teacher_message"]
    assert "一、学习动机" not in followup["teacher_message"]


def test_workflow_generates_chart_image_for_data_rich_board(monkeypatch: pytest.MonkeyPatch) -> None:
    package = build_initial_course_package()
    lesson = create_empty_lesson("数据趋势")
    package.lessons.append(lesson)
    content_html = """
    <h1>产品增长数据</h1>
    <h2>一、用户增长趋势</h2>
    <p>从年度数据看，2022 年活跃用户 120 万，2023 年活跃用户 180 万，2024 年活跃用户 260 万，整体呈持续增长趋势。</p>
    """.strip()
    chart_calls: list[dict[str, str]] = []

    monkeypatch.setattr(
        openai_course_ai,
        "generate_document_edit",
        lambda **kwargs: DocumentEditOutput(
            rationale="生成包含数据的板书",
            replacement_html=content_html,
        ),
    )
    monkeypatch.setattr(
        openai_course_ai,
        "generate_chart_image",
        lambda **kwargs: chart_calls.append(kwargs) or "data:image/png;base64,ZmFrZQ==",
    )

    result = course_workflow.invoke(
        {
            "lesson": lesson,
            "course_package": package,
            "request": ChatRequest(message="请生成一版用户增长数据板书"),
        }
    )

    assert result["document_updated"] is True
    assert chart_calls
    assert chart_calls[0]["chart_type"] == "折线图"
    assert "2024 年活跃用户 260 万" in chart_calls[0]["source_excerpt"]
    assert "AI 图表：折线图" in result["teacher_document"].content_html
    assert '<img src="data:image/png;base64,ZmFrZQ=="' in result["teacher_document"].content_html


def test_chart_fragment_rules_choose_pie_for_share_data() -> None:
    document = build_document(
        title="市场份额",
        content_html="""
<h1>市场份额</h1>
<h2>渠道占比</h2>
<p>本季度渠道占比数据为：线上 45%，门店 35%，代理商 20%。</p>
""",
    )

    fragments = extract_chart_data_fragments(document, request_message="生成图表")

    assert fragments
    assert fragments[0].chart_type == "饼图 / 环形图"


def test_workflow_records_section_followup_as_child_learning_need() -> None:
    package = build_initial_course_package()
    lesson = create_empty_lesson("开平方入门")
    package.lessons.append(lesson)
    lesson.board_document = build_document(
        title="开平方入门",
        content_html="""
<h1>开平方入门</h1>
<h2>一、乘法回顾</h2>
<p>乘法可以理解为相同数量的重复累加。</p>
<h2>二、根号开平方</h2>
<p>开平方是在问哪个非负数平方以后等于被开方数。</p>
""",
        document_id=lesson.board_document.id,
    )
    lesson.board_teaching_progress = BoardTeachingProgress(
        board_document_id=lesson.board_document.id,
        board_snapshot_hash=_board_snapshot_hash(lesson.board_document),
        current_section_index=1,
        completed_section_indexes=[1],
        waiting_for_continue=True,
    )

    result = course_workflow.invoke(
        {
            "lesson": lesson,
            "course_package": package,
            "request": ChatRequest(message="如果是负数被开方会怎么样？还有小数开方怎么算？"),
        }
    )

    needs = result["learning_requirement_sheet"].learning_need_checklist
    assert result["board_decision"].action == "no_change"
    assert result["document_updated"] is False
    assert any(need.startswith("2.1 ") and "负数" in need and "二、根号开平方" in need for need in needs)
    assert any(need.startswith("2.2 ") and "小数开方" in need for need in needs)
    assert "负数" in result["board_teaching_guide"].lecture_handout
    assert "小数开方" in result["teacher_message"]


def test_confirming_section_followup_appends_numbered_child_section() -> None:
    package = build_initial_course_package()
    lesson = create_empty_lesson("开平方入门")
    package.lessons.append(lesson)
    lesson.board_document = build_document(
        title="开平方入门",
        content_html="""
<h1>开平方入门</h1>
<h2>一、乘法回顾</h2>
<p>乘法可以理解为相同数量的重复累加。</p>
<h2>二、根号开平方</h2>
<p>开平方是在问哪个非负数平方以后等于被开方数。</p>
""",
        document_id=lesson.board_document.id,
    )
    lesson.board_teaching_progress = BoardTeachingProgress(
        board_document_id=lesson.board_document.id,
        board_snapshot_hash=_board_snapshot_hash(lesson.board_document),
        current_section_index=1,
        completed_section_indexes=[1],
        waiting_for_continue=True,
    )

    first = course_workflow.invoke(
        {
            "lesson": lesson,
            "course_package": package,
            "request": ChatRequest(message="如果是负数被开方会怎么样？"),
        }
    )
    lesson.learning_requirements = first["learning_requirement_sheet"]
    lesson.board_teaching_guide = first["board_teaching_guide"]

    result = course_workflow.invoke(
        {
            "lesson": lesson,
            "course_package": package,
            "request": ChatRequest(
                message="如果是负数被开方会怎么样？",
                board_edit_action="confirm",
                board_edit_topic=first["board_edit_prompt"].topic,
            ),
        }
    )

    assert result["board_decision"].action == "append_section"
    assert result["document_updated"] is False
    assert "负数被开方会怎么样" not in result["teacher_document"].content_text
    assert "没有写入板书" in result["teacher_message"]


def test_workflow_backfills_section_plans_for_legacy_board_teaching_guide() -> None:
    package = build_initial_course_package()
    lesson = create_empty_lesson("旧教案")
    package.lessons.append(lesson)
    lesson.board_document = build_document(
        title="旧教案",
        content_html="<h1>旧教案</h1><h2>第一节</h2><p>第一节正文。</p><h2>第二节</h2><p>第二节正文。</p>",
        document_id=lesson.board_document.id,
    )
    lesson.board_teaching_guide = BoardTeachingGuide(
        board_document_id=lesson.board_document.id,
        board_snapshot_hash=_board_snapshot_hash(lesson.board_document),
        board_title=lesson.board_document.title,
        selected_items=[
            BoardTeachingSelectedItem(
                excerpt="第一节正文。",
                source_heading="第一节",
                reason="旧版只存重点摘录。",
                order_index=1,
            )
        ],
        teacher_brief="旧版教案没有分节计划。",
    )

    result = course_workflow.invoke(
        {
            "lesson": lesson,
            "course_package": package,
            "request": ChatRequest(message="继续下一节", teaching_action="continue"),
        }
    )

    assert [plan.heading for plan in result["board_teaching_guide"].section_plans] == ["第一节", "第二节"]
    assert result["board_teaching_progress"].current_section_index == 0
    assert "第一节" in result["teacher_message"]


def test_workflow_teaches_chinese_numbered_chapter_from_reference_without_filler(tmp_path) -> None:
    package = build_initial_course_package()
    lesson = create_empty_lesson("资料导论")
    package.lessons.append(lesson)
    resource_path = tmp_path / "intro-notes.md"
    resource_path.write_text(
        "\n".join(
            [
                "# 第一章 概论",
                "第一章先建立三个问题：这门课研究什么、常见对象是什么、如何判断结果是否可靠。",
                "它还会说明基础概念、典型任务和评价方式之间的关系。",
                "# 第二章 方法",
                "第二章进入具体方法。",
            ]
        ),
        encoding="utf-8",
    )
    resource = build_resource_item(resource_path, "学习讲义.md")
    package.resources.append(resource)

    result = course_workflow.invoke(
        {
            "lesson": lesson,
            "course_package": package,
            "request": ChatRequest(
                message="为我讲第一章内容",
                resource_reference_action="confirm",
                resource_reference_resource_id=resource.id,
                resource_reference_chapter_id=resource.outline[0].id,
            ),
        }
    )

    assert result["needs_clarification"] is False
    assert result["board_decision"].action == "edit_board"
    assert result["selected_reference"] is not None
    assert result["selected_reference"].chapter_title == "第一章 概论"
    assert result["document_updated"] is True
    assert "第一章先建立三个问题" in result["board_teaching_guide"].lecture_handout
    assert "第一章先建立三个问题" in result["teacher_document"].content_text
    assert "请补入一个最小例子" not in result["board_teaching_guide"].lecture_handout
    assert "这门课研究什么" in result["teacher_message"]
    assert "顺手告诉我" not in result["teacher_message"]
    assert "工作项目" not in result["teacher_message"]


def test_workflow_turns_reference_chapter_into_polished_handout(
    monkeypatch: pytest.MonkeyPatch,
    tmp_path,
) -> None:
    monkeypatch.setattr(openai_course_ai, "generate_document_edit", lambda **kwargs: None)
    monkeypatch.setattr(openai_course_ai, "generate_teacher_message", lambda **kwargs: None)

    package = build_initial_course_package()
    lesson = create_empty_lesson("测试5")
    package.lessons.append(lesson)
    resource_path = tmp_path / "chapter-notes.md"
    resource_path.write_text(
        "\n".join(
            [
                "# 第7章 通用方法概要",
                "第7章通用方法概要 7.1 引言。本章关心如何把有限材料整理成可复用的方法。",
                "定义输入信息为原始材料，判断标准是能否支持解释、迁移和检查。",
                "（7-1）（7-2）（7-3）160 第7章通用方法概要 R(a) Remp(a) 这里有一些 PDF 抽取噪声。",
                "只照搬材料并不可靠，因为缺少边界条件时会误用结论。",
                "一致性要求材料、目标和输出结构能够互相支撑；进一步要讨论例子与反例。",
                "评价维度刻画方法适用范围，质量边界把目标、证据和检查问题联系起来。",
                "实践中可以通过分层讲解和反馈检查控制理解偏差。",
            ]
        ),
        encoding="utf-8",
    )
    resource = build_resource_item(resource_path, "通用方法讲义.md")
    package.resources.append(resource)

    result = course_workflow.invoke(
        {
            "lesson": lesson,
            "course_package": package,
            "request": ChatRequest(
                message="讲解第七章的内容",
                resource_reference_action="confirm",
                resource_reference_resource_id=resource.id,
                resource_reference_chapter_id=resource.outline[0].id,
            ),
        }
    )

    lecture_handout = result["board_teaching_guide"].lecture_handout
    assert result["board_decision"].action == "edit_board"
    assert result["document_updated"] is True
    assert result["teacher_document"].content_text
    assert "通用方法概要" in lecture_handout
    assert "输入信息" in lecture_handout
    assert "质量边界" in lecture_handout
    assert "PDF 抽取噪声" not in result["teacher_message"]
    assert len(lecture_handout) >= 450
    assert "（7-1）（7-2）（7-3）" not in result["teacher_message"]
    assert "有限材料" in result["teacher_message"] or "可复用的方法" in result["teacher_message"]


def test_workflow_turns_reference_chapter_into_detailed_handout(
    monkeypatch: pytest.MonkeyPatch,
    tmp_path,
) -> None:
    monkeypatch.setattr(openai_course_ai, "generate_document_edit", lambda **kwargs: None)
    monkeypatch.setattr(openai_course_ai, "generate_teacher_message", lambda **kwargs: None)

    package = build_initial_course_package()
    lesson = create_empty_lesson("测试8")
    package.lessons.append(lesson)
    resource_path = tmp_path / "structured-method.md"
    resource_path.write_text(
        "\n".join(
            [
                "# 第3章 结构化方法的建立",
                "第3章结构化方法的建立。本章的基础是先明确输入、处理步骤和评价标准。",
                "输入比较容易罗列，本章重点讨论如何把输入转成可操作的分析路径。",
                "这种先整理材料，再用检查标准进行判断的方法称作两步分析法。",
                "c0；）。max p（x16,1）。这里是 PDF 抽取噪声，不应该被照搬进板书。",
                "第一步是建立分类维度；第二步是用边界条件检查方案是否适用。",
                "开放问题包括证据不足、维度过多和目标变化，处理质量受材料数量、抽象层级和判断标准影响。",
            ]
        ),
        encoding="utf-8",
    )
    resource = build_resource_item(resource_path, "结构化方法讲义.md")
    package.resources.append(resource)

    result = course_workflow.invoke(
        {
            "lesson": lesson,
            "course_package": package,
            "request": ChatRequest(
                message="讲解第三章的内容",
                resource_reference_action="confirm",
                resource_reference_resource_id=resource.id,
                resource_reference_chapter_id=resource.outline[0].id,
            ),
        }
    )

    lecture_handout = result["board_teaching_guide"].lecture_handout
    assert result["board_decision"].action == "edit_board"
    assert result["document_updated"] is True
    assert result["board_teaching_guide"] is not None
    assert len(result["board_teaching_guide"].section_plans) >= 5
    assert len(lecture_handout) >= 450
    assert "结构化方法的建立" in lecture_handout
    assert "两步分析法" in lecture_handout
    assert "边界条件" in lecture_handout
    assert "c0；" not in result["teacher_message"]
    assert "p（x16" not in result["teacher_message"]
    assert "输入、处理步骤和评价标准" in result["teacher_message"] or "两步分析法" in result["teacher_message"]


def test_workflow_expands_important_reference_content(
    monkeypatch: pytest.MonkeyPatch,
    tmp_path,
) -> None:
    monkeypatch.setattr(openai_course_ai, "generate_document_edit", lambda **kwargs: None)
    monkeypatch.setattr(openai_course_ai, "generate_teacher_message", lambda **kwargs: None)

    package = build_initial_course_package()
    lesson = create_empty_lesson("资料讲解")
    package.lessons.append(lesson)
    resource_path = tmp_path / "case-notes.md"
    resource_path.write_text(
        "\n".join(
            [
                "# 案例复盘",
                "案例复盘是一类重要的资料学习任务，它通过背景、关键动作、结果和代价来还原事件结构。",
                "关键动作会改变原有角色的利益分配，使不同参与者获得新的机会或承担新的约束。",
                "制度或流程调整会削弱局部惯性，加强整体目标对具体行动的约束。",
                "从影响看，案例复盘既要说明收益，也要保留风险、代价和后续问题。",
            ]
        ),
        encoding="utf-8",
    )
    resource = build_resource_item(resource_path, "案例资料.md")
    package.resources.append(resource)

    result = course_workflow.invoke(
        {
            "lesson": lesson,
            "course_package": package,
            "request": ChatRequest(
                message="讲一下案例复盘",
                resource_reference_action="confirm",
                resource_reference_resource_id=resource.id,
                resource_reference_chapter_id=resource.outline[0].id,
            ),
        }
    )

    lecture_handout = result["board_teaching_guide"].lecture_handout
    assert result["board_decision"].action == "edit_board"
    assert result["document_updated"] is True
    assert "案例复盘" in lecture_handout
    assert "关键动作" in lecture_handout
    assert "风险" in lecture_handout or "代价" in lecture_handout
    assert "案例复盘" in result["teacher_message"]


def test_match_resources_uses_real_numbered_chapter_after_preface_and_toc(tmp_path) -> None:
    lesson = create_empty_lesson("资料学习")
    package = build_initial_course_package()
    package.lessons.append(lesson)
    resource_path = tmp_path / "structured-book.md"
    resource_path.write_text(
        "\n".join(
            [
                "# 前言",
                "这里是教材前言。",
                "# 目录",
                "第一章 概论。",
                "# 第一章 概论",
                "第一章正文讲对象、特征和判断方法。",
                "## 1.1对象与特征",
                "第一节说明基础任务。",
                "# 第二章 方法",
                "第二章进入具体方法。",
            ]
        ),
        encoding="utf-8",
    )
    package.resources.append(build_resource_item(resource_path, "结构化教材.md"))

    matches = match_resources(
        package,
        lesson,
        ChatRequest(message="讲第一章的内容"),
        effective_requirements(lesson),
    )

    assert matches
    assert matches[0].chapter_title == "第一章 概论"


def test_default_single_resource_skips_preface_and_toc(tmp_path) -> None:
    package = build_initial_course_package()
    lesson = create_empty_lesson("资料学习")
    package.lessons.append(lesson)
    resource_path = tmp_path / "book.md"
    resource_path.write_text(
        "\n".join(
            [
                "# 前言",
                "出版说明。",
                "# 目录",
                "第一章 概论。",
                "# 第一章 概论",
                "第一章先建立课程的整体地图。",
            ]
        ),
        encoding="utf-8",
    )
    package.resources.append(build_resource_item(resource_path, "学习资料.md"))

    result = course_workflow.invoke(
        {
            "lesson": lesson,
            "course_package": package,
            "request": ChatRequest(message="讲一下"),
        }
    )

    assert result["board_decision"].action == "no_change"
    assert result["reference_prompt"] is None
    assert result["selected_reference"] is not None
    assert result["selected_reference"].chapter_title == "第一章 概论"
    assert result["board_edit_prompt"] is None
    assert result["document_updated"] is False


def test_workflow_defaults_brief_followup_to_single_uploaded_resource(tmp_path) -> None:
    package = build_initial_course_package()
    lesson = create_empty_lesson("资料学习")
    package.lessons.append(lesson)
    resource_path = tmp_path / "single-notes.md"
    resource_path.write_text(
        "\n".join(
            [
                "# 第一章 概论",
                "第一章先建立课程的整体地图。",
                "# 第二章 深入",
                "第二章进入细节。",
            ]
        ),
        encoding="utf-8",
    )
    package.resources.append(build_resource_item(resource_path, "学习资料.md"))

    result = course_workflow.invoke(
        {
            "lesson": lesson,
            "course_package": package,
            "request": ChatRequest(message="讲一下"),
        }
    )

    assert result["needs_clarification"] is False
    assert result["board_decision"].action == "no_change"
    assert result["reference_prompt"] is None
    assert result["selected_reference"] is not None
    assert result["selected_reference"].resource_name == "学习资料.md"
    assert result["selected_reference"].chapter_title == "第一章 概论"
    assert result["board_edit_prompt"] is None
    assert result["document_updated"] is False


def test_empty_board_generates_directly_from_uploaded_chapter(tmp_path) -> None:
    package = build_initial_course_package()
    lesson = create_empty_lesson("资料学习")
    package.lessons.append(lesson)
    resource_path = tmp_path / "macro.md"
    resource_path.write_text(
        "\n".join(
            [
                "# 第一章 宏观经济学导论",
                "第一章先解释宏观经济学研究整体经济运行，包括 GDP、失业、通货膨胀和经济增长。",
                "课堂上要先说明总量指标为什么重要，再用一个家庭收入和全国总收入的类比帮助理解。",
                "# 第二章 国民收入核算",
                "第二章进入 GDP 的核算方法。",
            ]
        ),
        encoding="utf-8",
    )
    package.resources.append(build_resource_item(resource_path, "宏观经济学资料.md"))

    result = course_workflow.invoke(
        {
            "lesson": lesson,
            "course_package": package,
            "request": ChatRequest(message="讲解第一章节"),
        }
    )

    assert result["board_decision"].action == "edit_board"
    assert result["reference_prompt"] is None
    assert result["board_edit_prompt"] is None
    assert result["selected_reference"] is not None
    assert result["selected_reference"].chapter_title == "第一章 宏观经济学导论"
    assert result["document_updated"] is True
    assert "宏观经济学导论" in result["teacher_document"].content_text
    assert "GDP" in result["teacher_document"].content_text


def test_workflow_asks_which_file_when_multiple_resources_share_chapter_pointer(tmp_path) -> None:
    package = build_initial_course_package()
    lesson = create_empty_lesson("资料学习")
    package.lessons.append(lesson)
    algebra_path = tmp_path / "algebra.md"
    algebra_path.write_text("# 第一章 群\n第一章讲群的定义。", encoding="utf-8")
    calculus_path = tmp_path / "calculus.md"
    calculus_path.write_text("# 第一章 极限\n第一章讲极限。", encoding="utf-8")
    package.resources.append(build_resource_item(algebra_path, "抽象代数讲义.md"))
    package.resources.append(build_resource_item(calculus_path, "微积分讲义.md"))

    result = course_workflow.invoke(
        {
            "lesson": lesson,
            "course_package": package,
            "request": ChatRequest(message="我要学第一章内容"),
        }
    )

    assert result["needs_clarification"] is True
    assert result["board_decision"].action == "clarify_request"
    assert result["reference_prompt"] is None
    assert result["selected_reference"] is None
    assert "哪一份资料" in result["teacher_message"]
    assert "抽象代数讲义.md" in result["teacher_message"]
    assert "微积分讲义.md" in result["teacher_message"]


def test_workflow_uses_named_file_when_multiple_resources_are_uploaded(tmp_path) -> None:
    package = build_initial_course_package()
    lesson = create_empty_lesson("资料学习")
    package.lessons.append(lesson)
    algebra_path = tmp_path / "algebra.md"
    algebra_path.write_text("# 第一章 群\n第一章讲群的定义。", encoding="utf-8")
    linear_path = tmp_path / "linear.md"
    linear_path.write_text("# 第一章 矩阵\n第一章讲矩阵和线性方程组。", encoding="utf-8")
    package.resources.append(build_resource_item(algebra_path, "抽象代数讲义.md"))
    package.resources.append(build_resource_item(linear_path, "线性代数讲义.md"))

    result = course_workflow.invoke(
        {
            "lesson": lesson,
            "course_package": package,
            "request": ChatRequest(message="请讲线性代数讲义第一章"),
        }
    )

    assert result["needs_clarification"] is False
    assert result["board_decision"].action == "edit_board"
    assert result["reference_prompt"] is None
    assert result["selected_reference"] is not None
    assert result["selected_reference"].resource_name == "线性代数讲义.md"
    assert result["selected_reference"].chapter_title == "第一章 矩阵"
    assert result["document_updated"] is True
    assert "矩阵" in result["board_teaching_guide"].lecture_handout
    assert "矩阵" in result["teacher_document"].content_text


def test_workflow_can_answer_without_changing_the_board() -> None:
    package = build_initial_course_package()
    lesson = package.lessons[0]

    result = course_workflow.invoke(
        {
            "lesson": lesson,
            "course_package": package,
            "request": ChatRequest(message="请解释一下勾股定理的核心公式"),
        }
    )

    assert result["needs_clarification"] is False
    assert result["board_decision"].action == "no_change"
    assert result["document_updated"] is False
    assert result["board_teaching_guide"] is not None
    assert "勾股定理" in result["teacher_message"] or "直角三角形" in result["teacher_message"]


def test_workflow_answers_specific_board_followup_without_preference_clarification() -> None:
    package = build_initial_course_package()
    lesson = create_lesson("第7章 统计学习理论概要")
    lesson.board_document = build_document(
        title="第7章 统计学习理论概要：讲解板书",
        content_html="""
<h1>第7章 统计学习理论概要：讲解板书</h1>
<h2>一、本章定位</h2>
<p>经验风险小不等于真实风险小。机器学习真正追求的是推广能力，而不是只在训练集上表现好。</p>
<h2>二、机器学习问题的数学提法</h2>
<h3>1. 损失函数</h3>
<p>损失函数 L(y,f(x)) 衡量预测结果和真实结果之间的差距。分类问题里常见的是 0-1 损失：预测正确损失为 0，预测错误损失为 1。</p>
<p>损失函数的作用，是把“模型好不好”变成可以计算的量。</p>
<h3>2. 真实风险</h3>
<p>真实风险表示模型在总体分布上的平均损失。</p>
<h3>3. 经验风险</h3>
<p>经验风险是训练集上的平均损失，可以从样本直接计算。</p>
""",
    )
    package.lessons.append(lesson)

    result = course_workflow.invoke(
        {
            "lesson": lesson,
            "course_package": package,
            "request": ChatRequest(
                message="损失函数是什么？我不理解",
                conversation=[
                    ConversationTurn(role="user", content="直接展开讲解第7章内容"),
                    ConversationTurn(
                        role="assistant",
                        content="直接开讲，先抓主线。统计学习理论要回答训练误差什么时候能代表真实风险。",
                    ),
                ],
            ),
        }
    )

    assert result["needs_clarification"] is False
    assert result["board_decision"].action == "no_change"
    assert result["document_updated"] is False
    assert "损失函数" in result["teacher_message"]
    assert "预测结果" in result["teacher_message"] or "真实结果" in result["teacher_message"]


def test_workflow_direct_start_after_brief_clarification_generates_board_for_blank_lesson() -> None:
    package = build_initial_course_package()
    lesson = create_empty_lesson("法考 测试")
    package.lessons.append(lesson)

    result = course_workflow.invoke(
        {
            "lesson": lesson,
            "course_package": package,
            "request": ChatRequest(
                message="直接开讲",
                conversation=[
                    ConversationTurn(role="user", content="为我讲中华人民共和国民法典是什么"),
                    ConversationTurn(role="assistant", content="我可以先按入门节奏讲起来；你顺手告诉我，是为了考试、工作项目，还是日常兴趣？"),
                ],
            ),
        }
    )

    assert result["needs_clarification"] is False
    assert result["board_decision"].action == "no_change"
    assert result["document_updated"] is False
    assert result["teacher_document"].content_text == ""
    assert result["board_teaching_guide"] is not None
    assert result["board_teaching_guide"].lecture_handout
    assert result["board_edit_prompt"] is None


def test_workflow_starts_after_user_answers_from_zero_to_probe() -> None:
    package = build_initial_course_package()
    lesson = create_empty_lesson("在测试1")
    package.lessons.append(lesson)

    result = course_workflow.invoke(
        {
            "lesson": lesson,
            "course_package": package,
            "request": ChatRequest(
                message="从零开始",
                conversation=[
                    ConversationTurn(role="user", content="我要学什么是虚数"),
                    ConversationTurn(
                        role="assistant",
                        content="为了下一轮把例子和深度对准，你之前接触过这个主题吗，还是希望我从零开始？",
                    ),
                ],
            ),
        }
    )

    assert result["needs_clarification"] is False
    assert result["board_decision"].action == "no_change"
    assert result["document_updated"] is False
    assert "虚数" in result["teacher_message"]


def test_workflow_starts_when_user_delegates_after_option_prompt() -> None:
    package = build_initial_course_package()
    lesson = create_empty_lesson("在测试1")
    package.lessons.append(lesson)

    result = course_workflow.invoke(
        {
            "lesson": lesson,
            "course_package": package,
            "request": ChatRequest(
                message="都要，你自己看着办",
                conversation=[
                    ConversationTurn(role="user", content="我要学什么是虚数"),
                    ConversationTurn(
                        role="assistant",
                        content="虚数可以从直观图像讲，也可以从代数规则讲；你更想先解决哪一个卡点？",
                    ),
                ],
            ),
        }
    )

    assert result["needs_clarification"] is False
    assert result["board_decision"].action == "no_change"
    assert result["document_updated"] is False
    assert "虚数" in result["teacher_message"]


def test_workflow_formats_teacher_message_into_readable_paragraphs(monkeypatch: pytest.MonkeyPatch) -> None:
    package = build_initial_course_package()
    lesson = package.lessons[0]

    monkeypatch.setattr(
        openai_course_ai,
        "generate_teacher_message",
        lambda **kwargs: (
            "勾股定理先抓一件事，它说的不是死记公式，而是直角三角形三条边之间的稳定关系。"
            "为什么重要，因为你后面算距离、判定图形、做几何证明都会反复用到它。"
            "你可以把它理解成一把固定尺子，只要直角确定了，两条边一变，第三条边就被锁定了。"
            "最后你可以自己检查一下，3、4、5 为什么会刚好满足这个关系。"
        ),
    )

    result = course_workflow.invoke(
        {
            "lesson": lesson,
            "course_package": package,
            "request": ChatRequest(message="请解释一下勾股定理的核心公式"),
        }
    )

    assert "\n\n" in result["teacher_message"]
    assert "最后你可以自己检查一下" in result["teacher_message"]


def test_workflow_direct_edit_rewrites_only_selected_excerpt() -> None:
    package = build_initial_course_package()
    lesson = package.lessons[0]
    excerpt = "在直角三角形中，两条直角边的平方和，等于斜边的平方。"

    result = course_workflow.invoke(
        {
            "lesson": lesson,
            "course_package": package,
            "request": ChatRequest(
                message="把这一句讲得更易懂",
                interaction_mode="direct_edit",
                selection=SelectionRef(
                    kind="board",
                    lesson_id=lesson.id,
                    excerpt=excerpt,
                ),
            ),
        }
    )

    assert result["needs_clarification"] is False
    assert result["board_decision"].action == "edit_board"
    assert result["document_updated"] is False
    assert "换一种更好懂的说法" not in result["teacher_document"].content_text
    assert "没有写入板书" in result["teacher_message"]


def test_workflow_direct_edit_enhancement_preserves_original_excerpt() -> None:
    package = build_initial_course_package()
    lesson = package.lessons[0]
    lesson.board_document = build_document(
        title=lesson.title,
        content_text=(
            "题干：已知函数 f(x) 在区间上单调递增，求参数 a 的取值范围。\n"
            "解题方法：先求导，再根据导数符号分类讨论。\n"
            "课后提醒：注意端点条件。"
        ),
        document_id=lesson.board_document.id,
    )
    excerpt = "题干：已知函数 f(x) 在区间上单调递增，求参数 a 的取值范围。\n解题方法：先求导，再根据导数符号分类讨论。"

    result = course_workflow.invoke(
        {
            "lesson": lesson,
            "course_package": package,
            "request": ChatRequest(
                message="帮我把题干和解题方法写得更加完善全面",
                interaction_mode="direct_edit",
                selection=SelectionRef(
                    kind="board",
                    lesson_id=lesson.id,
                    excerpt=excerpt,
                ),
            ),
        }
    )

    assert result["needs_clarification"] is False
    assert result["board_decision"].action == "edit_board"
    assert result["document_updated"] is False
    assert "题干：已知函数 f(x) 在区间上单调递增，求参数 a 的取值范围。" in result["teacher_document"].content_text
    assert "解题方法：先求导，再根据导数符号分类讨论。" in result["teacher_document"].content_text
    assert "补充解析" not in result["teacher_document"].content_text
    assert "课后提醒：注意端点条件。" in result["teacher_document"].content_text
    assert "没有写入板书" in result["teacher_message"]


def test_workflow_direct_edit_new_page_appends_instead_of_replacing(monkeypatch: pytest.MonkeyPatch) -> None:
    package = build_initial_course_package()
    lesson = package.lessons[0]
    lesson.board_document = build_document(
        title="量化金融入门讲义",
        content_html=(
            "<h1>量化金融入门讲义：给第一次接触的人</h1>"
            "<p>量化金融是用数据、数学和程序辅助投资决策。</p>"
        ),
        document_id=lesson.board_document.id,
    )

    monkeypatch.setattr(
        openai_course_ai,
        "generate_document_edit",
        lambda **kwargs: DocumentEditOutput(
            rationale="append new page",
            replacement_html=(
                "<h1>量化金融中的数学工具方法</h1>"
                "<p>这一页讲蒙特卡洛方法、相关性和回归。</p>"
            ),
            replacement_text="量化金融中的数学工具方法\n这一页讲蒙特卡洛方法、相关性和回归。",
            teacher_talk_track="这一页我们继续补工具箱。",
            replace_whole=False,
            target_action="create_child_lesson",
        ),
    )

    result = course_workflow.invoke(
        {
            "lesson": lesson,
            "course_package": package,
            "request": ChatRequest(
                message="再为我新生成做一个页面，为我生成几个量化金融的数学工具知识方法，比如蒙特卡洛方法什么的",
                interaction_mode="direct_edit",
            ),
        }
    )

    assert result["board_decision"].action == "append_section"
    assert result["document_updated"] is True
    assert "量化金融入门讲义：给第一次接触的人" in result["teacher_document"].content_text
    assert "量化金融是用数据、数学和程序辅助投资决策。" in result["teacher_document"].content_text
    assert "量化金融中的数学工具方法" in result["teacher_document"].content_text
    assert result["teacher_document"].content_html.index("量化金融入门讲义") < result["teacher_document"].content_html.index(
        "量化金融中的数学工具方法"
    )


def test_append_request_recognizes_continue_new_chapter_wording() -> None:
    assert _is_append_document_request("续写一个新章节，如何解决过拟合？")


def test_expand_board_content_is_not_append_request() -> None:
    assert not _is_append_document_request("请扩展版书内容，细致讲解每个例子")
    assert not _is_append_document_request("把当前板书内容扩写得更详细")


def test_workflow_expands_existing_board_in_place_instead_of_appending_chapter(monkeypatch: pytest.MonkeyPatch) -> None:
    package = build_initial_course_package()
    lesson = create_empty_lesson("测试4")
    package.lessons.append(lesson)
    lesson.board_document = build_document(
        title="风险与收益",
        content_html=(
            "<h1>风险与收益</h1>"
            "<h2>一、什么是风险</h2>"
            "<p>风险不是一定亏钱，而是不确定性。比如同样买入一只股票，未来可能上涨，也可能下跌。</p>"
            "<h2>二、收益怎么理解</h2>"
            "<p>收益是投资结果相对本金的变化，可以用百分比来比较。</p>"
        ),
        document_id=lesson.board_document.id,
    )

    monkeypatch.setattr(
        openai_course_ai,
        "generate_document_edit",
        lambda **kwargs: DocumentEditOutput(
            rationale="bad appended chapter",
            replacement_html="<h2>补充章节</h2><p>这里补充一章风险和收益的内容。</p>",
            replacement_text="补充章节\n这里补充一章风险和收益的内容。",
            teacher_talk_track="继续补充。",
            replace_whole=False,
            target_action="append_section",
        ),
    )

    result = course_workflow.invoke(
        {
            "lesson": lesson,
            "course_package": package,
            "request": ChatRequest(message="请扩展版书内容，细致讲解每个例子"),
        }
    )

    content_text = result["teacher_document"].content_text
    content_html = result["teacher_document"].content_html
    assert result["board_decision"].action == "edit_board"
    assert result["document_updated"] is False
    assert "补充章节" not in content_text
    assert "风险不是一定亏钱" in content_text
    assert "收益是投资结果相对本金的变化" in content_text
    assert "展开说明" not in content_text
    assert content_html.index("一、什么是风险") < content_html.index("二、收益怎么理解")
    assert "没有写入板书" in result["teacher_message"]


def test_workflow_does_not_append_template_when_chapter_generation_model_unavailable() -> None:
    package = build_initial_course_package()
    lesson = create_empty_lesson("测试2")
    package.lessons.append(lesson)
    lesson.board_document = build_document(
        title="几个比较基础的量化数学知识",
        content_html=(
            "<h1>几个比较基础的量化数学知识</h1>"
            "<p>这一页先讲蒙特卡洛方法、相关性和回归。</p>"
        ),
        document_id=lesson.board_document.id,
    )
    lesson.board_teaching_guide = BoardTeachingGuide(
        board_document_id=lesson.board_document.id,
        board_snapshot_hash=_board_snapshot_hash(lesson.board_document),
        board_title=lesson.board_document.title,
        selected_items=[
            BoardTeachingSelectedItem(
                excerpt="续写一个新章节，如何解决过拟合",
                source_heading="补充章节",
                reason="stale guide from earlier bad append",
                mapped_needs=["续写一个新章节，如何解决过拟合"],
                teaching_role="main_idea",
                order_index=1,
            )
        ],
        teacher_brief="这次先抓“续写一个新章节，如何解决过拟合”这条主线。",
    )

    result = course_workflow.invoke(
        {
            "lesson": lesson,
            "course_package": package,
            "request": ChatRequest(message="续写一个新章节，如何解决过拟合？"),
        }
    )

    content_text = result["teacher_document"].content_text
    assert result["board_decision"].action == "append_section"
    assert result["document_updated"] is False
    assert "几个比较基础的量化数学知识" in content_text
    assert "蒙特卡洛方法、相关性和回归" in content_text
    assert "问题入口" not in content_text
    assert "补充章节：如何解决过拟合" not in content_text
    assert "练习任务" not in content_text
    assert "没有写入板书" in result["teacher_message"]


def test_workflow_rejects_low_value_ai_append_without_local_template(monkeypatch: pytest.MonkeyPatch) -> None:
    package = build_initial_course_package()
    lesson = create_empty_lesson("测试2")
    package.lessons.append(lesson)
    lesson.board_document = build_document(
        title="几个比较基础的量化数学知识",
        content_html=(
            "<h1>几个比较基础的量化数学知识</h1>"
            "<p>这一页先讲蒙特卡洛方法、相关性和回归。</p>"
        ),
        document_id=lesson.board_document.id,
    )

    monkeypatch.setattr(
        openai_course_ai,
        "generate_document_edit",
        lambda **kwargs: DocumentEditOutput(
            rationale="low value append",
            replacement_html=(
                "<h2>补充章节</h2>"
                "<p>这一节专门承接用户当前追问，把新问题接回原有主线。</p>"
                "<p>续写一个新章节，如何解决过拟合</p>"
            ),
            replacement_text=(
                "补充章节\n"
                "这一节专门承接用户当前追问，把新问题接回原有主线。\n"
                "续写一个新章节，如何解决过拟合"
            ),
            teacher_talk_track="直接开讲，先抓主线。",
            replace_whole=False,
            target_action="append_section",
        ),
    )

    result = course_workflow.invoke(
        {
            "lesson": lesson,
            "course_package": package,
            "request": ChatRequest(message="续写一个新章节，如何解决过拟合？"),
        }
    )

    content_text = result["teacher_document"].content_text
    assert result["board_decision"].action == "append_section"
    assert result["document_updated"] is False
    assert "补充章节：如何解决过拟合" not in content_text
    assert "问题入口" not in content_text
    assert "练习任务" not in content_text
    assert "用户当前追问" not in content_text
    assert "续写一个新章节" not in content_text
    assert "没有写入板书" in result["teacher_message"]


def test_workflow_rejects_too_short_ai_chapter_append_without_local_template(monkeypatch: pytest.MonkeyPatch) -> None:
    package = build_initial_course_package()
    lesson = create_empty_lesson("测试2")
    package.lessons.append(lesson)
    lesson.board_document = build_document(
        title="几个比较基础的量化数学知识",
        content_html=(
            "<h1>几个比较基础的量化数学知识</h1>"
            "<p>这一页先讲蒙特卡洛方法、相关性和回归。</p>"
        ),
        document_id=lesson.board_document.id,
    )

    monkeypatch.setattr(
        openai_course_ai,
        "generate_document_edit",
        lambda **kwargs: DocumentEditOutput(
            rationale="too short append",
            replacement_html="<h2>补充章节：如何解决过拟合</h2><p>可以用验证集、正则化和交叉验证来减少过拟合。</p>",
            replacement_text="补充章节：如何解决过拟合\n可以用验证集、正则化和交叉验证来减少过拟合。",
            teacher_talk_track="过拟合要看样本外效果。",
            replace_whole=False,
            target_action="append_section",
        ),
    )

    result = course_workflow.invoke(
        {
            "lesson": lesson,
            "course_package": package,
            "request": ChatRequest(message="续写一个新章节，如何解决过拟合？"),
        }
    )

    content_text = result["teacher_document"].content_text
    assert result["document_updated"] is False
    assert "可以用验证集、正则化和交叉验证来减少过拟合" not in content_text
    assert "例子拆解" not in content_text
    assert "参考答案" not in content_text
    assert "没有写入板书" in result["teacher_message"]


def test_workflow_preserves_existing_document_when_echoed_append_needs_model(monkeypatch: pytest.MonkeyPatch) -> None:
    package = build_initial_course_package()
    lesson = create_empty_lesson("测试2")
    package.lessons.append(lesson)
    lesson.board_document = build_document(
        title="几个比较基础的量化数学知识",
        content_html=(
            "<h1>几个比较基础的量化数学知识</h1>"
            "<p>这一页先讲蒙特卡洛方法、相关性和回归。</p>"
            "<h2>补充章节</h2>"
            "<p>续写一个新章节，如何解决过拟合</p>"
        ),
        document_id=lesson.board_document.id,
    )

    monkeypatch.setattr(
        openai_course_ai,
        "generate_document_edit",
        lambda **kwargs: None,
    )

    result = course_workflow.invoke(
        {
            "lesson": lesson,
            "course_package": package,
            "request": ChatRequest(message="续写一个新章节，如何解决过拟合？"),
        }
    )

    content_text = result["teacher_document"].content_text
    content_html = result["teacher_document"].content_html
    assert result["board_decision"].action == "append_section"
    assert result["document_updated"] is False
    assert "补充章节：如何解决过拟合" not in content_text
    assert "例子拆解" not in content_text
    assert content_html == lesson.board_document.content_html
    assert "没有写入板书" in result["teacher_message"]


def test_workflow_does_not_append_same_chapter_twice(monkeypatch: pytest.MonkeyPatch) -> None:
    package = build_initial_course_package()
    lesson = create_empty_lesson("测试2")
    package.lessons.append(lesson)
    lesson.board_document = build_document(
        title="几个比较基础的量化数学知识",
        content_html=(
            "<h1>几个比较基础的量化数学知识</h1>"
            "<p>这一页先讲蒙特卡洛方法、相关性和回归。</p>"
            "<h2>补充章节</h2>"
            "<p>续写一个新章节，如何解决过拟合</p>"
            "<h2>补充章节：如何解决过拟合</h2>"
            "<p>过拟合指模型在训练数据上表现很好，但一换到新数据、验证集或真实市场环境就明显变差。</p>"
            "<h3>一、先判断是不是过拟合</h3>"
            "<p>最直接的信号是训练集效果持续变好，但验证集或测试集效果停滞甚至变差。</p>"
        ),
        document_id=lesson.board_document.id,
    )

    monkeypatch.setattr(
        openai_course_ai,
        "generate_document_edit",
        lambda **kwargs: pytest.fail("duplicate append should not call Board AI again"),
    )

    result = course_workflow.invoke(
        {
            "lesson": lesson,
            "course_package": package,
            "request": ChatRequest(message="续写一个新章节，如何解决过拟合？"),
        }
    )

    assert result["board_decision"].action == "no_change"
    assert result["document_updated"] is False
    assert result["teacher_document"].content_text.count("补充章节：如何解决过拟合") == 1
    assert "如何解决过拟合" in result["teacher_message"]
    assert "续写一个新章节" not in result["teacher_message"]


def test_replace_selection_in_document_replaces_exact_block_without_nested_paragraphs() -> None:
    document = build_document(
        title="测试",
        content_html="<h1>测试</h1><p>原始题干</p><p>后续内容</p>",
    )

    next_document = replace_selection_in_document(
        document,
        selection_text="原始题干",
        replacement_text="原始题干\n\n补充说明：先圈出已知条件，再写解题步骤。",
    )

    assert "<p><p>" not in next_document.content_html
    assert "原始题干" in next_document.content_text
    assert "补充说明" in next_document.content_text
    assert "后续内容" in next_document.content_text


def test_replace_selection_in_document_preserves_rich_html_for_cross_block_selection() -> None:
    document = build_document(
        title="量化金融入门讲义",
        content_html=(
            "<h1>量化金融入门讲义：给第一次接触的人</h1>"
            "<p>开场正文。</p>"
            "<h2>九、量化金融和普通炒股有什么不同</h2>"
            "<p>上一节正文。</p>"
            "<h2>十、初学者最该优先掌握的学习顺序</h2>"
            "<p>如果你是零基础，建议按下面顺序学：</p>"
            "<ol>"
            "<li>先理解金融市场里有哪些资产；</li>"
            "<li>再理解收益率、风险、波动率、回撤这些基本指标；</li>"
            "<li>接着理解策略、回测、过拟合；</li>"
            "<li>然后学习最简单的统计知识，比如均值、方差、相关性；</li>"
            "<li>最后再进入编程和实际策略设计。</li>"
            "</ol>"
            "<p>很多人一开始就学复杂模型，结果概念混乱。</p>"
            "<h2>十一、一个入门练习</h2>"
            "<p>练习正文。</p>"
        ),
    )
    selection = (
        "十、初学者最该优先掌握的学习顺序 如果你是零基础，建议按下面顺序学： "
        "先理解金融市场里有哪些资产； 再理解收益率、风险、波动率、回撤这些基本指标； "
        "接着理解策略、回测、过拟合； 然后学习最简单的统计知识，比如均值、方差、相关性； "
        "最后再进入编程和实际策略设计。 很多人一开始就学复杂模型，结果概念混乱。"
    )

    updated = replace_selection_in_document(
        document,
        selection_text=selection,
        replacement_text="十、初学者最该优先掌握的学习顺序\n\n补充正文",
        replacement_html=(
            "<h2>十、初学者最该优先掌握的学习顺序</h2>"
            "<p>补充正文。</p>"
            "<h3>1. 为什么量化金融离不开一点数学</h3>"
            "<p>量化里的数学是为了把感觉变成可比较的东西。</p>"
        ),
    )

    assert updated.content_html.startswith("<h1>量化金融入门讲义")
    assert "<h2>九、量化金融和普通炒股有什么不同</h2>" in updated.content_html
    assert "<h2>十一、一个入门练习</h2>" in updated.content_html
    assert "<h3>1. 为什么量化金融离不开一点数学</h3>" in updated.content_html
    assert updated.content_html.count("<h2>十、初学者最该优先掌握的学习顺序</h2>") == 1
    assert "<p>量化金融入门讲义：给第一次接触的人</p>" not in updated.content_html


def test_workflow_generates_initial_scenario_document_for_blank_lesson() -> None:
    package = build_initial_course_package()
    lesson = create_empty_lesson("板书测试")
    package.lessons.append(lesson)

    result = course_workflow.invoke(
        {
            "lesson": lesson,
            "course_package": package,
            "request": ChatRequest(
                message=(
                    "嗨，我在练习客户访谈，当前水平是入门。"
                    "能不能给我生成一篇客户访谈的情景对话课文，"
                    "要覆盖开场、追问、确认需求和收尾复盘"
                ),
            ),
        }
    )

    assert result["board_decision"].action == "edit_board"
    assert result["document_updated"] is False
    assert result["teacher_document"].content_text == ""
    assert "没有写入板书" in result["teacher_message"]


def test_workflow_uses_fast_path_for_clear_generation_request(monkeypatch: pytest.MonkeyPatch) -> None:
    package = build_initial_course_package()
    lesson = create_empty_lesson("板书测试")
    package.lessons.append(lesson)

    monkeypatch.setattr(
        openai_course_ai,
        "assess_learning_requirements",
        lambda **kwargs: pytest.fail("clear generation request should skip PM AI assessment"),
    )
    monkeypatch.setattr(
        openai_course_ai,
        "generate_board_decision",
        lambda **kwargs: pytest.fail("clear generation request should skip board manager AI decision"),
    )
    monkeypatch.setattr(
        openai_course_ai,
        "generate_teacher_message",
        lambda **kwargs: pytest.fail("document generation with talk track should skip extra teacher AI call"),
    )
    monkeypatch.setattr(
        openai_course_ai,
        "generate_document_edit",
        lambda **kwargs: DocumentEditOutput(
            rationale="fast path",
            replacement_html="<h1>虚拟内存</h1><p>先理解地址空间，再理解页表和缺页异常。</p>",
            replacement_text="虚拟内存\n先理解地址空间，再理解页表和缺页异常。",
            teacher_talk_track="这节不用背定义，我们先抓住虚拟内存是在帮程序和物理内存之间做一层更灵活的映射。",
            replace_whole=True,
        ),
    )

    result = course_workflow.invoke(
        {
            "lesson": lesson,
            "course_package": package,
            "request": ChatRequest(
                message=(
                    "我是一名计算机专业学生，请给我生成一版虚拟内存的讲义，"
                    "重点讲地址空间、页表和缺页异常，并带一个入门例子"
                ),
            ),
        }
    )

    assert result["board_decision"].action == "edit_board"
    assert result["document_updated"] is True
    assert "虚拟内存" in result["teacher_document"].content_text
    assert "不用背定义" in result["teacher_message"]


def test_workflow_prompts_before_using_reference_when_one_candidate_is_clearly_best(tmp_path) -> None:
    package = build_initial_course_package()
    lesson = package.lessons[0]
    resource_path = tmp_path / "memory-notes.md"
    resource_path.write_text(
        "# 虚拟内存\n虚拟内存这一章主要解释地址空间、页表和缺页异常。\n\n## 地址转换\n地址转换依赖页表。",
        encoding="utf-8",
    )
    package.resources.append(build_resource_item(resource_path, "计算机系统导论笔记.md"))

    result = course_workflow.invoke(
        {
            "lesson": lesson,
            "course_package": package,
            "request": ChatRequest(message="请把虚拟内存这一节整理成更易懂的板书"),
        }
    )

    assert result["board_decision"].action == "await_reference_choice"
    assert result["reference_prompt"] is not None
    assert result["selected_reference"] is None
    assert result["reference_prompt"].chapter_title == "虚拟内存"
    assert "参考这章正文" in result["reference_prompt"].question


def test_workflow_compares_current_board_with_resource_directory_before_generating(tmp_path) -> None:
    package = build_initial_course_package()
    lesson = create_empty_lesson("资料整理课")
    lesson.board_document = build_document(
        title="案例复盘草稿",
        content_html="<h1>案例复盘草稿</h1><p>这节课准备讲背景、关键动作、影响和复盘问题。</p>",
    )
    package.lessons.append(lesson)
    resource_path = tmp_path / "case-notes.md"
    resource_path.write_text(
        "# 案例复盘\n案例复盘解释背景、关键动作、影响和复盘问题。",
        encoding="utf-8",
    )
    package.resources.append(build_resource_item(resource_path, "案例资料.md"))

    result = course_workflow.invoke(
        {
            "lesson": lesson,
            "course_package": package,
            "request": ChatRequest(message="请根据当前内容生成一版更完整的板书"),
        }
    )

    assert result["board_decision"].action == "await_reference_choice"
    assert result["reference_prompt"] is not None
    assert result["reference_prompt"].chapter_title == "案例复盘"


def test_workflow_prompts_for_reference_when_top_candidates_are_close(tmp_path) -> None:
    package = build_initial_course_package()
    lesson = package.lessons[0]
    first_path = tmp_path / "ring1.md"
    first_path.write_text(
        "# 环论\n## 环\n环这一节先解释加法群与乘法封闭。",
        encoding="utf-8",
    )
    second_path = tmp_path / "ring2.md"
    second_path.write_text(
        "# 抽象代数\n## 环\n环这一节重点讲单位元、零因子与理想。",
        encoding="utf-8",
    )
    package.resources.append(build_resource_item(first_path, "环论讲义.md"))
    package.resources.append(build_resource_item(second_path, "抽象代数讲义.md"))

    result = course_workflow.invoke(
        {
            "lesson": lesson,
            "course_package": package,
            "request": ChatRequest(message="请把环这一节整理成更易懂的板书"),
        }
    )

    assert result["board_decision"].action == "clarify_request"
    assert result["needs_clarification"] is True
    assert result["reference_prompt"] is None
    assert "哪一份资料" in result["teacher_message"]


def test_workflow_uses_selected_reference_after_user_confirms(tmp_path) -> None:
    package = build_initial_course_package()
    lesson = package.lessons[0]
    resource_path = tmp_path / "memory-notes.md"
    resource_path.write_text(
        "# 虚拟内存\n虚拟内存用于把程序看到的地址空间和物理内存解耦。\n\n## 页表\n页表记录虚拟页到物理页的映射。",
        encoding="utf-8",
    )
    resource = build_resource_item(resource_path, "计算机系统导论笔记.md")
    package.resources.append(resource)
    chapter = resource.outline[0]

    result = course_workflow.invoke(
        {
            "lesson": lesson,
            "course_package": package,
            "request": ChatRequest(
                message="请把虚拟内存这一节整理成更易懂的板书",
                resource_reference_action="confirm",
                resource_reference_resource_id=resource.id,
                resource_reference_chapter_id=chapter.id,
            ),
        }
    )

    assert result["board_decision"].action == "edit_board"
    assert result["selected_reference"] is not None
    assert result["selected_reference"].chunks
    assert "虚拟内存用于把程序看到的地址空间和物理内存解耦" in result["selected_reference"].full_text
    assert result["document_updated"] is True
    assert result["board_teaching_guide"] is not None


def test_build_resource_item_extracts_image_ocr_text(monkeypatch: pytest.MonkeyPatch, tmp_path) -> None:
    image_path = tmp_path / "math-problem.png"
    image_path.write_bytes(b"fake-image")
    extracted = "设正整数 n >= 2\n求证该数列满足条件"
    monkeypatch.setattr("app.services.resource_library.extract_image_text", lambda path: extracted)

    resource = build_resource_item(image_path, "数学题.png")

    assert resource.extracted_text_available is True
    assert resource.text_content == extracted
    assert resource.outline
    reference = extract_reference_context(resource, resource.outline[0].id, user_query="正整数数列")
    assert reference is not None
    assert "正整数" in reference.full_text


def test_build_resource_item_extracts_epub_outline_and_reference_context(tmp_path) -> None:
    resource_path = tmp_path / "macro.epub"
    chapter_html = """
<?xml version="1.0" encoding="utf-8"?>
<html xmlns="http://www.w3.org/1999/xhtml">
  <head><title>第一章 宏观经济学导论</title></head>
  <body>
    <h1>第一章 宏观经济学导论</h1>
    <p>宏观经济学研究整体经济运行，关注国内生产总值、通货膨胀、失业率和经济增长。</p>
    <p>学习这一章时，要先理解总量指标为什么能把千千万万个家庭、企业和政府行为连接起来。</p>
    <h2>1.1 国内生产总值</h2>
    <p>国内生产总值 GDP 衡量一定时期内一个经济体生产的最终产品和服务的市场价值。</p>
  </body>
</html>
""".strip()
    with ZipFile(resource_path, "w") as archive:
        archive.writestr(
            "META-INF/container.xml",
            """
<container version="1.0" xmlns="urn:oasis:names:tc:opendocument:xmlns:container">
  <rootfiles>
    <rootfile full-path="OEBPS/content.opf" media-type="application/oebps-package+xml"/>
  </rootfiles>
</container>
""".strip(),
        )
        archive.writestr(
            "OEBPS/content.opf",
            """
<package xmlns="http://www.idpf.org/2007/opf" version="3.0">
  <manifest>
    <item id="chapter1" href="chapter1.xhtml" media-type="application/xhtml+xml"/>
  </manifest>
  <spine>
    <itemref idref="chapter1"/>
  </spine>
</package>
""".strip(),
        )
        archive.writestr("OEBPS/chapter1.xhtml", chapter_html)

    resource = build_resource_item(resource_path, "曼昆宏观经济学.epub")

    assert resource.extracted_text_available is True
    assert resource.outline
    assert resource.outline[0].title == "第一章宏观经济学导论"
    reference = extract_reference_context(resource, resource.outline[0].id, user_query="讲解第一章")
    assert reference is not None
    assert reference.chapter_title == "第一章宏观经济学导论"
    assert "国内生产总值GDP" in reference.full_text
    assert "曼昆宏观经济学.epub" not in reference.chunks[0].excerpt


def test_epub_reference_prefers_body_chapter_over_toc_duplicate(tmp_path) -> None:
    package = build_initial_course_package()
    lesson = create_empty_lesson("测试11")
    package.lessons.append(lesson)
    resource_path = tmp_path / "macro-duplicate.epub"
    toc_html = """
<?xml version="1.0" encoding="utf-8"?>
<html xmlns="http://www.w3.org/1999/xhtml">
  <head><title>目录</title></head>
  <body>
    <h1>第2章宏观经济学的数据</h1>
    <p>【学习精要】</p>
    <p>【习题解析】</p>
    <p>【补充训练】</p>
  </body>
</html>
""".strip()
    chapter_html = """
<?xml version="1.0" encoding="utf-8"?>
<html xmlns="http://www.w3.org/1999/xhtml">
  <head><title>第2章宏观经济学的数据</title></head>
  <body>
    <h1>第2章宏观经济学的数据</h1>
    <p>【学习精要】国内生产总值 GDP 衡量一定时期内一个经济体生产的最终产品和服务的市场价值。</p>
    <p>本章还要区分名义 GDP、实际 GDP、GDP 平减指数、消费者价格指数 CPI 和失业率。</p>
    <p>这些指标共同回答宏观经济运行到底如何被计量，以及为什么同一个经济体会有产出、价格和就业三条观察线索。</p>
  </body>
</html>
""".strip()
    with ZipFile(resource_path, "w") as archive:
        archive.writestr(
            "META-INF/container.xml",
            """
<container version="1.0" xmlns="urn:oasis:names:tc:opendocument:xmlns:container">
  <rootfiles>
    <rootfile full-path="OEBPS/content.opf" media-type="application/oebps-package+xml"/>
  </rootfiles>
</container>
""".strip(),
        )
        archive.writestr(
            "OEBPS/content.opf",
            """
<package xmlns="http://www.idpf.org/2007/opf" version="3.0">
  <manifest>
    <item id="toc" href="toc.xhtml" media-type="application/xhtml+xml"/>
    <item id="chapter2" href="chapter2.xhtml" media-type="application/xhtml+xml"/>
  </manifest>
  <spine>
    <itemref idref="toc"/>
    <itemref idref="chapter2"/>
  </spine>
</package>
""".strip(),
        )
        archive.writestr("OEBPS/toc.xhtml", toc_html)
        archive.writestr("OEBPS/chapter2.xhtml", chapter_html)

    resource = build_resource_item(resource_path, "曼昆宏观经济学.epub")
    package.resources.append(resource)

    matches = match_resources(
        package,
        lesson,
        ChatRequest(message="讲解第二章内容"),
        effective_requirements(lesson),
    )

    assert matches
    reference = extract_reference_context(resource, matches[0].chapter_id, user_query="讲解第二章内容")
    assert reference is not None
    assert len(reference.full_text) > 120
    assert "国内生产总值GDP" in reference.full_text
    assert "GDP平减指数" in reference.full_text
    assert "【学习精要】\n【习题解析】\n【补充训练】" not in reference.full_text


def test_build_resource_item_extracts_docx_outline_and_reference_context(tmp_path) -> None:
    resource_path = tmp_path / "math_solutions.docx"
    document = DocxDocument()
    document.add_heading("函数压轴题答案", level=0)
    document.add_heading("题目分析", level=1)
    document.add_paragraph("先判断单调性，再看零点分布。")
    document.add_heading("参考答案", level=1)
    document.add_paragraph("设 f(x)=x^2-2x+1，然后分类讨论。")
    document.save(resource_path)

    resource = build_resource_item(resource_path, "math_solutions.docx")

    titles = [chapter.title for chapter in resource.outline]
    assert "题目分析" in titles
    chapter = next(chapter for chapter in resource.outline if chapter.title == "题目分析")
    reference = extract_reference_context(resource, chapter.id, user_query="单调性")
    assert reference is not None
    assert "先判断单调性" in reference.full_text


def _write_text_pdf(path, pages: list[list[str]]) -> None:
    pdf = canvas.Canvas(str(path), pagesize=letter)
    for lines in pages:
        y = 760
        for line in lines:
            pdf.drawString(72, y, line)
            y -= 18
        pdf.showPage()
    pdf.save()


def test_build_resource_item_extracts_pdf_toc_page_and_uses_offset_page_candidate(tmp_path) -> None:
    resource_path = tmp_path / "toc-offset.pdf"
    _write_text_pdf(
        resource_path,
        [
            ["Cover"],
            ["Contents", "Chapter 1 Overview 3", "Chapter 2 Details 6"],
            ["Preface", "This is not the chapter."],
            ["Front matter continues."],
            ["Chapter 1 Overview", "This chapter explains patterns, features, and classifiers."],
            ["Chapter 2 Details", "This chapter goes deeper."],
        ],
    )

    resource = build_resource_item(resource_path, "toc-offset.pdf")

    assert resource.outline
    chapter = resource.outline[0]
    assert chapter.title == "Chapter 1 Overview"
    assert chapter.page_start == 5
    assert "目录页 2" in chapter.summary
    reference = extract_reference_context(resource, chapter.id, user_query="patterns")
    assert reference is not None
    assert "patterns, features, and classifiers" in reference.full_text


def test_build_resource_item_uses_pdf_outline_ranges_by_hierarchy(tmp_path) -> None:
    resource_path = tmp_path / "csapp.pdf"
    writer = PdfWriter()
    for _ in range(10):
        writer.add_blank_page(width=200, height=200)
    chapter_one = writer.add_outline_item("Chapter 1", 0)
    writer.add_outline_item("Section 1.1", 1, parent=chapter_one)
    writer.add_outline_item("Section 1.2", 4, parent=chapter_one)
    chapter_two = writer.add_outline_item("Chapter 2", 6)
    writer.add_outline_item("Section 2.1", 7, parent=chapter_two)
    with resource_path.open("wb") as target:
        writer.write(target)

    resource = build_resource_item(resource_path, "csapp.pdf")

    chapter_one_outline = next(chapter for chapter in resource.outline if chapter.title == "Chapter 1")
    section_one_one = next(chapter for chapter in resource.outline if chapter.title == "Section 1.1")
    section_one_two = next(chapter for chapter in resource.outline if chapter.title == "Section 1.2")
    chapter_two_outline = next(chapter for chapter in resource.outline if chapter.title == "Chapter 2")

    assert chapter_one_outline.page_start == 1
    assert chapter_one_outline.page_end == 6
    assert section_one_one.page_start == 2
    assert section_one_one.page_end == 4
    assert section_one_two.page_start == 5
    assert section_one_two.page_end == 6
    assert chapter_two_outline.page_start == 7
    assert chapter_two_outline.page_end == 10


def test_build_resource_item_handles_same_page_outline_entries(tmp_path) -> None:
    resource_path = tmp_path / "same-page-outline.pdf"
    writer = PdfWriter()
    for _ in range(10):
        writer.add_blank_page(width=200, height=200)
    chapter_one = writer.add_outline_item("Chapter 1", 0)
    writer.add_outline_item("Section 1 overview", 1, parent=chapter_one)
    writer.add_outline_item("Section 1.1", 1, parent=chapter_one)
    writer.add_outline_item("Section 1.2", 4, parent=chapter_one)
    writer.add_outline_item("Chapter 2", 6)
    with resource_path.open("wb") as target:
        writer.write(target)

    resource = build_resource_item(resource_path, "same-page-outline.pdf")

    section_overview = next(chapter for chapter in resource.outline if chapter.title == "Section 1 overview")
    section_one_one = next(chapter for chapter in resource.outline if chapter.title == "Section 1.1")

    assert section_overview.page_start == 2
    assert section_overview.page_end == 4
    assert section_one_one.page_start == 2
    assert section_one_one.page_end == 4


def test_keywords_from_text_filters_common_english_function_words() -> None:
    keywords = _keywords_from_text("Virtual memory is the way a system maps virtual addresses to physical memory.")

    assert "virtual" in keywords
    assert "memory" in keywords
    assert "to" not in keywords
    assert "is" not in keywords


def test_match_resources_uses_current_board_scope_for_directory_matching(tmp_path) -> None:
    lesson = create_lesson("虚拟内存")
    package = build_initial_course_package()
    package.lessons.append(lesson)
    resource_path = tmp_path / "memory-notes.md"
    resource_path.write_text(
        "# 虚拟内存\n虚拟内存这一章主要解释地址空间、页表和缺页异常。\n\n## 地址转换\n地址转换依赖页表。",
        encoding="utf-8",
    )
    package.resources.append(build_resource_item(resource_path, "计算机系统导论笔记.md"))

    matches = match_resources(
        package,
        lesson,
        ChatRequest(message="把这一节整理得更易懂一点"),
        effective_requirements(lesson),
    )

    assert matches
    assert matches[0].chapter_title == "虚拟内存"


def test_match_resources_uses_requirement_theme_for_cross_language_directory_matching(tmp_path) -> None:
    lesson = create_lesson("虚拟内存")
    package = build_initial_course_package()
    package.lessons.append(lesson)
    resource_path = tmp_path / "csapp.md"
    resource_path.write_text(
        "# Virtual Memory\nVirtual memory explains address translation, page tables, TLBs, and page faults.",
        encoding="utf-8",
    )
    package.resources.append(build_resource_item(resource_path, "CSAPP notes.md"))

    requirements = effective_requirements(lesson)
    requirements.theme = "虚拟内存（板书版，适合 15–25 分钟讲解）"

    matches = match_resources(
        package,
        lesson,
        ChatRequest(message="请把这一节整理成更易懂的板书"),
        requirements,
    )

    assert matches
    assert matches[0].chapter_title == "Virtual Memory"
    assert matches[0].is_high_overlap is True


def test_match_resources_prioritizes_current_request_over_context_noise(tmp_path) -> None:
    lesson = create_lesson("进程与系统")
    lesson.board_document = build_document(
        title="进程与系统",
        content_html="<h1>进程与系统</h1><p>进程、进程、进程、并发、进程调度。</p>",
    )
    package = build_initial_course_package()
    package.lessons.append(lesson)

    vm_path = tmp_path / "vm.md"
    vm_path.write_text("# Virtual Memory\nAddress translation and page tables.", encoding="utf-8")
    proc_path = tmp_path / "proc.md"
    proc_path.write_text("# Concurrent Programming with Processes\nProcesses and concurrency.", encoding="utf-8")
    package.resources.append(build_resource_item(vm_path, "CSAPP VM.md"))
    package.resources.append(build_resource_item(proc_path, "CSAPP PROC.md"))

    requirements = effective_requirements(lesson)
    requirements.theme = "进程与系统"
    requirements.learning_goal = "理解进程调度，但当前这次请求是把虚拟内存这一节整理得更易懂。"

    matches = match_resources(
        package,
        lesson,
        ChatRequest(message="请把虚拟内存这一节整理成更易懂的板书"),
        requirements,
    )

    assert matches
    assert matches[0].chapter_title == "Virtual Memory"


def test_match_resources_understands_numeric_chapter_and_section_reference(tmp_path) -> None:
    lesson = create_empty_lesson("教材测试")
    package = build_initial_course_package()
    package.lessons.append(lesson)
    resource_path = tmp_path / "structured-notes.md"
    resource_path.write_text(
        "\n".join(
            [
                "# 第一章",
                "## 第一节",
                "内容 1.1",
                "## 第二节",
                "内容 1.2",
                "# 第二章",
                "## 第一节",
                "内容 2.1",
                "## 第二节",
                "内容 2.2",
                "# 第三章",
                "## 第一节",
                "内容 3.1",
                "## 第二节",
                "内容 3.2",
                "# 第四章",
                "## 第一节",
                "内容 4.1",
                "## 第二节",
                "内容 4.2",
                "# 第五章",
                "## 第一节",
                "内容 5.1",
                "## 第二节",
                "内容 5.2",
            ]
        ),
        encoding="utf-8",
    )
    package.resources.append(build_resource_item(resource_path, "structured-notes.md"))

    matches = match_resources(
        package,
        lesson,
        ChatRequest(message="请直接讲教材里的第5章第2节"),
        effective_requirements(lesson),
    )

    assert matches
    assert matches[0].chapter_title == "第二节"

    chinese_matches = match_resources(
        package,
        lesson,
        ChatRequest(message="请直接讲教材里的第五章第二节"),
        effective_requirements(lesson),
    )

    assert chinese_matches
    assert chinese_matches[0].chapter_title == "第二节"


def test_docx_import_export_roundtrip(tmp_path) -> None:
    document = build_document(
        title="法国咖啡厅点餐情景对话（含过去将来时）",
        content_html="""
<h1>法国咖啡厅点餐情景对话（含过去将来时）</h1>
<p>完整双语对话。</p>
<p>Je pensais que je prendrais seulement un café.</p>
        """.strip(),
    )
    target = tmp_path / "lesson.docx"

    export_docx(document, target)
    imported = import_docx(target)

    assert target.exists()
    assert "完整双语对话" in imported.content_text
    assert "Je pensais que je prendrais" in imported.content_text


def test_docx_export_writes_math_as_office_math(tmp_path) -> None:
    document = build_document(
        title="公式导出测试",
        content_html="""
<p>求 lim_{x→0} (sin x)/x。</p>
<p>lim_{x→a} f(x)/g(x) = lim_{x→a} f'(x)/g'(x)</p>
<p><span data-type="inline-math" data-latex="x^2"></span> 是平方。</p>
<p><span data-type="inline-math" data-latex="f'\\frac{x}{g}'(x)"></span> 是历史兼容。</p>
<div data-type="block-math" data-latex="\\frac{\\ln x}{1/x}"></div>
        """.strip(),
    )
    target = tmp_path / "math-export.docx"

    export_docx(document, target)
    with ZipFile(target) as package:
        document_xml = package.read("word/document.xml").decode("utf-8")
        media_files = [name for name in package.namelist() if name.startswith("word/media/") and name.endswith(".png")]

    assert media_files == []
    assert "<w:drawing>" not in document_xml
    assert "r:embed" not in document_xml
    assert "<m:oMath>" in document_xml
    assert document_xml.count("<m:f>") >= 3
    assert "<m:limLow>" in document_xml
    assert "<m:sSup>" in document_xml
    assert "lim_{" not in document_xml
    assert "\\frac" not in document_xml
    assert "<m:t>f'(x)</m:t>" in document_xml
    assert "<m:t>g'(x)</m:t>" in document_xml
    assert "<m:t>ln</m:t>" in document_xml


def test_replace_selection_preserves_page_settings() -> None:
    document = build_document(
        title="页面设置保留测试",
        content_text="第一段\n第二段",
        page_settings={
            "orientation": "landscape",
            "margin_preset": "narrow",
            "show_page_number": True,
            "header_text": "课堂讲义",
            "footer_text": "开放课堂",
        },
    )

    updated = replace_selection_in_document(
        document,
        selection_text="第二段",
        replacement_text="更新后的第二段",
    )

    assert updated.content_text.endswith("更新后的第二段")
    assert updated.page_settings.orientation == "landscape"
    assert updated.page_settings.margin_preset == "narrow"
    assert updated.page_settings.show_page_number is True
    assert updated.page_settings.header_text == "课堂讲义"
    assert updated.page_settings.footer_text == "开放课堂"


def test_docx_export_applies_basic_page_settings(tmp_path) -> None:
    document = build_document(
        title="导出版式测试",
        content_text="讲义正文",
        page_settings={
            "orientation": "landscape",
            "margin_preset": "wide",
            "header_text": "页眉示例",
            "footer_text": "页脚示例",
        },
    )
    target = tmp_path / "page-settings.docx"

    export_docx(document, target)
    exported = DocxDocument(target)
    section = exported.sections[0]

    assert section.orientation == WD_ORIENT.LANDSCAPE
    assert section.page_width > section.page_height
    assert exported.sections[0].header.paragraphs[0].text == "页眉示例"
    assert exported.sections[0].footer.paragraphs[0].text == "页脚示例"


def test_docx_export_keeps_page_break_nodes(tmp_path) -> None:
    document = build_document(
        title="分页测试",
        content_html='<p>第一页</p><div data-type="page-break" class="word-editor__page-break"></div><p>第二页</p>',
    )
    target = tmp_path / "page-break.docx"

    export_docx(document, target)

    with ZipFile(target) as archive:
        document_xml = archive.read("word/document.xml").decode("utf-8")
    assert '<w:br w:type="page"/>' in document_xml
