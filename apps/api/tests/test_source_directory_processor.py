from __future__ import annotations

import hashlib
import json
import sqlite3
import threading
import time
import zipfile
from concurrent.futures import ThreadPoolExecutor
from dataclasses import replace
from pathlib import Path
from types import SimpleNamespace

import pytest

from app.models import AIModelSelection, CoursePackage, SourceChapter, SourceIngestionRecord, SourceRange
from app.services import workspace_state
from app.services.source_directory_extractor import (
    DirectoryCandidate,
    DirectoryExtraction,
    _PdfHeadingLine,
    _PdfTocCandidate,
    _close_numeric_ranges,
    _map_pdf_toc_nodes,
    _parse_native_toc_text,
    _pdf_page_has_toc_heading,
    extract_directory,
)
from app.services.source_directory_processor import (
    CodexDirectoryNormalizer,
    DirectoryBatchDecision,
    DirectoryNodeDecision,
    DirectoryNormalizationResult,
    SourceDirectoryProcessingError,
    SourceDirectoryProcessor,
    _catalog_quality,
    _reclose_normalized_ranges,
    _validate_chapters,
)
from app.services.pdf_toc_parser import PdfTocExtraction, PdfTocNode
from app.services.source_evidence_store import SourceEvidenceStore
from app.services.source_ingestion_jobs import SourceIngestionJobStore
from app.services.source_ingestion_service import SourceIngestionError, SourceIngestionService
from app.services.source_structure_store import SourceStructureStore


class PassthroughNormalizer:
    def normalize(self, *, record, candidates, selection):
        return DirectoryNormalizationResult(
            candidates=tuple(candidates),
            turn_count=1 if candidates else 0,
            metadata={"test_adapter": "passthrough"},
        )


class FailingNormalizer:
    def normalize(self, *, record, candidates, selection):
        raise RuntimeError("catalog model failed")


def _record(path: Path, *, source_id: str = "source_directory") -> SourceIngestionRecord:
    content_hash = hashlib.sha256(path.read_bytes()).hexdigest()
    return SourceIngestionRecord(
        id=source_id,
        owner_user_id="user_directory",
        package_id="course_directory",
        title=path.stem,
        source_type="local_file",
        file_name=path.name,
        mime_type="application/pdf" if path.suffix == ".pdf" else "text/markdown",
        size_bytes=path.stat().st_size,
        status="parsing",
        metadata={"local_source_path": str(path), "content_hash": content_hash},
    )


def _model() -> AIModelSelection:
    return AIModelSelection(provider="openai_codex", model="catalog-test-model")


def _alternate_model() -> AIModelSelection:
    return AIModelSelection(provider="openai_codex", model="catalog-retry-model")


def _write_reordered_pptx(path: Path) -> None:
    presentation = """<?xml version="1.0" encoding="UTF-8"?>
<p:presentation xmlns:p="http://schemas.openxmlformats.org/presentationml/2006/main"
 xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">
  <p:sldIdLst>
    <p:sldId id="256" r:id="rId9"/>
    <p:sldId id="257" r:id="rId2"/>
  </p:sldIdLst>
</p:presentation>"""
    relationships = """<?xml version="1.0" encoding="UTF-8"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="rId2" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/slide" Target="slides/slide2.xml"/>
  <Relationship Id="rId9" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/slide" Target="slides/slide9.xml"/>
</Relationships>"""

    def slide(title: str) -> str:
        return f"""<?xml version="1.0" encoding="UTF-8"?>
<p:sld xmlns:p="http://schemas.openxmlformats.org/presentationml/2006/main"
 xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main">
  <p:cSld><p:spTree><p:sp><p:txBody><a:p><a:r><a:t>{title}</a:t></a:r></a:p></p:txBody></p:sp></p:spTree></p:cSld>
</p:sld>"""

    with zipfile.ZipFile(path, "w") as archive:
        archive.writestr("ppt/presentation.xml", presentation)
        archive.writestr("ppt/_rels/presentation.xml.rels", relationships)
        archive.writestr("ppt/slides/slide9.xml", slide("Playback first"))
        archive.writestr("ppt/slides/slide2.xml", slide("Playback second"))
        archive.writestr("ppt/slides/slide1.xml", slide("Orphan slide"))


def _write_epub_with_navigation(
    path: Path,
    *,
    navigation: str,
    document: str,
    document_name: str = "s1.xhtml",
) -> None:
    with zipfile.ZipFile(path, "w") as archive:
        archive.writestr(
            "META-INF/container.xml",
            "<container><rootfiles><rootfile full-path='OPS/content.opf'/></rootfiles></container>",
        )
        archive.writestr(
            "OPS/content.opf",
            "<package><manifest>"
            f"<item id='nav' href='nav.xhtml'/><item id='s1' href='{document_name}'/>"
            "</manifest><spine><itemref idref='s1'/></spine></package>",
        )
        archive.writestr("OPS/nav.xhtml", f"<nav><ol>{navigation}</ol></nav>")
        archive.writestr(f"OPS/{document_name}", document)


def test_markdown_extractor_keeps_only_heading_lines_and_inclusive_ranges(tmp_path: Path) -> None:
    path = tmp_path / "notes.md"
    path.write_text(
        "# First\n\nBODY_SECRET_ALPHA\n\n## Child\n\nBODY_SECRET_BETA\n\n# Second\nTail",
        encoding="utf-8",
    )

    extraction = extract_directory(_record(path), path)

    assert [candidate.title for candidate in extraction.candidates] == ["First", "Child", "Second"]
    assert all("BODY_SECRET" not in item.evidence[0].excerpt for item in extraction.candidates)
    assert extraction.candidates[0].source_range == SourceRange(
        kind="text_lines",
        start=1,
        end=8,
        display_label="Lines 1-8",
        metadata={"index_base": 1},
    )
    assert extraction.candidates[1].source_range.end == 8
    assert extraction.candidates[2].source_range.end == 10
    assert extraction.metadata["body_text_extracted"] is False


def test_docx_extractor_uses_document_xml_paragraph_order_including_tables(
    tmp_path: Path,
) -> None:
    from docx import Document

    path = tmp_path / "table-headings.docx"
    document = Document()
    document.add_paragraph("Preface outside chapter")
    table = document.add_table(rows=1, cols=1)
    table_heading = table.cell(0, 0).paragraphs[0]
    table_heading.text = "Table chapter"
    table_heading.style = "Heading 1"
    table.cell(0, 0).add_paragraph("TABLE_CELL_BODY_EVIDENCE")
    document.add_paragraph("AFTER_TABLE_BODY_EVIDENCE")
    document.add_heading("Next chapter", level=1)
    document.add_paragraph("NEXT_CHAPTER_BODY")
    document.save(path)

    extraction = extract_directory(_record(path), path)

    assert [candidate.title for candidate in extraction.candidates] == [
        "Table chapter",
        "Next chapter",
    ]
    assert extraction.candidates[0].source_range == SourceRange(
        kind="docx_paragraphs",
        start=1,
        end=3,
        display_label="Paragraphs 2-4",
        metadata={"index_base": 0},
    )
    assert extraction.candidates[1].source_range is not None
    assert extraction.candidates[1].source_range.start == 4
    assert extraction.candidates[1].source_range.end == 5
    assert extraction.metadata["paragraph_count"] == 6
    assert extraction.metadata["paragraph_sequence"] == "word_document_xml_v1"


def test_pptx_extractor_uses_relationship_backed_playback_order(tmp_path: Path) -> None:
    path = tmp_path / "reordered.pptx"
    _write_reordered_pptx(path)

    extraction = extract_directory(_record(path), path)

    assert [candidate.title for candidate in extraction.candidates] == [
        "Playback first",
        "Playback second",
    ]
    assert [candidate.source_range.start for candidate in extraction.candidates] == [1, 2]
    assert [candidate.source_range.end for candidate in extraction.candidates] == [1, 2]
    assert extraction.page_count == 2
    assert extraction.metadata["slide_sequence"] == "presentation_sldId_v1"


def test_pdf_outline_uses_physical_inclusive_pages_without_page_scan(tmp_path: Path) -> None:
    from pypdf import PdfWriter

    path = tmp_path / "outline.pdf"
    writer = PdfWriter()
    for _ in range(4):
        writer.add_blank_page(width=612, height=792)
    first = writer.add_outline_item("1 First", 0)
    writer.add_outline_item("1.1 Child", 1, parent=first)
    writer.add_outline_item("2 Second", 2)
    with path.open("wb") as handle:
        writer.write(handle)

    extraction = extract_directory(_record(path), path)

    assert extraction.inspected_page_count == 0
    assert extraction.ocr_page_count == 0
    assert [candidate.source_range.start for candidate in extraction.candidates] == [1, 2, 3]
    assert [candidate.source_range.end for candidate in extraction.candidates] == [2, 2, 4]
    assert all(candidate.mapping_status == "verified" for candidate in extraction.candidates)
    assert extraction.metadata["directory_source"] == "native_outline"
    assert extraction.metadata["body_text_extracted"] is False


def test_numeric_page_bookmarks_do_not_masquerade_as_chapter_outline(tmp_path: Path) -> None:
    from pypdf import PdfWriter

    path = tmp_path / "numeric-bookmarks.pdf"
    writer = PdfWriter()
    for page_index in range(12):
        writer.add_blank_page(width=612, height=792)
        writer.add_outline_item(str(page_index + 1), page_index)
    with path.open("wb") as handle:
        writer.write(handle)

    extraction = extract_directory(_record(path), path)

    assert extraction.candidates == ()
    assert extraction.metadata["directory_source"] == "heading_regions"
    assert extraction.inspected_page_count == 12


def test_pdf_toc_uses_numeric_page_navigation_without_body_scan(tmp_path: Path) -> None:
    import fitz
    from pypdf import PdfReader, PdfWriter

    native_path = tmp_path / "numeric-navigation-native.pdf"
    document = fitz.open()
    document.new_page(width=612, height=792)
    toc_page = document.new_page(width=612, height=792)
    toc_page.insert_text((72, 72), "Contents", fontsize=18)
    toc_page.insert_text((72, 110), "1 Alpha Section ........ 1", fontsize=12)
    toc_page.insert_text((72, 135), "2 Beta Section ........ 3", fontsize=12)
    alpha_page = document.new_page(width=612, height=792)
    alpha_page.insert_text((72, 72), "1 Alpha Section", fontsize=18)
    document.new_page(width=612, height=792)
    beta_page = document.new_page(width=612, height=792)
    beta_page.insert_text((72, 72), "2 Beta Section", fontsize=18)
    document.new_page(width=612, height=792)
    document.save(native_path)
    document.close()

    path = tmp_path / "numeric-navigation.pdf"
    writer = PdfWriter()
    for page in PdfReader(native_path).pages:
        writer.add_page(page)
    writer.add_outline_item("Contents", 1)
    for printed_page, physical_index in enumerate(range(2, 6), start=1):
        writer.add_outline_item(str(printed_page), physical_index)
    with path.open("wb") as handle:
        writer.write(handle)

    extraction = extract_directory(_record(path), path)

    assert [candidate.title for candidate in extraction.candidates] == [
        "1 Alpha Section",
        "2 Beta Section",
    ]
    assert [candidate.source_range.start for candidate in extraction.candidates] == [3, 5]
    assert [candidate.source_range.end for candidate in extraction.candidates] == [4, 6]
    assert extraction.inspected_page_count == 1
    assert extraction.ocr_page_count == 0
    assert extraction.metadata["toc_page_start"] == 2
    assert extraction.metadata["toc_page_end"] == 2
    assert extraction.metadata["printed_page_navigation"] is True
    assert extraction.metadata["heading_region_scan"] is False
    assert all(candidate.mapping_status == "verified" for candidate in extraction.candidates)


def test_image_only_pdf_outline_bounds_toc_ocr_before_numeric_page_navigation(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    import fitz
    from PIL import Image
    from pypdf import PdfReader, PdfWriter

    image_path = tmp_path / "page.png"
    Image.new("RGB", (64, 80), color="white").save(image_path)

    native_path = tmp_path / "image-only-native.pdf"
    document = fitz.open()
    for _ in range(26):
        page = document.new_page(width=612, height=792)
        page.insert_image(page.rect, filename=str(image_path))
    document.save(native_path)
    document.close()

    path = tmp_path / "image-only-navigation.pdf"
    writer = PdfWriter()
    for page in PdfReader(native_path).pages:
        writer.add_page(page)
    writer.add_outline_item("Contents", 20)
    for printed_page, physical_index in enumerate(range(21, 25), start=1):
        writer.add_outline_item(str(printed_page), physical_index)
    with path.open("wb") as handle:
        writer.write(handle)

    with fitz.open(path) as image_only_document:
        assert all(not page.get_text("text").strip() for page in image_only_document)

    ocr_calls: list[tuple[int, int]] = []

    def fake_toc_ocr(
        _path: Path,
        *,
        page_start: int,
        page_end: int,
    ) -> PdfTocExtraction:
        ocr_calls.append((page_start, page_end))
        return PdfTocExtraction(
            nodes=[
                PdfTocNode(
                    title="Chapter 1 Alpha",
                    number="1",
                    level=1,
                    printed_page=1,
                    toc_page=21,
                    confidence=0.9,
                ),
                PdfTocNode(
                    title="Chapter 2 Beta",
                    number="2",
                    level=1,
                    printed_page=3,
                    toc_page=21,
                    confidence=0.9,
                ),
            ],
            toc_page_start=page_start,
            toc_page_end=page_end,
        )

    monkeypatch.setattr(
        "app.services.source_directory_extractor.extract_pdf_toc_from_range",
        fake_toc_ocr,
    )

    extraction = extract_directory(_record(path), path)

    assert ocr_calls == [(21, 21)]
    assert extraction.inspected_page_count == 1
    assert extraction.ocr_page_count == 1
    assert extraction.metadata["toc_page_start"] == 21
    assert extraction.metadata["toc_page_end"] == 21
    assert extraction.metadata["printed_page_navigation"] is True
    assert extraction.metadata["heading_region_scan"] is False
    assert [candidate.source_range.start for candidate in extraction.candidates] == [22, 24]
    assert all(
        candidate.evidence[0].metadata["offset_support"] == 4
        for candidate in extraction.candidates
    )


def test_pdf_printed_page_offset_does_not_extrapolate_beyond_observed_navigation() -> None:
    nodes = [
        _PdfTocCandidate("Observed chapter", "1", 1, 1, 2, 0.9, "toc"),
        _PdfTocCandidate("Unobserved chapter", "", 1, 100, 2, 0.9, "toc"),
        _PdfTocCandidate("Heading matched chapter", "", 1, 101, 2, 0.9, "toc"),
    ]
    heading_cache = {
        150: [_PdfHeadingLine("Heading matched chapter", 150, 0.05, 18, "text")],
    }

    mapped = _map_pdf_toc_nodes(
        nodes,
        heading_cache,
        page_count=200,
        printed_page_map={1: 10, 2: 11, 3: 12, 4: 13},
    )

    assert mapped[0].mapping_status == "verified"
    assert mapped[0].source_range is not None
    assert mapped[0].source_range.start == 10
    assert mapped[0].evidence[0].metadata["mapping_method"] == "native_printed_page_navigation"
    assert mapped[1].mapping_status == "unmapped"
    assert mapped[1].source_range is None
    assert mapped[1].evidence[0].metadata["mapping_method"] == "unmapped"
    assert mapped[2].mapping_status == "verified"
    assert mapped[2].source_range is not None
    assert mapped[2].source_range.start == 150
    assert mapped[2].evidence[0].metadata["mapping_method"] == "body_heading_region_match"


def test_toc_layout_requests_trailing_column_pass_and_parses_leading_leaders(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    from app.services import pdf_toc_parser
    from app.services.image_ocr import OCRLineLayout, OCRPageLayout

    calls: list[bool] = []

    def fake_layouts(
        _path: Path,
        *,
        page_start: int,
        page_end: int,
        max_pages: int,
        trailing_column_pass: bool = False,
    ) -> list[OCRPageLayout]:
        assert (page_start, page_end, max_pages) == (21, 21, 24)
        calls.append(trailing_column_pass)
        return [
            OCRPageLayout(
                page_no=21,
                lines=[
                    OCRLineLayout(
                        "第2章 通用目录标题",
                        x=0.07,
                        y=0.65,
                        width=0.28,
                        height=0.018,
                    ),
                    OCRLineLayout(
                        "⋯143",
                        x=0.79,
                        y=0.65,
                        width=0.05,
                        height=0.012,
                    ),
                ],
            )
        ]

    monkeypatch.setattr(pdf_toc_parser, "extract_pdf_pages_layout", fake_layouts)

    extraction = pdf_toc_parser.extract_pdf_toc_from_range(
        tmp_path / "layout-contract.pdf",
        page_start=21,
        page_end=21,
    )

    assert calls == [True]
    assert [(node.title, node.printed_page) for node in extraction.nodes] == [
        ("第2章 通用目录标题", 143)
    ]


def test_toc_layout_splits_embedded_printed_page_from_scanned_title(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    from app.services import pdf_toc_parser
    from app.services.image_ocr import OCRLineLayout, OCRPageLayout

    monkeypatch.setattr(
        pdf_toc_parser,
        "extract_pdf_pages_layout",
        lambda *_args, **_kwargs: [
            OCRPageLayout(
                page_no=11,
                lines=[
                    OCRLineLayout(
                        "Section 2 General methods /48",
                        x=0.12,
                        y=0.65,
                        width=0.45,
                        height=0.018,
                    )
                ],
            )
        ],
    )

    extraction = pdf_toc_parser.extract_pdf_toc_from_range(
        tmp_path / "embedded-page.pdf",
        page_start=11,
        page_end=11,
    )

    assert [(node.title, node.printed_page) for node in extraction.nodes] == [
        ("Section 2 General methods", 48)
    ]


def test_pdf_layout_ocr_trailing_column_mode_is_explicit(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    from app.services import image_ocr

    path = tmp_path / "placeholder.pdf"
    path.write_bytes(b"placeholder")
    calls: list[list[str]] = []

    def fake_vision_payload(args: list[str], *, timeout: int) -> dict[str, object]:
        assert timeout >= 120
        calls.append(args)
        return {"pages": []}

    monkeypatch.setattr(image_ocr, "_run_vision_ocr_payload", fake_vision_payload)

    image_ocr.extract_pdf_pages_layout(
        path,
        page_start=1,
        page_end=1,
    )
    image_ocr.extract_pdf_pages_layout(
        path,
        page_start=1,
        page_end=1,
        trailing_column_pass=True,
    )

    assert calls[0][-1] == "1"
    assert calls[1][-1] == "trailing-column-lines"


def test_tesseract_tsv_preserves_distant_page_number_as_separate_layout_line() -> None:
    from app.services import image_ocr

    payload = "\n".join(
        [
            "level\tpage_num\tblock_num\tpar_num\tline_num\tword_num\tleft\ttop\twidth\theight\tconf\ttext",
            "1\t1\t0\t0\t0\t0\t0\t0\t1000\t1200\t-1\t",
            "5\t1\t1\t1\t1\t1\t100\t180\t110\t40\t95\tChapter",
            "5\t1\t1\t1\t1\t2\t225\t180\t70\t40\t94\tOne",
            "5\t1\t1\t1\t1\t3\t820\t180\t35\t40\t96\t12",
        ]
    )

    page = image_ocr._parse_tesseract_tsv(payload, page_no=7)

    assert page.page_no == 7
    assert [line.text for line in page.lines] == ["Chapter One", "12"]
    assert page.lines[0].x == pytest.approx(0.1)
    assert page.lines[1].x == pytest.approx(0.82)


def test_pdf_layout_ocr_falls_back_to_tesseract_when_vision_is_unavailable(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    from app.services import image_ocr
    from app.services.image_ocr import OCRLineLayout, OCRPageLayout

    path = tmp_path / "scan.pdf"
    path.write_bytes(b"scan")
    expected = [OCRPageLayout(page_no=3, lines=[OCRLineLayout("Chapter 1", x=0.1, y=0.9)])]
    monkeypatch.setattr(image_ocr, "_run_vision_ocr_payload", lambda *_args, **_kwargs: None)
    monkeypatch.setattr(
        image_ocr,
        "_extract_pdf_pages_layout_with_tesseract",
        lambda *_args, **_kwargs: expected,
    )

    layouts = image_ocr.extract_pdf_pages_layout(
        path,
        page_start=3,
        page_end=3,
    )

    assert layouts == expected


def test_scanned_toc_probe_uses_repeated_structural_rows_in_leading_batches(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    from app.services import pdf_toc_parser
    from app.services.pdf_toc_parser import PdfTocExtraction, PdfTocNode

    calls: list[tuple[int, int]] = []

    def fake_extract(_path: Path, *, page_start: int, page_end: int) -> PdfTocExtraction:
        calls.append((page_start, page_end))
        if page_start < 9:
            return PdfTocExtraction(toc_page_start=page_start, toc_page_end=page_end)
        return PdfTocExtraction(
            nodes=[
                PdfTocNode(title="Chapter 1 Foundations", printed_page=1, toc_page=11),
                PdfTocNode(title="Section 1 Concepts", printed_page=3, toc_page=11),
                PdfTocNode(title="Chapter 2 Practice", printed_page=9, toc_page=12),
                PdfTocNode(title="Section 2 Review", printed_page=12, toc_page=12),
            ],
            toc_page_start=page_start,
            toc_page_end=page_end,
        )

    monkeypatch.setattr(pdf_toc_parser, "extract_pdf_toc_from_range", fake_extract)

    extraction = pdf_toc_parser.probe_pdf_toc_from_leading_pages(
        tmp_path / "scan.pdf",
        page_count=200,
        max_probe_pages=48,
    )

    assert calls == [(1, 8), (9, 16)]
    assert extraction.toc_page_start == 11
    assert extraction.toc_page_end == 12
    assert len(extraction.nodes) == 4


def test_semantic_outline_titles_may_contain_question_marks(tmp_path: Path) -> None:
    from pypdf import PdfWriter

    path = tmp_path / "semantic-bookmarks.pdf"
    writer = PdfWriter()
    for _ in range(3):
        writer.add_blank_page(width=612, height=792)
    writer.add_outline_item("Why? Motivating the API", 0)
    writer.add_outline_item("What About Directories?", 1)
    writer.add_outline_item("Summary!", 2)
    with path.open("wb") as handle:
        writer.write(handle)

    extraction = extract_directory(_record(path), path)

    assert [candidate.title for candidate in extraction.candidates] == [
        "Why? Motivating the API",
        "What About Directories?",
        "Summary!",
    ]
    assert extraction.metadata["directory_source"] == "native_outline"
    assert extraction.inspected_page_count == 0


def test_pdf_outline_collision_demotes_unmatched_destination_titles(tmp_path: Path) -> None:
    import fitz
    from pypdf import PdfReader, PdfWriter

    native_path = tmp_path / "collision-native.pdf"
    document = fitz.open()
    first_page = document.new_page(width=612, height=792)
    first_page.insert_text((72, 90), "Alpha Section", fontsize=18)
    document.new_page(width=612, height=792)
    document.save(native_path)
    document.close()

    path = tmp_path / "collision.pdf"
    writer = PdfWriter()
    for page in PdfReader(native_path).pages:
        writer.add_page(page)
    writer.add_outline_item("Alpha Section", 0)
    writer.add_outline_item("Beta Section", 0)
    writer.add_outline_item("Gamma Section", 1)
    with path.open("wb") as handle:
        writer.write(handle)

    extraction = extract_directory(_record(path), path)

    assert extraction.inspected_page_count == 1
    assert extraction.ocr_page_count == 0
    assert [candidate.mapping_status for candidate in extraction.candidates] == [
        "partial",
        "unmapped",
        "verified",
    ]
    assert extraction.candidates[1].source_range is None
    assert extraction.candidates[1].metadata["reported_destination_page"] == 1
    assert extraction.candidates[2].source_range == SourceRange(
        kind="pdf_pages",
        start=2,
        end=2,
        display_label="PDF p. 2",
        metadata={"index_base": 1, "physical_pages": True},
    )


def test_pdf_outline_collision_keeps_legitimate_same_page_headings(tmp_path: Path) -> None:
    import fitz
    from pypdf import PdfReader, PdfWriter

    native_path = tmp_path / "same-page-native.pdf"
    document = fitz.open()
    first_page = document.new_page(width=612, height=792)
    first_page.insert_text((72, 90), "Alpha Section", fontsize=18)
    first_page.insert_text((72, 180), "Beta Section", fontsize=16)
    document.new_page(width=612, height=792)
    document.save(native_path)
    document.close()

    path = tmp_path / "same-page.pdf"
    writer = PdfWriter()
    for page in PdfReader(native_path).pages:
        writer.add_page(page)
    writer.add_outline_item("Alpha Section", 0)
    writer.add_outline_item("Beta Section", 0)
    writer.add_outline_item("Gamma Section", 1)
    with path.open("wb") as handle:
        writer.write(handle)

    extraction = extract_directory(_record(path), path)

    assert extraction.inspected_page_count == 1
    assert extraction.ocr_page_count == 0
    assert [candidate.mapping_status for candidate in extraction.candidates] == [
        "verified",
        "verified",
        "verified",
    ]
    assert [candidate.source_range.start for candidate in extraction.candidates] == [1, 1, 2]
    assert [candidate.source_range.end for candidate in extraction.candidates] == [1, 1, 2]


def test_pdf_outline_collision_with_empty_text_fails_closed_without_ocr(tmp_path: Path) -> None:
    from pypdf import PdfWriter

    path = tmp_path / "empty-collision.pdf"
    writer = PdfWriter()
    writer.add_blank_page(width=612, height=792)
    writer.add_blank_page(width=612, height=792)
    writer.add_outline_item("Alpha Section", 0)
    writer.add_outline_item("Beta Section", 0)
    writer.add_outline_item("Gamma Section", 1)
    with path.open("wb") as handle:
        writer.write(handle)

    extraction = extract_directory(_record(path), path)

    assert extraction.inspected_page_count == 1
    assert extraction.ocr_page_count == 0
    assert [candidate.mapping_status for candidate in extraction.candidates] == [
        "unmapped",
        "unmapped",
        "verified",
    ]
    assert extraction.candidates[0].source_range is None
    assert extraction.candidates[1].source_range is None


def test_pdf_without_toc_reads_only_heading_regions(tmp_path: Path) -> None:
    import fitz

    path = tmp_path / "headings.pdf"
    document = fitz.open()
    for heading in ("1 First", "1.1 Child", "2 Second"):
        page = document.new_page(width=612, height=792)
        page.insert_text((72, 48), heading, fontsize=20)
        page.insert_text((72, 300), "BODY_SECRET_MUST_NOT_BECOME_DIRECTORY", fontsize=12)
    document.save(path)
    document.close()

    extraction = extract_directory(_record(path), path)

    assert [candidate.title for candidate in extraction.candidates] == ["1 First", "1.1 Child", "2 Second"]
    assert all("BODY_SECRET" not in candidate.title for candidate in extraction.candidates)
    assert extraction.inspected_page_count == 3
    assert extraction.metadata["directory_source"] == "heading_regions"
    assert extraction.metadata["body_text_extracted"] is False


def test_pdf_heading_fallback_keeps_multiple_same_page_sections(tmp_path: Path) -> None:
    import fitz

    path = tmp_path / "same-page-headings.pdf"
    document = fitz.open()
    page = document.new_page(width=612, height=792)
    page.insert_text((72, 42), "1 Parent", fontsize=18)
    page.insert_text((72, 92), "1.1 First child", fontsize=16)
    page.insert_text((72, 142), "1.2 Second child", fontsize=16)
    page.insert_text((72, 400), "BODY_SECRET_MUST_NOT_BECOME_DIRECTORY", fontsize=12)
    document.save(path)
    document.close()

    extraction = extract_directory(_record(path), path)

    assert [candidate.title for candidate in extraction.candidates] == [
        "1 Parent",
        "1.1 First child",
        "1.2 Second child",
    ]
    assert all(candidate.source_range is not None for candidate in extraction.candidates)
    assert all(candidate.source_range.start == 1 for candidate in extraction.candidates if candidate.source_range)
    assert all(candidate.source_range.end == 1 for candidate in extraction.candidates if candidate.source_range)


def test_pdf_sparse_hidden_text_still_uses_heading_region_ocr(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    import fitz

    path = tmp_path / "sparse-hidden-text.pdf"
    document = fitz.open()
    page = document.new_page(width=612, height=792)
    page.insert_text((72, 42), "2024", fontsize=8)
    document.save(path)
    document.close()
    monkeypatch.setattr(
        "app.services.source_directory_extractor.extract_image_text",
        lambda _path: "1 Recovered heading",
    )

    extraction = extract_directory(_record(path), path)

    assert extraction.ocr_page_count == 1
    assert [candidate.title for candidate in extraction.candidates] == [
        "1 Recovered heading"
    ]


def test_numeric_range_is_not_verified_across_unmapped_logical_boundary() -> None:
    def mapped(key: str, start: int) -> DirectoryCandidate:
        return DirectoryCandidate(
            local_key=key,
            title=key,
            level=1,
            source_range=SourceRange(
                kind="pdf_pages",
                start=start,
                end=start,
                metadata={"index_base": 1, "physical_pages": True},
            ),
            mapping_status="verified",
            confidence=0.95,
        )

    candidates = [
        mapped("first", 10),
        DirectoryCandidate(
            local_key="unmapped-successor",
            title="Unmapped successor",
            level=1,
            mapping_status="unmapped",
        ),
        mapped("later", 30),
    ]

    closed = _close_numeric_ranges(candidates, maximum=40, kind="pdf_pages")

    assert closed[0].mapping_status == "partial"
    assert closed[0].source_range is not None
    assert closed[0].source_range.end == 10
    assert closed[0].metadata["range_boundary_local_key"] == "unmapped-successor"
    assert closed[1].mapping_status == "unmapped"
    assert closed[2].mapping_status == "verified"


def test_split_native_toc_lines_are_parsed_without_body_text() -> None:
    heading_lines = [
        _PdfHeadingLine("目", 5, 0.03, 20, "pdf_heading_region_text"),
        _PdfHeadingLine("录", 5, 0.06, 20, "pdf_heading_region_text"),
    ]
    nodes = _parse_native_toc_text(
        "目\n录\n第一章 Random events\n１\n……\n第一节 Definitions\n２\n",
        toc_page=5,
    )

    assert _pdf_page_has_toc_heading(heading_lines) is True
    assert [(node.title, node.printed_page, node.level) for node in nodes] == [
        ("第一章 Random events", 1, 1),
        ("第一节 Definitions", 2, 2),
    ]


def test_verified_printed_page_offset_overrides_ambiguous_heading_match() -> None:
    nodes = [
        _PdfTocCandidate("Chapter 1 Alpha", "1", 1, 1, 2, 0.9, "toc"),
        _PdfTocCandidate("Common heading", "", 2, 2, 2, 0.9, "toc"),
        _PdfTocCandidate("Chapter 2 Beta", "2", 1, 3, 2, 0.9, "toc"),
    ]
    heading_cache = {
        8: [_PdfHeadingLine("Chapter 1 Alpha", 8, 0.05, 18, "text")],
        10: [_PdfHeadingLine("Chapter 2 Beta", 10, 0.05, 18, "text")],
        20: [_PdfHeadingLine("Common heading", 20, 0.05, 18, "text")],
    }

    mapped = _map_pdf_toc_nodes(nodes, heading_cache, page_count=30)

    assert [candidate.source_range.start for candidate in mapped if candidate.source_range] == [8, 9, 10]
    assert mapped[1].evidence[0].metadata["mapping_method"] == "verified_printed_to_physical_offset"


def test_normalized_hierarchy_recloses_ranges_after_candidate_rejection() -> None:
    def candidate(key: str, level: int, start: int, end: int) -> DirectoryCandidate:
        return DirectoryCandidate(
            local_key=key,
            title=key,
            level=level,
            source_range=SourceRange(
                kind="pdf_pages",
                start=start,
                end=end,
                metadata={"index_base": 1, "physical_pages": True},
            ),
            mapping_status="verified",
        )

    original = (
        candidate("parent", 1, 1, 2),
        candidate("rejected-running-header", 1, 3, 3),
        candidate("normalized-child", 1, 4, 6),
    )
    normalized = (original[0], replace(original[2], level=2))

    closed = _reclose_normalized_ranges(
        normalized,
        extraction=DirectoryExtraction(
            candidates=original,
            page_count=6,
            metadata={"format": "pdf"},
        ),
    )

    assert closed[0].source_range is not None
    assert closed[0].source_range.end == 6
    assert closed[1].source_range is not None
    assert closed[1].source_range.end == 6


def test_verified_parent_range_encloses_descendants_on_next_peer_page() -> None:
    def candidate(key: str, level: int, start: int) -> DirectoryCandidate:
        return DirectoryCandidate(
            local_key=key,
            title=key,
            level=level,
            source_range=SourceRange(
                kind="pdf_pages",
                start=start,
                end=start,
                display_label=f"PDF p. {start}",
                metadata={"index_base": 1, "physical_pages": True},
            ),
            mapping_status="verified",
        )

    closed = _reclose_normalized_ranges(
        (
            candidate("parent", 1, 70),
            candidate("first-child", 2, 70),
            candidate("last-child", 2, 71),
            candidate("next-peer", 1, 71),
        ),
        extraction=DirectoryExtraction(
            candidates=(),
            page_count=80,
            metadata={"format": "pdf"},
        ),
    )

    assert closed[0].mapping_status == "verified"
    assert closed[0].source_range is not None
    assert closed[0].source_range.end == 71
    assert closed[0].source_range.display_label == "PDF pp. 70-71"
    assert closed[0].source_range.end == closed[3].source_range.start


def test_verified_parent_is_demoted_instead_of_crossing_next_peer() -> None:
    def candidate(key: str, level: int, start: int, end: int) -> DirectoryCandidate:
        return DirectoryCandidate(
            local_key=key,
            title=key,
            level=level,
            source_range=SourceRange(
                kind="pdf_pages",
                start=start,
                end=end,
                metadata={"index_base": 1, "physical_pages": True},
            ),
            mapping_status="verified",
            confidence=0.95,
        )

    enclosed = _reclose_normalized_ranges(
        (
            candidate("parent", 1, 70, 70),
            candidate("child", 2, 72, 72),
            candidate("next-peer", 1, 71, 80),
        ),
        extraction=DirectoryExtraction(
            candidates=(),
            page_count=80,
            metadata={"format": "pdf"},
        ),
    )

    assert enclosed[0].mapping_status == "partial"
    assert enclosed[0].source_range is not None
    assert enclosed[0].source_range.end == 70
    assert enclosed[0].metadata["range_boundary_status"] == "descendant_crosses_successor"
    assert enclosed[0].metadata["range_boundary_local_key"] == "next-peer"


def test_chapter_validator_rejects_verified_child_outside_verified_parent() -> None:
    parent = SourceChapter(
        id="parent",
        package_id="course_directory",
        source_ingestion_id="source_directory",
        title="Parent",
        order_index=0,
        range=SourceRange(kind="pdf_pages", start=70, end=70),
        mapping_status="verified",
    )
    child = SourceChapter(
        id="child",
        package_id="course_directory",
        source_ingestion_id="source_directory",
        parent_id="parent",
        title="Child",
        level=2,
        order_index=1,
        range=SourceRange(kind="pdf_pages", start=71, end=71),
        mapping_status="verified",
    )

    with pytest.raises(SourceDirectoryProcessingError, match="outside its verified parent"):
        _validate_chapters((parent, child))


def test_epub_anchor_ranges_are_not_subject_to_numeric_parent_enclosure() -> None:
    parent = SourceChapter(
        id="parent",
        package_id="course_directory",
        source_ingestion_id="source_directory",
        title="Parent",
        order_index=0,
        range=SourceRange(
            kind="epub_spine",
            start=1,
            end=1,
            start_anchor="parent",
            end_anchor="next",
        ),
        mapping_status="verified",
    )
    child = SourceChapter(
        id="child",
        package_id="course_directory",
        source_ingestion_id="source_directory",
        parent_id="parent",
        title="Child",
        level=2,
        order_index=1,
        range=SourceRange(
            kind="epub_spine",
            start=2,
            end=2,
            start_anchor="child",
            end_anchor="next-child",
        ),
        mapping_status="verified",
    )

    _validate_chapters((parent, child))


def test_epub_bad_fragment_is_unmapped_without_authoritative_range(tmp_path: Path) -> None:
    path = tmp_path / "bad-fragment.epub"
    _write_epub_with_navigation(
        path,
        navigation="<li><a href='s1.xhtml#missing'>Missing anchor</a></li>",
        document="<html><body><h1 id='present'>Present</h1></body></html>",
    )

    extraction = extract_directory(_record(path), path)
    candidate = extraction.candidates[0]

    assert candidate.mapping_status == "unmapped"
    assert candidate.source_range is None
    assert candidate.metadata["navigation_provenance"] == "native"
    assert candidate.metadata["native_level"] == 1
    assert candidate.metadata["fragment_validation"] == "missing_id_or_name"
    assert extraction.metadata["body_text_extracted"] is False


def test_epub_url_encoded_fragment_is_decoded_and_verified(tmp_path: Path) -> None:
    path = tmp_path / "encoded-fragment.epub"
    _write_epub_with_navigation(
        path,
        navigation="<li><a href='s1.xhtml#section%20one'>Encoded anchor</a></li>",
        document="<html><body><a name='section one'></a></body></html>",
    )

    extraction = extract_directory(_record(path), path)
    candidate = extraction.candidates[0]

    assert candidate.mapping_status == "verified"
    assert candidate.source_range is not None
    assert candidate.source_range.start_anchor == "section one"
    assert candidate.source_locator == "epub:OPS/s1.xhtml#section one"
    assert extraction.metadata["anchor_validation"] == "xhtml_id_name_attributes_only"


def test_epub_extensionless_content_document_can_verify_fragment(tmp_path: Path) -> None:
    path = tmp_path / "extensionless-content.epub"
    _write_epub_with_navigation(
        path,
        navigation="<li><a href='content#target'>Target</a></li>",
        document="<html><body><section id='target'>Body is not retained.</section></body></html>",
        document_name="content",
    )

    extraction = extract_directory(_record(path), path)

    assert extraction.candidates[0].mapping_status == "verified"
    assert extraction.candidates[0].source_range is not None
    assert extraction.candidates[0].source_range.start_anchor == "target"
    assert extraction.metadata["body_text_extracted"] is False


def test_epub_native_navigation_truncation_is_explicit(
    tmp_path: Path,
    monkeypatch,
) -> None:
    path = tmp_path / "truncated-navigation.epub"
    _write_epub_with_navigation(
        path,
        navigation=(
            "<li><a href='s1.xhtml#one'>One</a></li>"
            "<li><a href='s1.xhtml#two'>Two</a></li>"
            "<li><a href='s1.xhtml#three'>Three</a></li>"
        ),
        document="<html><body><i id='one'></i><i id='two'></i><i id='three'></i></body></html>",
    )
    monkeypatch.setattr("app.services.source_directory_extractor.MAX_DIRECTORY_NODES", 2)

    extraction = extract_directory(_record(path), path)

    assert len(extraction.candidates) == 2
    assert extraction.metadata["native_navigation_count"] == 3
    assert extraction.metadata["published_navigation_count"] == 2
    assert extraction.metadata["navigation_truncated"] is True
    assert extraction.metadata["navigation_node_limit"] == 2
    assert any("catalog is partial" in warning for warning in extraction.warnings)
    assert extraction.candidates[1].mapping_status == "verified"
    assert extraction.candidates[1].source_range is not None
    assert extraction.candidates[1].source_range.end_anchor == "three"

    reclosed = _reclose_normalized_ranges(
        extraction.candidates,
        extraction=extraction,
    )
    assert reclosed[1].source_range is not None
    assert reclosed[1].source_range.end == extraction.candidates[1].source_range.end
    assert reclosed[1].source_range.end_anchor == "three"

    quality = _catalog_quality(
        (
            SourceChapter(
                id="chapter-one",
                package_id="course-directory",
                source_ingestion_id="source-epub",
                title="One",
                order_index=0,
                range=reclosed[0].source_range,
                mapping_status="verified",
            ),
            SourceChapter(
                id="chapter-two",
                package_id="course-directory",
                source_ingestion_id="source-epub",
                title="Two",
                order_index=1,
                range=reclosed[1].source_range,
                mapping_status="verified",
            ),
        ),
        catalog_complete=False,
    )
    assert quality.level == "partially_verified"
    assert quality.confidence == 0.9
    assert any("只发布部分导航" in item for item in quality.diagnostics)


def test_epub_navigation_truncation_demotes_ranges_left_open_by_lookahead(
    tmp_path: Path,
    monkeypatch,
) -> None:
    path = tmp_path / "nested-truncated-navigation.epub"
    _write_epub_with_navigation(
        path,
        navigation=(
            "<li><a href='s1.xhtml#one'>One</a><ol>"
            "<li><a href='s1.xhtml#two'>Two</a><ol>"
            "<li><a href='s1.xhtml#three'>Three</a></li>"
            "</ol></li></ol></li>"
        ),
        document="<html><body><i id='one'></i><i id='two'></i><i id='three'></i></body></html>",
    )
    monkeypatch.setattr("app.services.source_directory_extractor.MAX_DIRECTORY_NODES", 2)

    extraction = extract_directory(_record(path), path)

    assert len(extraction.candidates) == 2
    assert [candidate.mapping_status for candidate in extraction.candidates] == [
        "partial",
        "partial",
    ]
    assert all(
        candidate.metadata["range_boundary_status"] == "navigation_truncated"
        for candidate in extraction.candidates
    )


def test_html_ranges_store_last_heading_and_next_heading_text_boundary(tmp_path: Path) -> None:
    path = tmp_path / "reference.html"
    path.write_text(
        "<h1 id='parent'>Parent</h1><p>Parent body</p>"
        "<h2 id='child'>Child</h2><p>Child body</p>"
        "<h1 id='next'>Next</h1><p>Next body</p>",
        encoding="utf-8",
    )

    extraction = extract_directory(_record(path), path)
    parent, child, following = extraction.candidates

    assert parent.source_range is not None
    assert parent.source_range.end == 1
    assert parent.source_range.end_anchor == "next"
    assert parent.source_range.metadata["end_heading_ordinal"] == 2
    assert child.source_range is not None
    assert child.source_range.end == 1
    assert child.source_range.end_anchor == "next"
    assert following.source_range is not None
    assert following.source_range.end == 2
    assert following.source_range.end_anchor == ""


def test_epub_same_spine_sections_use_next_anchor_as_text_boundary(tmp_path: Path) -> None:
    path = tmp_path / "reference.epub"
    with zipfile.ZipFile(path, "w") as archive:
        archive.writestr(
            "META-INF/container.xml",
            "<container><rootfiles><rootfile full-path='OPS/content.opf'/></rootfiles></container>",
        )
        archive.writestr(
            "OPS/content.opf",
            "<package><manifest>"
            "<item id='nav' href='nav.xhtml'/><item id='s1' href='s1.xhtml'/>"
            "<item id='s2' href='s2.xhtml'/></manifest>"
            "<spine><itemref idref='s1'/><itemref idref='s2'/></spine></package>",
        )
        archive.writestr(
            "OPS/nav.xhtml",
            "<nav><ol><li><a href='s1.xhtml#parent'>Parent</a>"
            "<ol><li><a href='s1.xhtml#child'>Child</a></li></ol></li>"
            "<li><a href='s1.xhtml#sibling'>Sibling</a></li>"
            "<li><a href='s2.xhtml#next'>Next</a></li></ol></nav>",
        )
        archive.writestr(
            "OPS/s1.xhtml",
            "<h1 id='parent'>Parent</h1><h2 id='child'>Child</h2>"
            "<h1 id='sibling'>Sibling</h1>",
        )
        archive.writestr("OPS/s2.xhtml", "<h1 id='next'>Next</h1>")

    extraction = extract_directory(_record(path), path)
    parent, child, sibling, following = extraction.candidates

    assert parent.source_range is not None
    assert parent.source_range.start == 0
    assert parent.source_range.end == 0
    assert parent.source_range.end_anchor == "sibling"
    assert child.source_range is not None
    assert child.source_range.end_anchor == "sibling"
    assert sibling.source_range is not None
    assert sibling.source_range.end == 1
    assert sibling.source_range.end_anchor == "next"
    assert following.source_range is not None
    assert following.source_range.start == 1
    assert following.source_range.end == 1


def test_xlsx_uses_workbook_relationships_for_reordered_noncontiguous_sheets(
    tmp_path: Path,
) -> None:
    path = tmp_path / "reordered.xlsx"
    with zipfile.ZipFile(path, "w") as archive:
        archive.writestr(
            "xl/workbook.xml",
            "<workbook xmlns='http://schemas.openxmlformats.org/spreadsheetml/2006/main' "
            "xmlns:r='http://schemas.openxmlformats.org/officeDocument/2006/relationships'>"
            "<sheets>"
            "<sheet name='Second in storage' sheetId='1' r:id='rSecond'/>"
            "<sheet name='First in storage' sheetId='2' r:id='rFirst'/>"
            "</sheets></workbook>",
        )
        archive.writestr(
            "xl/_rels/workbook.xml.rels",
            "<Relationships xmlns='http://schemas.openxmlformats.org/package/2006/relationships'>"
            "<Relationship Id='rFirst' Target='worksheets/sheet2.xml'/>"
            "<Relationship Id='rSecond' Target='worksheets/sheet7.xml'/>"
            "</Relationships>",
        )
        archive.writestr(
            "xl/worksheets/sheet2.xml",
            "<worksheet xmlns='http://schemas.openxmlformats.org/spreadsheetml/2006/main'>"
            "<dimension ref='A1:A4'/></worksheet>",
        )
        archive.writestr(
            "xl/worksheets/sheet7.xml",
            "<worksheet xmlns='http://schemas.openxmlformats.org/spreadsheetml/2006/main'>"
            "<dimension ref='A1:A17'/></worksheet>",
        )

    extraction = extract_directory(_record(path), path)

    assert [candidate.title for candidate in extraction.candidates] == [
        "Second in storage",
        "First in storage",
    ]
    assert [candidate.source_range.end for candidate in extraction.candidates if candidate.source_range] == [
        17,
        4,
    ]
    assert [
        candidate.source_range.metadata["sheet_path"]
        for candidate in extraction.candidates
        if candidate.source_range
    ] == ["xl/worksheets/sheet7.xml", "xl/worksheets/sheet2.xml"]


def test_processor_publishes_catalog_without_chunks_or_visuals(tmp_path: Path) -> None:
    path = tmp_path / "source.md"
    path.write_text("# One\nBody one\n## Child\nBody child\n# Two\nBody two", encoding="utf-8")
    database = tmp_path / "openclass.sqlite3"
    store = SourceStructureStore(database)
    processor = SourceDirectoryProcessor(
        store=store,
        normalizer_factory=lambda _record: PassthroughNormalizer(),
    )

    structure = processor.process(record=_record(path), path=path, catalog_model=_model())
    catalog = store.get_catalog_view(source=_record(path))

    assert structure.strategy == "codex_directory_v1"
    assert structure.catalog_version == 1
    assert structure.chunk_count == 0
    assert structure.visual_count == 0
    assert catalog.catalog_version == 1
    assert len(catalog.chapters) == 3
    assert all(chapter.range is not None for chapter in catalog.chapters)
    assert all(chapter.catalog_version == 1 for chapter in catalog.chapters)
    with sqlite3.connect(database) as connection:
        assert connection.execute("SELECT COUNT(*) FROM source_chunks").fetchone()[0] == 0
        assert connection.execute("SELECT COUNT(*) FROM source_visual_assets").fetchone()[0] == 0
        run = connection.execute(
            "SELECT status, turn_count, chapter_count FROM source_catalog_runs"
        ).fetchone()
    assert run == ("succeeded", 1, 3)


def test_processor_rejects_stale_metadata_fingerprint_before_extraction(tmp_path: Path) -> None:
    path = tmp_path / "source.md"
    path.write_text("# One\nBody", encoding="utf-8")
    record = _record(path).model_copy(
        update={
            "metadata": {
                **_record(path).metadata,
                "content_hash": "0" * 64,
            }
        }
    )

    class MustNotNormalize:
        def normalize(self, **_kwargs):
            raise AssertionError("A stale fingerprint must be rejected before extraction and Codex")

    store = SourceStructureStore(tmp_path / "openclass.sqlite3")
    processor = SourceDirectoryProcessor(
        store=store,
        normalizer_factory=lambda _record: MustNotNormalize(),
    )

    with pytest.raises(SourceDirectoryProcessingError, match="fingerprint"):
        processor.process(record=record, path=path, catalog_model=_model())

    assert store.get_structure(
        owner_user_id=record.owner_user_id,
        package_id=record.package_id,
        source_id=record.id,
    ) is None
    assert [run.status for run in store.list_catalog_runs(
        owner_user_id=record.owner_user_id,
        package_id=record.package_id,
        source_id=record.id,
    )] == ["failed"]


def test_processor_detects_file_change_before_publish_and_preserves_catalog(tmp_path: Path) -> None:
    path = tmp_path / "source.md"
    path.write_text("# One\nBody\n# Two\nBody", encoding="utf-8")
    store = SourceStructureStore(tmp_path / "openclass.sqlite3")
    record = _record(path)
    first = SourceDirectoryProcessor(
        store=store,
        normalizer_factory=lambda _record: PassthroughNormalizer(),
    ).process(record=record, path=path, catalog_model=_model())

    class MutatingNormalizer(PassthroughNormalizer):
        def normalize(self, *, record, candidates, selection):
            result = super().normalize(record=record, candidates=candidates, selection=selection)
            path.write_text("# Replaced while cataloging\nNew body", encoding="utf-8")
            return result

    racing = SourceDirectoryProcessor(
        store=store,
        normalizer_factory=lambda _record: MutatingNormalizer(),
    )

    with pytest.raises(SourceDirectoryProcessingError, match="changed while"):
        racing.process(record=record, path=path, catalog_model=_model())

    preserved = store.get_structure(
        owner_user_id=record.owner_user_id,
        package_id=record.package_id,
        source_id=record.id,
    )
    assert preserved is not None
    assert preserved.id == first.id
    assert preserved.catalog_version == 1


def test_post_publish_progress_failure_does_not_mark_committed_catalog_failed(
    tmp_path: Path,
) -> None:
    path = tmp_path / "source.md"
    path.write_text("# One\nBody\n# Two\nBody", encoding="utf-8")
    store = SourceStructureStore(tmp_path / "openclass.sqlite3")
    processor = SourceDirectoryProcessor(
        store=store,
        normalizer_factory=lambda _record: PassthroughNormalizer(),
    )

    def progress(phase: str, _value: int) -> None:
        if phase == "catalog_ready":
            raise RuntimeError("progress transport closed after commit")

    published = processor.process(
        record=_record(path),
        path=path,
        catalog_model=_model(),
        progress_callback=progress,
    )

    assert published.status == "ready"
    assert published.catalog_version == 1
    assert [run.status for run in store.list_catalog_runs(
        owner_user_id="user_directory",
        package_id="course_directory",
        source_id="source_directory",
    )] == ["succeeded"]


def test_structured_upload_bypasses_open_notebook_and_full_indexing(
    tmp_path: Path,
    monkeypatch,
) -> None:
    database = tmp_path / "openclass.sqlite3"
    source_store = SourceEvidenceStore(database)
    structure_store = SourceStructureStore(database)
    processor = SourceDirectoryProcessor(
        store=structure_store,
        normalizer_factory=lambda _record: PassthroughNormalizer(),
    )

    class RejectingOpenNotebookAdapter:
        api_url = "http://notebook.test"

        def create_notebook(self, **_kwargs):
            raise AssertionError("Directory uploads must not create an OpenNotebook notebook")

        def upload_file_source(self, **_kwargs):
            raise AssertionError("Directory uploads must not enter OpenNotebook")

        def get_command(self, _command_id):
            raise AssertionError("Directory uploads must not poll OpenNotebook")

    monkeypatch.setattr(workspace_state, "UPLOAD_DIR", tmp_path / "uploads")
    service = SourceIngestionService(
        adapter=RejectingOpenNotebookAdapter(),
        source_backend="open_notebook",
        store=source_store,
        job_store=SourceIngestionJobStore(database),
        structure_store=structure_store,
        directory_processor=processor,
    )
    package = CoursePackage(id="course_directory", title="Directory", summary="", lessons=[])

    queued = service.queue_file_source(
        owner_user_id="user_directory",
        package=package,
        file_name="source.md",
        content=b"# One\nBody one\n## Child\nBody child",
        mime_type="text/markdown",
        catalog_model=_model(),
    )
    completed = service.process_file_source(
        owner_user_id="user_directory",
        package_id=package.id,
        source_id=queued.id,
    )
    listed = service.list_sources(
        owner_user_id="user_directory",
        package_id=package.id,
    )[0]
    rebuilt = service.rebuild_catalog(
        owner_user_id="user_directory",
        package_id=package.id,
        source_id=queued.id,
    )
    assert rebuilt is not None
    catalog = structure_store.get_catalog_view(source=rebuilt)

    assert queued.status == "parsing"
    assert queued.metadata["catalog_pipeline"] == "codex_directory_v1"
    assert queued.metadata["adapter"] == "codex_directory_v1"
    assert completed.status == "ready"
    assert completed.open_notebook_source_id == ""
    assert listed.status == "ready"
    assert catalog.catalog_version == 2
    with sqlite3.connect(database) as connection:
        assert connection.execute("SELECT COUNT(*) FROM source_chunks").fetchone()[0] == 0
        assert connection.execute("SELECT COUNT(*) FROM source_visual_assets").fetchone()[0] == 0


def test_directory_processing_is_serial_per_source_and_reuses_completed_upload(
    tmp_path: Path,
    monkeypatch,
) -> None:
    database = tmp_path / "openclass.sqlite3"
    source_store = SourceEvidenceStore(database)
    structure_store = SourceStructureStore(database)

    class TrackingNormalizer(PassthroughNormalizer):
        def __init__(self) -> None:
            self._lock = threading.Lock()
            self.calls = 0
            self.active = 0
            self.max_active = 0

        def reset(self) -> None:
            with self._lock:
                self.calls = 0
                self.active = 0
                self.max_active = 0

        def normalize(self, *, record, candidates, selection):
            with self._lock:
                self.calls += 1
                self.active += 1
                self.max_active = max(self.max_active, self.active)
            try:
                time.sleep(0.08)
                return super().normalize(
                    record=record,
                    candidates=candidates,
                    selection=selection,
                )
            finally:
                with self._lock:
                    self.active -= 1

    normalizer = TrackingNormalizer()
    processor = SourceDirectoryProcessor(
        store=structure_store,
        normalizer_factory=lambda _record: normalizer,
    )
    monkeypatch.setattr(workspace_state, "UPLOAD_DIR", tmp_path / "uploads")
    service = SourceIngestionService(
        source_backend="native",
        store=source_store,
        job_store=SourceIngestionJobStore(database),
        structure_store=structure_store,
        directory_processor=processor,
    )
    package = CoursePackage(id="course_directory", title="Directory", summary="", lessons=[])
    queued = service.queue_file_source(
        owner_user_id="user_directory",
        package=package,
        file_name="source.md",
        content=b"# One\nBody\n# Two\nBody",
        mime_type="text/markdown",
        catalog_model=_model(),
    )

    start_processing = threading.Barrier(2)

    def process_upload() -> SourceIngestionRecord:
        start_processing.wait()
        return service.process_file_source(
            owner_user_id="user_directory",
            package_id=package.id,
            source_id=queued.id,
        )

    with ThreadPoolExecutor(max_workers=2) as executor:
        completed = list(executor.map(lambda _index: process_upload(), range(2)))

    assert [record.status for record in completed] == ["ready", "ready"]
    assert normalizer.calls == 1
    assert normalizer.max_active == 1
    first_catalog = structure_store.get_structure(
        owner_user_id="user_directory",
        package_id=package.id,
        source_id=queued.id,
    )
    assert first_catalog is not None
    assert first_catalog.catalog_version == 1

    normalizer.reset()
    start_rebuild = threading.Barrier(2)

    def rebuild_catalog() -> SourceIngestionRecord | None:
        start_rebuild.wait()
        return service.rebuild_catalog(
            owner_user_id="user_directory",
            package_id=package.id,
            source_id=queued.id,
        )

    with ThreadPoolExecutor(max_workers=2) as executor:
        rebuilt = list(executor.map(lambda _index: rebuild_catalog(), range(2)))

    assert all(record is not None and record.status == "ready" for record in rebuilt)
    assert normalizer.calls == 2
    assert normalizer.max_active == 1
    final_catalog = structure_store.get_structure(
        owner_user_id="user_directory",
        package_id=package.id,
        source_id=queued.id,
    )
    assert final_catalog is not None
    assert final_catalog.catalog_version == 3


def test_deleting_source_waits_for_active_directory_processing_and_removes_final_catalog(
    tmp_path: Path,
    monkeypatch,
) -> None:
    database = tmp_path / "openclass.sqlite3"
    source_store = SourceEvidenceStore(database)
    structure_store = SourceStructureStore(database)
    normalizer_started = threading.Event()
    release_normalizer = threading.Event()
    remove_attempted = threading.Event()

    class BlockingNormalizer(PassthroughNormalizer):
        def normalize(self, *, record, candidates, selection):
            normalizer_started.set()
            assert release_normalizer.wait(timeout=5)
            return super().normalize(
                record=record,
                candidates=candidates,
                selection=selection,
            )

    monkeypatch.setattr(workspace_state, "UPLOAD_DIR", tmp_path / "uploads")
    service = SourceIngestionService(
        source_backend="native",
        store=source_store,
        job_store=SourceIngestionJobStore(database),
        structure_store=structure_store,
        directory_processor=SourceDirectoryProcessor(
            store=structure_store,
            normalizer_factory=lambda _record: BlockingNormalizer(),
        ),
    )
    package = CoursePackage(id="course_directory", title="Directory", summary="", lessons=[])
    queued = service.queue_file_source(
        owner_user_id="user_directory",
        package=package,
        file_name="source.md",
        content=b"# One\nBody\n# Two\nBody",
        mime_type="text/markdown",
        catalog_model=_model(),
    )

    def process_source() -> SourceIngestionRecord:
        return service.process_file_source(
            owner_user_id="user_directory",
            package_id=package.id,
            source_id=queued.id,
        )

    def remove_source() -> SourceIngestionRecord | None:
        remove_attempted.set()
        return service.remove_source(
            owner_user_id="user_directory",
            package_id=package.id,
            source_id=queued.id,
        )

    with ThreadPoolExecutor(max_workers=2) as executor:
        processing = executor.submit(process_source)
        assert normalizer_started.wait(timeout=5)
        removing = executor.submit(remove_source)
        assert remove_attempted.wait(timeout=5)
        assert not removing.done()
        release_normalizer.set()
        assert processing.result(timeout=5).status == "ready"
        assert removing.result(timeout=5) is not None

    assert source_store.get_source(
        owner_user_id="user_directory",
        package_id=package.id,
        source_id=queued.id,
    ) is None
    assert structure_store.get_structure(
        owner_user_id="user_directory",
        package_id=package.id,
        source_id=queued.id,
    ) is None
    assert not Path(str(queued.metadata["local_source_path"])).exists()


def test_stale_directory_task_cannot_recreate_an_already_deleted_source(
    tmp_path: Path,
    monkeypatch,
) -> None:
    database = tmp_path / "openclass.sqlite3"
    source_store = SourceEvidenceStore(database)
    job_store = SourceIngestionJobStore(database)
    structure_store = SourceStructureStore(database)
    monkeypatch.setattr(workspace_state, "UPLOAD_DIR", tmp_path / "uploads")
    service = SourceIngestionService(
        source_backend="native",
        store=source_store,
        job_store=job_store,
        structure_store=structure_store,
        directory_processor=SourceDirectoryProcessor(
            store=structure_store,
            normalizer_factory=lambda _record: PassthroughNormalizer(),
        ),
    )
    package = CoursePackage(id="course_directory", title="Directory", summary="", lessons=[])
    queued = service.queue_file_source(
        owner_user_id="user_directory",
        package=package,
        file_name="source.md",
        content=b"# One\nBody",
        mime_type="text/markdown",
        catalog_model=_model(),
    )
    stale_job = job_store.latest_for_source(
        owner_user_id="user_directory",
        package_id=package.id,
        source_id=queued.id,
    )
    assert stale_job is not None
    assert service.remove_source(
        owner_user_id="user_directory",
        package_id=package.id,
        source_id=queued.id,
    ) is not None

    with pytest.raises(SourceIngestionError, match="removed"):
        service._save_and_index(queued, stale_job)

    assert source_store.get_source(
        owner_user_id="user_directory",
        package_id=package.id,
        source_id=queued.id,
    ) is None
    assert structure_store.get_structure(
        owner_user_id="user_directory",
        package_id=package.id,
        source_id=queued.id,
    ) is None


def test_retry_directory_source_uses_the_current_catalog_model(
    tmp_path: Path,
    monkeypatch,
) -> None:
    database = tmp_path / "openclass.sqlite3"
    source_store = SourceEvidenceStore(database)
    structure_store = SourceStructureStore(database)
    observed_selections: list[AIModelSelection] = []

    class RecordingNormalizer:
        def normalize(self, *, record, candidates, selection):
            observed_selections.append(selection)
            return DirectoryNormalizationResult(
                candidates=tuple(candidates),
                turn_count=1 if candidates else 0,
                metadata={"test_adapter": "recording"},
            )

    monkeypatch.setattr(workspace_state, "UPLOAD_DIR", tmp_path / "uploads")
    service = SourceIngestionService(
        source_backend="native",
        store=source_store,
        job_store=SourceIngestionJobStore(database),
        structure_store=structure_store,
        directory_processor=SourceDirectoryProcessor(
            store=structure_store,
            normalizer_factory=lambda _record: RecordingNormalizer(),
        ),
    )
    package = CoursePackage(id="course_directory", title="Directory", summary="", lessons=[])
    queued = service.queue_file_source(
        owner_user_id="user_directory",
        package=package,
        file_name="source.md",
        content=b"# One\nBody",
        mime_type="text/markdown",
        catalog_model=_model(),
    )
    source_store.save_source(
        queued.model_copy(
            update={
                "status": "failed",
                "error": "Sign in with ChatGPT/Codex to use subscription models.",
            }
        )
    )

    retried = service.retry_source(
        owner_user_id="user_directory",
        package_id=package.id,
        source_id=queued.id,
        catalog_model=_alternate_model(),
    )

    assert retried is not None
    assert retried.status == "ready"
    assert retried.metadata["catalog_model"] == _alternate_model().model_dump(mode="json")
    assert observed_selections == [_alternate_model()]
    structure = structure_store.get_structure(
        owner_user_id="user_directory",
        package_id=package.id,
        source_id=queued.id,
    )
    assert structure is not None
    assert structure.catalog_model == "catalog-retry-model"


@pytest.mark.parametrize(
    ("operation", "legacy_source"),
    [
        ("retry", False),
        ("rebuild", False),
        ("retry", True),
        ("rebuild", True),
    ],
)
def test_deleted_source_cannot_be_resurrected_after_catalog_operation_read(
    operation: str,
    legacy_source: bool,
    tmp_path: Path,
    monkeypatch,
) -> None:
    database = tmp_path / "openclass.sqlite3"
    source_store = SourceEvidenceStore(database)
    job_store = SourceIngestionJobStore(database)
    structure_store = SourceStructureStore(database)
    monkeypatch.setattr(workspace_state, "UPLOAD_DIR", tmp_path / "uploads")
    service = SourceIngestionService(
        source_backend="native",
        store=source_store,
        job_store=job_store,
        structure_store=structure_store,
        directory_processor=SourceDirectoryProcessor(
            store=structure_store,
            normalizer_factory=lambda _record: PassthroughNormalizer(),
        ),
    )
    package = CoursePackage(id="course_directory", title="Directory", summary="", lessons=[])
    queued = service.queue_file_source(
        owner_user_id="user_directory",
        package=package,
        file_name="source.md",
        content=b"# One\nBody",
        mime_type="text/markdown",
        catalog_model=_model(),
    )
    if legacy_source:
        legacy_metadata = dict(queued.metadata)
        legacy_metadata.pop("catalog_pipeline", None)
        legacy_metadata["adapter"] = "openclass_native"
        queued = source_store.save_source(
            queued.model_copy(update={"metadata": legacy_metadata})
        )
    original_get_source = source_store.get_source
    operation_read_source = threading.Event()
    release_operation = threading.Event()
    paused_once = threading.Event()

    def controlled_get_source(*, owner_user_id: str, package_id: str, source_id: str):
        result = original_get_source(
            owner_user_id=owner_user_id,
            package_id=package_id,
            source_id=source_id,
        )
        if threading.current_thread().name.startswith("catalog-race") and not paused_once.is_set():
            paused_once.set()
            operation_read_source.set()
            assert release_operation.wait(timeout=5)
        return result

    monkeypatch.setattr(source_store, "get_source", controlled_get_source)

    def run_operation() -> SourceIngestionRecord | None:
        if operation == "retry":
            return service.retry_source(
                owner_user_id="user_directory",
                package_id=package.id,
                source_id=queued.id,
            )
        return service.rebuild_catalog(
            owner_user_id="user_directory",
            package_id=package.id,
            source_id=queued.id,
        )

    with ThreadPoolExecutor(max_workers=1, thread_name_prefix="catalog-race") as executor:
        pending = executor.submit(run_operation)
        assert operation_read_source.wait(timeout=5)
        assert service.remove_source(
            owner_user_id="user_directory",
            package_id=package.id,
            source_id=queued.id,
        ) is not None
        release_operation.set()
        assert pending.result(timeout=5) is None

    assert original_get_source(
        owner_user_id="user_directory",
        package_id=package.id,
        source_id=queued.id,
    ) is None
    assert structure_store.get_structure(
        owner_user_id="user_directory",
        package_id=package.id,
        source_id=queued.id,
    ) is None
    assert job_store.latest_for_source(
        owner_user_id="user_directory",
        package_id=package.id,
        source_id=queued.id,
    ) is None


def test_failed_rebuild_preserves_previous_catalog_version(tmp_path: Path) -> None:
    path = tmp_path / "source.md"
    path.write_text("# One\nBody\n# Two\nBody", encoding="utf-8")
    store = SourceStructureStore(tmp_path / "openclass.sqlite3")
    first = SourceDirectoryProcessor(
        store=store,
        normalizer_factory=lambda _record: PassthroughNormalizer(),
    ).process(record=_record(path), path=path, catalog_model=_model())
    failing = SourceDirectoryProcessor(
        store=store,
        normalizer_factory=lambda _record: FailingNormalizer(),
    )

    with pytest.raises(SourceDirectoryProcessingError, match="catalog model failed"):
        failing.process(record=_record(path), path=path, catalog_model=_model())

    preserved = store.get_structure(
        owner_user_id="user_directory",
        package_id="course_directory",
        source_id="source_directory",
    )
    assert preserved is not None
    assert preserved.id == first.id
    assert preserved.catalog_version == 1
    runs = store.list_catalog_runs(
        owner_user_id="user_directory",
        package_id="course_directory",
        source_id="source_directory",
    )
    assert [run.status for run in runs] == ["failed", "succeeded"]


def test_codex_normalizer_executes_bounded_batches_serially(monkeypatch) -> None:
    calls: list[int] = []
    runtime_settings: list[tuple[str | None, str | None, bool]] = []
    progress_updates: list[tuple[str, int]] = []

    def fake_parse(self, **kwargs):
        packet = json.loads(kwargs["user_prompt"].split("\n", 1)[1])
        calls.append(packet["batch_index"])
        runtime_settings.append(
            (
                kwargs["reasoning_effort"],
                kwargs["service_tier"],
                kwargs["service_tier_is_set"],
            )
        )
        return SimpleNamespace(
            output_parsed=DirectoryBatchDecision(
                batch_hash=packet["batch_hash"],
                decisions=[
                    DirectoryNodeDecision(
                        local_key=node["local_key"],
                        keep=True,
                        title=node["title"],
                        number=node["number"],
                        level=node["level"],
                    )
                    for node in packet["nodes"]
                ],
            )
        )

    monkeypatch.setattr(
        "app.services.source_directory_processor.CodexAppServerTextClient.parse",
        fake_parse,
    )
    candidates = [
        DirectoryCandidate(
            local_key=f"node-{index}",
            title=f"Node {index}",
            order_index=index,
            source_locator=f"text:line:{index + 1}",
            source_range=SourceRange(
                kind="text_lines",
                start=index + 1,
                end=index + 1,
                metadata={"index_base": 1},
            ),
            mapping_status="verified",
            confidence=1.0,
        )
        for index in range(241)
    ]

    result = CodexDirectoryNormalizer(
        user_id="user_directory",
        progress_callback=lambda phase, value: progress_updates.append((phase, value)),
    ).normalize(
        record=SourceIngestionRecord(
            id="source_batches",
            owner_user_id="user_directory",
            package_id="course_directory",
            title="Batches",
            source_type="local_file",
            file_name="batches.md",
            mime_type="text/markdown",
        ),
        candidates=candidates,
        selection=AIModelSelection(
            provider="openai_codex",
            model="catalog-test-model",
            reasoning_effort="high",
            service_tier="priority",
        ),
    )

    assert calls == [0, 1, 2]
    assert runtime_settings == [("high", "priority", True)] * 3
    assert progress_updates == [
        ("normalizing_directory", 69),
        ("normalizing_directory", 75),
        ("normalizing_directory", 80),
    ]
    assert result.turn_count == 3
    assert len(result.candidates) == 241


def test_codex_normalizer_preserves_native_epub_levels_across_batches(monkeypatch) -> None:
    calls: list[int] = []

    def fake_parse(self, **kwargs):
        packet = json.loads(kwargs["user_prompt"].split("\n", 1)[1])
        calls.append(packet["batch_index"])
        return SimpleNamespace(
            output_parsed=DirectoryBatchDecision(
                batch_hash=packet["batch_hash"],
                decisions=[
                    DirectoryNodeDecision(
                        local_key=node["local_key"],
                        keep=False,
                        title=node["title"],
                        number=node["number"],
                        level=1,
                    )
                    for node in packet["nodes"]
                ],
            )
        )

    monkeypatch.setattr(
        "app.services.source_directory_processor.CodexAppServerTextClient.parse",
        fake_parse,
    )
    monkeypatch.setattr("app.services.source_directory_processor.MAX_CODEX_BATCH_NODES", 2)
    candidates = [
        DirectoryCandidate(
            local_key=f"epub-{index}",
            title=f"Node {index}",
            level=native_level,
            order_index=index,
            source_locator=f"epub:OPS/s1.xhtml#node-{index}",
            source_range=SourceRange(
                kind="epub_spine",
                start=0,
                end=0,
                container="OPS/s1.xhtml",
                start_anchor=f"node-{index}",
            ),
            mapping_status="verified",
            confidence=1.0,
            metadata={
                "navigation_provenance": "native",
                "hierarchy_locked": True,
                "native_level": native_level,
            },
        )
        for index, native_level in enumerate((1, 2, 3))
    ]

    result = CodexDirectoryNormalizer(user_id="user_directory").normalize(
        record=SourceIngestionRecord(
            id="source_epub_batches",
            owner_user_id="user_directory",
            package_id="course_directory",
            title="EPUB batches",
            source_type="local_file",
            file_name="batches.epub",
            mime_type="application/epub+zip",
        ),
        candidates=candidates,
        selection=_model(),
    )

    assert calls == [0, 1]
    assert [candidate.level for candidate in result.candidates] == [1, 2, 3]
    assert [candidate.source_locator for candidate in result.candidates] == [
        candidate.source_locator for candidate in candidates
    ]
