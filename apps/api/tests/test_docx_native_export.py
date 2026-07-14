from __future__ import annotations

import xml.etree.ElementTree as ET
from zipfile import ZipFile

import pytest
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Cm

from app.services.docx_quality_check import DocxExportQualityError, assert_docx_export_quality, inspect_docx_export
from app.services.docx_styles import apply_textbook_docx_styles
from app.services.latex_to_omml import append_omml_math
from app.services.rich_document import build_document, export_docx

_NS = {
    "m": "http://schemas.openxmlformats.org/officeDocument/2006/math",
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
}


def _document_root(path):
    with ZipFile(path) as archive:
        return ET.fromstring(archive.read("word/document.xml"))


def _math_text(root: ET.Element) -> str:
    return "".join(node.text or "" for node in root.findall(".//m:t", _NS))


def test_latex_to_omml_converts_core_math_structures(tmp_path) -> None:
    document = Document()
    paragraph = document.add_paragraph()
    append_omml_math(
        paragraph,
        r"\|P\|=\max_{1\le i\le n}\Delta x_i,\quad S(P,\xi)=\sum_{i=1}^{n}f(\xi_i)",
        display=True,
    )
    export_path = tmp_path / "math.docx"
    document.save(export_path)

    root = _document_root(export_path)
    math_text = _math_text(root)

    assert root.findall(".//m:oMathPara", _NS)
    assert root.findall(".//m:sSubSup", _NS)
    assert "‖P‖" in math_text
    assert "ξ" in math_text
    assert "∑" in math_text


def test_docx_quality_check_rejects_raw_latex_text(tmp_path) -> None:
    document = Document()
    document.add_paragraph(r"公式残留：\frac{1}{2}")
    export_path = tmp_path / "bad.docx"
    document.save(export_path)

    with pytest.raises(DocxExportQualityError):
        assert_docx_export_quality(export_path)


def test_docx_quality_check_accepts_native_export(tmp_path) -> None:
    document = build_document(
        title="Doc",
        content_text=(
            "# 标题\n\n"
            "设 \\(f\\) 在 \\([a,b]\\) 上可积。\n\n"
            "$$\\Phi(x)=\\int_a^x f(t)\\,dt,\\quad \\forall x\\in[a,b]$$"
        ),
    )
    export_path = tmp_path / "good.docx"

    export_docx(document, export_path)
    report = inspect_docx_export(export_path)

    assert report.passed
    assert report.omml_count > 0
    assert report.display_omml_count > 0


def test_docx_styles_create_textbook_styles() -> None:
    document = Document()

    apply_textbook_docx_styles(document)

    assert document.styles["OpenClass Body Text"]
    assert document.styles["OpenClass Compact"]
    assert document.styles["OpenClass Formula"]
    assert document.styles["OpenClass Preformatted"]
    assert document.styles["OpenClass Code Line"]
    assert document.styles["OpenClass Code Caption"]
    assert document.styles["OpenClass Figure Caption"]
    assert document.styles["Heading 1"].font.bold is True
    title = document.styles["Title"]
    assert title.element.get_or_add_pPr().find(qn("w:pBdr")) is None
    assert title.element.get_or_add_rPr().find(qn("w:szCs")).get(qn("w:val")) == "36"


def test_docx_export_renders_code_listing_without_math_or_lost_indentation(tmp_path) -> None:
    document = build_document(
        title="4.2 Structured document",
        content_text=(
            "# 4.2 Structured document\n\n"
            "Use `list_sum` and keep `\"1010\"` as code.\n\n"
            "```python title=\"Accumulator\"\n"
            "def list_sum(values):\n"
            "    total = 0\n"
            "    for value in values:\n"
            "        total += value\n"
            "    return total\n"
            "```"
        ),
    )
    export_path = tmp_path / "code-listing.docx"

    export_docx(document, export_path)
    report = inspect_docx_export(export_path)
    root = _document_root(export_path)
    code_rows = [
        "".join(node.text or "" for node in row.findall(".//w:t", _NS))
        for row in root.findall(".//w:tr", _NS)
    ]
    body_text = "".join(node.text or "" for node in root.findall(".//w:t", _NS))
    math_text = _math_text(root)
    table_layout = root.find(".//w:tblLayout", _NS)

    assert report.passed
    assert report.code_listing_count == 1
    assert report.code_line_count == 5
    assert report.code_math_count == 0
    assert report.duplicate_title_count == 0
    assert report.page_number_field_count == 1
    assert table_layout is not None
    assert table_layout.attrib[qn("w:type")] == "fixed"
    assert code_rows == [
        "1def list_sum(values):",
        "2    total = 0",
        "3    for value in values:",
        "4        total += value",
        "5    return total",
    ]
    assert "代码清单 4-1　Accumulator" in body_text
    assert "list_sum" not in math_text
    assert '"1010"' in body_text


def test_docx_export_renders_structured_diagram_as_numbered_figure(tmp_path) -> None:
    document = build_document(
        title="2.3 Process",
        content_text=(
            "## 2.3.1 Process stages\n\n"
            "```openclass-diagram\n"
            "{\n"
            '  "caption": "State transition",\n'
            '  "direction": "down",\n'
            '  "nodes": [\n'
            '    {"id": "start", "label": "Initial state"},\n'
            '    {"id": "finish", "label": "Final state"}\n'
            "  ],\n"
            '  "edges": [{"from": "start", "to": "finish", "label": "advance"}]\n'
            "}\n"
            "```"
        ),
    )
    export_path = tmp_path / "diagram.docx"

    export_docx(document, export_path)
    report = inspect_docx_export(export_path)
    root = _document_root(export_path)
    visible_text = "".join(node.text or "" for node in root.findall(".//w:t", _NS))
    extent = root.find(".//wp:extent", _NS)

    assert report.passed
    assert report.figure_count == 1
    assert report.figure_caption_count == 1
    assert report.media_count == 1
    assert extent is not None
    assert int(extent.attrib["cy"]) <= Cm(11.5)
    assert "图 2-1　State transition" in visible_text
    assert "openclass-diagram" not in visible_text
