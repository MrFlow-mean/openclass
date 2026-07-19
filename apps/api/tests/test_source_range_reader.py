from __future__ import annotations

import hashlib
import json
import zipfile
from pathlib import Path
from types import SimpleNamespace

import pytest
from pypdf import PdfWriter
from reportlab.pdfgen import canvas

from app.models import (
    LearningSourceGrounding,
    LearningSourceReference,
    Lesson,
    RetrievalEvidence,
    SelectionRef,
    SourceChapter,
    SourceIngestionRecord,
    SourceRange,
    SourceStructure,
    SourceStructureView,
)
from app.services import codex_chat, source_grounded_board, source_range_reader
from app.services.lesson_factory import build_requirements
from app.services.source_range_reader import SourceRangeReadError, read_verified_source_range


OWNER_ID = "user_catalog_reader"
PACKAGE_ID = "package_catalog_reader"


def _sha256(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


def _write_reordered_pptx(path: Path) -> None:
    presentation = """<?xml version="1.0" encoding="UTF-8"?>
<p:presentation xmlns:p="http://schemas.openxmlformats.org/presentationml/2006/main"
 xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">
  <p:sldIdLst>
    <p:sldId id="256" r:id="rId9"/>
    <p:sldId id="257" r:id="rId2"/>
  </p:sldIdLst>
</p:presentation>"""
    relationships = """<?xml version="1.0" encoding="UTF-8"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="rId2" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/slide" Target="slides/slide2.xml"/>
  <Relationship Id="rId9" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/slide" Target="slides/slide9.xml"/>
</Relationships>"""

    def slide(title: str, body: str) -> str:
        return f"""<?xml version="1.0" encoding="UTF-8"?>
<p:sld xmlns:p="http://schemas.openxmlformats.org/presentationml/2006/main"
 xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main">
  <p:cSld><p:spTree><p:sp><p:txBody>
    <a:p><a:r><a:t>{title}</a:t></a:r></a:p>
    <a:p><a:r><a:t>{body}</a:t></a:r></a:p>
  </p:txBody></p:sp></p:spTree></p:cSld>
</p:sld>"""

    with zipfile.ZipFile(path, "w") as archive:
        archive.writestr("ppt/presentation.xml", presentation)
        archive.writestr("ppt/_rels/presentation.xml.rels", relationships)
        archive.writestr(
            "ppt/slides/slide9.xml",
            slide("Playback first", "FIRST_PLAYBACK_BODY_EVIDENCE"),
        )
        archive.writestr(
            "ppt/slides/slide2.xml",
            slide("Playback second", "SECOND_PLAYBACK_BODY"),
        )
        archive.writestr(
            "ppt/slides/slide1.xml",
            slide("Orphan slide", "ORPHAN_BODY"),
        )


def _catalog_objects(
    path: Path,
    source_range: SourceRange,
    *,
    owner_user_id: str = OWNER_ID,
) -> tuple[SourceIngestionRecord, SourceStructure, SourceChapter, SelectionRef]:
    content_hash = _sha256(path)
    source = SourceIngestionRecord(
        id="source_catalog_reader",
        owner_user_id=owner_user_id,
        package_id=PACKAGE_ID,
        title="Catalog source",
        source_type="local_file",
        file_name=path.name,
        mime_type=("application/pdf" if path.suffix.lower() == ".pdf" else "text/plain"),
        size_bytes=path.stat().st_size,
        status="ready",
        metadata={
            "content_hash": content_hash,
            "catalog_pipeline": "codex_directory_v1",
        },
    )
    structure = SourceStructure(
        id="structure_catalog_reader",
        owner_user_id=owner_user_id,
        package_id=PACKAGE_ID,
        source_ingestion_id=source.id,
        status="ready",
        strategy="codex_directory_v1",
        catalog_version=3,
        source_content_hash=content_hash,
        catalog_schema_version="codex_directory_v1",
        metadata={"catalog_pipeline": "codex_directory_v1"},
    )
    chapter = SourceChapter(
        id="chapter_catalog_reader",
        owner_user_id=owner_user_id,
        package_id=PACKAGE_ID,
        source_ingestion_id=source.id,
        title="Selected chapter",
        path=["Selected chapter"],
        range=source_range,
        mapping_status="verified",
        source_content_hash=content_hash,
        catalog_version=3,
        confidence=0.95,
    )
    selection = SelectionRef(
        kind="source",
        excerpt="Catalog source / Selected chapter",
        source_ingestion_id=source.id,
        source_chapter_id=chapter.id,
        source_chapter_title=chapter.title,
        heading_path=chapter.path,
        source_range=source_range,
        catalog_version=3,
        source_content_hash=content_hash,
    )
    return source, structure, chapter, selection


def test_text_range_reads_only_selected_lines_without_creating_chunk_ids(
    monkeypatch: pytest.MonkeyPatch,
    tmp_path: Path,
) -> None:
    path = tmp_path / "reference.txt"
    path.write_text(
        "outside-before\nselected alpha evidence\nselected beta evidence\noutside-after\n",
        encoding="utf-8",
    )
    source_range = SourceRange(
        kind="text_lines",
        start=2,
        end=3,
        display_label="lines 2-3",
        metadata={"index_base": 1},
    )
    source, structure, chapter, selection = _catalog_objects(path, source_range)
    monkeypatch.setattr(source_range_reader, "_source_path", lambda _source: path)

    result = read_verified_source_range(
        owner_user_id=OWNER_ID,
        package_id=PACKAGE_ID,
        source=source,
        structure=structure,
        chapter=chapter,
        selection=selection,
    )

    assert len(result.evidence_items) == 1
    evidence = result.evidence_items[0]
    assert "selected alpha evidence" in evidence.expanded_text
    assert "selected beta evidence" in evidence.expanded_text
    assert "outside-before" not in evidence.expanded_text
    assert "outside-after" not in evidence.expanded_text
    assert evidence.chunk_ids == []
    assert evidence.metadata["source_range"]["end"] == 3
    assert evidence.metadata["range_end_inclusive"] is True
    assert evidence.metadata["source_content_hash"] == _sha256(path)


def test_docx_range_reads_table_cell_paragraphs_in_catalog_sequence(
    monkeypatch: pytest.MonkeyPatch,
    tmp_path: Path,
) -> None:
    from docx import Document

    path = tmp_path / "table-range.docx"
    document = Document()
    document.add_paragraph("PREFACE_OUTSIDE_RANGE")
    table = document.add_table(rows=1, cols=1)
    table_heading = table.cell(0, 0).paragraphs[0]
    table_heading.text = "Table chapter"
    table_heading.style = "Heading 1"
    table.cell(0, 0).add_paragraph("TABLE_CELL_BODY_EVIDENCE")
    document.add_paragraph("AFTER_TABLE_BODY_EVIDENCE")
    document.add_heading("Next chapter", level=1)
    document.add_paragraph("NEXT_CHAPTER_BODY")
    document.save(path)
    source_range = SourceRange(
        kind="docx_paragraphs",
        start=1,
        end=3,
        display_label="Paragraphs 2-4",
        metadata={"index_base": 0},
    )
    source, structure, chapter, selection = _catalog_objects(path, source_range)
    monkeypatch.setattr(source_range_reader, "_source_path", lambda _source: path)

    result = read_verified_source_range(
        owner_user_id=OWNER_ID,
        package_id=PACKAGE_ID,
        source=source,
        structure=structure,
        chapter=chapter,
        selection=selection,
    )

    text = "\n".join(item.expanded_text for item in result.evidence_items)
    assert "Table chapter" in text
    assert "TABLE_CELL_BODY_EVIDENCE" in text
    assert "AFTER_TABLE_BODY_EVIDENCE" in text
    assert "PREFACE_OUTSIDE_RANGE" not in text
    assert "Next chapter" not in text
    assert "NEXT_CHAPTER_BODY" not in text


def test_pptx_range_maps_logical_playback_position_to_relationship_target(
    monkeypatch: pytest.MonkeyPatch,
    tmp_path: Path,
) -> None:
    path = tmp_path / "reordered-range.pptx"
    _write_reordered_pptx(path)
    source_range = SourceRange(
        kind="ppt_slides",
        start=1,
        end=1,
        display_label="Slide 1",
        metadata={"index_base": 1},
    )
    source, structure, chapter, selection = _catalog_objects(path, source_range)
    monkeypatch.setattr(source_range_reader, "_source_path", lambda _source: path)

    result = read_verified_source_range(
        owner_user_id=OWNER_ID,
        package_id=PACKAGE_ID,
        source=source,
        structure=structure,
        chapter=chapter,
        selection=selection,
    )

    text = "\n".join(item.expanded_text for item in result.evidence_items)
    assert "Playback first" in text
    assert "FIRST_PLAYBACK_BODY_EVIDENCE" in text
    assert "Playback second" not in text
    assert "SECOND_PLAYBACK_BODY" not in text
    assert "Orphan slide" not in text
    assert "ORPHAN_BODY" not in text


@pytest.mark.parametrize(
    ("mutation", "message_fragment"),
    [
        ("range", "资料范围已失效或被修改"),
        ("version", "目录引用的版本已经失效"),
        ("hash", "文件指纹与目录引用不一致"),
        ("chapter", "章节标识已失效或被修改"),
    ],
)
def test_reader_rejects_forged_or_stale_reference_fields(
    monkeypatch: pytest.MonkeyPatch,
    tmp_path: Path,
    mutation: str,
    message_fragment: str,
) -> None:
    path = tmp_path / "reference.txt"
    path.write_text("line one evidence\nline two evidence\nline three evidence\n", encoding="utf-8")
    source_range = SourceRange(
        kind="text_lines",
        start=1,
        end=2,
        metadata={"index_base": 1},
    )
    source, structure, chapter, selection = _catalog_objects(path, source_range)
    if mutation == "range":
        selection.source_range = SourceRange(
            kind="text_lines",
            start=1,
            end=3,
            metadata={"index_base": 1},
        )
    elif mutation == "version":
        selection.catalog_version = 2
    elif mutation == "hash":
        selection.source_content_hash = "f" * 64
    elif mutation == "chapter":
        selection.source_chapter_id = "chapter_forged"
    monkeypatch.setattr(source_range_reader, "_source_path", lambda _source: path)

    with pytest.raises(SourceRangeReadError, match=message_fragment):
        read_verified_source_range(
            owner_user_id=OWNER_ID,
            package_id=PACKAGE_ID,
            source=source,
            structure=structure,
            chapter=chapter,
            selection=selection,
        )


def test_reader_rejects_cross_user_source_before_opening_file(
    monkeypatch: pytest.MonkeyPatch,
    tmp_path: Path,
) -> None:
    path = tmp_path / "reference.txt"
    path.write_text("selected source evidence\n", encoding="utf-8")
    source_range = SourceRange(kind="text_lines", start=1, end=1)
    source, structure, chapter, selection = _catalog_objects(
        path,
        source_range,
        owner_user_id="another_user",
    )
    opened = False

    def unexpected_path(_source):
        nonlocal opened
        opened = True
        return path

    monkeypatch.setattr(source_range_reader, "_source_path", unexpected_path)
    with pytest.raises(SourceRangeReadError, match="无权读取"):
        read_verified_source_range(
            owner_user_id=OWNER_ID,
            package_id=PACKAGE_ID,
            source=source,
            structure=structure,
            chapter=chapter,
            selection=selection,
        )
    assert opened is False


def test_reader_rejects_file_replacement_after_catalog_publication(
    monkeypatch: pytest.MonkeyPatch,
    tmp_path: Path,
) -> None:
    path = tmp_path / "reference.txt"
    path.write_text("original selected source evidence\n", encoding="utf-8")
    source_range = SourceRange(kind="text_lines", start=1, end=1)
    source, structure, chapter, selection = _catalog_objects(path, source_range)
    path.write_text("replacement selected source evidence\n", encoding="utf-8")
    monkeypatch.setattr(source_range_reader, "_source_path", lambda _source: path)

    with pytest.raises(SourceRangeReadError, match="文件内容已经变化"):
        read_verified_source_range(
            owner_user_id=OWNER_ID,
            package_id=PACKAGE_ID,
            source=source,
            structure=structure,
            chapter=chapter,
            selection=selection,
        )


def test_native_pdf_reads_only_inclusive_selected_physical_pages(
    monkeypatch: pytest.MonkeyPatch,
    tmp_path: Path,
) -> None:
    path = tmp_path / "native.pdf"
    document = canvas.Canvas(str(path))
    for page_no in range(1, 5):
        document.drawString(72, 720, f"UNIQUE_PAGE_{page_no} durable selected evidence")
        document.showPage()
    document.save()
    source_range = SourceRange(
        kind="pdf_pages",
        start=2,
        end=3,
        display_label="PDF pp. 2-3",
        metadata={"index_base": 1},
    )
    source, structure, chapter, selection = _catalog_objects(path, source_range)
    monkeypatch.setattr(source_range_reader, "_source_path", lambda _source: path)
    monkeypatch.setattr(
        source_range_reader,
        "extract_pdf_pages_layout",
        lambda *_args, **_kwargs: (_ for _ in ()).throw(
            AssertionError("native selected pages must not invoke OCR")
        ),
    )

    result = read_verified_source_range(
        owner_user_id=OWNER_ID,
        package_id=PACKAGE_ID,
        source=source,
        structure=structure,
        chapter=chapter,
        selection=selection,
    )

    combined = "\n".join(item.expanded_text for item in result.evidence_items)
    assert "UNIQUE_PAGE_2" in combined
    assert "UNIQUE_PAGE_3" in combined
    assert "UNIQUE_PAGE_1" not in combined
    assert "UNIQUE_PAGE_4" not in combined
    assert result.evidence_items[0].metadata["page_start"] == 2
    assert result.evidence_items[0].metadata["page_end_inclusive"] == 3


def test_pdf_sparse_watermark_text_layer_does_not_suppress_selected_page_ocr(
    monkeypatch: pytest.MonkeyPatch,
    tmp_path: Path,
) -> None:
    path = tmp_path / "watermarked-scan.pdf"
    document = canvas.Canvas(str(path))
    document.drawString(240, 24, "CONFIDENTIAL 2024")
    document.showPage()
    document.save()
    source_range = SourceRange(
        kind="pdf_pages",
        start=1,
        end=1,
        display_label="PDF p. 1",
        metadata={"index_base": 1},
    )
    source, structure, chapter, selection = _catalog_objects(path, source_range)
    monkeypatch.setattr(source_range_reader, "_source_path", lambda _source: path)
    monkeypatch.setattr(source_range_reader, "extract_pdf_pages_layout", lambda *_args, **_kwargs: [])
    monkeypatch.setattr(
        source_range_reader,
        "extract_pdf_pages_text",
        lambda *_args, **_kwargs: "OCR recovered the selected scanned page body evidence.",
    )

    result = read_verified_source_range(
        owner_user_id=OWNER_ID,
        package_id=PACKAGE_ID,
        source=source,
        structure=structure,
        chapter=chapter,
        selection=selection,
    )

    assert result.evidence_items[0].metadata["retrieval_mode"] == "on_demand_pdf_ocr"
    assert "selected scanned page body" in result.evidence_items[0].expanded_text


def test_scanned_pdf_ocrs_only_selected_pages_in_serial_bounded_batches(
    monkeypatch: pytest.MonkeyPatch,
    tmp_path: Path,
) -> None:
    path = tmp_path / "scanned.pdf"
    writer = PdfWriter()
    for _ in range(12):
        writer.add_blank_page(width=612, height=792)
    with path.open("wb") as handle:
        writer.write(handle)
    source_range = SourceRange(
        kind="pdf_pages",
        start=2,
        end=10,
        display_label="PDF pp. 2-10",
        metadata={"index_base": 1},
    )
    source, structure, chapter, selection = _catalog_objects(path, source_range)
    monkeypatch.setattr(source_range_reader, "_source_path", lambda _source: path)
    monkeypatch.setattr(source_range_reader, "extract_pdf_pages_layout", lambda *_args, **_kwargs: [])
    ocr_calls: list[tuple[int, int, int]] = []

    def fake_ocr(_path: Path, *, page_start: int, page_end: int, max_pages: int) -> str:
        ocr_calls.append((page_start, page_end, max_pages))
        return f"OCR selected evidence for physical pages {page_start} through {page_end}."

    monkeypatch.setattr(source_range_reader, "extract_pdf_pages_text", fake_ocr)

    result = read_verified_source_range(
        owner_user_id=OWNER_ID,
        package_id=PACKAGE_ID,
        source=source,
        structure=structure,
        chapter=chapter,
        selection=selection,
    )

    assert ocr_calls == [(2, 9, 8), (10, 10, 1)]
    assert all(call[0] >= 2 and call[1] <= 10 for call in ocr_calls)
    assert [item.metadata["batch_index"] for item in result.evidence_items] == [1, 2]
    assert all(item.metadata["retrieval_mode"] == "on_demand_pdf_ocr" for item in result.evidence_items)


def test_json_bracketed_array_path_reads_the_selected_node(
    monkeypatch: pytest.MonkeyPatch,
    tmp_path: Path,
) -> None:
    path = tmp_path / "reference.json"
    path.write_text(
        json.dumps(
            {
                "items": [
                    {"name": "outside first object"},
                    {"name": "selected bracket-index evidence"},
                ]
            }
        ),
        encoding="utf-8",
    )
    source_range = SourceRange(
        kind="structured_path",
        path=["items", "[1]"],
        display_label="/items/[1]",
        metadata={"syntax": "json_pointer"},
    )
    source, structure, chapter, selection = _catalog_objects(path, source_range)
    monkeypatch.setattr(source_range_reader, "_source_path", lambda _source: path)

    result = read_verified_source_range(
        owner_user_id=OWNER_ID,
        package_id=PACKAGE_ID,
        source=source,
        structure=structure,
        chapter=chapter,
        selection=selection,
    )

    text = result.evidence_items[0].expanded_text
    assert "selected bracket-index evidence" in text
    assert "outside first object" not in text


def test_xml_indexed_sibling_path_reads_the_requested_occurrence(
    monkeypatch: pytest.MonkeyPatch,
    tmp_path: Path,
) -> None:
    path = tmp_path / "reference.xml"
    path.write_text(
        "<root><item>outside first item</item>"
        "<item><value>selected second-item evidence</value></item></root>",
        encoding="utf-8",
    )
    source_range = SourceRange(
        kind="structured_path",
        path=["root", "item[2]"],
        display_label="/root/item[2]",
        metadata={"syntax": "indexed_xml_path"},
    )
    source, structure, chapter, selection = _catalog_objects(path, source_range)
    monkeypatch.setattr(source_range_reader, "_source_path", lambda _source: path)

    result = read_verified_source_range(
        owner_user_id=OWNER_ID,
        package_id=PACKAGE_ID,
        source=source,
        structure=structure,
        chapter=chapter,
        selection=selection,
    )

    text = result.evidence_items[0].expanded_text
    assert "selected second-item evidence" in text
    assert "outside first item" not in text


def test_epub_container_validates_start_item_and_numeric_end_spans_parent_range(
    monkeypatch: pytest.MonkeyPatch,
    tmp_path: Path,
) -> None:
    path = tmp_path / "reference.epub"
    with zipfile.ZipFile(path, "w") as archive:
        archive.writestr(
            "META-INF/container.xml",
            "<container><rootfiles><rootfile full-path='OPS/content.opf'/></rootfiles></container>",
        )
        archive.writestr(
            "OPS/content.opf",
            "<package><manifest>"
            "<item id='s1' href='s1.xhtml'/><item id='s2' href='s2.xhtml'/>"
            "<item id='s3' href='s3.xhtml'/></manifest>"
            "<spine><itemref idref='s1'/><itemref idref='s2'/><itemref idref='s3'/></spine>"
            "</package>",
        )
        archive.writestr(
            "OPS/s1.xhtml",
            "<html><body><h1 id='start'>Parent start</h1><p>first spine evidence</p></body></html>",
        )
        archive.writestr(
            "OPS/s2.xhtml",
            "<html><body><p>middle spine evidence</p></body></html>",
        )
        archive.writestr(
            "OPS/s3.xhtml",
            "<html><body><p>final spine evidence</p></body></html>",
        )
    source_range = SourceRange(
        kind="epub_spine",
        start=0,
        end=2,
        container="OPS/s1.xhtml",
        start_anchor="start",
        display_label="OPS/s1.xhtml-OPS/s3.xhtml",
        metadata={"index_base": 0, "href": "OPS/s1.xhtml"},
    )
    source, structure, chapter, selection = _catalog_objects(path, source_range)
    monkeypatch.setattr(source_range_reader, "_source_path", lambda _source: path)

    result = read_verified_source_range(
        owner_user_id=OWNER_ID,
        package_id=PACKAGE_ID,
        source=source,
        structure=structure,
        chapter=chapter,
        selection=selection,
    )

    text = "\n".join(item.expanded_text for item in result.evidence_items)
    assert "first spine evidence" in text
    assert "middle spine evidence" in text
    assert "final spine evidence" in text


def test_epub_end_anchor_keeps_prefix_of_boundary_spine(
    monkeypatch: pytest.MonkeyPatch,
    tmp_path: Path,
) -> None:
    path = tmp_path / "anchored-boundary.epub"
    with zipfile.ZipFile(path, "w") as archive:
        archive.writestr(
            "META-INF/container.xml",
            "<container><rootfiles><rootfile full-path='OPS/content.opf'/></rootfiles></container>",
        )
        archive.writestr(
            "OPS/content.opf",
            "<package><manifest>"
            "<item id='s1' href='s1.xhtml'/><item id='s2' href='s2.xhtml'/>"
            "</manifest><spine><itemref idref='s1'/><itemref idref='s2'/></spine>"
            "</package>",
        )
        archive.writestr(
            "OPS/s1.xhtml",
            "<html><body><h1 id='start'>Start</h1><p>first spine evidence</p></body></html>",
        )
        archive.writestr(
            "OPS/s2.xhtml",
            "<html><body><p>prefix belongs to previous chapter</p>"
            "<h1 id='next'>Next</h1><p>next chapter evidence</p></body></html>",
        )
    source_range = SourceRange(
        kind="epub_spine",
        start=0,
        end=1,
        container="OPS/s1.xhtml",
        start_anchor="start",
        end_anchor="next",
        display_label="EPUB spine 0-1",
        metadata={"index_base": 0, "href": "OPS/s1.xhtml"},
    )
    source, structure, chapter, selection = _catalog_objects(path, source_range)
    monkeypatch.setattr(source_range_reader, "_source_path", lambda _source: path)

    result = read_verified_source_range(
        owner_user_id=OWNER_ID,
        package_id=PACKAGE_ID,
        source=source,
        structure=structure,
        chapter=chapter,
        selection=selection,
    )

    text = "\n".join(item.expanded_text for item in result.evidence_items)
    assert "first spine evidence" in text
    assert "prefix belongs to previous chapter" in text
    assert "next chapter evidence" not in text


def test_dom_synthetic_heading_ordinals_and_whole_document_ranges_are_readable(
    monkeypatch: pytest.MonkeyPatch,
    tmp_path: Path,
) -> None:
    path = tmp_path / "reference.html"
    path.write_text(
        "<html><body><h1>First heading</h1><p>selected first-section evidence</p>"
        "<h2>Second heading</h2><p>outside second-section evidence</p></body></html>",
        encoding="utf-8",
    )
    selected_range = SourceRange(
        kind="dom_anchor",
        container=path.name,
        start_anchor="heading-1",
        end_anchor="heading-2",
        display_label="#heading-1",
        metadata={"heading_ordinal": 0},
    )
    source, structure, chapter, selection = _catalog_objects(path, selected_range)
    monkeypatch.setattr(source_range_reader, "_source_path", lambda _source: path)

    selected = read_verified_source_range(
        owner_user_id=OWNER_ID,
        package_id=PACKAGE_ID,
        source=source,
        structure=structure,
        chapter=chapter,
        selection=selection,
    )

    selected_text = selected.evidence_items[0].expanded_text
    assert "selected first-section evidence" in selected_text
    assert "outside second-section evidence" not in selected_text

    whole_range = SourceRange(
        kind="dom_anchor",
        container=path.name,
        display_label="Whole document",
    )
    whole_source, whole_structure, whole_chapter, whole_selection = _catalog_objects(
        path,
        whole_range,
    )
    whole = read_verified_source_range(
        owner_user_id=OWNER_ID,
        package_id=PACKAGE_ID,
        source=whole_source,
        structure=whole_structure,
        chapter=whole_chapter,
        selection=whole_selection,
    )
    whole_text = whole.evidence_items[0].expanded_text
    assert "selected first-section evidence" in whole_text
    assert "outside second-section evidence" in whole_text


def test_empty_structured_path_reads_json_root_range(
    monkeypatch: pytest.MonkeyPatch,
    tmp_path: Path,
) -> None:
    path = tmp_path / "root.json"
    path.write_text('{"rootEvidence":"selected whole JSON document"}', encoding="utf-8")
    source_range = SourceRange(
        kind="structured_path",
        path=[],
        display_label="/",
    )
    source, structure, chapter, selection = _catalog_objects(path, source_range)
    monkeypatch.setattr(source_range_reader, "_source_path", lambda _source: path)

    result = read_verified_source_range(
        owner_user_id=OWNER_ID,
        package_id=PACKAGE_ID,
        source=source,
        structure=structure,
        chapter=chapter,
        selection=selection,
    )

    assert "selected whole JSON document" in result.evidence_items[0].expanded_text


def test_grounded_board_uses_range_reader_and_freezes_traceable_bundle(
    monkeypatch: pytest.MonkeyPatch,
    tmp_path: Path,
) -> None:
    path = tmp_path / "reference.txt"
    path.write_text(
        "outside before\nselected grounded evidence\nselected worked evidence\noutside after\n",
        encoding="utf-8",
    )
    source_range = SourceRange(
        kind="text_lines",
        start=2,
        end=3,
        display_label="lines 2-3",
        metadata={"index_base": 1},
    )
    source, structure, chapter, selection = _catalog_objects(path, source_range)
    lesson = Lesson.model_construct(id="lesson_catalog_reader")
    saved_bundles = []
    monkeypatch.setattr(source_range_reader, "_source_path", lambda _source: path)
    monkeypatch.setattr(
        source_grounded_board.workspace_state,
        "load_workspace_for_user",
        lambda _user_id: object(),
    )
    monkeypatch.setattr(
        source_grounded_board.workspace_state,
        "find_lesson_package",
        lambda _workspace, _lesson_id: (SimpleNamespace(id=PACKAGE_ID), lesson),
    )
    monkeypatch.setattr(
        source_grounded_board.source_evidence_store,
        "get_source",
        lambda **_kwargs: source,
    )
    monkeypatch.setattr(
        source_grounded_board.source_evidence_store,
        "save_bundle",
        lambda bundle: saved_bundles.append(bundle) or bundle,
    )
    monkeypatch.setattr(
        source_grounded_board.source_structure_store,
        "get_structure_view",
        lambda **_kwargs: SourceStructureView(
            source=source,
            structure=structure,
            chapters=[chapter],
            chunks=[],
            visuals=[],
        ),
    )
    monkeypatch.setattr(
        source_grounded_board.source_structure_store,
        "get_catalog_chapter",
        lambda **_kwargs: (structure, chapter),
    )
    monkeypatch.setattr(
        source_grounded_board.source_structure_store,
        "chapter_evidence_by_id",
        lambda **_kwargs: (_ for _ in ()).throw(
            AssertionError("new catalogs must not read legacy SourceChunk evidence")
        ),
    )
    monkeypatch.setattr(
        source_grounded_board.source_structure_store,
        "visual_evidence_for_scope",
        lambda **_kwargs: (_ for _ in ()).throw(
            AssertionError("new catalogs must not read a global visual index")
        ),
    )

    plan = source_grounded_board.resolve_source_grounded_board_plan(
        owner_user_id=OWNER_ID,
        lesson=lesson,
        selection=selection,
        query="Generate from the selected range",
    )

    assert plan is not None
    frozen = plan.requirement.source_grounding.frozen_evidence
    assert "selected grounded evidence" in frozen[0].expanded_text
    assert "outside before" not in frozen[0].expanded_text
    assert saved_bundles[0].metadata["catalog_pipeline"] == "codex_directory_v1"
    assert saved_bundles[0].metadata["catalog_version"] == 3
    assert saved_bundles[0].metadata["source_range"]["start"] == 2
    assert saved_bundles[0].evidence_items[0].chunk_ids == []


def test_long_on_demand_range_is_summarized_serially_without_source_chunks(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    source_range = SourceRange(
        kind="pdf_pages",
        start=2,
        end=25,
        display_label="PDF pp. 2-25",
    )
    evidence = [
        RetrievalEvidence(
            id=f"range_evidence_{index}",
            source_ingestion_id="source_catalog_reader",
            source_title="Catalog source",
            chapter_id="chapter_catalog_reader",
            section_path=["Selected chapter"],
            page_range=f"PDF pp. {page_start}-{page_start + 7}",
            expanded_text=character * 120_000,
            token_count=30_000,
            metadata={
                "retrieval_mode": "on_demand_pdf_ocr",
                "source_range": source_range.model_dump(mode="json"),
                "source_locator": f"pdf:pages:{page_start}-{page_start + 7}",
                "catalog_version": 3,
                "source_content_hash": "a" * 64,
            },
        )
        for index, (page_start, character) in enumerate(((2, "a"), (10, "汉")))
    ]
    requirement = build_requirements("On-demand source batch test")
    requirement.source_grounding = LearningSourceGrounding(
        requested_by_user=True,
        confirmation_status="confirmed",
        confirmed_bundle_id="bundle_catalog_reader",
        confirmed_references=[
            LearningSourceReference(
                evidence_bundle_id="bundle_catalog_reader",
                source_ingestion_id="source_catalog_reader",
                source_chapter_id="chapter_catalog_reader",
                chunk_ids=[],
            )
        ],
        frozen_evidence=evidence,
    )
    monkeypatch.setattr(
        codex_chat.source_structure_store,
        "source_chunks_by_ids",
        lambda **_kwargs: (_ for _ in ()).throw(
            AssertionError("on-demand ranges must not reconstruct global SourceChunk rows")
        ),
    )

    class FakeAdapter:
        def __init__(self) -> None:
            self.calls: list[dict[str, object]] = []

        def parse_structured(self, **kwargs):
            payload = json.loads(kwargs["user_prompt"])
            self.calls.append(payload)
            evidence_ids = [item["evidence_id"] for item in payload["evidence"]]
            return SimpleNamespace(
                output_parsed=kwargs["schema"](
                    summary=f"Serial summary for {','.join(evidence_ids)}"
                ),
                activity=[],
            )

        def analyze_image_batch(self, **_kwargs):
            raise AssertionError("this test has no visual evidence")

    adapter = FakeAdapter()
    prepared, image_inputs = codex_chat._prepare_source_generation_inputs(
        adapter=adapter,
        requirement=requirement,
        owner_user_id=OWNER_ID,
        is_cancelled=None,
        on_activity=None,
    )

    assert [call["batch_index"] for call in adapter.calls] == list(
        range(len(adapter.calls))
    )
    called_evidence_ids = [
        call["evidence"][0]["evidence_id"] for call in adapter.calls
    ]
    assert called_evidence_ids.count("range_evidence_0") == 4
    assert called_evidence_ids.count("range_evidence_1") == 12
    assert all(
        len(json.dumps(call, ensure_ascii=False).encode("utf-8")) < 36_000
        for call in adapter.calls
    )
    summaries = prepared.source_grounding.frozen_evidence
    assert len(summaries) == len(adapter.calls)
    assert all(item.chunk_ids == [] for item in summaries)
    assert summaries[0].metadata["covered_page_ranges"] == ["PDF pp. 2-9"]
    assert summaries[4].metadata["covered_source_locators"] == ["pdf:pages:10-17"]
    assert summaries[0].metadata["source_provenance"][0]["text_part_count"] == 4
    assert summaries[4].metadata["source_provenance"][0]["text_part_count"] == 12
    assert sum(len(item.expanded_text.encode("utf-8")) for item in summaries) <= 48_000
    assert image_inputs == []
