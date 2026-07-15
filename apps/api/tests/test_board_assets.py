from __future__ import annotations

import base64
import io
from zipfile import ZipFile

import pytest
from fastapi.testclient import TestClient
from docx import Document as DocxDocument
from PIL import Image

from app.models import BoardDocument, UserView
from app.services import board_asset_store as board_asset_store_module
from app.services import docx_exporter, html_document_export
from app.services.board_asset_store import BoardAssetError, BoardAssetStore
from app.services.html_document_export import HtmlExportBudgetError, standalone_html
from app.services.rich_document import (
    export_docx,
    html_to_tiptap_doc,
    rebuild_document_from_content_json,
    tiptap_doc_to_html,
)


PNG_BYTES = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNk+A8AAQUBAScY42YAAAAASUVORK5CYII="
)
OWNER = UserView(
    id="board_asset_owner",
    email="owner@example.com",
    role="user",
    created_at="2026-01-01T00:00:00+00:00",
)
OTHER_USER = OWNER.model_copy(update={"id": "board_asset_other", "email": "other@example.com"})


@pytest.fixture
def asset_store(tmp_path) -> BoardAssetStore:
    return BoardAssetStore(tmp_path / "openclass.sqlite3", tmp_path / "board-assets")


def _put_asset(store: BoardAssetStore, *, owner_user_id: str = OWNER.id):
    return store.put_bytes(
        owner_user_id=owner_user_id,
        lesson_id="lesson_board_asset",
        document_id="document_board_asset",
        source_visual_id="visual_board_asset",
        content=PNG_BYTES,
        mime_type="image/png",
        file_name="visual.png",
    )


def _visual_document(asset_id: str) -> BoardDocument:
    content_html = (
        '<p>图片前文</p>'
        '<section data-type="resource-visual-block" '
        f'data-board-asset-id="{asset_id}" '
        'data-visual-id="visual_board_asset" '
        'data-source-ingestion-id="source_board_asset" '
        'data-source-chapter-id="chapter_board_asset" '
        'data-source-title="测试资料" data-visual-kind="chart" '
        'data-caption="增长趋势" data-source="测试资料" '
        'data-source-locator="page:7" data-page-no="7" '
        'data-page-range="第 7 页" data-slide-no="3" data-sheet-name="数据" '
        'data-original-alt="增长趋势原图">'
        '<img src="javascript:alert(1)" onerror="alert(2)">'
        '</section>'
        '<p>图片后文</p>'
    )
    return BoardDocument(
        title="板书图片导出",
        content_text="图片前文\n增长趋势\n图片后文",
        content_html=content_html,
        content_json=html_to_tiptap_doc(content_html),
    )


def test_board_asset_store_deduplicates_per_owner_and_tracks_references(
    asset_store: BoardAssetStore,
) -> None:
    first = _put_asset(asset_store)
    second = asset_store.put_bytes(
        owner_user_id=OWNER.id,
        lesson_id="lesson_second",
        content=PNG_BYTES,
        mime_type="application/octet-stream",
        source_visual_id="visual_second",
    )
    other_owner = _put_asset(asset_store, owner_user_id=OTHER_USER.id)

    assert second.id == first.id
    assert other_owner.id != first.id
    assert asset_store.read_bytes(first.id, OTHER_USER.id) is None
    assert {ref.lesson_id for ref in asset_store.references_for_lesson(
        owner_user_id=OWNER.id,
        lesson_id="lesson_second",
    )} == {"lesson_second"}
    assert asset_store.resolve_path(first).read_bytes() == PNG_BYTES


def test_board_asset_store_rejects_invalid_mime_size_and_pixel_budget(
    asset_store: BoardAssetStore,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    with pytest.raises(BoardAssetError, match="do not match"):
        asset_store.put_bytes(
            owner_user_id=OWNER.id,
            lesson_id="lesson_invalid_mime",
            content=PNG_BYTES,
            mime_type="image/jpeg",
        )

    monkeypatch.setattr(board_asset_store_module, "MAX_BOARD_ASSET_BYTES", len(PNG_BYTES) - 1)
    with pytest.raises(BoardAssetError, match="25 MiB"):
        _put_asset(asset_store)

    monkeypatch.setattr(board_asset_store_module, "MAX_BOARD_ASSET_BYTES", 25 * 1024 * 1024)
    monkeypatch.setattr(board_asset_store_module, "MAX_BOARD_ASSET_PIXELS", 0)
    with pytest.raises(BoardAssetError, match="pixel budget"):
        _put_asset(asset_store)


def test_board_asset_store_limits_animated_decoded_pixels(
    asset_store: BoardAssetStore,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    frames = [Image.new("RGB", (10, 10), color=(index * 40, 0, 0)) for index in range(3)]
    output = io.BytesIO()
    frames[0].save(output, format="GIF", save_all=True, append_images=frames[1:], duration=50)
    monkeypatch.setattr(board_asset_store_module, "MAX_BOARD_ASSET_PIXELS", 250)

    with pytest.raises(BoardAssetError, match="pixel budget"):
        asset_store.put_bytes(
            owner_user_id=OWNER.id,
            lesson_id="lesson_animated",
            content=output.getvalue(),
            mime_type="image/gif",
        )


def test_board_asset_store_rejects_tampered_storage_key(asset_store: BoardAssetStore) -> None:
    record = _put_asset(asset_store)
    with asset_store._connect() as conn:
        with conn:
            conn.execute(
                "UPDATE board_assets SET storage_key = ? WHERE id = ?",
                ("../../outside.png", record.id),
            )

    assert asset_store.read_bytes(record.id, OWNER.id) is None


def test_board_asset_content_route_is_owner_scoped_and_supports_etag(
    asset_store: BoardAssetStore,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    import app.main as main_module
    from app.routers import auth as auth_router
    from app.routers import documents as documents_router

    record = _put_asset(asset_store)
    monkeypatch.setattr(documents_router, "get_board_asset_store", lambda: asset_store)
    main_module.app.dependency_overrides[auth_router.current_user] = lambda: OWNER
    try:
        with TestClient(main_module.app) as client:
            response = client.get(f"/api/board-assets/{record.id}/content")
            assert response.status_code == 200
            assert response.content == PNG_BYTES
            assert response.headers["content-type"] == "image/png"
            assert response.headers["cache-control"].startswith("private")
            assert response.headers["content-disposition"].startswith("inline;")
            assert response.headers["x-content-type-options"] == "nosniff"
            assert response.headers["etag"] == f'"{record.content_hash}"'

            cached = client.get(
                f"/api/board-assets/{record.id}/content",
                headers={"If-None-Match": response.headers["etag"]},
            )
            assert cached.status_code == 304
            assert cached.content == b""

            main_module.app.dependency_overrides[auth_router.current_user] = lambda: OTHER_USER
            denied = client.get(f"/api/board-assets/{record.id}/content")
            assert denied.status_code == 404
    finally:
        main_module.app.dependency_overrides.clear()

    assert documents_router._unique_export_path("docx") != documents_router._unique_export_path("docx")


def test_resource_visual_fields_survive_html_round_trip(asset_store: BoardAssetStore) -> None:
    record = _put_asset(asset_store)
    initial_json = _visual_document(record.id).content_json
    regenerated_html = tiptap_doc_to_html(initial_json)
    regenerated_json = html_to_tiptap_doc(regenerated_html)
    attrs = regenerated_json["content"][1]["attrs"]

    assert attrs == {
        "marker": "",
        "assetId": record.id,
        "visualId": "visual_board_asset",
        "sourceIngestionId": "source_board_asset",
        "sourceChapterId": "chapter_board_asset",
        "sourceTitle": "测试资料",
        "kind": "chart",
        "caption": "增长趋势",
        "source": "测试资料",
        "sourceLocator": "page:7",
        "pageNo": 7,
        "pageRange": "第 7 页",
        "slideNo": 3,
        "sheetName": "数据",
        "recreationKind": "original",
        "recreationStatus": "original_only",
        "recreationConfidence": "0.00",
        "recreationNote": "",
        "recreationHtml": "",
        "originalSrc": "",
        "originalAlt": "增长趋势原图",
        "originalInitiallyCollapsed": False,
    }


def test_rebuild_from_content_json_preserves_markdown_structure(asset_store: BoardAssetStore) -> None:
    record = _put_asset(asset_store)
    base = _visual_document(record.id)
    content_json = {
        "type": "doc",
        "content": [
            {"type": "heading", "attrs": {"level": 1}, "content": [{"type": "text", "text": "标题"}]},
            {
                "type": "bulletList",
                "content": [
                    {
                        "type": "listItem",
                        "content": [
                            {"type": "paragraph", "content": [{"type": "text", "text": "列表项"}]}
                        ],
                    }
                ],
            },
            base.content_json["content"][1],
        ],
    }

    rebuilt = rebuild_document_from_content_json(base, content_json)

    assert rebuilt.content_text.startswith("# 标题\n\n- 列表项")
    assert "资料图表：增长趋势" in rebuilt.content_text
    assert record.id not in rebuilt.content_text


def test_docx_and_html_exports_embed_authenticated_board_asset(
    asset_store: BoardAssetStore,
    monkeypatch: pytest.MonkeyPatch,
    tmp_path,
) -> None:
    record = _put_asset(asset_store)
    document = _visual_document(record.id)
    monkeypatch.setattr(docx_exporter, "get_board_asset_store", lambda: asset_store)

    docx_path = tmp_path / "board.docx"
    export_docx(document, docx_path, owner_user_id=OWNER.id)
    with ZipFile(docx_path) as archive:
        media = [name for name in archive.namelist() if name.startswith("word/media/")]
        assert len(media) == 1
        assert archive.read(media[0]) == PNG_BYTES

    exported_html = standalone_html(
        document,
        owner_user_id=OWNER.id,
        asset_store=asset_store,
    )
    encoded = base64.b64encode(PNG_BYTES).decode("ascii")
    assert f'data:image/png;base64,{encoded}' in exported_html
    assert 'alt="增长趋势原图"' in exported_html
    assert "javascript:alert" not in exported_html
    assert "onerror=" not in exported_html

    missing_html = standalone_html(
        document,
        owner_user_id=OTHER_USER.id,
        asset_store=asset_store,
    )
    assert "图片内容不可用（板书资产缺失）" in missing_html


def test_html_export_bounds_repeated_asset_expansion_and_resolves_once(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    document = _visual_document("basset_repeated")
    visual_node = document.content_json["content"][1]
    document.content_json["content"] = [visual_node, visual_node]
    resolver_calls = 0

    def resolver(_asset_id: str) -> tuple[str, bytes]:
        nonlocal resolver_calls
        resolver_calls += 1
        return "image/png", PNG_BYTES

    one_embedded_size = len("data:image/png;base64,") + 4 * ((len(PNG_BYTES) + 2) // 3)
    monkeypatch.setattr(
        html_document_export,
        "_MAX_HTML_EXPORT_EMBEDDED_IMAGE_BYTES",
        one_embedded_size,
    )

    with pytest.raises(HtmlExportBudgetError, match="embedded-image limit"):
        standalone_html(document, asset_resolver=resolver)

    assert resolver_calls == 1


def test_docx_converts_webp_and_constrains_picture_width(
    asset_store: BoardAssetStore,
    monkeypatch: pytest.MonkeyPatch,
    tmp_path,
) -> None:
    output = io.BytesIO()
    Image.new("RGB", (1000, 100), color=(25, 100, 180)).save(output, format="WEBP")
    record = asset_store.put_bytes(
        owner_user_id=OWNER.id,
        lesson_id="lesson_webp",
        content=output.getvalue(),
        mime_type="image/webp",
    )
    monkeypatch.setattr(docx_exporter, "get_board_asset_store", lambda: asset_store)
    export_path = tmp_path / "webp.docx"

    export_docx(_visual_document(record.id), export_path, owner_user_id=OWNER.id)

    exported = DocxDocument(export_path)
    assert len(exported.inline_shapes) == 1
    shape = exported.inline_shapes[0]
    section = exported.sections[0]
    assert shape.width <= section.page_width - section.left_margin - section.right_margin
    with ZipFile(export_path) as archive:
        media = [name for name in archive.namelist() if name.startswith("word/media/")]
        assert len(media) == 1
        assert archive.read(media[0]).startswith(b"\x89PNG\r\n\x1a\n")


def test_html_export_uses_tag_attribute_and_url_allowlists() -> None:
    document = BoardDocument(
        title="不可信 HTML",
        content_text="安全正文",
        content_html=(
            '<form action="javascript:alert(1)"><a href=javascript:alert(2)>危险链接</a></form>'
            '<svg><a xlink:href="javascript:alert(3)">SVG 链接</a></svg>'
            '<p style="color:red;background-image:url(javascript:alert(4))" onclick="alert(5)">安全正文</p>'
        ),
    )

    exported = standalone_html(document)

    assert "javascript:" not in exported
    assert "<form" not in exported
    assert "<svg" not in exported
    assert "xlink:" not in exported
    assert "onclick" not in exported
    assert "background-image" not in exported
    assert 'style="color:red"' in exported

    trusted_json_document = document.model_copy(
        update={
            "content_json": {
                "type": "doc",
                "content": [
                    {"type": "paragraph", "content": [{"type": "text", "text": "可信 JSON 正文"}]}
                ],
            },
            "content_html": "<p>不应采用的 HTML 正文</p>",
        }
    )
    trusted_export = standalone_html(trusted_json_document)
    assert "可信 JSON 正文" in trusted_export
    assert "不应采用的 HTML 正文" not in trusted_export
