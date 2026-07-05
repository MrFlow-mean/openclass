from __future__ import annotations

import html
import hashlib
import mimetypes
import posixpath
import re
import zipfile
from xml.etree import ElementTree as ET
from pathlib import Path

from pypdf import PdfReader
from docx import Document as DocxDocument

from app.models import (
    LibraryChapter,
    ResourceContextChunk,
    ResourceLibraryItem,
    ResourcePageStructure,
    ResourceReferenceContext,
    ResourceSourceUnit,
)
from app.services.epub_toc import extract_epub_toc_entries
from app.services.image_ocr import extract_image_text, extract_pdf_pages_text
from app.services.rag_anything_adapter import parse_with_rag_anything
from app.services.resource_page_structure import (
    build_page_structure_from_source_units,
    build_pdf_page_structure,
    enrich_source_units_with_page_structure,
    physical_page_candidates_for_printed_page,
)
from app.services.source_ingestion import mark_local_file_ready
from app.services.resource_visual_evidence import select_resource_visual_evidence


_PDF_TEXT_SUMMARY_LIMIT = 140
_PDF_LOCATOR_SEPARATOR = " || "
_PARSER_ARTIFACT_TITLE_PATTERN = re.compile(
    r"^(?:text|image|img|table|equation|formula|figure|page)\d{3,}$",
    re.IGNORECASE,
)


def _is_parser_artifact_title(title: str) -> bool:
    compact = re.sub(r"[\s_-]+", "", title).strip("：:").lower()
    return bool(_PARSER_ARTIFACT_TITLE_PATTERN.fullmatch(compact))


def _clean_outline_title(title: str) -> str:
    return re.sub(r"\s+", " ", title).strip()


def _heading_parts_from_value(value: object) -> list[str]:
    if value is None:
        return []
    if isinstance(value, str):
        raw_parts = re.split(r"\s*(?:>|/)\s*", value) if (">" in value or "/" in value) else [value]
        return [_clean_outline_title(part) for part in raw_parts if _clean_outline_title(part)]
    if isinstance(value, list):
        parts: list[str] = []
        for item in value:
            parts.extend(_heading_parts_from_value(item))
        return parts
    return []


def _readable_heading_path(unit: ResourceSourceUnit) -> list[str]:
    raw_parts = list(unit.heading_path)
    if not raw_parts:
        for key in ("heading_path", "headings", "section_path", "title_path", "breadcrumbs", "breadcrumb"):
            raw_parts = _heading_parts_from_value(unit.metadata.get(key))
            if raw_parts:
                break
    cleaned: list[str] = []
    for part in raw_parts:
        title = _clean_outline_title(str(part))
        if title and not _is_parser_artifact_title(title):
            cleaned.append(title[:120])
    return cleaned[:6]


def _normalize_extracted_text(text: str) -> str:
    cleaned = text.replace("\x00", "").replace("\r\n", "\n")
    cleaned = re.sub(r"(?<=[\u4e00-\u9fff])\s+(?=[\u4e00-\u9fff])", "", cleaned)
    cleaned = re.sub(r"(?<=[0-9])\s+(?=[0-9])", "", cleaned)
    cleaned = re.sub(r"(?<=[\u4e00-\u9fff])\s+(?=[0-9A-Za-z])", "", cleaned)
    cleaned = re.sub(r"(?<=[0-9A-Za-z])\s+(?=[\u4e00-\u9fff])", "", cleaned)
    cleaned = re.sub(r"\s+([,，.。!?！？；;：:])", r"\1", cleaned)
    cleaned = re.sub(r"([，。！？；：])\s+", r"\1", cleaned)
    cleaned = re.sub(r"([,.!?;:])\s+(?=[\u4e00-\u9fff])", r"\1", cleaned)
    cleaned = re.sub(r"[ \t]{2,}", " ", cleaned)
    cleaned = re.sub(r"\n{3,}", "\n\n", cleaned)
    return cleaned.strip()


def _summary_snippet(text: str, *, limit: int = _PDF_TEXT_SUMMARY_LIMIT) -> str:
    compact = re.sub(r"\s+", " ", _normalize_extracted_text(text)).strip()
    return compact[:limit].strip(" ，,。") if compact else ""


def _read_pdf_text_window(
    reader: PdfReader,
    *,
    page_start: int,
    page_end: int,
    max_pages: int = 6,
    max_nonempty_pages: int | None = None,
) -> str:
    start_index = max(page_start - 1, 0)
    end_index = min(max(page_end, page_start), len(reader.pages))
    extracted: list[str] = []
    nonempty_pages = 0
    scanned_pages = 0
    for page_index in range(start_index, end_index):
        if scanned_pages >= max_pages:
            break
        scanned_pages += 1
        try:
            text = reader.pages[page_index].extract_text() or ""
        except Exception:
            continue
        text = _normalize_extracted_text(text)
        if not text:
            continue
        extracted.append(text)
        nonempty_pages += 1
        if max_nonempty_pages is not None and nonempty_pages >= max_nonempty_pages:
            break
    return "\n".join(extracted).strip()


def _pdf_locator_hint(
    title: str,
    *,
    source: str,
    toc_page: int | None = None,
    printed_page: int | None = None,
    actual_page: int | None = None,
) -> str:
    parts = [title, f"source={source}"]
    if toc_page is not None:
        parts.append(f"toc_page={toc_page}")
    if printed_page is not None:
        parts.append(f"printed_page={printed_page}")
    if actual_page is not None:
        parts.append(f"actual_page={actual_page}")
    return _PDF_LOCATOR_SEPARATOR.join(parts)


def _pdf_locator_value(locator_hint: str | None, key: str) -> int | None:
    if not locator_hint:
        return None
    for part in locator_hint.split(_PDF_LOCATOR_SEPARATOR):
        name, sep, value = part.partition("=")
        if sep and name.strip() == key:
            try:
                return int(value.strip())
            except ValueError:
                return None
    return None


def _pdf_locator_source(locator_hint: str | None) -> str | None:
    if not locator_hint:
        return None
    for part in locator_hint.split(_PDF_LOCATOR_SEPARATOR):
        name, sep, value = part.partition("=")
        if sep and name.strip() == "source":
            return value.strip()
    return None


def _chapter(
    title: str,
    summary: str,
    keywords: list[str],
    level: int = 1,
    *,
    locator_hint: str | None = None,
    order_index: int = 0,
    scan_strategy: str = "outline_only",
    page_start: int | None = None,
    page_end: int | None = None,
) -> LibraryChapter:
    page_range = None
    if page_start and page_end and page_end >= page_start:
        page_range = f"{page_start}-{page_end}" if page_end > page_start else str(page_start)
    elif page_start:
        page_range = str(page_start)

    return LibraryChapter(
        title=title,
        summary=summary,
        keywords=keywords,
        level=level,
        locator_hint=locator_hint or title,
        order_index=order_index,
        scan_strategy=scan_strategy,  # type: ignore[arg-type]
        page_start=page_start,
        page_end=page_end,
        page_range=page_range,
    )


def _attach_outline_hierarchy(chapters: list[LibraryChapter]) -> list[LibraryChapter]:
    stack: list[LibraryChapter] = []
    enriched: list[LibraryChapter] = []
    for chapter in chapters:
        while stack and stack[-1].level >= chapter.level:
            stack.pop()
        parent = stack[-1] if stack else None
        path = [*(parent.path if parent else []), chapter.title]
        enriched_chapter = chapter.model_copy(
            update={
                "parent_id": parent.id if parent else None,
                "parent_title": parent.title if parent else None,
                "path": path,
            }
        )
        enriched.append(enriched_chapter)
        stack.append(enriched_chapter)
    return enriched


def _markdown_sections(text: str) -> list[dict[str, object]]:
    lines = text.splitlines()
    headings: list[tuple[int, int, str]] = []
    for index, raw_line in enumerate(lines):
        line = raw_line.strip()
        if not line or not line.startswith("#"):
            continue
        level = len(line) - len(line.lstrip("#"))
        title = _clean_outline_title(line[level:])
        if title and not _is_parser_artifact_title(title):
            headings.append((index, level, title))

    sections: list[dict[str, object]] = []
    for index, (line_number, level, title) in enumerate(headings):
        end = len(lines)
        for next_line_number, next_level, _ in headings[index + 1 :]:
            if next_level <= level:
                end = next_line_number
                break
        content = "\n".join(lines[line_number + 1 : end]).strip()
        sections.append(
            {
                "title": title,
                "level": level,
                "content": content,
                "order_index": index,
            }
        )
    return sections


def _extract_markdown_outline(text: str) -> list[LibraryChapter]:
    chapters: list[LibraryChapter] = []
    for section in _markdown_sections(text):
        title = str(section["title"])
        if _is_parser_artifact_title(title):
            continue
        content = str(section["content"])
        summary_seed = re.sub(r"\s+", " ", content).strip()[:90]
        summary = summary_seed or f"来自资料标题“{title}”的章节摘要待进一步展开。"
        chapters.append(
            _chapter(
                title=title,
                summary=summary,
                keywords=_keywords_from_text(f"{title}\n{content}") or [token.lower() for token in re.findall(r"[A-Za-z\u4e00-\u9fff]+", title)[:5]],
                level=int(section["level"]),
                locator_hint=title,
                order_index=int(section["order_index"]),
                scan_strategy="heading_section",
            )
        )
    return chapters


def _keywords_from_text(text: str) -> list[str]:
    stopwords = {
        "the",
        "and",
        "for",
        "with",
        "that",
        "this",
        "from",
        "of",
        "to",
        "in",
        "on",
        "is",
        "are",
        "as",
        "by",
        "be",
        "or",
        "an",
        "at",
        "into",
        "about",
        "lesson",
        "chapter",
        "section",
        "一个",
        "一些",
        "我们",
        "你们",
        "什么",
        "以及",
        "当前",
        "这个",
        "那个",
        "可以",
        "通过",
    }
    counts: dict[str, int] = {}
    for token in re.findall(r"[A-Za-z\u4e00-\u9fff]{2,}", text.lower()):
        if token in stopwords or token.isdigit():
            continue
        counts[token] = counts.get(token, 0) + 1
    return [token for token, _ in sorted(counts.items(), key=lambda item: item[1], reverse=True)[:8]]


def _generic_chapter_from_text(title: str, text: str, *, summary_prefix: str) -> LibraryChapter:
    normalized_text = _normalize_extracted_text(text)
    snippet = re.sub(r"\s+", " ", normalized_text[:4000]).strip()[:120] or f"围绕“{title}”补充资料入口。"
    return _chapter(
        title=title,
        summary=f"{summary_prefix}{snippet}",
        keywords=_keywords_from_text(normalized_text) or [token.lower() for token in re.findall(r"[A-Za-z\u4e00-\u9fff]+", title)[:5]],
        locator_hint=title,
        order_index=0,
        scan_strategy="fulltext_match",
    )


def _ai_generated_outline(original_name: str, text: str) -> list[LibraryChapter]:
    normalized_text = _normalize_extracted_text(text)
    if len(normalized_text) < 80:
        return []
    try:
        from app.services.openai_course_ai import openai_course_ai

        generated = openai_course_ai.generate_resource_outline(
            resource_name=original_name,
            extracted_text=normalized_text,
        )
    except Exception:
        return []
    if generated is None:
        return []

    chapters: list[LibraryChapter] = []
    seen_titles: set[str] = set()
    for index, item in enumerate(generated.chapters):
        title = _clean_outline_title(item.title)
        summary = re.sub(r"\s+", " ", item.summary).strip()
        if not title or _is_parser_artifact_title(title) or title.lower() in seen_titles:
            continue
        seen_titles.add(title.lower())
        chapters.append(
            _chapter(
                title=title[:80],
                summary=summary[:220] or f"从资料“{original_name}”生成的目录入口。",
                keywords=[
                    keyword[:40]
                    for keyword in (item.keywords or _keywords_from_text(f"{title}\n{summary}"))[:8]
                    if keyword.strip()
                ],
                level=max(1, min(item.level, 4)),
                locator_hint=title,
                order_index=index,
                scan_strategy="fulltext_match",
            )
        )
    return chapters


def _read_text_file(file_path: Path) -> str:
    return file_path.read_text(encoding="utf-8", errors="ignore")


def _docx_items(file_path: Path) -> list[dict[str, object]]:
    source = DocxDocument(file_path)
    items: list[dict[str, object]] = []

    for paragraph in source.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        style_name = (paragraph.style.name if paragraph.style else "").lower()
        level = 0
        if "heading 1" in style_name or "title" in style_name:
            level = 1
        elif "heading 2" in style_name:
            level = 2
        elif "heading 3" in style_name:
            level = 3
        items.append({"text": text, "level": level})

    for table in source.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                items.append({"text": " | ".join(cells), "level": 0})

    return items


def _read_docx_text(file_path: Path) -> str:
    return "\n".join(str(item["text"]) for item in _docx_items(file_path))


def _docx_sections(file_path: Path) -> list[dict[str, object]]:
    items = _docx_items(file_path)
    headings = [(index, int(item["level"]), str(item["text"])) for index, item in enumerate(items) if int(item["level"]) > 0]
    sections: list[dict[str, object]] = []

    for index, (item_index, level, title) in enumerate(headings):
        end = len(items)
        for next_item_index, next_level, _ in headings[index + 1 :]:
            if next_level <= level:
                end = next_item_index
                break
        content = "\n".join(str(item["text"]) for item in items[item_index + 1 : end]).strip()
        sections.append(
            {
                "title": title,
                "level": level,
                "content": content,
                "order_index": index,
            }
        )
    return sections


def _extract_docx_outline(file_path: Path) -> list[LibraryChapter]:
    chapters: list[LibraryChapter] = []
    for section in _docx_sections(file_path):
        title = str(section["title"])
        content = str(section["content"])
        summary_seed = re.sub(r"\s+", " ", content).strip()[:90]
        summary = summary_seed or f"来自资料标题“{title}”的章节摘要待进一步展开。"
        chapters.append(
            _chapter(
                title=title,
                summary=summary,
                keywords=_keywords_from_text(f"{title}\n{content}") or [token.lower() for token in re.findall(r"[A-Za-z\u4e00-\u9fff]+", title)[:5]],
                level=int(section["level"]),
                locator_hint=title,
                order_index=int(section["order_index"]),
                scan_strategy="heading_section",
            )
        )
    return chapters


def _extract_docx_section_text(file_path: Path, chapter: LibraryChapter) -> str:
    sections = _docx_sections(file_path)
    target = next(
        (
            section
            for section in sections
            if str(section["title"]) == (chapter.locator_hint or chapter.title)
        ),
        None,
    )
    if target is None:
        return _read_docx_text(file_path)
    content = str(target["content"]).strip()
    if content:
        return content
    return str(target["title"])


_EPUB_CHINESE_DIGITS = {
    "零": 0,
    "〇": 0,
    "一": 1,
    "二": 2,
    "两": 2,
    "三": 3,
    "四": 4,
    "五": 5,
    "六": 6,
    "七": 7,
    "八": 8,
    "九": 9,
}
_EPUB_CHINESE_UNITS = {"十": 10, "百": 100}
_EPUB_NUMBER_PATTERN = r"(?:\d+|[一二三四五六七八九十百〇零两]+)"


def _parse_epub_outline_number(value: str | None) -> int | None:
    cleaned = (value or "").strip()
    if not cleaned:
        return None
    if cleaned.isdigit():
        return int(cleaned)

    total = 0
    current = 0
    seen = False
    for char in cleaned:
        if char in _EPUB_CHINESE_DIGITS:
            current = _EPUB_CHINESE_DIGITS[char]
            seen = True
            continue
        unit = _EPUB_CHINESE_UNITS.get(char)
        if unit is None:
            return None
        total += (current or 1) * unit
        current = 0
        seen = True
    if not seen:
        return None
    return total + current


def _extract_epub_requested_outline_reference(text: str) -> tuple[int | None, int | None]:
    match = re.search(
        rf"第\s*({_EPUB_NUMBER_PATTERN})\s*章(?:\s*第\s*({_EPUB_NUMBER_PATTERN})\s*[节讲部分])?",
        text,
    )
    if match:
        chapter_no = _parse_epub_outline_number(match.group(1))
        section_no = _parse_epub_outline_number(match.group(2)) if match.group(2) else None
        return chapter_no, section_no
    english = re.search(r"\bchapter\s*(\d+)\s*(?:section\s*(\d+))?\b", text, flags=re.IGNORECASE)
    if english:
        return int(english.group(1)), int(english.group(2)) if english.group(2) else None
    dotted = re.search(r"\b(\d+)\.(\d+)\b", text)
    if dotted:
        return int(dotted.group(1)), int(dotted.group(2))
    return None, None


def _epub_title_outline_reference(title: str) -> tuple[int | None, int | None]:
    cleaned = title.strip()
    chapter = re.search(rf"第\s*({_EPUB_NUMBER_PATTERN})\s*章", cleaned)
    if chapter:
        return _parse_epub_outline_number(chapter.group(1)), None
    dotted = re.search(r"^\s*(\d+)\s*[.．]\s*(\d+)", cleaned)
    if dotted:
        return int(dotted.group(1)), int(dotted.group(2))
    english = re.search(r"\bchapter\s*(\d+)\b", cleaned, flags=re.IGNORECASE)
    if english:
        return int(english.group(1)), None
    return None, None


def _epub_is_html_path(path: str) -> bool:
    lowered = path.lower()
    return lowered.endswith((".xhtml", ".html", ".htm"))


def _decode_epub_bytes(data: bytes) -> str:
    for encoding in ("utf-8", "utf-16", "gb18030", "latin-1"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="ignore")


def _epub_text_from_html(raw_html: str) -> str:
    cleaned = re.sub(r"(?is)<(script|style|head|nav)\b.*?</\1>", "\n", raw_html)
    cleaned = re.sub(r"(?is)<br\b[^>]*>", "\n", cleaned)
    cleaned = re.sub(r"(?is)</?(h[1-6]|p|div|section|article|li|tr|td|th|blockquote)\b[^>]*>", "\n", cleaned)
    cleaned = re.sub(r"(?is)<[^>]+>", "", cleaned)
    text = html.unescape(cleaned)
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    return _normalize_extracted_text("\n".join(lines))


def _epub_fragment_text(fragment: str) -> str:
    return _epub_text_from_html(fragment).strip()


def _epub_html_title(raw_html: str, path: str) -> str:
    for match in re.finditer(r"(?is)<h([1-6])\b[^>]*>(.*?)</h\1>", raw_html):
        title = _epub_fragment_text(match.group(2))
        if title:
            return title[:120]
    title_match = re.search(r"(?is)<title\b[^>]*>(.*?)</title>", raw_html)
    if title_match:
        title = _epub_fragment_text(title_match.group(1))
        if title:
            return title[:120]
    return Path(path).stem.replace("_", " ").replace("-", " ").strip() or path


def _epub_rootfile_path(archive: zipfile.ZipFile) -> str | None:
    try:
        container_xml = archive.read("META-INF/container.xml")
    except KeyError:
        return None
    try:
        root = ET.fromstring(container_xml)
    except ET.ParseError:
        return None
    rootfile = root.find(".//{*}rootfile")
    if rootfile is None:
        return None
    full_path = rootfile.attrib.get("full-path")
    return full_path.strip() if full_path else None


def _epub_reading_order_paths(archive: zipfile.ZipFile) -> list[str]:
    rootfile_path = _epub_rootfile_path(archive)
    if not rootfile_path:
        return []
    try:
        opf_root = ET.fromstring(archive.read(rootfile_path))
    except (KeyError, ET.ParseError):
        return []

    base_dir = posixpath.dirname(rootfile_path)
    manifest: dict[str, tuple[str, str, str]] = {}
    for item in opf_root.findall(".//{*}manifest/{*}item"):
        item_id = item.attrib.get("id", "").strip()
        href = item.attrib.get("href", "").strip()
        media_type = item.attrib.get("media-type", "").strip()
        properties = item.attrib.get("properties", "").strip()
        if not item_id or not href:
            continue
        path = posixpath.normpath(posixpath.join(base_dir, href))
        manifest[item_id] = (path, media_type, properties)

    ordered: list[str] = []
    for itemref in opf_root.findall(".//{*}spine/{*}itemref"):
        item_id = itemref.attrib.get("idref", "").strip()
        path, media_type, properties = manifest.get(item_id, ("", "", ""))
        if not path:
            continue
        if "nav" in properties or "toc" in properties:
            continue
        if media_type in {"application/xhtml+xml", "text/html"} or _epub_is_html_path(path):
            ordered.append(path)
    return [path for path in ordered if path in archive.namelist()]


def _epub_html_items(file_path: Path) -> list[dict[str, object]]:
    try:
        with zipfile.ZipFile(file_path) as archive:
            paths = _epub_reading_order_paths(archive)
            if not paths:
                paths = [
                    path
                    for path in archive.namelist()
                    if _epub_is_html_path(path) and not re.search(r"(?:^|/)(?:nav|toc|cover)\.", path, flags=re.IGNORECASE)
                ]
            items: list[dict[str, object]] = []
            for order_index, path in enumerate(paths):
                try:
                    raw_html = _decode_epub_bytes(archive.read(path))
                except KeyError:
                    continue
                text = _epub_text_from_html(raw_html)
                if not text:
                    continue
                items.append(
                    {
                        "path": path,
                        "title": _epub_html_title(raw_html, path),
                        "text": text,
                        "order_index": order_index,
                    }
                )
            return items
    except (zipfile.BadZipFile, OSError):
        return []


def _is_epub_separator_title(title: str) -> bool:
    compact = re.sub(r"\s+", "", title).lower()
    return compact in {"封面", "版权", "目录", "目次", "前言", "序", "绪言", "contents", "cover", "titlepage"}


def _looks_like_epub_heading(line: str) -> bool:
    cleaned = line.strip()
    if len(cleaned) > 90:
        return False
    if _looks_like_reference_heading(cleaned):
        return True
    if re.match(r"^(?:chapter\s+\d+|\d+\s*[.．]\s*\d+|\d+\s+[A-Za-z\u4e00-\u9fff])", cleaned, flags=re.IGNORECASE):
        return True
    return False


def _epub_sections(file_path: Path) -> list[dict[str, object]]:
    items = _epub_html_items(file_path)
    sections: list[dict[str, object]] = []
    for item in items:
        text = str(item["text"])
        lines = [line.strip() for line in text.splitlines() if line.strip()]
        if not lines:
            continue
        heading_indexes = [index for index, line in enumerate(lines) if _looks_like_epub_heading(line)]
        if not heading_indexes:
            title = str(item["title"]).strip() or Path(str(item["path"])).stem
            if _is_epub_separator_title(title) and len(lines) > 1:
                title = next((line for line in lines if len(line) <= 80 and not _is_epub_separator_title(line)), title)
            content = "\n".join(lines).strip()
            if content:
                sections.append(
                    {
                        "title": title[:120],
                        "content": content,
                        "level": 1,
                        "order_index": int(item["order_index"]),
                    }
                )
            continue

        for local_index, line_index in enumerate(heading_indexes):
            end = heading_indexes[local_index + 1] if local_index + 1 < len(heading_indexes) else len(lines)
            title = lines[line_index].strip()
            content = "\n".join(lines[line_index + 1 : end]).strip()
            if _is_epub_separator_title(title) and not content:
                continue
            level = _toc_entry_level(title)
            sections.append(
                {
                    "title": title[:120],
                    "content": content or title,
                    "level": level,
                    "order_index": len(sections),
                }
            )
    return sections


def _read_epub_text(file_path: Path) -> str:
    return _normalize_extracted_text("\n\n".join(str(item["text"]) for item in _epub_html_items(file_path)))


def _extract_epub_outline(file_path: Path) -> list[LibraryChapter]:
    toc_entries = extract_epub_toc_entries(file_path)
    if toc_entries:
        chapters: list[LibraryChapter] = []
        for index, entry in enumerate(toc_entries):
            title = entry.title
            page_start = entry.page_start
            chapter = _chapter(
                title=title,
                summary=f"来自 EPUB 目录的章节入口：{title}。",
                keywords=_keywords_from_text(title),
                level=entry.level,
                locator_hint=title,
                order_index=index,
                scan_strategy="outline_only",
                page_start=page_start,
                page_end=page_start,
            )
            stable_seed = f"{file_path.name}\n{index}\n{title}\n{entry.source}"
            stable_id = f"chapter_epub_{hashlib.sha256(stable_seed.encode('utf-8')).hexdigest()[:16]}"
            chapters.append(chapter.model_copy(update={"id": stable_id}))
        return chapters

    chapters: list[LibraryChapter] = []
    for section in _epub_sections(file_path):
        title = str(section["title"]).strip()
        if not title or _is_epub_separator_title(title):
            continue
        content = str(section["content"])
        summary_seed = re.sub(r"\s+", " ", content).strip()[:120]
        summary = summary_seed or f"来自 EPUB 标题“{title}”的章节摘要待进一步展开。"
        chapters.append(
            _chapter(
                title=title,
                summary=summary,
                keywords=_keywords_from_text(f"{title}\n{content}"),
                level=int(section["level"]),
                locator_hint=title,
                order_index=int(section["order_index"]),
                scan_strategy="heading_section",
            )
        )
    return chapters


def _epub_section_with_children(sections: list[dict[str, object]], start_index: int) -> str:
    target = sections[start_index]
    target_level = int(target["level"])

    def section_text(section: dict[str, object]) -> str:
        title = str(section["title"]).strip()
        content = str(section["content"]).strip()
        if content.startswith(title):
            return content
        return f"{title}\n{content}".strip()

    parts = [section_text(target)]
    for section in sections[start_index + 1 :]:
        if int(section["level"]) <= target_level:
            break
        parts.append(section_text(section))
    return "\n\n".join(part for part in parts if part.strip())


def _generic_outline_marker_count(text: str) -> int:
    count = 0
    for line in text.splitlines():
        compact = re.sub(r"\s+", "", line)
        if not 2 <= len(compact) <= 36:
            continue
        if re.match(r"^[【\[\(（].{1,34}[】\]\)）]$", compact):
            count += 1
            continue
        if re.match(r"^(?:第?[一二三四五六七八九十百千万\d]+[章节部分讲课、.．)]|[A-Za-z][.)])", compact):
            count += 1
    return count


def _continuous_explanatory_sentence_count(text: str) -> int:
    segments = re.split(r"[。！？!?；;]\s*", text)
    return sum(1 for segment in segments if len(re.sub(r"\s+", "", segment)) >= 18)


def _body_text_density(text: str) -> float:
    compact = re.sub(r"\s+", "", text)
    if not compact:
        return 0.0
    body_chars = 0
    for line in text.splitlines():
        stripped = line.strip()
        if not stripped:
            continue
        line_compact = re.sub(r"\s+", "", stripped)
        if len(line_compact) <= 36 and _generic_outline_marker_count(stripped):
            continue
        body_chars += len(line_compact)
    return body_chars / len(compact)


def _epub_section_body_score(sections: list[dict[str, object]], index: int) -> tuple[int, int]:
    text = _epub_section_with_children(sections, index)
    compact = re.sub(r"\s+", "", text)
    outline_markers = _generic_outline_marker_count(text)
    explanatory_sentences = _continuous_explanatory_sentence_count(text)
    body_density = _body_text_density(text)
    if len(compact) <= 80 and outline_markers >= 2 and explanatory_sentences == 0:
        return (-1, len(compact))
    score = min(len(compact), 2000)
    if outline_markers >= 2 and explanatory_sentences <= 1:
        score -= 200
    if body_density < 0.45:
        score -= 120
    return (score, len(compact))


def _extract_epub_section_text(file_path: Path, chapter: LibraryChapter, user_query: str) -> tuple[str, str]:
    sections = _epub_sections(file_path)
    if not sections:
        return chapter.title, _read_epub_text(file_path)

    requested_chapter_no, requested_section_no = _extract_epub_requested_outline_reference(user_query)
    if requested_chapter_no is not None:
        exact_candidates: list[int] = []
        fallback_candidates: list[int] = []
        for index, section in enumerate(sections):
            chapter_no, section_no = _epub_title_outline_reference(str(section["title"]))
            if chapter_no != requested_chapter_no:
                continue
            if requested_section_no is not None and section_no == requested_section_no:
                exact_candidates.append(index)
            if requested_section_no is None and section_no is None:
                exact_candidates.append(index)
            fallback_candidates.append(index)
        candidates = exact_candidates or fallback_candidates
        if candidates:
            best_index = max(candidates, key=lambda candidate: _epub_section_body_score(sections, candidate))
            section = sections[best_index]
            return str(section["title"]), _epub_section_with_children(sections, best_index)

    target_title = (chapter.locator_hint or chapter.title).strip()
    title_candidates = [
        index
        for index, section in enumerate(sections)
        if str(section["title"]).strip() == target_title
    ]
    if title_candidates:
        best_index = max(title_candidates, key=lambda candidate: _epub_section_body_score(sections, candidate))
        section = sections[best_index]
        return str(section["title"]), _epub_section_with_children(sections, best_index)

    first_index = next(
        (
            index
            for index, section in enumerate(sections)
            if int(section["level"]) == 1 and not _is_epub_separator_title(str(section["title"]))
        ),
        0,
    )
    section = sections[first_index]
    return str(section["title"]), _epub_section_with_children(sections, first_index)


def _looks_like_reference_heading(line: str) -> bool:
    cleaned = line.strip()
    return bool(
        re.match(r"^(?:第\s*[0-9一二三四五六七八九十百〇零两]+\s*章|[0-9]+\s*[.．]\s*[0-9]+)", cleaned)
        or re.match(r"^[一二三四五六七八九十]+[、.．]\s*", cleaned)
    )


def _looks_like_page_artifact(line: str) -> bool:
    cleaned = re.sub(r"\s+", "", line.strip())
    if not cleaned:
        return True
    if cleaned.isdigit() and len(cleaned) <= 3:
        return True
    if re.fullmatch(r"第[0-9一二三四五六七八九十百〇零两]+章(?:概论|绪论)?", cleaned):
        return True
    return False


def _join_reference_lines(lines: list[str]) -> str:
    text = "".join(line.strip() for line in lines if line.strip())
    return re.sub(r"\s+", " ", text).strip()


def _reference_text_passages(text: str) -> list[str]:
    lines = [line.strip() for line in _normalize_extracted_text(text).splitlines() if line.strip()]
    passages: list[str] = []
    current: list[str] = []

    def flush() -> None:
        nonlocal current
        passage = _join_reference_lines(current)
        if len(passage) >= 8:
            passages.append(passage)
        current = []

    for line in lines:
        if _looks_like_page_artifact(line):
            continue
        if _looks_like_reference_heading(line):
            flush()
            passages.append(line)
            continue
        if re.match(r"^[•·\-—]", line):
            flush()
            passages.append(line)
            continue

        current.append(line)
        joined = _join_reference_lines(current)
        if re.search(r"[。！？!?]$", line) and len(joined) >= 80:
            flush()
        elif len(joined) >= 260:
            flush()
    flush()

    fallback = [
        segment.strip()
        for segment in re.split(r"\n{2,}|(?<=[。！？.!?])\s+", text)
        if len(segment.strip()) >= 8
    ]
    candidates = passages or fallback
    unique: list[str] = []
    seen: set[str] = set()
    for passage in candidates:
        cleaned = re.sub(r"\s+", " ", passage).strip()
        if len(cleaned) < 8:
            continue
        key = cleaned[:80]
        if key in seen:
            continue
        seen.add(key)
        unique.append(cleaned)
    return unique


def _rank_passages(text: str, query: str, *, anchor: str | None = None) -> list[str]:
    paragraphs = _reference_text_passages(text)
    if not paragraphs:
        compact = re.sub(r"\s+", " ", text).strip()
        return [compact] if compact else []

    query_terms = [term for term in _keywords_from_text(f"{query}\n{anchor or ''}") if len(term) >= 2]
    scored: list[tuple[int, str]] = []
    for paragraph in paragraphs:
        score = 0
        lowered = paragraph.lower()
        for term in query_terms:
            if term.lower() in lowered:
                score += 2
        if anchor and anchor.lower() in lowered:
            score += 3
        if score:
            scored.append((score, paragraph))
    if not scored:
        return paragraphs[:4]
    return [paragraph for _, paragraph in sorted(scored, key=lambda item: item[0], reverse=True)[:4]]


def _build_teaching_hint(chapter_title: str, excerpt: str) -> str:
    focus = _keywords_from_text(excerpt)[:3]
    if focus:
        return f"讲解时先用自己的话串起 {', '.join(focus)}，再回到“{chapter_title}”的主线。"
    return f"讲解时先概括这段在“{chapter_title}”里解决了什么问题，再给一个更口语化解释。"


def _child_chapters(resource: ResourceLibraryItem, chapter: LibraryChapter) -> list[LibraryChapter]:
    children = [
        candidate
        for candidate in resource.outline
        if candidate.parent_id == chapter.id and not _looks_like_page_artifact(candidate.title)
    ]
    if children:
        return children[:8]

    descendants: list[LibraryChapter] = []
    started = False
    for candidate in sorted(resource.outline, key=lambda item: item.order_index):
        if candidate.id == chapter.id:
            started = True
            continue
        if not started:
            continue
        if candidate.level <= chapter.level:
            break
        if not _looks_like_page_artifact(candidate.title):
            descendants.append(candidate)
    return descendants[:8]


def _outline_chunk(chapter: LibraryChapter, children: list[LibraryChapter]) -> ResourceContextChunk | None:
    if not children:
        return None
    titles = [child.title.strip() for child in children if child.title.strip()]
    if not titles:
        return None
    outline = " -> ".join(titles[:6])
    return ResourceContextChunk(
        title=f"{chapter.title} / 目录主线",
        excerpt=f"这一章可以按目录顺序来讲：{outline}。",
        teaching_hint="先把目录讲成学习地图，再展开正文里的定义、例子和系统流程。",
    )


def _generic_teaching_points(
    *,
    chapter: LibraryChapter,
    children: list[LibraryChapter],
    text: str,
) -> list[str]:
    keywords = _keywords_from_text(f"{chapter.title}\n{text}")[:5]
    points = [
        f"先说明“{chapter.title}”这一节在资料结构中要解决的核心问题。",
        "把抽取到的关键术语、材料证据或推理步骤组织成一条可复述的学习主线。",
        "优先解释概念之间的关系、适用条件和容易混淆的边界，而不是照搬原文段落。",
        "讲解时配一个最小例子、对比或检查问题，用来验证学习者是否能迁移。",
    ]
    if children:
        child_titles = "、".join(child.title for child in children[:4])
        points.insert(1, f"参考子目录顺序组织讲解：{child_titles}。")
    if keywords:
        points.insert(2, f"围绕 {', '.join(keywords[:3])} 的关系展开，不把关键词拆成孤立卡片。")
    return points


def _build_reference_context(
    resource: ResourceLibraryItem,
    chapter: LibraryChapter,
    query: str,
    raw_text: str,
) -> ResourceReferenceContext | None:
    normalized_text = _normalize_extracted_text(raw_text)
    compact = re.sub(r"\s+", " ", normalized_text).strip()
    if not compact:
        return None
    if not resource.extracted_text_available and len(compact) < 320:
        return None

    children = _child_chapters(resource, chapter)
    passages = _rank_passages(normalized_text[:12000], query, anchor=chapter.title)
    chunks = [
        ResourceContextChunk(
            title=f"{chapter.title} / 参考片段 {index}",
            excerpt=passage[:420],
            teaching_hint=_build_teaching_hint(chapter.title, passage),
        )
        for index, passage in enumerate(passages[:3], start=1)
    ]
    outline = _outline_chunk(chapter, children)
    if outline is not None:
        chunks.insert(0, outline)

    teaching_points = _generic_teaching_points(chapter=chapter, children=children, text=compact)
    unique_points: list[str] = []
    seen_points: set[str] = set()
    for point in teaching_points:
        if point in seen_points:
            continue
        seen_points.add(point)
        unique_points.append(point)

    if children:
        child_titles = "、".join(child.title for child in children[:5])
        summary = f"《{resource.name}》的《{chapter.title}》包含这些讲解入口：{child_titles}。"
    else:
        summary = f"《{resource.name}》的《{chapter.title}》可以作为本次讲解参考。"
    summary = (
        f"{summary}"
        "下面的上下文会优先保留本章结构、关键定义和可用于课堂解释的片段。"
    )
    return ResourceReferenceContext(
        resource_id=resource.id,
        chapter_id=chapter.id,
        resource_name=resource.name,
        chapter_title=chapter.title,
        summary=summary,
        teaching_points=unique_points[:6],
        chunks=chunks,
        visual_evidence=select_resource_visual_evidence(resource, chapter, query=query, max_items=2),
        full_text=normalized_text,
    )


def _extract_markdown_section_text(file_path: Path, chapter: LibraryChapter) -> str:
    text = _read_text_file(file_path)
    sections = _markdown_sections(text)
    target = next(
        (
            section
            for section in sections
            if str(section["title"]) == (chapter.locator_hint or chapter.title)
        ),
        None,
    )
    if target is None:
        return text[:4000]
    content = str(target["content"]).strip()
    if content:
        return content
    return str(target["title"])


def _extract_pdf_chapter_text(
    file_path: Path,
    chapter: LibraryChapter,
    query: str,
    *,
    page_structure: ResourcePageStructure | None = None,
) -> str:
    reader = PdfReader(str(file_path))
    candidate_pages = _pdf_page_candidates(chapter, len(reader.pages), page_structure=page_structure)
    best_text = ""
    best_score = -1
    locator_source = _pdf_locator_source(chapter.locator_hint)
    trusted_locator = locator_source == "pdf_outline"
    for page_start in candidate_pages:
        page_end = min(chapter.page_end or page_start + 3, len(reader.pages))
        raw_text = _read_pdf_text_window(
            reader,
            page_start=page_start,
            page_end=page_end,
            max_pages=max(1, min(12, page_end - page_start + 1)),
        )
        if not raw_text:
            raw_text = extract_pdf_pages_text(
                file_path,
                page_start=page_start,
                page_end=page_end,
                max_pages=max(1, min(6, page_end - page_start + 1)),
            ) or ""
        raw_text = _normalize_extracted_text(raw_text)
        if not raw_text:
            continue

        score = _chapter_text_match_score(raw_text, chapter, query)
        if score > best_score:
            best_score = score
            best_text = raw_text
        if score >= 2 or trusted_locator:
            return raw_text

    if best_text and best_score > 0:
        return best_text

    searched_text = _find_pdf_text_by_keywords(reader, chapter, query)
    if searched_text:
        return searched_text

    if best_text and trusted_locator:
        return best_text

    if chapter.page_start or locator_source == "toc_page":
        return ""

    joined = _read_pdf_text_window(reader, page_start=1, page_end=min(3, len(reader.pages)), max_pages=3)
    if joined:
        return joined

    # 只有没有页码定位时，才退回到前几页相关片段，避免把前言错塞进正文章节。
    fallback: list[str] = []
    for page in reader.pages[: min(5, len(reader.pages))]:
        try:
            fallback.append(_normalize_extracted_text(page.extract_text() or ""))
        except Exception:
            continue
    fallback_text = "\n".join(fallback)
    passages = _rank_passages(fallback_text, query, anchor=chapter.title)
    return "\n\n".join(passages)


def _pdf_page_candidates(
    chapter: LibraryChapter,
    total_pages: int,
    *,
    page_structure: ResourcePageStructure | None = None,
) -> list[int]:
    printed_page = _pdf_locator_value(chapter.locator_hint, "printed_page")
    raw_candidates = [
        chapter.page_start,
        _pdf_locator_value(chapter.locator_hint, "actual_page"),
        *physical_page_candidates_for_printed_page(page_structure, printed_page),
        printed_page,
    ]
    toc_page = _pdf_locator_value(chapter.locator_hint, "toc_page")
    if toc_page and printed_page:
        raw_candidates.extend([toc_page + printed_page, toc_page + printed_page - 1])

    candidates: list[int] = []
    for candidate in raw_candidates:
        if candidate is None or candidate < 1 or candidate > total_pages:
            continue
        for nearby in (candidate, candidate - 2, candidate - 1, candidate + 1, candidate + 2):
            if 1 <= nearby <= total_pages and nearby not in candidates:
                candidates.append(nearby)
    return candidates


def _chapter_text_match_score(text: str, chapter: LibraryChapter, query: str) -> int:
    compact_text = re.sub(r"\s+", "", text).lower()
    compact_title = re.sub(r"\s+", "", chapter.title).lower()
    score = 0
    if compact_title and compact_title in compact_text:
        score += 4
    for path_item in chapter.path:
        compact_path_item = re.sub(r"\s+", "", path_item).lower()
        if compact_path_item and compact_path_item in compact_text:
            score += 2
    for keyword in _keywords_from_text(f"{chapter.title}\n{' '.join(chapter.keywords)}\n{query}")[:8]:
        if re.sub(r"\s+", "", keyword.lower()) in compact_text:
            score += 1
    return score


def _find_pdf_text_by_keywords(reader: PdfReader, chapter: LibraryChapter, query: str) -> str:
    scored_pages: list[tuple[int, int, str]] = []
    for page_index, page in enumerate(reader.pages):
        try:
            text = _normalize_extracted_text(page.extract_text() or "")
        except Exception:
            continue
        if not text:
            continue
        score = _chapter_text_match_score(text, chapter, query)
        if score:
            scored_pages.append((score, -page_index, text))
    if not scored_pages:
        return ""
    scored_pages.sort(reverse=True)
    return scored_pages[0][2]


def _outline_entries_to_chapters(
    entries: list[tuple[str, int, int | None]],
    total_pages: int,
    *,
    reader: PdfReader | None = None,
) -> list[LibraryChapter]:
    chapters: list[LibraryChapter] = []
    for index, (title, level, page_start) in enumerate(entries):
        page_end = None
        if page_start:
            page_end = total_pages
            for _, candidate_level, candidate_page in entries[index + 1 :]:
                if not candidate_page or candidate_level > level:
                    continue
                if candidate_page <= page_start:
                    # Some PDFs place a section label and its first subsection on the
                    # same page. Keep scanning until we find the next real page break.
                    continue
                page_end = max(page_start, candidate_page - 1)
                break
        page_label = None
        if page_start:
            page_label = str(page_start) if page_end == page_start or not page_end else f"{page_start}-{page_end}"
        summary = (
            f"PDF 页 {page_label} 已按目录定位；引用时将读取该页范围正文。"
            if page_label
            else f"PDF 目录项“{title}”被收录进课程资料库。"
        )
        if reader is not None and page_start:
            window_text = _read_pdf_text_window(
                reader,
                page_start=page_start,
                page_end=page_end or page_start,
                max_pages=6,
                max_nonempty_pages=1,
            )
            snippet = _summary_snippet(window_text)
            if snippet:
                summary = f"PDF 页 {page_label} 内容摘要：{snippet}"
        chapters.append(
            _chapter(
                title=title,
                summary=summary,
                keywords=[token.lower() for token in re.findall(r"[A-Za-z\u4e00-\u9fff]+", title)[:5]],
                level=level,
                locator_hint=_pdf_locator_hint(title, source="pdf_outline", actual_page=page_start) if page_start else title,
                order_index=index,
                scan_strategy="page_window" if page_start else "outline_only",
                page_start=page_start,
                page_end=page_end,
            )
        )
    return chapters


def _toc_entry_level(title: str) -> int:
    cleaned = title.strip()
    if re.match(r"^(?:第\s*[0-9一二三四五六七八九十百〇零两]+\s*章|chapter\s+\d+)\b", cleaned, flags=re.IGNORECASE):
        return 1
    if re.match(r"^\d+\s*[.．]\s*\d+", cleaned):
        return 2
    return 1


def _parse_toc_entries(text: str) -> list[tuple[str, int, int]]:
    entries: list[tuple[str, int, int]] = []
    for raw_line in text.splitlines():
        line = re.sub(r"\s+", " ", raw_line).strip(" .·•\t")
        if not line or line in {"目录", "目 录", "contents", "Contents"}:
            continue
        line = re.sub(r"[.．·•…]{2,}", " ", line)
        match = re.search(r"(?P<title>.+?)\s+(?P<page>\d{1,4})$", line)
        if not match:
            continue
        title = match.group("title").strip(" .·•…")
        if len(title) < 2 or not re.search(r"(?:第\s*[0-9一二三四五六七八九十百〇零两]+\s*章|\d+\s*[.．]\s*\d+|chapter\s+\d+)", title, re.IGNORECASE):
            continue
        page_number = int(match.group("page"))
        entries.append((title, _toc_entry_level(title), page_number))
    return entries


def _looks_like_toc_page(text: str) -> bool:
    compact = re.sub(r"\s+", "", text).lower()
    return "目录" in compact or "contents" in compact or len(_parse_toc_entries(text)) >= 2


def _extract_pdf_toc_text_pages(reader: PdfReader, file_path: Path, *, max_pages: int = 20) -> list[tuple[int, str]]:
    toc_pages: list[tuple[int, str]] = []
    in_toc = False
    for page_number in range(1, min(max_pages, len(reader.pages)) + 1):
        text = _read_pdf_text_window(reader, page_start=page_number, page_end=page_number, max_pages=1)
        if not text and page_number <= 12:
            text = extract_pdf_pages_text(file_path, page_start=page_number, page_end=page_number, max_pages=1) or ""
            text = _normalize_extracted_text(text)
        if not text:
            if in_toc:
                break
            continue
        entry_count = len(_parse_toc_entries(text))
        if _looks_like_toc_page(text) or (in_toc and entry_count):
            toc_pages.append((page_number, text))
            in_toc = True
            continue
        if in_toc:
            break
    return toc_pages


def _resolve_toc_entry_actual_page(
    reader: PdfReader,
    *,
    title: str,
    toc_page: int,
    printed_page: int,
    page_structure: ResourcePageStructure | None = None,
) -> int | None:
    candidates = [
        *physical_page_candidates_for_printed_page(page_structure, printed_page),
        printed_page,
        toc_page + printed_page,
        toc_page + printed_page - 1,
    ]
    candidates = list(dict.fromkeys(candidates))
    for candidate in candidates:
        if candidate < 1 or candidate > len(reader.pages):
            continue
        window_text = _read_pdf_text_window(
            reader,
            page_start=candidate,
            page_end=min(candidate + 1, len(reader.pages)),
            max_pages=2,
        )
        if not window_text:
            continue
        pseudo_chapter = LibraryChapter(title=title, summary="", keywords=_keywords_from_text(title))
        if _chapter_text_match_score(window_text, pseudo_chapter, title) > 0:
            return candidate
    return next((candidate for candidate in candidates if 1 <= candidate <= len(reader.pages)), None)


def _toc_entries_to_chapters(
    reader: PdfReader,
    toc_pages: list[tuple[int, str]],
    *,
    page_structure: ResourcePageStructure | None = None,
) -> list[LibraryChapter]:
    raw_entries: list[tuple[str, int, int, int]] = []
    for toc_page, toc_text in toc_pages:
        for title, level, printed_page in _parse_toc_entries(toc_text):
            raw_entries.append((title, level, printed_page, toc_page))
    if not raw_entries:
        return []

    chapters: list[LibraryChapter] = []
    for index, (title, level, printed_page, toc_page) in enumerate(raw_entries):
        actual_page = _resolve_toc_entry_actual_page(
            reader,
            title=title,
            toc_page=toc_page,
            printed_page=printed_page,
            page_structure=page_structure,
        )
        page_end = None
        next_entry = next(
            (
                candidate
                for candidate in raw_entries[index + 1 :]
                if candidate[1] <= level and candidate[2] > printed_page
            ),
            None,
        )
        if actual_page and next_entry:
            next_actual_page = _resolve_toc_entry_actual_page(
                reader,
                title=next_entry[0],
                toc_page=next_entry[3],
                printed_page=next_entry[2],
                page_structure=page_structure,
            )
            if next_actual_page and next_actual_page > actual_page:
                page_end = next_actual_page - 1
        elif actual_page:
            page_end = min(actual_page + 3, len(reader.pages))

        page_label = str(actual_page) if actual_page and (not page_end or page_end == actual_page) else (
            f"{actual_page}-{page_end}" if actual_page and page_end else None
        )
        summary = (
            f"PDF 目录页 {toc_page} 标注页码 {printed_page}；引用时会尝试实际页、目录页偏移和全文检索定位正文。"
        )
        if page_label:
            summary = f"PDF 页 {page_label} 已由目录页 {toc_page} 的页码 {printed_page} 定位；引用时会再次校验正文。"
        chapters.append(
            _chapter(
                title=title,
                summary=summary,
                keywords=_keywords_from_text(title),
                level=level,
                locator_hint=_pdf_locator_hint(
                    title,
                    source="toc_page",
                    toc_page=toc_page,
                    printed_page=printed_page,
                    actual_page=actual_page,
                ),
                order_index=index,
                scan_strategy="page_window" if actual_page else "fulltext_match",
                page_start=actual_page,
                page_end=page_end,
            )
        )
    return chapters


def extract_outline(
    file_path: Path,
    original_name: str,
    mime_type: str,
) -> tuple[list[LibraryChapter], bool, str | None, ResourcePageStructure | None]:
    if mime_type.startswith("image/"):
        generic_title = Path(original_name).stem
        extracted_text = extract_image_text(file_path)
        if extracted_text:
            ai_outline = _ai_generated_outline(original_name, extracted_text)
            if ai_outline:
                return ai_outline, True, extracted_text, None
            return (
                [
                    _generic_chapter_from_text(
                        generic_title,
                        extracted_text,
                        summary_prefix="从图片中识别到的文字摘要：",
                    )
                ],
                True,
                extracted_text,
                None,
            )
        return (
            [
                _chapter(
                    title=generic_title,
                    summary=f"已上传图片资料“{original_name}”，可作为当前课程的视觉参考。",
                    keywords=[token.lower() for token in re.findall(r"[A-Za-z\u4e00-\u9fff]+", generic_title)[:6]],
                    locator_hint=generic_title,
                    order_index=0,
                )
            ],
            False,
            None,
            None,
        )

    if mime_type in {"text/plain", "text/markdown"} or file_path.suffix.lower() in {".md", ".txt"}:
        text = _read_text_file(file_path)
        outline = _extract_markdown_outline(text)
        if outline:
            return outline, True, text, None
        ai_outline = _ai_generated_outline(original_name, text)
        if ai_outline:
            return ai_outline, True, text, None
        return (
            [
                _generic_chapter_from_text(
                    Path(original_name).stem,
                    text,
                    summary_prefix="从文本资料中抽取到的内容摘要：",
                )
            ],
            True,
            text,
            None,
        )

    if file_path.suffix.lower() == ".docx":
        text = _read_docx_text(file_path)
        outline = _extract_docx_outline(file_path)
        if outline:
            return outline, True, text, None
        ai_outline = _ai_generated_outline(original_name, text)
        if ai_outline:
            return ai_outline, True, text, None
        return (
            [
                _generic_chapter_from_text(
                    Path(original_name).stem,
                    text,
                    summary_prefix="从 Word 资料中抽取到的内容摘要：",
                )
            ],
            True,
            text,
            None,
        )

    if file_path.suffix.lower() == ".epub":
        text = _read_epub_text(file_path)
        outline = _extract_epub_outline(file_path)
        if outline:
            return outline, True, text[:200000] if text else None, None
        if text:
            ai_outline = _ai_generated_outline(original_name, text)
            if ai_outline:
                return ai_outline, True, text[:200000], None
            return (
                [
                    _generic_chapter_from_text(
                        Path(original_name).stem,
                        text,
                        summary_prefix="从 EPUB 资料中抽取到的内容摘要：",
                    )
                ],
                True,
                text[:200000],
                None,
            )

    if file_path.suffix.lower() == ".pdf":
        reader = PdfReader(str(file_path))
        page_structure = build_pdf_page_structure(reader)
        if reader.outline:
            entries: list[tuple[str, int, int | None]] = []

            def _walk_outline(items: list, level: int = 1) -> None:
                for item in items:
                    if isinstance(item, list):
                        _walk_outline(item, level + 1)
                        continue
                    title = str(getattr(item, "title", item))
                    page_start = None
                    try:
                        page_start = reader.get_destination_page_number(item) + 1
                    except Exception:
                        page_start = None
                    entries.append((title, level, page_start))

            _walk_outline(list(reader.outline))
            chapters = _outline_entries_to_chapters(entries, len(reader.pages), reader=reader)
            if chapters:
                return chapters, True, None, page_structure

        toc_pages = _extract_pdf_toc_text_pages(reader, file_path)
        toc_chapters = _toc_entries_to_chapters(reader, toc_pages, page_structure=page_structure)
        if toc_chapters:
            return toc_chapters, True, None, page_structure
        extracted_text = []
        for page in reader.pages[:2]:
            try:
                extracted_text.append(page.extract_text() or "")
            except Exception:
                continue
        joined_text = "\n".join(extracted_text).strip()
        if joined_text:
            ai_outline = _ai_generated_outline(original_name, joined_text)
            if ai_outline:
                return ai_outline, True, None, page_structure
            return (
                [
                    _generic_chapter_from_text(
                        Path(original_name).stem,
                        joined_text,
                        summary_prefix="从 PDF 前几页抽取到的内容摘要：",
                    )
                ],
                True,
                None,
                page_structure,
            )
    generic_title = Path(original_name).stem
    return (
        [
            _chapter(
                title=generic_title,
                summary=f"当前资料尚未提取出显式目录，先以“{generic_title}”作为入口。",
                keywords=[token.lower() for token in re.findall(r"[A-Za-z\u4e00-\u9fff]+", generic_title)[:5]],
                locator_hint=generic_title,
                order_index=0,
            )
        ],
        False,
        None,
        None,
    )


def _outline_from_source_units(original_name: str, source_units: list[ResourceSourceUnit]) -> list[LibraryChapter]:
    chapters: list[LibraryChapter] = []
    chapters_by_path: dict[tuple[str, ...], LibraryChapter] = {}
    for unit in sorted(source_units, key=lambda item: item.order_index):
        heading_path = _readable_heading_path(unit)
        if not heading_path:
            continue
        for depth in range(1, len(heading_path) + 1):
            path = tuple(heading_path[:depth])
            if path in chapters_by_path:
                continue
            title = path[-1]
            snippet = _summary_snippet(unit.text) if depth == len(heading_path) else ""
            summary = snippet or f"来自资料“{Path(original_name).stem}”的目录入口：{' > '.join(path)}。"
            chapter = _chapter(
                title=title,
                summary=summary,
                keywords=_keywords_from_text(f"{' '.join(path)}\n{unit.text}"),
                level=depth,
                locator_hint=f"heading:{' > '.join(path)}",
                order_index=len(chapters),
                scan_strategy="heading_section",
                page_start=unit.page_no,
                page_end=unit.page_no,
            )
            chapters_by_path[path] = chapter
            chapters.append(chapter)
    return chapters


def _outline_from_source_text(original_name: str, text: str) -> list[LibraryChapter]:
    outline = _extract_markdown_outline(text)
    if outline:
        return outline
    ai_outline = _ai_generated_outline(original_name, text)
    if ai_outline:
        return ai_outline
    return [
        _generic_chapter_from_text(
            Path(original_name).stem,
            text,
            summary_prefix="从资料解析结果中抽取到的内容摘要：",
        )
    ]


def _concept_index_from_outline(outline: list[LibraryChapter]) -> dict[str, list[str]]:
    concept_index: dict[str, list[str]] = {}
    for chapter in outline:
        path_keywords = _keywords_from_text(" ".join(chapter.path))
        for keyword in [*chapter.keywords, *path_keywords]:
            concept_index.setdefault(keyword, []).append(chapter.id)
    return concept_index


def _outline_looks_like_body_fragments(outline: list[LibraryChapter]) -> bool:
    if not outline:
        return True
    structured_count = 0
    fragment_count = 0
    artifact_count = 0
    for chapter in outline[:20]:
        title = chapter.title.strip()
        if _is_parser_artifact_title(title):
            artifact_count += 1
            fragment_count += 1
            continue
        if _looks_like_reference_heading(title):
            structured_count += 1
            continue
        compact = re.sub(r"\s+", "", title)
        punctuation_count = len(re.findall(r"[,，;；:：/*%{}=]|\\.\\.\\.", title))
        if len(compact) >= 34 or punctuation_count >= 2:
            fragment_count += 1
    if artifact_count >= 2:
        return True
    return structured_count <= 1 and fragment_count >= max(2, min(5, len(outline[:20]) // 2))


def resource_with_epub_catalog_outline(resource: ResourceLibraryItem) -> ResourceLibraryItem:
    if not resource.source_path or Path(resource.source_path).suffix.lower() != ".epub":
        return resource
    if not _outline_looks_like_body_fragments(resource.outline):
        return resource
    file_path = Path(resource.source_path)
    if not file_path.exists():
        return resource
    outline = _attach_outline_hierarchy(_extract_epub_outline(file_path))
    if not outline:
        return resource
    return resource.model_copy(
        update={
            "outline": outline,
            "concept_index": _concept_index_from_outline(outline),
        }
    )


def _source_units_from_native_text(original_name: str, text_content: str | None) -> list[ResourceSourceUnit]:
    if not text_content:
        return []
    suffix = Path(original_name).suffix.lower().lstrip(".") or "document"
    return [
        ResourceSourceUnit(
            content_type="text",
            text=text_content[:500000],
            source_locator=f"native:{suffix}:fulltext",
            order_index=0,
        )
    ]


def _native_provider_from_warnings(warnings: list[str]) -> tuple[str, str]:
    fallback = any("used native parser" in warning.lower() for warning in warnings)
    if fallback:
        return "native_fallback", "Parsed by native OpenClass parser after RAG-Anything fallback."
    return "native", "Parsed by native OpenClass parser."


def build_resource_item(file_path: Path, original_name: str) -> ResourceLibraryItem:
    mime_type = mimetypes.guess_type(original_name)[0] or "application/octet-stream"
    parser_provider = "native"
    parser_artifacts_path: str | None = None
    parser_message = "Parsed by native OpenClass parser."
    parse_warnings: list[str] = []
    source_units: list[ResourceSourceUnit] = []
    page_structure: ResourcePageStructure | None = None

    rag_attempt = parse_with_rag_anything(file_path, original_name, mime_type)
    parse_warnings.extend(rag_attempt.warnings)
    if rag_attempt.result is not None:
        rag_result = rag_attempt.result
        extracted = True
        text_content = rag_result.text_content
        page_structure = build_page_structure_from_source_units(rag_result.source_units)
        source_units = enrich_source_units_with_page_structure(rag_result.source_units, page_structure)
        outline = _outline_from_source_units(original_name, source_units) or _outline_from_source_text(
            original_name,
            rag_result.text_content,
        )
        parser_provider = rag_result.parser_provider
        parser_artifacts_path = rag_result.parser_artifacts_path
        parser_message = rag_result.parser_message
        parse_warnings.extend(rag_result.warnings)
    else:
        outline, extracted, text_content, page_structure = extract_outline(file_path, original_name, mime_type)
        source_units = _source_units_from_native_text(original_name, text_content)
        parser_provider, parser_message = _native_provider_from_warnings(parse_warnings)

    outline = _attach_outline_hierarchy(outline)
    concept_index = _concept_index_from_outline(outline)

    resource = ResourceLibraryItem(
        name=original_name,
        mime_type=mime_type,
        resource_type="image" if mime_type.startswith("image/") else "document",
        size_bytes=file_path.stat().st_size,
        outline=outline,
        concept_index=concept_index,
        extracted_text_available=extracted,
        text_content=text_content,
        source_path=str(file_path),
        parser_provider=parser_provider,
        parser_artifacts_path=parser_artifacts_path,
        parser_message=parser_message,
        parse_warnings=parse_warnings,
        source_units=source_units,
        page_structure=page_structure,
    )
    return mark_local_file_ready(resource, str(file_path))


def extract_reference_context(
    resource: ResourceLibraryItem,
    chapter_id: str,
    *,
    user_query: str,
) -> ResourceReferenceContext | None:
    chapter = next((candidate for candidate in resource.outline if candidate.id == chapter_id), None)
    if chapter is None:
        return None

    if resource.parser_provider.startswith("raganything") and resource.text_content and resource.source_units:
        return _build_reference_context(
            resource,
            chapter,
            user_query,
            raw_text=resource.text_content,
        )

    if resource.text_content and resource.resource_type == "image":
        return _build_reference_context(
            resource,
            chapter,
            user_query,
            raw_text=resource.text_content,
        )

    if not resource.source_path:
        return _build_reference_context(
            resource,
            chapter,
            user_query,
            raw_text=resource.text_content or f"{chapter.title}\n{chapter.summary}\n{' '.join(chapter.keywords)}",
        )

    file_path = Path(resource.source_path)
    if not file_path.exists():
        if resource.text_content:
            return _build_reference_context(
                resource,
                chapter,
                user_query,
                raw_text=resource.text_content,
            )
        return None

    suffix = file_path.suffix.lower()
    raw_text = ""
    if resource.mime_type in {"text/plain", "text/markdown"} or suffix in {".md", ".txt"}:
        if chapter.scan_strategy == "heading_section":
            raw_text = _extract_markdown_section_text(file_path, chapter)
        else:
            raw_text = _read_text_file(file_path)
    elif suffix == ".docx":
        if chapter.scan_strategy == "heading_section":
            raw_text = _extract_docx_section_text(file_path, chapter)
        else:
            raw_text = _read_docx_text(file_path)
    elif suffix == ".epub":
        chapter_title, raw_text = _extract_epub_section_text(file_path, chapter, user_query)
        if chapter_title and chapter_title != chapter.title:
            chapter = chapter.model_copy(
                update={
                    "title": chapter_title,
                    "summary": _summary_snippet(raw_text, limit=180)
                    or f"EPUB 章节“{chapter_title}”可作为本次讲解参考。",
                    "keywords": _keywords_from_text(f"{chapter_title}\n{raw_text}"),
                    "locator_hint": chapter_title,
                }
            )
    elif suffix == ".pdf":
        raw_text = _extract_pdf_chapter_text(
            file_path,
            chapter,
            user_query,
            page_structure=resource.page_structure,
        )
    else:
        raw_text = resource.text_content or f"{chapter.title}\n{chapter.summary}\n{' '.join(chapter.keywords)}"

    return _build_reference_context(resource, chapter, user_query, raw_text)
