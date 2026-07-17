from __future__ import annotations

import hashlib
import html
import posixpath
import re
from collections import Counter
from dataclasses import dataclass, field
from html.parser import HTMLParser
from pathlib import Path
from typing import Any, Callable
from urllib.parse import unquote

from app.models import (
    SourceChapter,
    SourceChunk,
    SourceIngestionRecord,
    SourceStructure,
    SourceVisualAsset,
)
from app.services import workspace_state
from app.services.ai_logging import ai_usage_logger
from app.services.image_ocr import (
    extract_image_text,
    extract_pdf_pages_layout,
    extract_pdf_pages_text,
    ordered_ocr_lines,
)
from app.services.pdf_toc_parser import (
    PdfOutlineAnchor,
    PdfTocNode,
    extract_pdf_toc,
    extract_pdf_toc_from_range,
    is_toc_heading,
    normalize_toc_text,
    parse_structural_heading,
)
from app.services.native_source_index import source_chunk_text_hash
from app.services.source_chapter_identity import stable_source_chapter_id
from app.services.source_archive import SafeSourceArchive
from app.services.source_ingestion_jobs import SourceIngestionCoordinator
from app.services.source_structure_store import SourceStructureStore, source_structure_store
from app.services.source_structure_quality import (
    evaluate_source_structure_quality,
    meaningful_text_character_count,
)
from app.services.source_visual_extraction import (
    CURRENT_SOURCE_VISUAL_INDEX_VERSION,
    SourceVisualExtractionResult,
    SourceVisualExtractor,
    source_visual_extractor,
)
from app.services.source_visual_storage import source_visual_staging
from app.services.source_xml import parse_untrusted_xml

CHUNK_CHAR_LIMIT = 1800
CHUNK_CHAR_OVERLAP = 160
CURRENT_SOURCE_STRUCTURE_INDEX_VERSION = 8
SourceIndexProgressCallback = Callable[[str, int], None]


def _report_progress(
    callback: SourceIndexProgressCallback | None,
    phase: str,
    progress: int,
) -> None:
    if callback is not None:
        callback(phase, max(0, min(100, progress)))


@dataclass
class PageText:
    page_no: int
    text: str
    start_offset: int = 0
    end_offset: int = 0
    content_start_offset: int | None = None


@dataclass
class DetectedChapter:
    title: str
    number: str = ""
    level: int = 1
    source_locator: str = ""
    start_offset: int | None = None
    end_offset: int | None = None
    page_start: int | None = None
    page_end: int | None = None
    confidence: float = 0.0
    verified: bool = False
    metadata: dict[str, Any] = field(default_factory=dict)


@dataclass
class ParsedSourceDocument:
    text: str
    chapters: list[DetectedChapter] = field(default_factory=list)
    pages: list[PageText] = field(default_factory=list)
    strategy: str = "linear_text"
    warnings: list[str] = field(default_factory=list)
    metadata: dict[str, Any] = field(default_factory=dict)


@dataclass(frozen=True)
class _EpubNavigationItem:
    target: str
    fragment: str
    label: str
    order: int
    level: int

    @property
    def source_locator(self) -> str:
        suffix = f"#{self.fragment}" if self.fragment else ""
        return f"epub:{self.target}{suffix}"


class SourceStructureIndexer:
    def __init__(
        self,
        *,
        store: SourceStructureStore = source_structure_store,
        visual_extractor: SourceVisualExtractor = source_visual_extractor,
        coordinator: SourceIngestionCoordinator | None = None,
    ) -> None:
        self.store = store
        self.visual_extractor = visual_extractor
        self.coordinator = coordinator or store.coordinator

    def ensure_structure(
        self,
        record: SourceIngestionRecord,
        *,
        progress_callback: SourceIndexProgressCallback | None = None,
    ) -> SourceStructure | None:
        current = self.store.get_structure(
            owner_user_id=record.owner_user_id,
            package_id=record.package_id,
            source_id=record.id,
        )
        if current and current.status in {"ready", "linear_only", "failed"}:
            structure_is_current = not source_structure_needs_upgrade(current)
            visuals_are_current = (
                current.visual_index_version >= CURRENT_SOURCE_VISUAL_INDEX_VERSION
            )
            if structure_is_current and visuals_are_current:
                _report_progress(progress_callback, "persisting", 96)
                return current
            if current.status in {"ready", "linear_only"}:
                if structure_is_current and not _record_supports_visual_index(record):
                    _report_progress(progress_callback, "persisting", 96)
                    return current
                # Reuse the public rebuild boundary for lazy upgrades.  Besides
                # preserving the existing observable contract, this keeps text,
                # chapter, chunk, and visual identities on one atomic save path.
                # Structure-only upgrades preserve the current visual index so
                # a visual rebuild cannot block newer chapter parsing.
                if not structure_is_current and visuals_are_current:
                    if progress_callback is None:
                        return self.rebuild_structure(record, preserve_existing_visuals=True)
                    return self.rebuild_structure(
                        record,
                        preserve_existing_visuals=True,
                        progress_callback=progress_callback,
                    )
                if progress_callback is None:
                    return self.rebuild_structure(record)
                return self.rebuild_structure(record, progress_callback=progress_callback)
        if progress_callback is None:
            return self.rebuild_structure(record)
        return self.rebuild_structure(record, progress_callback=progress_callback)

    def rebuild_structure(
        self,
        record: SourceIngestionRecord,
        *,
        preserve_existing_visuals: bool = False,
        progress_callback: SourceIndexProgressCallback | None = None,
    ) -> SourceStructure:
        _report_progress(progress_callback, "waiting_for_worker", 24)
        weight = self.coordinator.processing_weight(
            size_bytes=record.size_bytes,
            source_type=record.source_type,
        )
        with self.coordinator.processing_slot(weight=weight):
            return self._rebuild_structure(
                record,
                preserve_existing_visuals=preserve_existing_visuals,
                progress_callback=progress_callback,
            )

    def _rebuild_structure(
        self,
        record: SourceIngestionRecord,
        *,
        preserve_existing_visuals: bool = False,
        progress_callback: SourceIndexProgressCallback | None = None,
    ) -> SourceStructure:
        previous_visuals: list[SourceVisualAsset] = []
        previous_chapters: list[SourceChapter] = []
        if preserve_existing_visuals:
            previous_view = self.store.get_structure_view(source=record, chunk_limit=0)
            previous = previous_view.structure
            previous_visuals = previous_view.visuals
            previous_chapters = previous_view.chapters
        else:
            previous = self.store.get_structure(
                owner_user_id=record.owner_user_id,
                package_id=record.package_id,
                source_id=record.id,
            )
        building = SourceStructure(
            owner_user_id=record.owner_user_id,
            package_id=record.package_id,
            source_ingestion_id=record.id,
            status="building",
            strategy="linear_text",
            visual_index_status="pending",
            visual_index_version=CURRENT_SOURCE_VISUAL_INDEX_VERSION,
            metadata={
                "source_title": record.title,
                "mime_type": record.mime_type,
                "structure_index_version": CURRENT_SOURCE_STRUCTURE_INDEX_VERSION,
            },
        )
        with source_visual_staging():
            return self._rebuild_structure_staged(
                record,
                previous=previous,
                building=building,
                preserve_existing_visuals=preserve_existing_visuals,
                previous_visuals=previous_visuals,
                previous_chapters=previous_chapters,
                progress_callback=progress_callback,
            )

    def _rebuild_structure_staged(
        self,
        record: SourceIngestionRecord,
        *,
        previous: SourceStructure | None,
        building: SourceStructure,
        preserve_existing_visuals: bool = False,
        previous_visuals: list[SourceVisualAsset] | None = None,
        previous_chapters: list[SourceChapter] | None = None,
        progress_callback: SourceIndexProgressCallback | None = None,
    ) -> SourceStructure:
        extracted_storage_keys: list[str] = []
        try:
            _report_progress(progress_callback, "parsing", 25)
            parsed = (
                self._parse_record(record, progress_callback=progress_callback)
                if progress_callback is not None
                else self._parse_record(record)
            )
            _report_progress(progress_callback, "mapping_structure", 58)
            chapters = self._chapters_for_record(record, parsed)
            quality_result = evaluate_source_structure_quality(
                chapters=chapters,
                text=parsed.text,
                strategy=parsed.strategy,
                metadata=parsed.metadata,
            )
            chapters = quality_result.chapters
            _report_progress(progress_callback, "building_chunks", 63)
            chunks = self._chunks_for_record(record, parsed, chapters)
            _report_progress(progress_callback, "extracting_visuals", 70)
            if preserve_existing_visuals and previous is not None:
                visual_result = SourceVisualExtractionResult(
                    status=(
                        previous.visual_index_status
                        if previous.visual_index_status in {"ready", "partial"}
                        else "ready"
                    ),
                    visuals=_reanchor_existing_visuals(
                        previous_visuals or [],
                        structure=building,
                        chapters=chapters,
                        chunks=chunks,
                        previous_chapters=previous_chapters or [],
                    ),
                )
            else:
                try:
                    visual_kwargs: dict[str, Any] = {
                        "record": record,
                        "path": _local_source_path(record),
                        "structure": building,
                        "chapters": chapters,
                        "chunks": chunks,
                    }
                    if progress_callback is not None:
                        visual_kwargs["progress_callback"] = lambda completed, total: _report_progress(
                            progress_callback,
                            "extracting_visuals",
                            70 + round(22 * completed / max(1, total)),
                        )
                    visual_result = self.visual_extractor.extract(**visual_kwargs)
                    if (
                        visual_result.status == "failed"
                        and previous is not None
                        and previous.status in {"ready", "linear_only"}
                        and previous.visual_index_status in {"ready", "partial"}
                    ):
                        previous_view = self.store.get_structure_view(
                            source=record,
                            chunk_limit=0,
                        )
                        visual_result = SourceVisualExtractionResult(
                            status=previous.visual_index_status,
                            visuals=_reanchor_existing_visuals(
                                previous_view.visuals,
                                structure=building,
                                chapters=chapters,
                                chunks=chunks,
                                previous_chapters=previous_view.chapters,
                            ),
                            warnings=list(
                                dict.fromkeys(
                                    [
                                        *visual_result.warnings,
                                        "视觉索引重建失败，已保留上一次可用的视觉索引。",
                                    ]
                                )
                            ),
                        )
                except Exception as visual_exc:
                    if previous is not None and previous.status in {"ready", "linear_only"}:
                        raise
                    ai_usage_logger.log_event(
                        "source_visual_initial_index_failed",
                        owner_user_id=record.owner_user_id,
                        package_id=record.package_id,
                        source_ingestion_id=record.id,
                        error=str(visual_exc),
                    )
                    visual_result = SourceVisualExtractionResult(
                        status="failed",
                        warnings=[
                            "资料文本已建立索引，但视觉索引失败；重新构建后才能使用图表证据。"
                        ],
                    )
            _report_progress(progress_callback, "persisting", 94)
            extracted_storage_keys = [
                visual.storage_key for visual in visual_result.visuals if visual.storage_key
            ]
            has_verified_toc = any(chapter.anchor_status == "verified" for chapter in chapters)
            status = "ready" if has_verified_toc else "linear_only"
            confidence = quality_result.quality.confidence
            structure = building.model_copy(
                update={
                    "status": status,
                    "strategy": parsed.strategy,
                    "confidence": confidence,
                    "quality": quality_result.quality,
                    "warnings": list(
                        dict.fromkeys(
                            [
                                *parsed.warnings,
                                *quality_result.warnings,
                                *visual_result.warnings,
                            ]
                        )
                    ),
                    "visual_count": len(visual_result.visuals),
                    "visual_index_status": visual_result.status,
                    "visual_index_version": CURRENT_SOURCE_VISUAL_INDEX_VERSION,
                    "metadata": {
                        **building.metadata,
                        **parsed.metadata,
                        "text_length": len(parsed.text),
                        "chapter_node_count": len(chapters),
                        "verified_chapter_count": sum(
                            chapter.anchor_status == "verified" for chapter in chapters
                        ),
                        "unverified_chapter_count": sum(
                            chapter.anchor_status == "unverified" for chapter in chapters
                        ),
                        "visual_count": len(visual_result.visuals),
                        "verified_visual_count": sum(
                            visual.anchor_status == "verified"
                            for visual in visual_result.visuals
                        ),
                        "unverified_visual_count": sum(
                            visual.anchor_status == "unverified"
                            for visual in visual_result.visuals
                        ),
                        "visual_index_version": CURRENT_SOURCE_VISUAL_INDEX_VERSION,
                    },
                }
            )
            return self.store.save_structure_bundle(
                structure=structure,
                chapters=chapters,
                chunks=chunks,
                visuals=visual_result.visuals,
            )
        except Exception as exc:  # pragma: no cover - defensive boundary around third-party parsers
            self.store.cleanup_unreferenced_visual_assets(extracted_storage_keys)
            if previous is not None and previous.status in {"ready", "linear_only"}:
                return self.store.record_rebuild_failure(structure=previous, error=str(exc))
            failed = building.model_copy(
                update={
                    "status": "failed",
                    "error": str(exc),
                    "warnings": ["资料结构索引失败，请检查文件格式后重试。"],
                }
            )
            return self.store.save_structure_bundle(structure=failed, chapters=[], chunks=[])

    def _parse_record(
        self,
        record: SourceIngestionRecord,
        *,
        progress_callback: SourceIndexProgressCallback | None = None,
    ) -> ParsedSourceDocument:
        local_path = _local_source_path(record)
        if not local_path:
            return ParsedSourceDocument(
                text="",
                strategy="linear_text",
                warnings=["未找到资料原文件，请重新导入后再建立索引。"],
                metadata={"source_type": record.source_type, "missing_local_source_path": True},
            )
        suffix = local_path.suffix.lower()
        if suffix == ".epub" or _looks_like_epub(record.mime_type):
            return _parse_epub(local_path)
        if suffix == ".pdf" or record.mime_type == "application/pdf":
            return _parse_pdf(local_path, progress_callback=progress_callback)
        if suffix in {".docx", ".doc"} or "wordprocessingml" in record.mime_type:
            return _parse_docx(local_path)
        if suffix == ".pptx" or "presentationml" in record.mime_type:
            return _parse_pptx(local_path)
        if suffix in {".xlsx", ".xls"} or "spreadsheetml" in record.mime_type:
            return _parse_xlsx(local_path)
        if suffix in {".html", ".htm"} or record.mime_type == "text/html":
            return _parse_html_document(local_path)
        if record.mime_type.startswith("image/") or suffix in {".png", ".jpg", ".jpeg", ".webp", ".gif"}:
            return _parse_image(local_path)
        if record.mime_type.startswith(("audio/", "video/")):
            return _parse_text_document(local_path, prefer_markdown=False)
        return _parse_text_document(local_path, prefer_markdown=suffix in {".md", ".markdown"})

    def _chapters_for_record(self, record: SourceIngestionRecord, parsed: ParsedSourceDocument) -> list[SourceChapter]:
        chapters: list[SourceChapter] = []
        level_stack: list[SourceChapter] = []
        semantic_occurrences: Counter[tuple[tuple[str, ...], str, str, int]] = Counter()
        for index, chapter in enumerate(parsed.chapters):
            is_verified = chapter.verified and chapter.start_offset is not None
            end_offset = chapter.end_offset
            if is_verified and end_offset is None:
                next_chapter = next(
                    (
                        candidate
                        for candidate in parsed.chapters[index + 1 :]
                        if candidate.verified and candidate.start_offset is not None
                    ),
                    None,
                )
                end_offset = next_chapter.start_offset if next_chapter else len(parsed.text)
            excerpt = (
                _compact(parsed.text[chapter.start_offset or 0 : end_offset or len(parsed.text)], 360)
                if is_verified
                else ""
            )
            title = _clean_label(chapter.title)
            number = chapter.number or _number_from_title(title)
            level = max(1, chapter.level)
            while level_stack and level_stack[-1].level >= level:
                level_stack.pop()
            parent = level_stack[-1] if level_stack else None
            normalized_number = _normalize_chapter_number(number)
            semantic_key = (
                tuple(parent.path if parent else ()),
                normalized_number,
                _normalize_for_match(title),
                level,
            )
            semantic_occurrence = semantic_occurrences[semantic_key]
            semantic_occurrences[semantic_key] += 1
            source_chapter = SourceChapter(
                id=stable_source_chapter_id(
                    source_ingestion_id=record.id,
                    parent_path=parent.path if parent else (),
                    normalized_number=normalized_number,
                    title=title,
                    level=level,
                    source_locator=chapter.source_locator,
                    order_index=semantic_occurrence,
                ),
                owner_user_id=record.owner_user_id,
                package_id=record.package_id,
                source_ingestion_id=record.id,
                parent_id=parent.id if parent else None,
                number=number,
                normalized_number=normalized_number,
                title=title,
                level=level,
                path=[*(parent.path if parent else []), title],
                order_index=index,
                source_locator=chapter.source_locator,
                body_start_offset=chapter.start_offset if is_verified else None,
                body_end_offset=end_offset if is_verified else None,
                page_start=chapter.page_start if is_verified else None,
                page_end=chapter.page_end if is_verified else None,
                anchor_status="verified" if is_verified else "unverified",
                confidence=chapter.confidence,
                excerpt=excerpt,
                metadata={
                    **chapter.metadata,
                    "semantic_identity_version": 2,
                    "semantic_occurrence": semantic_occurrence,
                },
            )
            chapters.append(source_chapter)
            level_stack.append(source_chapter)
        return chapters

    def _chunks_for_record(
        self,
        record: SourceIngestionRecord,
        parsed: ParsedSourceDocument,
        chapters: list[SourceChapter],
    ) -> list[SourceChunk]:
        if not parsed.text.strip():
            return []
        chapter_ranges = [
            (
                chapter.id,
                chapter.body_start_offset,
                chapter.body_end_offset,
                chapter.page_start,
                chapter.page_end,
                chapter.level,
            )
            for chapter in chapters
            if chapter.anchor_status == "verified"
            and chapter.body_start_offset is not None
            and chapter.body_end_offset is not None
        ]
        verified_chapter_starts = sorted(
            {
                chapter.body_start_offset
                for chapter in chapters
                if chapter.anchor_status == "verified"
                and chapter.body_start_offset is not None
                and chapter.body_start_offset >= 0
            }
        )
        # Do not create a tiny preamble chunk before the document's first
        # heading (some HTML parsers report that heading a few characters in).
        # Every later verified chapter start is a hard boundary, including a
        # child section whose parent range continues beyond it.
        chapter_boundaries = verified_chapter_starts[1:]
        chunks: list[SourceChunk] = []
        cursor = 0
        order_index = 0
        text_length = len(parsed.text)
        while cursor < text_length:
            next_chapter_boundary = next(
                (boundary for boundary in chapter_boundaries if boundary > cursor),
                None,
            )
            ends_at_chapter_boundary = bool(
                next_chapter_boundary is not None
                and next_chapter_boundary <= cursor + CHUNK_CHAR_LIMIT
            )
            end = min(
                text_length,
                cursor + CHUNK_CHAR_LIMIT,
                next_chapter_boundary if ends_at_chapter_boundary else text_length,
            )
            if end < text_length and not ends_at_chapter_boundary:
                boundary = parsed.text.rfind("\n\n", cursor + CHUNK_CHAR_LIMIT // 2, end)
                if boundary > cursor:
                    end = boundary
            chunk_text = parsed.text[cursor:end].strip()
            if chunk_text:
                chapter_id, chapter_page_start, chapter_page_end = _chapter_for_chunk(
                    cursor,
                    end,
                    chapter_ranges,
                )
                physical_page_start, physical_page_end = _page_range_for_offsets(
                    parsed.pages,
                    cursor,
                    end,
                )
                page_start = physical_page_start or chapter_page_start
                page_end = physical_page_end or chapter_page_end
                text_hash = source_chunk_text_hash(chunk_text)
                chunks.append(
                    SourceChunk(
                        id=_stable_source_chunk_id(
                            source_ingestion_id=record.id,
                            chapter_id=chapter_id,
                            order_index=order_index,
                            text_hash=text_hash,
                        ),
                        owner_user_id=record.owner_user_id,
                        package_id=record.package_id,
                        source_ingestion_id=record.id,
                        chapter_id=chapter_id,
                        order_index=order_index,
                        source_locator=_locator_for_offset(parsed.pages, cursor),
                        text=chunk_text,
                        start_offset=cursor,
                        end_offset=end,
                        page_start=page_start,
                        page_end=page_end,
                        token_count=_estimate_tokens(chunk_text),
                        metadata={"text_hash": text_hash},
                    )
                )
                order_index += 1
            if end >= text_length:
                break
            cursor = (
                end
                if ends_at_chapter_boundary
                else max(end - CHUNK_CHAR_OVERLAP, cursor + 1)
            )
        return chunks

def _stable_source_chunk_id(
    *,
    source_ingestion_id: str,
    chapter_id: str | None,
    order_index: int,
    text_hash: str,
) -> str:
    identity = "\x1f".join(
        (source_ingestion_id, chapter_id or "", str(order_index), text_hash)
    )
    digest = hashlib.sha256(identity.encode("utf-8")).hexdigest()[:24]
    return f"sourcechunk_{digest}"


def _local_source_path(record: SourceIngestionRecord) -> Path | None:
    raw_path = record.metadata.get("local_source_path")
    if not isinstance(raw_path, str) or not raw_path.strip():
        return _find_legacy_upload_path(record)
    path = Path(raw_path).expanduser()
    if path.exists() and path.is_file():
        return path
    return _find_legacy_upload_path(record)


def _find_legacy_upload_path(record: SourceIngestionRecord) -> Path | None:
    upload_dir = workspace_state.UPLOAD_DIR
    if not upload_dir.exists():
        return None
    file_name = Path(record.file_name or record.title).name.strip()
    if not file_name:
        return None
    for candidate in upload_dir.rglob("*"):
        if not candidate.is_file():
            continue
        if candidate.name == file_name or candidate.name.endswith(f"_{file_name}"):
            return candidate
    return None


def _looks_like_epub(mime_type: str) -> bool:
    return mime_type == "application/epub+zip" or "epub" in mime_type.lower()


def _record_supports_visual_index(record: SourceIngestionRecord) -> bool:
    suffix = Path(record.file_name or record.title).suffix.lower()
    mime_type = record.mime_type.split(";", 1)[0].strip().lower()
    return bool(
        suffix
        in {
            ".pdf",
            ".docx",
            ".pptx",
            ".xlsx",
            ".epub",
            ".html",
            ".htm",
            ".md",
            ".markdown",
            ".csv",
            ".png",
            ".jpg",
            ".jpeg",
            ".webp",
            ".gif",
            ".tif",
            ".tiff",
            ".bmp",
            ".svg",
        }
        or mime_type
        in {
            "application/pdf",
            "application/epub+zip",
            "text/html",
            "text/markdown",
            "text/x-markdown",
            "text/csv",
        }
        or mime_type.startswith("image/")
        or "wordprocessingml.document" in mime_type
        or "presentationml.presentation" in mime_type
        or "spreadsheetml.sheet" in mime_type
    )


def _parse_text_document(path: Path, *, prefer_markdown: bool) -> ParsedSourceDocument:
    raw = path.read_bytes()
    try:
        text = raw.decode("utf-8")
    except UnicodeDecodeError:
        text = raw.decode("utf-8", errors="replace")
    chapters = _headings_from_markdown(text) if prefer_markdown or _looks_like_markdown(text) else []
    return ParsedSourceDocument(
        text=text,
        chapters=chapters,
        strategy="markdown_heading" if chapters else "linear_text",
        metadata={"parser": "text"},
    )


def _parse_docx(path: Path) -> ParsedSourceDocument:
    try:
        from docx import Document
    except Exception as exc:  # pragma: no cover - dependency guard
        raise RuntimeError("python-docx is required to parse DOCX source structure.") from exc
    # python-docx performs its own ZIP reads; preflight the package with the
    # shared archive limits before handing it to that parser.
    with SafeSourceArchive(path):
        pass
    document = Document(str(path))
    parts: list[str] = []
    chapters: list[DetectedChapter] = []
    offset = 0
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        style_name = (paragraph.style.name if paragraph.style is not None else "") or ""
        level = _docx_heading_level(style_name)
        if level:
            number = _number_from_title(text)
            chapters.append(
                DetectedChapter(
                    title=text,
                    number=number,
                    level=level,
                    source_locator=f"docx:paragraph:{len(parts)}",
                    start_offset=offset,
                    confidence=0.86,
                    verified=True,
                    metadata={"source": "docx_heading", "style": style_name},
                )
            )
        parts.append(text)
        offset += len(text) + 2
    full_text = "\n\n".join(parts)
    _close_chapter_ranges(chapters, len(full_text))
    return ParsedSourceDocument(
        text=full_text,
        chapters=chapters,
        strategy="docx_heading" if chapters else "linear_text",
        metadata={"parser": "docx"},
    )


def _parse_html_document(path: Path) -> ParsedSourceDocument:
    html_text = path.read_text(encoding="utf-8", errors="replace")
    text, chapters = _html_text_and_headings(html_text, source_name=path.name)
    for chapter in chapters:
        chapter.source_locator = f"html:{path.name}"
        chapter.verified = True
        chapter.confidence = 0.86
        chapter.metadata = {"source": "html_heading", "file": path.name}
    _close_chapter_ranges(chapters, len(text))
    return ParsedSourceDocument(
        text=text,
        chapters=chapters,
        strategy="markdown_heading" if chapters else "linear_text",
        metadata={"parser": "html"},
    )


def _parse_image(path: Path) -> ParsedSourceDocument:
    text = extract_image_text(path) or ""
    return ParsedSourceDocument(
        text=text,
        strategy="linear_text",
        warnings=[] if text else ["图片中没有识别到可索引文字。"],
        metadata={"parser": "vision_ocr", "ocr": True},
    )


def _parse_pptx(path: Path) -> ParsedSourceDocument:
    with SafeSourceArchive(path) as archive:
        slide_names = sorted(
            (name for name in archive.namelist() if re.fullmatch(r"ppt/slides/slide\d+\.xml", name)),
            key=lambda name: int(re.search(r"(\d+)", Path(name).stem).group(1)),
        )
        parts: list[str] = []
        chapters: list[DetectedChapter] = []
        pages: list[PageText] = []
        offset = 0
        for slide_no, name in enumerate(slide_names, start=1):
            root = parse_untrusted_xml(archive.read(name))
            texts = [str(node.text or "").strip() for node in root.iter() if node.tag.endswith("}t") and str(node.text or "").strip()]
            if not texts:
                continue
            slide_text = "\n".join(texts)
            prefix = f"\n\n[Slide {slide_no}]\n"
            start = offset
            parts.append(prefix + slide_text)
            offset += len(prefix) + len(slide_text)
            pages.append(
                PageText(
                    page_no=slide_no,
                    text=slide_text,
                    start_offset=start,
                    end_offset=offset,
                    content_start_offset=start + len(prefix),
                )
            )
            chapters.append(
                DetectedChapter(
                    title=texts[0],
                    level=1,
                    source_locator=f"pptx:slide:{slide_no}",
                    start_offset=start + len(prefix),
                    page_start=slide_no,
                    page_end=slide_no + 1,
                    confidence=0.9,
                    verified=True,
                    metadata={"source": "pptx_slide", "slide": slide_no},
                )
            )
    # Keep the exact assembled coordinate space. Trimming here would shift every
    # slide offset away from the text persisted in the index.
    full_text = "".join(parts)
    _close_chapter_ranges(chapters, len(full_text))
    return ParsedSourceDocument(
        text=full_text,
        chapters=chapters,
        pages=pages,
        strategy="linear_text",
        metadata={"parser": "pptx", "slide_count": len(pages)},
    )


def _parse_xlsx(path: Path) -> ParsedSourceDocument:
    with SafeSourceArchive(path) as archive:
        names = archive.namelist()
        shared: list[str] = []
        if "xl/sharedStrings.xml" in names:
            root = parse_untrusted_xml(archive.read("xl/sharedStrings.xml"))
            for item in root.iter():
                if item.tag.endswith("}si"):
                    shared.append("".join(str(node.text or "") for node in item.iter() if node.tag.endswith("}t")))
        sheet_names = sorted(
            (name for name in names if re.fullmatch(r"xl/worksheets/sheet\d+\.xml", name)),
            key=lambda name: int(re.search(r"(\d+)", Path(name).stem).group(1)),
        )
        parts: list[str] = []
        chapters: list[DetectedChapter] = []
        offset = 0
        for sheet_no, name in enumerate(sheet_names, start=1):
            root = parse_untrusted_xml(archive.read(name))
            rows: list[str] = []
            for row in (node for node in root.iter() if node.tag.endswith("}row")):
                values: list[str] = []
                for cell in (node for node in row if node.tag.endswith("}c")):
                    cell_type = cell.attrib.get("t", "")
                    value_node = next((node for node in cell.iter() if node.tag.endswith("}v")), None)
                    if cell_type == "inlineStr":
                        value = "".join(str(node.text or "") for node in cell.iter() if node.tag.endswith("}t"))
                    else:
                        value = str(value_node.text or "") if value_node is not None else ""
                        if cell_type == "s" and value.isdigit() and int(value) < len(shared):
                            value = shared[int(value)]
                    values.append(value.strip())
                if any(values):
                    rows.append("\t".join(values).rstrip())
            sheet_text = "\n".join(rows)
            if not sheet_text:
                continue
            label = f"Sheet {sheet_no}"
            prefix = f"\n\n[{label}]\n"
            start = offset
            parts.append(prefix + sheet_text)
            offset += len(prefix) + len(sheet_text)
            chapters.append(
                DetectedChapter(
                    title=label,
                    level=1,
                    source_locator=f"xlsx:sheet:{sheet_no}",
                    start_offset=start + len(prefix),
                    confidence=1.0,
                    verified=True,
                    metadata={"source": "xlsx_sheet", "sheet": sheet_no},
                )
            )
    # Keep the exact assembled coordinate space; chapter offsets above include
    # the synthetic sheet prefix.
    full_text = "".join(parts)
    _close_chapter_ranges(chapters, len(full_text))
    return ParsedSourceDocument(
        text=full_text,
        chapters=chapters,
        strategy="linear_text",
        metadata={"parser": "xlsx", "sheet_count": len(chapters)},
    )


def _assemble_pdf_pages(page_texts: list[tuple[int, str]]) -> tuple[str, list[PageText]]:
    parts: list[str] = []
    pages: list[PageText] = []
    offset = 0
    for page_no, text in page_texts:
        prefix = f"\n\n[Page {page_no}]\n"
        page_text = f"{prefix}{text}"
        start = offset
        parts.append(page_text)
        offset += len(page_text)
        pages.append(
            PageText(
                page_no=page_no,
                text=text,
                start_offset=start,
                end_offset=offset,
                content_start_offset=start + len(prefix),
            )
        )
    return "".join(parts), pages


def _parse_pdf(
    path: Path,
    *,
    progress_callback: SourceIndexProgressCallback | None = None,
) -> ParsedSourceDocument:
    try:
        from pypdf import PdfReader
    except Exception as exc:  # pragma: no cover - dependency guard
        raise RuntimeError("pypdf is required to parse PDF source structure.") from exc
    reader = PdfReader(str(path))
    pages: list[PageText] = []
    parts: list[str] = []
    offset = 0
    page_count = len(reader.pages)
    last_progress = -1
    for page_index, page in enumerate(reader.pages):
        text = (page.extract_text() or "").strip()
        page_text = f"\n\n[Page {page_index + 1}]\n{text}"
        start = offset
        parts.append(page_text)
        offset += len(page_text)
        pages.append(
            PageText(
                page_no=page_index + 1,
                text=text,
                start_offset=start,
                end_offset=offset,
                content_start_offset=start + len(f"\n\n[Page {page_index + 1}]\n"),
            )
        )
        page_progress = 25 + round(30 * (page_index + 1) / max(1, page_count))
        if page_progress != last_progress:
            _report_progress(progress_callback, "reading_pages", page_progress)
            last_progress = page_progress
    full_text = "".join(parts)
    native_meaningful_character_count = sum(
        meaningful_text_character_count(page.text) for page in pages
    )
    ocr_attempted = False
    ocr_used = False
    ocr_page_count = 0
    ocr_replaced_page_count = 0
    ocr_page_mapping_preserved = True
    warnings: list[str] = []
    if native_meaningful_character_count < max(40, len(pages) * 8) and pages:
        ocr_attempted = True
        ocr_page_limit = min(len(pages), 200)
        ocr_layouts = extract_pdf_pages_layout(
            path,
            page_start=1,
            page_end=len(pages),
            max_pages=ocr_page_limit,
        )
        if ocr_layouts:
            ocr_page_count = len(ocr_layouts)
            ocr_text_by_page = {
                layout.page_no: "\n".join(
                    line.text
                    for line in ordered_ocr_lines(layout.lines)
                    if line.text.strip()
                ).strip()
                for layout in ocr_layouts
            }
            merged_page_texts: list[tuple[int, str]] = []
            for page in pages:
                ocr_page_text = ocr_text_by_page.get(page.page_no, "")
                if meaningful_text_character_count(
                    ocr_page_text
                ) > meaningful_text_character_count(page.text):
                    merged_page_texts.append((page.page_no, ocr_page_text))
                    ocr_replaced_page_count += 1
                else:
                    merged_page_texts.append((page.page_no, page.text))
            if ocr_replaced_page_count:
                full_text, pages = _assemble_pdf_pages(merged_page_texts)
                ocr_used = True
            if len(reader.pages) > ocr_page_limit:
                warnings.append(
                    f"扫描 PDF 仅尝试前 {ocr_page_limit} 页的 OCR；"
                    "其余页面保留原文字层。"
                )
        else:
            ocr_text = extract_pdf_pages_text(
                path,
                page_start=1,
                page_end=len(pages),
                max_pages=ocr_page_limit,
            )
        if not ocr_layouts and ocr_text:
            fallback_ocr_text = ocr_text.strip()
            if native_meaningful_character_count:
                full_text = (
                    f"{full_text}\n\n[OCR text without page mapping]\n"
                    f"{fallback_ocr_text}"
                )
            else:
                full_text = fallback_ocr_text
                pages = []
            ocr_used = True
            ocr_page_count = ocr_page_limit
            ocr_page_mapping_preserved = False
            warnings.append(
                "扫描 PDF 已提取文字，但 OCR 未返回逐页坐标；"
                "目录节点不会据此冒充精确页段。"
            )
    outline_chapters = _pdf_outline_chapters(reader, pages, full_text)
    chapters = outline_chapters
    strategy = "pdf_outline" if outline_chapters else "linear_text"
    toc_metadata: dict[str, Any] = {}
    if outline_chapters:
        extraction = extract_pdf_toc(
            path,
            outline=[
                PdfOutlineAnchor(
                    title=chapter.title,
                    page_no=chapter.page_start,
                    level=chapter.level,
                    metadata=chapter.metadata,
                )
                for chapter in outline_chapters
                if chapter.page_start is not None
            ],
            page_count=len(pages),
        )
        warnings.extend(extraction.warnings)
        if extraction.nodes:
            toc_chapters = _chapters_from_pdf_toc(extraction.nodes, pages)
            chapters = _merge_pdf_navigation(outline_chapters, toc_chapters, extraction.nodes)
            strategy = "pdf_merged_toc"
            toc_metadata = {
                "toc_page_start": extraction.toc_page_start,
                "toc_page_end": extraction.toc_page_end,
                "printed_page_offset": extraction.printed_page_offset,
                "printed_page_mapping_support": extraction.mapping_support,
                "ocr_toc_node_count": len(extraction.nodes),
            }
    elif not outline_chapters:
        detected_toc_pages = _detected_pdf_toc_pages(pages)
        if detected_toc_pages:
            extraction = extract_pdf_toc_from_range(
                path,
                page_start=min(page.page_no for page in detected_toc_pages),
                page_end=max(page.page_no for page in detected_toc_pages),
            )
            warnings.extend(extraction.warnings)
            if extraction.nodes:
                printed_page_offset, mapping_support = _verify_pdf_toc_nodes(
                    extraction.nodes,
                    pages,
                )
                layout_chapters = [
                    chapter
                    for chapter in _chapters_from_pdf_toc(extraction.nodes, pages)
                    if chapter.verified
                ]
                if any(chapter.verified for chapter in layout_chapters):
                    chapters = layout_chapters
                    strategy = "pdf_layout_toc"
                    toc_metadata = {
                        "toc_page_start": extraction.toc_page_start,
                        "toc_page_end": extraction.toc_page_end,
                        "printed_page_offset": printed_page_offset,
                        "printed_page_mapping_support": mapping_support,
                        "ocr_toc_node_count": len(extraction.nodes),
                    }
    if not any(chapter.verified for chapter in chapters):
        toc_fallback = _pdf_toc_chapters(pages, full_text)
        if toc_fallback:
            chapters = toc_fallback
            strategy = "pdf_toc"
        elif not chapters:
            strategy = "linear_text"
    _close_pdf_navigation_ranges(chapters, pages, len(full_text))
    return ParsedSourceDocument(
        text=full_text,
        chapters=chapters,
        pages=pages,
        strategy=strategy,
        warnings=warnings,
        metadata={
            "parser": "pdf",
            "page_count": len(reader.pages),
            "ocr_attempted": ocr_attempted,
            "ocr": ocr_used,
            "ocr_page_count": ocr_page_count,
            "ocr_replaced_page_count": ocr_replaced_page_count,
            "ocr_page_mapping_preserved": ocr_page_mapping_preserved,
            **toc_metadata,
        },
    )


def _parse_epub(path: Path) -> ParsedSourceDocument:
    with SafeSourceArchive(path) as archive:
        names = archive.namelist()
        spine_items = _epub_spine_items(archive)
        html_names = spine_items or sorted(
            name for name in names if name.lower().endswith((".xhtml", ".html", ".htm"))
        )
        docs: list[tuple[str, str, list[DetectedChapter], dict[str, int]]] = []
        parts: list[str] = []
        offset_by_name: dict[str, int] = {}
        offset = 0
        heading_chapters: list[DetectedChapter] = []
        for name in html_names:
            try:
                raw = archive.read(name)
            except KeyError:
                continue
            text, headings, _visual_refs, anchors = _html_text_headings_and_visuals(
                raw.decode("utf-8", errors="replace"),
                source_name=name,
            )
            if not text.strip():
                continue
            prefix = f"\n\n[{name}]\n"
            offset_by_name[name] = offset + len(prefix)
            parts.append(prefix + text)
            for heading in headings:
                heading.start_offset = offset + len(prefix) + (heading.start_offset or 0)
                heading.source_locator = f"epub:{name}"
                heading.verified = True
                heading.confidence = 0.82
                heading.metadata = {"source": "epub_heading", "file": name}
                heading_chapters.append(heading)
            docs.append((name, text, headings, anchors))
            offset += len(prefix) + len(text)
        # EPUB navigation and fragment offsets use the assembled document
        # coordinate space, including each file prefix.
        full_text = "".join(parts)
        nav_items = _epub_navigation_items(archive, names)
        nav_chapters = _chapters_from_epub_nav(nav_items, docs, offset_by_name)
        if nav_chapters:
            _close_epub_navigation_ranges(nav_chapters, len(full_text))
            return ParsedSourceDocument(
                text=full_text,
                chapters=nav_chapters,
                strategy="epub_navigation",
                metadata={"parser": "epub", "navigation_items": len(nav_chapters)},
            )
        _close_chapter_ranges(heading_chapters, len(full_text))
        return ParsedSourceDocument(
            text=full_text,
            chapters=heading_chapters,
            strategy="epub_heading" if heading_chapters else "linear_text",
            warnings=[] if heading_chapters else ["EPUB 未发现可验证导航目录或标题结构。"],
            metadata={"parser": "epub"},
        )


def _epub_spine_items(archive: SafeSourceArchive) -> list[str]:
    try:
        container = parse_untrusted_xml(archive.read("META-INF/container.xml"))
    except Exception:
        return []
    rootfile = ""
    for element in container.iter():
        if element.tag.endswith("rootfile"):
            rootfile = element.attrib.get("full-path", "")
            break
    if not rootfile:
        return []
    base = str(Path(rootfile).parent)
    if base == ".":
        base = ""
    try:
        opf = parse_untrusted_xml(archive.read(rootfile))
    except Exception:
        return []
    manifest: dict[str, str] = {}
    spine_ids: list[str] = []
    for element in opf.iter():
        tag = element.tag.split("}")[-1]
        if tag == "item":
            item_id = element.attrib.get("id", "")
            href = element.attrib.get("href", "")
            if item_id and href:
                manifest[item_id] = f"{base}/{href}".lstrip("/")
        elif tag == "itemref":
            item_id = element.attrib.get("idref", "")
            if item_id:
                spine_ids.append(item_id)
    return [manifest[item_id] for item_id in spine_ids if item_id in manifest]


def _epub_navigation_items(
    archive: SafeSourceArchive,
    names: list[str],
) -> list[_EpubNavigationItem]:
    items: list[_EpubNavigationItem] = []
    nav_names = [name for name in names if re.search(r"(^|/)(nav|toc)\.(xhtml|html|htm)$", name, re.I)]
    for name in nav_names:
        try:
            text = archive.read(name).decode("utf-8", errors="replace")
        except Exception:
            continue
        parser = _NavLinkParser()
        parser.feed(text)
        base = str(Path(name).parent)
        if base == ".":
            base = ""
        for href, label, level in parser.links:
            if not label.strip() or not href.strip():
                continue
            target, fragment = _epub_navigation_target(base, href)
            if target:
                items.append(
                    _EpubNavigationItem(
                        target=target,
                        fragment=fragment,
                        label=_clean_label(label),
                        order=len(items),
                        level=level,
                    )
                )
    if items:
        return _dedupe_nav_items(items)
    for name in [entry for entry in names if entry.lower().endswith(".ncx")]:
        try:
            root = parse_untrusted_xml(archive.read(name))
        except Exception:
            continue
        base = str(Path(name).parent)
        if base == ".":
            base = ""
        nav_map = next(
            (element for element in root.iter() if element.tag.split("}")[-1] == "navMap"),
            None,
        )
        if nav_map is None:
            continue

        def visit(parent: Any, level: int = 1) -> None:
            for point in parent:
                if point.tag.split("}")[-1] != "navPoint":
                    continue
                label = ""
                src = ""
                for child in point:
                    child_tag = child.tag.split("}")[-1]
                    if child_tag == "navLabel":
                        label_node = next(
                            (
                                descendant
                                for descendant in child.iter()
                                if descendant.tag.split("}")[-1] == "text"
                                and descendant.text
                            ),
                            None,
                        )
                        label = label_node.text if label_node is not None else ""
                    elif child_tag == "content":
                        src = child.attrib.get("src", "")
                if label and src:
                    target, fragment = _epub_navigation_target(base, src)
                    if target:
                        items.append(
                            _EpubNavigationItem(
                                target=target,
                                fragment=fragment,
                                label=_clean_label(label),
                                order=len(items),
                                level=level,
                            )
                        )
                visit(point, level + 1)

        visit(nav_map)
    return _dedupe_nav_items(items)


def _chapters_from_epub_nav(
    nav_items: list[_EpubNavigationItem],
    docs: list[tuple[str, str, list[DetectedChapter], dict[str, int]]],
    offset_by_name: dict[str, int],
) -> list[DetectedChapter]:
    docs_by_name = {name: text for name, text, _headings, _anchors in docs}
    anchors_by_name = {name: anchors for name, _text, _headings, anchors in docs}
    chapters: list[DetectedChapter] = []
    for item in nav_items:
        text = docs_by_name.get(item.target, "")
        base_offset = offset_by_name.get(item.target)
        local_index = _find_title_offset(text, item.label)
        fragment_offset = anchors_by_name.get(item.target, {}).get(item.fragment)
        start_offset = (
            base_offset + fragment_offset
            if base_offset is not None and fragment_offset is not None
            else base_offset + local_index
            if base_offset is not None and local_index >= 0
            else None
        )
        if fragment_offset is not None:
            confidence = 0.98
            anchor_source = "epub_fragment"
        elif local_index >= 0:
            confidence = 0.95
            anchor_source = "body_title_match"
        elif base_offset is not None:
            confidence = 0.62
            anchor_source = "target_file_only"
        else:
            confidence = 0.5
            anchor_source = "target_missing"
        number = _number_from_title(item.label)
        numbered_level = len(number.split(".")) if number else 1
        chapters.append(
            DetectedChapter(
                title=item.label,
                number=number,
                level=max(1, item.level, numbered_level),
                source_locator=item.source_locator,
                start_offset=start_offset,
                confidence=confidence,
                verified=start_offset is not None,
                metadata={
                    "source": "epub_navigation",
                    "file": item.target,
                    "fragment": item.fragment,
                    "nav_order": item.order,
                    "navigation_level": item.level,
                    "direct_anchor_verified": start_offset is not None,
                    "anchor_source": anchor_source,
                },
            )
        )
    return chapters


def _epub_navigation_target(base: str, href: str) -> tuple[str, str]:
    target_path, separator, fragment = html.unescape(href).partition("#")
    decoded_path = unquote(target_path.strip())
    if not decoded_path:
        return ("", fragment.strip() if separator else "")
    target = posixpath.normpath(posixpath.join(base, decoded_path)).lstrip("/")
    if target == ".." or target.startswith("../"):
        return ("", "")
    return (target, unquote(fragment.strip()) if separator else "")


@dataclass
class _HTMLVisualRef:
    src: str
    alt: str = ""
    title: str = ""
    caption: str = ""
    text_offset: int = 0
    paragraph_index: int | None = None


class _TextHeadingParser(HTMLParser):
    def __init__(self) -> None:
        super().__init__()
        self.parts: list[str] = []
        self.headings: list[DetectedChapter] = []
        self._heading_tag: str | None = None
        self._heading_text: list[str] = []
        self._heading_start = 0
        self.visuals: list[_HTMLVisualRef] = []
        self._figure_visual_start: int | None = None
        self._figcaption_text: list[str] | None = None
        self._paragraph_index = 0
        self.anchors: dict[str, int] = {}

    @property
    def text_offset(self) -> int:
        return sum(len(part) for part in self.parts)

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        attrs_dict = {key: value or "" for key, value in attrs}
        for anchor in (attrs_dict.get("id", ""), attrs_dict.get("name", "")):
            if anchor.strip() and anchor.strip() not in self.anchors:
                self.anchors[anchor.strip()] = self.text_offset
        if tag in {"p", "div", "section", "article", "br", "li", "tr"}:
            self.parts.append("\n")
        if tag in {"p", "div", "section", "article", "li"}:
            self._paragraph_index += 1
        if tag == "figure":
            self._figure_visual_start = len(self.visuals)
        if tag == "figcaption":
            self._figcaption_text = []
        if tag == "img":
            src = attrs_dict.get("src", "").strip()
            if src:
                self.visuals.append(
                    _HTMLVisualRef(
                        src=src,
                        alt=attrs_dict.get("alt", "").strip(),
                        title=attrs_dict.get("title", "").strip(),
                        text_offset=self.text_offset,
                        paragraph_index=self._paragraph_index,
                    )
                )
        if re.fullmatch(r"h[1-6]", tag):
            self.parts.append("\n\n")
            self._heading_tag = tag
            self._heading_text = []
            self._heading_start = self.text_offset

    def handle_endtag(self, tag: str) -> None:
        if tag == "figcaption" and self._figcaption_text is not None:
            caption = _clean_label(" ".join(self._figcaption_text))
            start = self._figure_visual_start if self._figure_visual_start is not None else len(self.visuals)
            for visual in self.visuals[start:]:
                visual.caption = caption
            self._figcaption_text = None
        if tag == "figure":
            self._figure_visual_start = None
        if tag == self._heading_tag:
            text = _clean_label(" ".join(self._heading_text))
            if text:
                level = int(tag[1])
                self.headings.append(
                    DetectedChapter(
                        title=text,
                        number=_number_from_title(text),
                        level=level,
                        start_offset=self._heading_start,
                    )
                )
            self._heading_tag = None
            self._heading_text = []
            self.parts.append("\n")
        elif tag in {"p", "div", "section", "article", "li", "tr"}:
            self.parts.append("\n")

    def handle_data(self, data: str) -> None:
        text = html.unescape(data)
        if not text.strip():
            return
        if self._heading_tag:
            self._heading_text.append(text.strip())
        if self._figcaption_text is not None:
            self._figcaption_text.append(text.strip())
        self.parts.append(text.strip() + " ")


class _NavLinkParser(HTMLParser):
    def __init__(self) -> None:
        super().__init__()
        self.links: list[tuple[str, str, int]] = []
        self._href: str | None = None
        self._text: list[str] = []
        self._list_depth = 0
        self._link_level = 1

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        if tag in {"ol", "ul"}:
            self._list_depth += 1
        if tag != "a":
            return
        attrs_dict = {key: value or "" for key, value in attrs}
        self._href = attrs_dict.get("href", "")
        self._text = []
        self._link_level = max(1, self._list_depth)

    def handle_endtag(self, tag: str) -> None:
        if tag == "a" and self._href:
            self.links.append((self._href, " ".join(self._text), self._link_level))
        if tag == "a":
            self._href = None
            self._text = []
        if tag in {"ol", "ul"}:
            self._list_depth = max(0, self._list_depth - 1)

    def handle_data(self, data: str) -> None:
        if self._href is not None:
            self._text.append(data.strip())


def _html_text_and_headings(content: str, *, source_name: str) -> tuple[str, list[DetectedChapter]]:
    text, headings, _visuals, _anchors = _html_text_headings_and_visuals(
        content,
        source_name=source_name,
    )
    return text, headings


def _html_text_headings_and_visuals(
    content: str,
    *,
    source_name: str,
) -> tuple[str, list[DetectedChapter], list[_HTMLVisualRef], dict[str, int]]:
    parser = _TextHeadingParser()
    parser.feed(content)
    raw_text = "".join(parser.parts)
    text = _normalize_text(raw_text)

    def normalized_offset(raw_offset: int) -> int:
        return len(_normalize_text(raw_text[: max(0, raw_offset)]))

    for heading in parser.headings:
        heading.source_locator = f"html:{source_name}"
        heading.start_offset = normalized_offset(heading.start_offset or 0)
    for visual in parser.visuals:
        visual.text_offset = normalized_offset(visual.text_offset)
    anchors = {
        anchor: normalized_offset(raw_offset)
        for anchor, raw_offset in parser.anchors.items()
    }
    return text, parser.headings, parser.visuals, anchors


def _headings_from_markdown(text: str) -> list[DetectedChapter]:
    chapters: list[DetectedChapter] = []
    for match in re.finditer(r"^(#{1,6})\s+(.+?)\s*$", text, flags=re.M):
        title = _clean_label(match.group(2))
        if not title:
            continue
        chapters.append(
            DetectedChapter(
                title=title,
                number=_number_from_title(title),
                level=len(match.group(1)),
                source_locator=f"markdown:line:{text[: match.start()].count(chr(10)) + 1}",
                start_offset=match.start(),
                confidence=0.88,
                verified=True,
                metadata={"source": "markdown_heading"},
            )
        )
    _close_chapter_ranges(chapters, len(text))
    return chapters


def _pdf_outline_chapters(reader: Any, pages: list[PageText], full_text: str) -> list[DetectedChapter]:
    try:
        outline = reader.outline
    except Exception:
        return []
    flattened: list[tuple[Any, int]] = []

    def visit(items: Any, level: int = 1) -> None:
        if not isinstance(items, list):
            flattened.append((items, level))
            return
        for item in items:
            if isinstance(item, list):
                visit(item, level + 1)
            else:
                flattened.append((item, level))

    visit(outline)
    chapters: list[DetectedChapter] = []
    for item, level in flattened:
        title = _clean_label(getattr(item, "title", "") or str(item))
        if not title:
            continue
        try:
            page_index = reader.get_destination_page_number(item)
        except Exception:
            page_index = None
        page = pages[page_index] if isinstance(page_index, int) and 0 <= page_index < len(pages) else None
        # An outline title often also appears in the printed table of contents.
        # When the PDF destination cannot be resolved, a global title search can
        # therefore turn a TOC occurrence into a false body anchor with a huge
        # chapter range. Preserve the navigation node, but do not verify it until
        # a physical page destination is available.
        verified = page is not None
        start_offset = None
        title_match_verified = False
        if page is not None:
            local_title_offset = _find_title_offset(page.text, title)
            title_match_verified = local_title_offset >= 0
            page_content_offset = _page_content_start_offset(page, full_text=full_text)
            start_offset = (
                page_content_offset + local_title_offset
                if local_title_offset >= 0
                else page_content_offset
            )
        chapters.append(
            DetectedChapter(
                title=title,
                number=_number_from_title(title),
                level=level,
                source_locator=f"pdf:outline:{page.page_no if page else ''}",
                start_offset=start_offset,
                page_start=page.page_no if page else None,
                confidence=0.93 if title_match_verified else 0.76 if verified else 0.55,
                verified=verified,
                metadata={
                    "source": "pdf_outline",
                    "verification": "destination_page" if verified else "destination_unresolved",
                    "anchor_source": (
                        "pdf_destination_title"
                        if title_match_verified
                        else "pdf_destination_page" if verified else "destination_unresolved"
                    ),
                    "title_match_verified": title_match_verified,
                },
            )
        )
    return chapters


def _chapters_from_pdf_toc(nodes: list[PdfTocNode], pages: list[PageText]) -> list[DetectedChapter]:
    chapters: list[DetectedChapter] = []
    for node in nodes:
        page = (
            pages[node.physical_page - 1]
            if node.physical_page is not None and 1 <= node.physical_page <= len(pages)
            else None
        )
        verified = node.verified and page is not None
        local_offset = _find_title_offset(page.text, node.title) if page is not None else -1
        chapters.append(
            DetectedChapter(
                title=node.title,
                number=node.number or _number_from_title(node.title),
                level=node.level,
                source_locator=f"pdf:toc-page:{node.toc_page}:printed:{node.printed_page}",
                start_offset=(
                    (page.content_start_offset if page.content_start_offset is not None else page.start_offset)
                    + local_offset
                    if verified and page and local_offset >= 0
                    else (
                        page.content_start_offset
                        if verified and page and page.content_start_offset is not None
                        else page.start_offset if verified and page else None
                    )
                ),
                page_start=page.page_no if verified and page else None,
                confidence=node.confidence,
                verified=verified,
                metadata={
                    **node.metadata,
                    "anchor_source": (
                        "pdf_toc_body_title"
                        if local_offset >= 0
                        else "pdf_toc_page_mapping" if verified else "toc_candidate"
                    ),
                    "title_match_verified": local_offset >= 0,
                },
            )
        )
    return chapters


def _verify_pdf_toc_nodes(nodes: list[PdfTocNode], pages: list[PageText]) -> tuple[int | None, int]:
    if not nodes:
        return None, 0
    for index, node in enumerate(nodes):
        if node.printed_page > 0 or index == 0:
            continue
        parent = nodes[index - 1]
        if parent.printed_page > 0 and node.level > parent.level:
            node.printed_page = parent.printed_page
            node.metadata = {**node.metadata, "printed_page_inferred": True}
    last_toc_page = max(node.toc_page for node in nodes)
    body_pages = [page for page in pages if page.page_no > last_toc_page]
    direct_matches: dict[int, PageText] = {}
    offset_votes: Counter[int] = Counter()
    for index, node in enumerate(nodes):
        for page in body_pages:
            if _find_title_offset(page.text, node.title) < 0:
                continue
            direct_matches[index] = page
            if node.printed_page > 0:
                offset_votes[page.page_no - node.printed_page] += 1
            break
    printed_page_offset: int | None = None
    mapping_support = 0
    if offset_votes:
        printed_page_offset, mapping_support = offset_votes.most_common(1)[0]
        if mapping_support < 2:
            printed_page_offset = None
    pages_by_number = {page.page_no: page for page in body_pages}
    for index, node in enumerate(nodes):
        direct_page = direct_matches.get(index)
        if node.printed_page <= 0 and direct_page is not None and printed_page_offset is not None:
            node.printed_page = max(1, direct_page.page_no - printed_page_offset)
            node.metadata = {**node.metadata, "printed_page_inferred": True}
        mapped_page = (
            pages_by_number.get(node.printed_page + printed_page_offset)
            if printed_page_offset is not None and node.printed_page > 0
            else None
        )
        page = mapped_page or direct_page
        if page is None:
            continue
        canonical_title = _canonical_structural_title_from_body(page.text, node)
        if canonical_title and canonical_title != node.title:
            node.metadata = {
                **node.metadata,
                "ocr_title": node.title,
                "title_canonicalized_from_body": True,
            }
            node.title = canonical_title
        node.physical_page = page.page_no
        node.verified = True
        node.confidence = 0.9 if mapped_page is not None else 0.82
        node.metadata = {
            **node.metadata,
            "verification": (
                "verified_printed_page_mapping" if mapped_page is not None else "body_title_match"
            ),
            "printed_page_offset": printed_page_offset,
            "printed_page_mapping_support": mapping_support,
        }
    return printed_page_offset, mapping_support


def _canonical_structural_title_from_body(text: str, node: PdfTocNode) -> str | None:
    node_marker = parse_structural_heading(node.title)
    if node_marker is None:
        return None
    candidates: list[str] = []
    for raw_line in text.splitlines():
        line = _clean_label(raw_line)
        if not line or len(line) > 180:
            continue
        marker = parse_structural_heading(line)
        if marker is None:
            continue
        if marker.kind != node_marker.kind or marker.level != node_marker.level:
            continue
        if marker.number and node_marker.number and marker.number.lower() != node_marker.number.lower():
            continue
        candidates.append(line)
    if not candidates:
        return None
    normalized_node = _normalize_for_match(node.title)
    exact = next(
        (candidate for candidate in candidates if _normalize_for_match(candidate) == normalized_node),
        None,
    )
    if exact is not None:
        return exact
    return max(candidates, key=len)


def _merge_pdf_navigation(
    outline: list[DetectedChapter],
    toc_chapters: list[DetectedChapter],
    toc_nodes: list[PdfTocNode],
) -> list[DetectedChapter]:
    matched_outline_pairs = {
        (str(node.metadata.get("outline_title") or ""), int(node.metadata.get("outline_page") or 0))
        for node in toc_nodes
        if node.metadata.get("outline_title") and node.metadata.get("outline_page")
    }
    matched_root_pages = {
        chapter.page_start
        for chapter in toc_chapters
        if chapter.level == 1 and chapter.page_start is not None
    }
    unmatched_outline = [
        chapter
        for chapter in outline
        if (chapter.title, chapter.page_start or 0) not in matched_outline_pairs
        and not (chapter.level == 1 and chapter.page_start in matched_root_pages)
    ]
    first_body_page = min(
        (chapter.page_start for chapter in toc_chapters if chapter.page_start is not None),
        default=max((node.toc_page for node in toc_nodes), default=0) + 1,
    )
    prefix = [chapter for chapter in unmatched_outline if (chapter.page_start or 0) < first_body_page]
    suffix = [chapter for chapter in unmatched_outline if (chapter.page_start or 0) >= first_body_page]
    return prefix + toc_chapters + suffix


def _close_pdf_navigation_ranges(
    chapters: list[DetectedChapter],
    pages: list[PageText],
    text_length: int,
) -> None:
    for index, chapter in enumerate(chapters):
        if not chapter.verified or chapter.start_offset is None or chapter.page_start is None:
            continue
        boundary = next(
            (
                candidate
                for candidate in chapters[index + 1 :]
                if candidate.verified
                and candidate.start_offset is not None
                and candidate.page_start is not None
                and candidate.level <= chapter.level
            ),
            None,
        )
        chapter.end_offset = boundary.start_offset if boundary else text_length
        if boundary and boundary.page_start is not None:
            chapter.page_end = max(chapter.page_start + 1, boundary.page_start)
        else:
            chapter.page_end = pages[-1].page_no + 1 if pages else chapter.page_start + 1


def _pdf_toc_chapters(pages: list[PageText], full_text: str) -> list[DetectedChapter]:
    toc_pages = _detected_pdf_toc_pages(pages)
    if not toc_pages:
        return []
    body_pages = [page for page in pages if page.page_no > max(toc.page_no for toc in toc_pages)]
    candidates: list[tuple[str, int, PageText]] = []
    seen: set[tuple[str, int]] = set()
    for toc_page in toc_pages[:6]:
        for line in toc_page.text.splitlines():
            parsed = _parse_toc_line(line)
            if not parsed:
                continue
            title, printed_page = parsed
            key = (_normalize_for_match(title), printed_page)
            if key in seen:
                continue
            seen.add(key)
            candidates.append((title, printed_page, toc_page))

    direct_matches: dict[tuple[str, int], tuple[PageText, int]] = {}
    offset_votes: Counter[int] = Counter()
    for title, printed_page, _toc_page in candidates:
        for page in body_pages:
            local_offset = _find_title_offset(page.text, title)
            if local_offset < 0:
                continue
            direct_matches[(_normalize_for_match(title), printed_page)] = (page, local_offset)
            offset_votes[page.page_no - printed_page] += 1
            break

    printed_page_offset: int | None = None
    mapping_support = 0
    if offset_votes:
        printed_page_offset, mapping_support = offset_votes.most_common(1)[0]
        if mapping_support < 2:
            printed_page_offset = None

    pages_by_number = {page.page_no: page for page in body_pages}
    chapters: list[DetectedChapter] = []
    for title, printed_page, toc_page in candidates:
        direct_match = direct_matches.get((_normalize_for_match(title), printed_page))
        mapped_page = (
            pages_by_number.get(printed_page + printed_page_offset)
            if printed_page_offset is not None
            else None
        )
        matched_page = mapped_page or (direct_match[0] if direct_match else None)
        local_offset = _find_title_offset(matched_page.text, title) if matched_page else -1
        if local_offset < 0 and direct_match and matched_page is direct_match[0]:
            local_offset = direct_match[1]
        title_offset = (
            (
                matched_page.content_start_offset
                if matched_page.content_start_offset is not None
                else matched_page.start_offset
            )
            + local_offset
            if matched_page and local_offset >= 0
            else (
                matched_page.content_start_offset
                if matched_page and matched_page.content_start_offset is not None
                else matched_page.start_offset if matched_page else None
            )
        )
        verified_by_mapping = mapped_page is not None and printed_page_offset is not None
        verified = verified_by_mapping or direct_match is not None
        chapters.append(
            DetectedChapter(
                title=title,
                number=_number_from_title(title),
                level=_toc_title_level(title),
                source_locator=f"pdf:toc-page:{toc_page.page_no}:printed:{printed_page}",
                start_offset=title_offset,
                page_start=matched_page.page_no if matched_page else None,
                confidence=0.9 if verified_by_mapping else 0.82 if verified else 0.62,
                verified=verified,
                metadata={
                    "source": "pdf_toc",
                    "printed_page": printed_page,
                    "verification": (
                        "verified_printed_page_mapping"
                        if verified_by_mapping
                        else "body_title_match" if direct_match else "toc_candidate"
                    ),
                    "printed_page_offset": printed_page_offset,
                    "printed_page_mapping_support": mapping_support,
                    "anchor_source": (
                        "pdf_toc_body_title"
                        if local_offset >= 0
                        else "pdf_toc_page_mapping" if verified else "toc_candidate"
                    ),
                    "title_match_verified": local_offset >= 0,
                },
            )
        )
    return chapters


def _parse_toc_line(line: str) -> tuple[str, int] | None:
    raw = normalize_toc_text(html.unescape(line or "")).strip()
    if len(raw) < 4 or len(raw) > 240:
        return None
    page_match = re.search(
        r"[（(\[]?\s*(\d{1,4}|[IVXLCDM]{1,8})\s*[）)\]]?\s*([.．。·•⋯…\s]*)$",
        raw,
        flags=re.I,
    )
    if not page_match:
        return None
    prefix = raw[: page_match.start()]
    leader_match = re.search(r"(?:[.．。·•⋯…]|\s){2,}$", prefix)
    trailing_leader = page_match.group(2)
    has_leader_before = bool(
        leader_match and re.search(r"[.．。·•⋯…]", leader_match.group(0))
    )
    has_leader_after = bool(re.search(r"[.．。·•⋯…]", trailing_leader))
    if not has_leader_before and not has_leader_after:
        return None
    title = _clean_label(prefix[: leader_match.start()] if leader_match else prefix)
    if not title or re.fullmatch(r"\d+(?:\.\d+)*", title):
        return None
    page_token = page_match.group(1)
    printed_page = int(page_token) if page_token.isdigit() else _roman_numeral_value(page_token)
    return (title, printed_page) if printed_page > 0 else None


def _toc_title_level(title: str) -> int:
    cleaned = _clean_label(title)
    structural_marker = parse_structural_heading(cleaned)
    if structural_marker is not None:
        return structural_marker.level
    number = _number_from_title(cleaned)
    if number and "." in number:
        return min(6, len(number.split(".")))
    if re.match(r"^第\s*[0-9一二三四五六七八九十百千零〇两]+\s*节", cleaned):
        return 2
    return 1


def _roman_numeral_value(value: str) -> int:
    values = {"I": 1, "V": 5, "X": 10, "L": 50, "C": 100, "D": 500, "M": 1000}
    total = 0
    previous = 0
    for character in reversed(value.upper()):
        current = values.get(character, 0)
        total += -current if current < previous else current
        previous = max(previous, current)
    return total


def _looks_like_toc_page(text: str) -> bool:
    lines = [line for line in text.splitlines() if line.strip()]
    toc_heading = any(is_toc_heading(line) for line in lines[:12])
    parsed_rows = sum(1 for line in lines if _parse_toc_line(line))
    structural_rows = sum(1 for line in lines if parse_structural_heading(line) is not None)
    return toc_heading or parsed_rows >= 3 or (structural_rows >= 4 and parsed_rows >= 2)


def _detected_pdf_toc_pages(pages: list[PageText]) -> list[PageText]:
    candidates = pages[:30]
    if not candidates:
        return []

    explicit_start = next(
        (
            index
            for index, page in enumerate(candidates)
            if any(
                is_toc_heading(line)
                for line in [line for line in page.text.splitlines() if line.strip()][:12]
            )
        ),
        None,
    )
    if explicit_start is None:
        explicit_start = next(
            (index for index, page in enumerate(candidates) if _looks_like_toc_page(page.text)),
            None,
        )
    if explicit_start is None:
        return []

    detected = [candidates[explicit_start]]
    for page in candidates[explicit_start + 1 :]:
        lines = [line for line in page.text.splitlines() if line.strip()]
        parsed_rows = sum(1 for line in lines if _parse_toc_line(line))
        structural_rows = sum(1 for line in lines if parse_structural_heading(line) is not None)
        if parsed_rows < 3 or structural_rows < 2:
            break
        detected.append(page)
    return detected


def _docx_heading_level(style_name: str) -> int | None:
    match = re.search(r"(?:heading|标题)\s*(\d+)", style_name, flags=re.I)
    if not match:
        return None
    return max(1, min(6, int(match.group(1))))


def _looks_like_markdown(text: str) -> bool:
    return bool(re.search(r"^#{1,6}\s+\S", text, flags=re.M))


def _close_chapter_ranges(chapters: list[DetectedChapter], text_length: int) -> None:
    ordered_chapters = sorted(
        chapters,
        key=lambda chapter: chapter.start_offset if chapter.start_offset is not None else text_length,
    )
    for index, chapter in enumerate(ordered_chapters):
        if chapter.start_offset is None:
            continue
        next_chapter = ordered_chapters[index + 1] if index + 1 < len(ordered_chapters) else None
        if chapter.end_offset is None:
            chapter.end_offset = next_chapter.start_offset if next_chapter else text_length
        if chapter.page_start is not None and chapter.page_end is None:
            next_page_start = (
                next_chapter.page_start
                if next_chapter and next_chapter.page_start is not None
                else None
            )
            chapter.page_end = max(
                chapter.page_start + 1,
                next_page_start or chapter.page_start + 1,
            )


def _close_epub_navigation_ranges(
    chapters: list[DetectedChapter],
    text_length: int,
) -> None:
    for index, chapter in enumerate(chapters):
        subtree_end = next(
            (
                candidate_index
                for candidate_index in range(index + 1, len(chapters))
                if chapters[candidate_index].level <= chapter.level
            ),
            len(chapters),
        )
        descendants = chapters[index + 1 : subtree_end]
        first_descendant_anchor = next(
            (
                candidate
                for candidate in descendants
                if candidate.verified and candidate.start_offset is not None
            ),
            None,
        )
        if chapter.start_offset is None and first_descendant_anchor is not None:
            chapter.start_offset = first_descendant_anchor.start_offset
            chapter.verified = True
            chapter.confidence = max(
                chapter.confidence,
                min(0.9, first_descendant_anchor.confidence),
            )
            chapter.metadata = {
                **chapter.metadata,
                "anchor_source": "first_verified_descendant",
            }
        if chapter.start_offset is None:
            continue
        next_scope_anchor = next(
            (
                candidate.start_offset
                for candidate in chapters[subtree_end:]
                if candidate.verified and candidate.start_offset is not None
            ),
            None,
        )
        chapter.end_offset = max(
            chapter.start_offset,
            next_scope_anchor if next_scope_anchor is not None else text_length,
        )
        if descendants:
            chapter.metadata = {
                **chapter.metadata,
                "navigation_container": True,
                "range_source": "navigation_subtree",
            }


def _chapter_for_chunk(
    start: int,
    end: int,
    chapter_ranges: list[tuple[str, int, int, int | None, int | None, int]],
) -> tuple[str | None, int | None, int | None]:
    candidates: list[
        tuple[bool, int, int, int, str, int | None, int | None]
    ] = []
    midpoint = start + max(0, end - start - 1) // 2
    for chapter_id, chapter_start, chapter_end, page_start, page_end, level in chapter_ranges:
        overlap = max(0, min(end, chapter_end) - max(start, chapter_start))
        if overlap <= 0:
            continue
        contains_midpoint = chapter_start <= midpoint < chapter_end
        candidates.append(
            (
                contains_midpoint,
                chapter_start if contains_midpoint else overlap,
                level,
                overlap,
                chapter_id,
                page_start,
                page_end,
            )
        )
    if not candidates:
        return (None, None, None)
    # Prefer the most recently started range containing the chunk midpoint.
    # This selects the nearest nested section instead of an older, overly broad
    # range. If no range contains the midpoint, fall back to greatest overlap.
    best = max(candidates, key=lambda candidate: candidate[:4])
    return (best[4], best[5], best[6])


def _page_range_for_offsets(
    pages: list[PageText],
    start: int,
    end: int,
) -> tuple[int | None, int | None]:
    overlapping = [
        page
        for page in pages
        if page.end_offset > start and page.start_offset < end
    ]
    if not overlapping:
        return (None, None)
    return (overlapping[0].page_no, overlapping[-1].page_no + 1)


def source_structure_needs_upgrade(structure: SourceStructure) -> bool:
    try:
        version = int(structure.metadata.get("structure_index_version") or 0)
    except (TypeError, ValueError):
        version = 0
    return version < CURRENT_SOURCE_STRUCTURE_INDEX_VERSION


def _reanchor_existing_visuals(
    visuals: list[SourceVisualAsset],
    *,
    structure: SourceStructure,
    chapters: list[SourceChapter],
    chunks: list[SourceChunk],
    previous_chapters: list[SourceChapter] | None = None,
) -> list[SourceVisualAsset]:
    verified_chapters = [
        chapter for chapter in chapters if chapter.anchor_status == "verified"
    ]
    chapter_by_id = {chapter.id: chapter for chapter in verified_chapters}
    previous_chapter_by_id = {
        chapter.id: chapter for chapter in (previous_chapters or [])
    }
    ordered_chunks = sorted(chunks, key=lambda chunk: chunk.order_index)
    reanchored: list[SourceVisualAsset] = []
    for visual in visuals:
        chapter = chapter_by_id.get(visual.chapter_id or "")
        reanchor_method = "chapter_id" if chapter is not None else ""
        old_chapter = previous_chapter_by_id.get(visual.chapter_id or "")
        if chapter is None and old_chapter is not None:
            semantic_signature = _chapter_semantic_signature(old_chapter)
            semantic_candidates = [
                candidate
                for candidate in verified_chapters
                if _chapter_semantic_signature(candidate) == semantic_signature
            ]
            if len(semantic_candidates) == 1:
                chapter = semantic_candidates[0]
                reanchor_method = "semantic_chapter"
            elif len(semantic_candidates) > 1:
                old_semantic_siblings = sorted(
                    (
                        candidate
                        for candidate in (previous_chapters or [])
                        if _chapter_semantic_signature(candidate)
                        == semantic_signature
                    ),
                    key=lambda candidate: candidate.order_index,
                )
                old_occurrence = next(
                    (
                        index
                        for index, candidate in enumerate(old_semantic_siblings)
                        if candidate.id == old_chapter.id
                    ),
                    -1,
                )
                ordered_candidates = sorted(
                    semantic_candidates,
                    key=lambda candidate: candidate.order_index,
                )
                if 0 <= old_occurrence < len(ordered_candidates):
                    chapter = ordered_candidates[old_occurrence]
                    reanchor_method = "semantic_occurrence"
        if chapter is None and visual.source_locator:
            locator_candidates = [
                candidate
                for candidate in verified_chapters
                if _source_locator_prefix_matches(
                    visual.source_locator,
                    candidate.source_locator,
                )
            ]
            if len(locator_candidates) == 1:
                chapter = locator_candidates[0]
                reanchor_method = "source_locator"
        page_no = visual.slide_no or visual.page_start
        if chapter is None and page_no is not None:
            page_chapters = [
                candidate
                for candidate in verified_chapters
                if candidate.page_start is not None
                and candidate.page_start <= page_no
                and (candidate.page_end or candidate.page_start + 1) > page_no
            ]
            if page_chapters:
                chapter = max(
                    page_chapters,
                    key=lambda candidate: (candidate.level, candidate.order_index),
                )
                reanchor_method = "physical_page"
        page_chunks = []
        if page_no is not None:
            page_chunks = [
                chunk
                for chunk in ordered_chunks
                if chunk.page_start is not None
                and chunk.page_start <= page_no
                and (chunk.page_end or chunk.page_start + 1) > page_no
            ]
        if chapter is not None:
            chapter_chunks = [
                chunk for chunk in ordered_chunks if chunk.chapter_id == chapter.id
            ]
            if chapter_chunks:
                chapter_page_chunks = [
                    chunk for chunk in page_chunks if chunk.chapter_id == chapter.id
                ]
                page_chunks = chapter_page_chunks or chapter_chunks
        before_chunk_id = page_chunks[0].id if page_chunks else None
        after_chunk_id = page_chunks[-1].id if page_chunks else None
        surrounding_text = "\n\n".join(
            dict.fromkeys(chunk.text.strip() for chunk in page_chunks if chunk.text.strip())
        )[:2000]
        has_grounded_anchor = bool(
            visual.metadata.get("standalone_image")
            or chapter is not None
            or page_chunks
        )
        reanchored.append(
            visual.model_copy(
                update={
                    "structure_id": structure.id,
                    "chapter_id": chapter.id if chapter else None,
                    "before_chunk_id": before_chunk_id,
                    "after_chunk_id": after_chunk_id,
                    "surrounding_text": surrounding_text,
                    "anchor_status": (
                        "verified"
                        if visual.anchor_status == "verified" and has_grounded_anchor
                        else "unverified"
                    ),
                    "metadata": {
                        **visual.metadata,
                        "reanchored_after_structure_upgrade": True,
                        "structure_reanchor_method": reanchor_method or "unresolved",
                    },
                }
            )
        )
    return reanchored


def _chapter_semantic_signature(
    chapter: SourceChapter,
) -> tuple[tuple[str, ...], str, str, int]:
    path = tuple(
        _normalize_for_match(part) for part in chapter.path if part.strip()
    )
    return (
        path,
        _normalize_chapter_number(chapter.normalized_number or chapter.number),
        _normalize_for_match(chapter.title),
        chapter.level,
    )


def _source_locator_prefix_matches(first: str, second: str) -> bool:
    return bool(
        first
        and second
        and (
            first == second
            or first.startswith(f"{second}:")
            or second.startswith(f"{first}:")
        )
    )


def _locator_for_offset(pages: list[PageText], offset: int) -> str:
    for page in pages:
        if page.start_offset <= offset <= page.end_offset:
            return f"page:{page.page_no}"
    return ""


def _page_content_start_offset(page: PageText, *, full_text: str | None = None) -> int:
    if page.content_start_offset is not None:
        return page.content_start_offset
    marker = f"\n\n[Page {page.page_no}]\n"
    if full_text is not None and full_text.startswith(marker, page.start_offset):
        return page.start_offset + len(marker)
    return page.start_offset


def _find_title_offset(text: str, title: str) -> int:
    if not text or not title:
        return -1
    normalized_title = "".join(character.lower() for character in title if not character.isspace())
    if not normalized_title:
        return -1
    normalized_text_parts: list[str] = []
    original_offsets: list[int] = []
    for offset, character in enumerate(text):
        if character.isspace():
            continue
        lowered = character.lower()
        normalized_text_parts.append(lowered)
        original_offsets.extend([offset] * len(lowered))
    normalized_text = "".join(normalized_text_parts)
    index = normalized_text.find(normalized_title)
    if index < 0:
        return -1
    return original_offsets[index]


def _number_from_title(title: str) -> str:
    cleaned = _clean_label(title)
    structural_marker = parse_structural_heading(cleaned)
    if structural_marker is not None and structural_marker.number:
        return structural_marker.number
    patterns = [
        r"^(\d+(?:\.\d+){0,8})(?=\s|[.:：、-]|$)",
        r"^chapter\s+(\d+(?:\.\d+){0,8})(?=\s|[.:：、-]|$)",
        r"^第\s*(\d+)\s*[章节]",
    ]
    lowered = cleaned.lower()
    for pattern in patterns:
        match = re.search(pattern, lowered, flags=re.I)
        if match:
            return match.group(1)
    return ""


def _normalize_chapter_number(number: str) -> str:
    parts = [part for part in re.split(r"\.+", number.strip()) if part != ""]
    normalized: list[str] = []
    for part in parts:
        if part.isdigit():
            normalized.append(str(int(part)))
        else:
            normalized.append(part.lower())
    return ".".join(normalized)


def _clean_label(value: str) -> str:
    return re.sub(r"\s+", " ", html.unescape(value or "")).strip()


def _normalize_text(value: str) -> str:
    lines = [line.strip() for line in value.splitlines()]
    return "\n".join(line for line in lines if line)


def _normalize_for_match(value: str) -> str:
    return re.sub(r"\s+", " ", _clean_label(value)).lower()


def _compact(text: str, limit: int) -> str:
    compacted = re.sub(r"\s+", " ", text).strip()
    return compacted if len(compacted) <= limit else compacted[: limit - 1].rstrip() + "…"


def _estimate_tokens(text: str) -> int:
    stripped = text.strip()
    if not stripped:
        return 0
    ascii_chars = sum(1 for char in stripped if ord(char) < 128)
    non_ascii_chars = len(stripped) - ascii_chars
    return max(1, ascii_chars // 4 + non_ascii_chars // 2)


def _dedupe_nav_items(
    items: list[_EpubNavigationItem],
) -> list[_EpubNavigationItem]:
    seen: set[tuple[str, str, str]] = set()
    deduped: list[_EpubNavigationItem] = []
    for item in items:
        key = (item.target, item.fragment, item.label)
        if key in seen:
            continue
        seen.add(key)
        deduped.append(item)
    return deduped


source_structure_indexer = SourceStructureIndexer()
