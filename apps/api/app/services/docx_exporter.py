from __future__ import annotations

import io
from collections.abc import Callable
from pathlib import Path

from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.image.exceptions import UnrecognizedImageError

from app.models import BoardDocument
from app.services.board_asset_store import board_asset_id_from_url, get_board_asset_store
from app.services.rich_document import core as rd
from app.services.docx_quality_check import assert_docx_export_quality
from app.services.docx_styles import apply_textbook_docx_styles


def export_docx(
    document: BoardDocument,
    path: Path,
    *,
    owner_user_id: str = "",
    asset_resolver: Callable[[str], tuple[str, bytes] | None] | None = None,
) -> Path:
    target = DocxDocument()
    apply_textbook_docx_styles(target)
    rd._apply_page_settings(target, document.page_settings)
    target.add_heading(document.title, level=0)
    current_page_units = 4.0

    parser = rd._DocxBlockParser()
    content_html = (document.content_html or "").strip()
    content_text = (document.content_text or "").strip()
    if _json_contains_resource_visual(document.content_json):
        content_html = rd.tiptap_doc_to_html(document.content_json)
    elif content_html and content_text and rd._html_has_visible_raw_math_text(content_html):
        content_html = rd.text_to_html(content_text)
    elif content_html:
        content_html = rd._repair_suspicious_math_html(content_html)
    else:
        content_html = rd.text_to_html(content_text)
    parser.feed(content_html)
    parser._flush()
    blocks = parser.blocks or [("p", [("text", line)], {}) for line in document.content_text.splitlines() if line.strip()]
    blocks = rd._normalize_fenced_docx_blocks(blocks)

    for tag, fragments, attrs in blocks:
        text = rd._fragment_text(fragments)
        if tag == "h1":
            paragraph = target.add_heading("", level=1)
            paragraph.paragraph_format.keep_with_next = True
            rd._append_fragments(paragraph, fragments)
            current_page_units = rd._advance_page_units(current_page_units, rd._estimated_block_units(tag, text))
        elif tag == "h2":
            paragraph = target.add_heading("", level=2)
            paragraph.paragraph_format.keep_with_next = True
            rd._append_fragments(paragraph, fragments)
            current_page_units = rd._advance_page_units(current_page_units, rd._estimated_block_units(tag, text))
        elif tag == "h3":
            paragraph = target.add_heading("", level=3)
            paragraph.paragraph_format.keep_with_next = True
            rd._append_fragments(paragraph, fragments)
            current_page_units = rd._advance_page_units(current_page_units, rd._estimated_block_units(tag, text))
        elif tag == "li":
            rd._add_fragment_paragraph(target, fragments, style="List Bullet")
            current_page_units = rd._advance_page_units(current_page_units, rd._estimated_block_units(tag, text))
        elif tag == "blockquote":
            rd._add_fragment_paragraph(target, fragments, style="Intense Quote")
            current_page_units = rd._advance_page_units(current_page_units, rd._estimated_block_units(tag, text))
        elif tag == "img":
            src = str(attrs.get("src") or "").strip()
            asset_id = str(attrs.get("asset_id") or "").strip() or board_asset_id_from_url(src)
            image_bytes = (
                _board_asset_bytes(asset_id, owner_user_id, asset_resolver)
                or rd._decode_data_uri(src)
            )
            if image_bytes and _add_docx_picture(target, image_bytes):
                pass
            elif text:
                target.add_paragraph(f"[图片不可用] {text}")
            current_page_units = rd._advance_page_units(current_page_units, 10)
        elif tag == "pageBreak":
            target.add_page_break()
            current_page_units = 0
        elif tag == "table":
            rows = attrs.get("rows")
            if isinstance(rows, list):
                current_page_units = rd._maybe_page_break_before_table(target, rows, current_page_units)
                rd._add_fragment_table(target, rows)
                current_page_units = rd._advance_page_units(current_page_units, rd._estimated_table_units(rows))
        elif tag == "pre":
            rd._add_preformatted_paragraph(target, text)
            current_page_units = rd._advance_page_units(current_page_units, rd._estimated_block_units(tag, text))
        elif tag == "math" or (len(fragments) == 1 and fragments[0][0] == "math"):
            paragraph = target.add_paragraph(style="OpenClass Formula")
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            rd._append_fragments(paragraph, fragments, auto_math=False, display_math=True)
            current_page_units = rd._advance_page_units(current_page_units, rd._estimated_block_units("math", text))
        else:
            fenced_text = rd._fenced_code_text(text)
            if fenced_text is not None:
                rd._add_preformatted_paragraph(target, fenced_text)
                current_page_units = rd._advance_page_units(current_page_units, rd._estimated_block_units("pre", fenced_text))
                continue
            formula_latex = rd._formula_only_latex(text)
            if formula_latex:
                paragraph = target.add_paragraph(style="OpenClass Formula")
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                rd._append_math(paragraph, formula_latex, display=True)
                current_page_units = rd._advance_page_units(current_page_units, rd._estimated_block_units("math", text))
            else:
                rd._add_fragment_paragraph(target, fragments, style="OpenClass Body Text")
                current_page_units = rd._advance_page_units(current_page_units, rd._estimated_block_units(tag, text))

    path.parent.mkdir(parents=True, exist_ok=True)
    target.save(path)
    assert_docx_export_quality(path)
    return path


def _board_asset_bytes(
    asset_id: str,
    owner_user_id: str,
    asset_resolver: Callable[[str], tuple[str, bytes] | None] | None,
) -> bytes | None:
    if not asset_id:
        return None
    if asset_resolver is not None:
        resolved = asset_resolver(asset_id)
        return resolved[1] if resolved is not None else None
    if not owner_user_id:
        return None
    stored = get_board_asset_store().read_bytes(asset_id, owner_user_id)
    return stored[1] if stored is not None else None


def _json_contains_resource_visual(value: object) -> bool:
    if isinstance(value, dict):
        if value.get("type") == "resourceVisualBlock":
            return True
        return any(_json_contains_resource_visual(item) for item in value.values())
    if isinstance(value, list):
        return any(_json_contains_resource_visual(item) for item in value)
    return False


def _add_docx_picture(target: DocxDocument, image_bytes: bytes) -> bool:
    try:
        shape = target.add_picture(io.BytesIO(image_bytes))
    except UnrecognizedImageError:
        converted = _convert_picture_to_png(image_bytes)
        if converted is None:
            return False
        try:
            shape = target.add_picture(io.BytesIO(converted))
        except UnrecognizedImageError:
            return False

    section = target.sections[-1]
    available_width = int(section.page_width - section.left_margin - section.right_margin)
    if available_width > 0 and shape.width > available_width:
        original_width = int(shape.width)
        original_height = int(shape.height)
        shape.width = available_width
        shape.height = max(1, round(original_height * available_width / original_width))
    return True


def _convert_picture_to_png(image_bytes: bytes) -> bytes | None:
    try:
        from PIL import Image

        with Image.open(io.BytesIO(image_bytes)) as image:
            image.seek(0)
            converted = image.convert("RGBA" if "A" in image.getbands() else "RGB")
            output = io.BytesIO()
            converted.save(output, format="PNG", optimize=True)
            return output.getvalue()
    except Exception:
        return None
