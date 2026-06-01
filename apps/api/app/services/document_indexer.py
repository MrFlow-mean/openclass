from __future__ import annotations

import hashlib
import mimetypes
import os
import re
import sqlite3
import threading
import time
from dataclasses import dataclass
from pathlib import Path

from docx import Document as DocxDocument
from pypdf import PdfReader

from app.models import now_iso
from app.services.document_index_store import DocumentBlockRecord, DocumentIndexStore, DocumentPageRecord
from app.services.image_ocr import extract_image_text, extract_pdf_page_texts


PDF_OCR_MAX_PAGES = int(os.getenv("OPENCLASS_DOCUMENT_OCR_MAX_PAGES", "80"))
INDEXER_POLL_SECONDS = float(os.getenv("OPENCLASS_DOCUMENT_INDEXER_POLL_SECONDS", "1.5"))
BLOCK_TARGET_CHARS = 1000
BLOCK_MAX_CHARS = 1500
MAX_BLOCKS_PER_RESOURCE = 900


@dataclass(frozen=True)
class ParsedDocumentIndex:
    pages: list[DocumentPageRecord]
    blocks: list[DocumentBlockRecord]
    status: str
    message: str


_worker_started_paths: set[str] = set()
_worker_lock = threading.Lock()


def enqueue_resource_index(database_path: Path, resource_id: str) -> None:
    store = DocumentIndexStore()
    with _connect(database_path) as conn:
        with conn:
            store.create_schema(conn)
            store.enqueue(conn, resource_id)
            conn.execute(
                """
                UPDATE resources
                SET index_status = 'queued',
                    index_message = ?,
                    index_updated_at = ?
                WHERE id = ?
                """,
                ("等待后台解析资料", now_iso(), resource_id),
            )


def delete_resource_index(database_path: Path, resource_id: str) -> None:
    store = DocumentIndexStore()
    with _connect(database_path) as conn:
        with conn:
            store.create_schema(conn)
            store.delete_resource(conn, resource_id)


def start_document_index_worker(database_path: Path) -> None:
    key = str(database_path.resolve(strict=False))
    with _worker_lock:
        if key in _worker_started_paths:
            return
        _worker_started_paths.add(key)
    thread = threading.Thread(
        target=_worker_loop,
        args=(database_path,),
        name="openclass-document-indexer",
        daemon=True,
    )
    thread.start()


def index_next_resource(database_path: Path) -> str | None:
    store = DocumentIndexStore()
    with _connect(database_path) as conn:
        with conn:
            store.create_schema(conn)
            resource_id = store.claim_next_job(conn)
            if resource_id:
                conn.execute(
                    """
                    UPDATE resources
                    SET index_status = 'processing',
                        index_message = ?,
                        index_updated_at = ?
                    WHERE id = ?
                    """,
                    ("正在解析资料并建立索引", now_iso(), resource_id),
                )
    if not resource_id:
        return None
    index_resource_now(database_path, resource_id)
    return resource_id


def index_resource_now(database_path: Path, resource_id: str) -> None:
    store = DocumentIndexStore()
    with _connect(database_path) as conn:
        resource = conn.execute("SELECT * FROM resources WHERE id = ?", (resource_id,)).fetchone()
    if resource is None:
        return

    source_path = Path(str(resource["source_path"] or ""))
    if not source_path.exists():
        _mark_failed(database_path, resource_id, "源文件不存在")
        return

    try:
        parsed = build_document_index(
            resource_id=resource_id,
            source_path=source_path,
            resource_name=str(resource["name"]),
            mime_type=str(resource["mime_type"]),
        )
    except Exception as exc:
        _mark_failed(database_path, resource_id, str(exc) or exc.__class__.__name__)
        return

    with _connect(database_path) as conn:
        with conn:
            store.create_schema(conn)
            store.replace_index(conn, resource_id=resource_id, pages=parsed.pages, blocks=parsed.blocks)
            store.finish_job(conn, resource_id, status=parsed.status)
            conn.execute(
                """
                UPDATE resources
                SET extracted_text_available = ?,
                    index_status = ?,
                    index_message = ?,
                    index_updated_at = ?,
                    page_count = ?,
                    indexed_block_count = ?
                WHERE id = ?
                """,
                (
                    int(parsed.status == "ready"),
                    parsed.status,
                    parsed.message,
                    now_iso(),
                    len(parsed.pages),
                    len(parsed.blocks),
                    resource_id,
                ),
            )


def build_document_index(
    *,
    resource_id: str,
    source_path: Path,
    resource_name: str,
    mime_type: str | None = None,
) -> ParsedDocumentIndex:
    detected_mime = mime_type or mimetypes.guess_type(resource_name)[0] or "application/octet-stream"
    suffix = source_path.suffix.lower()
    if detected_mime.startswith("image/"):
        return _index_image(resource_id, source_path, resource_name)
    if suffix == ".pdf" or detected_mime == "application/pdf":
        return _index_pdf(resource_id, source_path)
    if suffix == ".docx":
        return _index_docx(resource_id, source_path, resource_name)
    if suffix in {".md", ".markdown"} or detected_mime == "text/markdown":
        return _index_plain_text(resource_id, source_path, resource_name, markdown=True)
    if suffix == ".txt" or detected_mime.startswith("text/"):
        return _index_plain_text(resource_id, source_path, resource_name, markdown=False)
    return ParsedDocumentIndex([], [], "failed", "暂不支持这种文件类型")


def _index_pdf(resource_id: str, source_path: Path) -> ParsedDocumentIndex:
    reader = PdfReader(str(source_path))
    labels = _page_labels(reader)
    pages: list[DocumentPageRecord] = []
    for page_number, page in enumerate(reader.pages, start=1):
        try:
            raw_text = page.extract_text() or ""
        except Exception:
            raw_text = ""
        text = _normalize_text(raw_text)
        printed_page = _printed_page_number(
            actual_page=page_number,
            label=labels.get(page_number),
            text=text,
        )
        if text:
            pages.append(
                DocumentPageRecord(
                    resource_id=resource_id,
                    page_number=page_number,
                    printed_page=printed_page,
                    text=text,
                    text_hash=_hash_text(text),
                    text_source="source_file",
                )
            )

    if not pages and PDF_OCR_MAX_PAGES > 0:
        for ocr_page in extract_pdf_page_texts(source_path, max_pages=PDF_OCR_MAX_PAGES):
            text = _normalize_text(ocr_page.text)
            if not text:
                continue
            printed_page = _printed_page_number(
                actual_page=ocr_page.page_number,
                label=labels.get(ocr_page.page_number),
                text=text,
            )
            pages.append(
                DocumentPageRecord(
                    resource_id=resource_id,
                    page_number=ocr_page.page_number,
                    printed_page=printed_page,
                    text=text,
                    text_hash=_hash_text(text),
                    text_source="ocr",
                )
            )

    blocks = _blocks_from_pages(resource_id, pages)
    if blocks:
        source_note = "OCR" if all(page.text_source == "ocr" for page in pages) else "PDF 文本"
        return ParsedDocumentIndex(pages, blocks, "ready", f"已完成索引：{len(pages)} 页，{len(blocks)} 个正文块（{source_note}）")
    return ParsedDocumentIndex(pages, [], "no_text", "没有抽取到可用于定位的正文")


def _index_docx(resource_id: str, source_path: Path, resource_name: str) -> ParsedDocumentIndex:
    document = DocxDocument(source_path)
    heading_path = [Path(resource_name).stem]
    blocks: list[DocumentBlockRecord] = []
    buffer: list[str] = []
    order = 0

    def flush() -> None:
        nonlocal order, buffer
        text = _normalize_text("\n".join(buffer))
        buffer = []
        if not text:
            return
        for part in _split_text(text):
            blocks.append(_block(resource_id, order, part, heading_path=list(heading_path), text_source="source_file"))
            order += 1

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        style_name = (paragraph.style.name if paragraph.style else "").lower()
        if "heading" in style_name or "title" in style_name:
            flush()
            level = _heading_level(style_name)
            heading_path[:] = heading_path[: max(level - 1, 0)]
            heading_path.append(text[:120])
        else:
            buffer.append(text)
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                buffer.append(" | ".join(cells))
    flush()
    return _text_index_result(blocks)


def _index_plain_text(resource_id: str, source_path: Path, resource_name: str, *, markdown: bool) -> ParsedDocumentIndex:
    text = _normalize_text(source_path.read_text(encoding="utf-8", errors="ignore"))
    title = Path(resource_name).stem
    blocks: list[DocumentBlockRecord] = []
    if markdown:
        sections = _markdown_sections(text)
        for section in sections:
            for part in _split_text(section["text"]):
                blocks.append(
                    _block(
                        resource_id,
                        len(blocks),
                        part,
                        heading_path=section["heading_path"],
                        block_type="section",
                    )
                )
    if not blocks:
        for part in _split_text(text):
            blocks.append(_block(resource_id, len(blocks), part, heading_path=[title]))
    return _text_index_result(blocks)


def _index_image(resource_id: str, source_path: Path, resource_name: str) -> ParsedDocumentIndex:
    text = _normalize_text(extract_image_text(source_path) or "")
    if not text:
        return ParsedDocumentIndex([], [], "no_text", "图片 OCR 没有识别到正文")
    page = DocumentPageRecord(
        resource_id=resource_id,
        page_number=1,
        printed_page=1,
        text=text,
        text_hash=_hash_text(text),
        text_source="ocr",
    )
    blocks = [
        _block(
            resource_id,
            index,
            part,
            heading_path=[Path(resource_name).stem],
            page_start=1,
            page_end=1,
            printed_page_start=1,
            printed_page_end=1,
            text_source="ocr",
        )
        for index, part in enumerate(_split_text(text))
    ]
    return ParsedDocumentIndex([page], blocks, "ready", f"已完成图片 OCR 索引：{len(blocks)} 个正文块")


def _blocks_from_pages(resource_id: str, pages: list[DocumentPageRecord]) -> list[DocumentBlockRecord]:
    blocks: list[DocumentBlockRecord] = []
    heading_path: list[str] = []
    for page in pages:
        page_heading = _first_heading(page.text)
        if page_heading:
            heading_path = [page_heading]
        for text in _split_text(_strip_running_page_number(page.text)):
            if len(blocks) >= MAX_BLOCKS_PER_RESOURCE:
                return blocks
            blocks.append(
                _block(
                    resource_id,
                    len(blocks),
                    text,
                    heading_path=heading_path or [f"第 {page.page_number} 页"],
                    page_start=page.page_number,
                    page_end=page.page_number,
                    printed_page_start=page.printed_page,
                    printed_page_end=page.printed_page,
                    text_source=page.text_source,
                )
            )
    return blocks


def _block(
    resource_id: str,
    order_index: int,
    text: str,
    *,
    heading_path: list[str],
    page_start: int | None = None,
    page_end: int | None = None,
    printed_page_start: int | None = None,
    printed_page_end: int | None = None,
    block_type: str = "paragraph",
    text_source: str = "source_file",
    confidence: float = 1.0,
) -> DocumentBlockRecord:
    text_hash = _hash_text(text)
    seed = f"{resource_id}:{order_index}:{page_start}:{text_hash}"
    return DocumentBlockRecord(
        resource_id=resource_id,
        block_id=f"dblk_{hashlib.sha1(seed.encode('utf-8')).hexdigest()[:12]}",
        order_index=order_index,
        page_start=page_start,
        page_end=page_end,
        printed_page_start=printed_page_start,
        printed_page_end=printed_page_end,
        heading_path=heading_path,
        text=text,
        text_hash=text_hash,
        keywords=_keywords(text),
        block_type=block_type,
        text_source=text_source,
        confidence=confidence,
    )


def _text_index_result(blocks: list[DocumentBlockRecord]) -> ParsedDocumentIndex:
    if blocks:
        return ParsedDocumentIndex([], blocks, "ready", f"已完成索引：{len(blocks)} 个正文块")
    return ParsedDocumentIndex([], [], "no_text", "没有抽取到可用于定位的正文")


def _worker_loop(database_path: Path) -> None:
    while True:
        try:
            indexed = index_next_resource(database_path)
        except Exception:
            indexed = None
        if indexed is None:
            time.sleep(INDEXER_POLL_SECONDS)


def _mark_failed(database_path: Path, resource_id: str, message: str) -> None:
    store = DocumentIndexStore()
    with _connect(database_path) as conn:
        with conn:
            store.create_schema(conn)
            store.finish_job(conn, resource_id, status="failed", error=message)
            conn.execute(
                """
                UPDATE resources
                SET index_status = 'failed',
                    index_message = ?,
                    index_updated_at = ?
                WHERE id = ?
                """,
                (message[:500], now_iso(), resource_id),
            )


def _connect(database_path: Path) -> sqlite3.Connection:
    conn = sqlite3.connect(database_path, timeout=10)
    conn.row_factory = sqlite3.Row
    conn.execute("PRAGMA busy_timeout = 5000")
    conn.execute("PRAGMA journal_mode = WAL")
    conn.execute("PRAGMA synchronous = NORMAL")
    return conn


def _normalize_text(text: str) -> str:
    cleaned = text.replace("\x00", "").replace("\r\n", "\n")
    cleaned = re.sub(r"(?<=[\u4e00-\u9fff])\s+(?=[\u4e00-\u9fff])", "", cleaned)
    cleaned = re.sub(r"[ \t]{2,}", " ", cleaned)
    cleaned = re.sub(r"\n{3,}", "\n\n", cleaned)
    return cleaned.strip()


def _split_text(text: str) -> list[str]:
    normalized = _normalize_text(text)
    if not normalized:
        return []
    paragraphs = [part.strip() for part in re.split(r"\n\s*\n+", normalized) if part.strip()]
    if not paragraphs:
        paragraphs = [line.strip() for line in normalized.splitlines() if line.strip()]
    chunks: list[str] = []
    current = ""
    for paragraph in paragraphs:
        for part in _split_long(paragraph):
            candidate = part if not current else f"{current}\n\n{part}"
            if len(candidate) <= BLOCK_TARGET_CHARS or not current:
                current = candidate
            else:
                chunks.append(current)
                current = part
    if current:
        chunks.append(current)
    return chunks[:MAX_BLOCKS_PER_RESOURCE]


def _split_long(text: str) -> list[str]:
    if len(text) <= BLOCK_MAX_CHARS:
        return [text]
    sentences = [part.strip() for part in re.split(r"(?<=[。！？!?；;])\s*", text) if part.strip()]
    if len(sentences) <= 1:
        return [text[index : index + BLOCK_MAX_CHARS].strip() for index in range(0, len(text), BLOCK_MAX_CHARS)]
    chunks: list[str] = []
    current = ""
    for sentence in sentences:
        candidate = sentence if not current else f"{current}{sentence}"
        if len(candidate) <= BLOCK_TARGET_CHARS or not current:
            current = candidate
        else:
            chunks.append(current)
            current = sentence
    if current:
        chunks.append(current)
    return chunks


def _page_labels(reader: PdfReader) -> dict[int, str]:
    try:
        labels = list(reader.page_labels)
    except Exception:
        return {}
    return {index: str(label) for index, label in enumerate(labels, start=1)}


def _arabic_page_number(value: str | None) -> int | None:
    if not value:
        return None
    compact = value.strip()
    if re.fullmatch(r"\d{1,5}", compact):
        return int(compact)
    return None


def _printed_page_from_text(text: str) -> int | None:
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    for line in [*lines[:5], *lines[-8:]]:
        match = re.fullmatch(r"(?:第\s*)?(\d{1,5})(?:\s*页)?", line)
        if match:
            return int(match.group(1))
    return None


def _printed_page_number(*, actual_page: int, label: str | None, text: str) -> int | None:
    text_number = _printed_page_from_text(text)
    label_number = _arabic_page_number(label)
    if label_number is not None and label_number != actual_page:
        return label_number
    return text_number if text_number is not None else label_number


def _strip_running_page_number(text: str) -> str:
    lines = [line for line in text.splitlines()]
    if lines and _printed_page_from_text(lines[0]):
        lines = lines[1:]
    if lines and _printed_page_from_text(lines[-1]):
        lines = lines[:-1]
    return "\n".join(lines).strip()


def _first_heading(text: str) -> str | None:
    for line in [line.strip() for line in text.splitlines() if line.strip()][:12]:
        if len(line) > 120:
            continue
        if re.match(r"^(第\s*\d+\s*[章节]|(?:\d{1,3}[.．]){1,4}\d{0,3}\s*\S+|\d{1,3}\s+\S+)", line):
            return line
    return None


def _heading_level(style_name: str) -> int:
    match = re.search(r"heading\s*(\d+)", style_name)
    if match:
        return max(1, min(6, int(match.group(1))))
    return 1


def _markdown_sections(text: str) -> list[dict[str, object]]:
    lines = text.splitlines()
    heading_path: list[str] = []
    sections: list[dict[str, object]] = []
    buffer: list[str] = []

    def flush() -> None:
        body = "\n".join(buffer).strip()
        if body:
            sections.append({"heading_path": list(heading_path), "text": body})
        buffer.clear()

    for line in lines:
        match = re.match(r"^(#{1,6})\s+(.+)$", line.strip())
        if match:
            flush()
            level = len(match.group(1))
            heading_path = heading_path[: level - 1]
            heading_path.append(match.group(2).strip()[:120])
        else:
            buffer.append(line)
    flush()
    return sections


def _keywords(text: str) -> list[str]:
    counts: dict[str, int] = {}
    for token in re.findall(r"[A-Za-z0-9\u4e00-\u9fff]{2,}", text.lower()):
        counts[token] = counts.get(token, 0) + 1
    return [token for token, _ in sorted(counts.items(), key=lambda item: item[1], reverse=True)[:12]]


def _hash_text(text: str) -> str:
    compact = re.sub(r"\s+", " ", text).strip()
    return hashlib.sha1(compact.encode("utf-8")).hexdigest()[:16]
