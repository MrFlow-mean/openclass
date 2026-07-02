from __future__ import annotations

import html
import io
import re
from functools import lru_cache
from html.parser import HTMLParser
from pathlib import Path
from typing import Any

from docx import Document as DocxDocument
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm

from app.models import BoardDocument, DocumentPageSettings


EMPTY_TIPTAP_DOC: dict[str, Any] = {"type": "doc", "content": [{"type": "paragraph"}]}
InlineFragment = tuple[str, str]
TableRows = list[list[list[InlineFragment]]]
DocxBlock = tuple[str, list[InlineFragment], dict[str, Any]]

_CJK_RE = re.compile(r"[\u3400-\u9fff]")
_MATH_SIGNAL_RE = re.compile(
    r"\\(?:begin|end|frac|dfrac|tfrac|sqrt|lim|sum|prod|int|sin|cos|tan|ln|log|exp|to|left|right|leftarrow|rightarrow|leftrightarrow|Leftarrow|Rightarrow|Leftrightarrow|Longleftarrow|Longrightarrow|Longleftrightarrow|infty|cdot|times|div|leq?|geq?|approx|neq?|pm|sim|in|notin|mid|subseteq?|supseteq?|cup|cap|mathbb|mathcal|mathfrak|mathbf|mathrm|operatorname|text|dots|cdots|ldots|vdots|partial|nabla|forall|exists|alpha|beta|gamma|delta|epsilon|varepsilon|zeta|eta|theta|iota|kappa|lambda|mu|xi|pi|rho|varrho|sigma|tau|upsilon|phi|varphi|chi|psi|omega|Gamma|Delta|Theta|Lambda|Xi|Pi|Sigma|Phi|Psi|Omega)\b"
    r"|[_^]"
    r"|[=<>≤≥≈≠]"
    r"|[A-Za-z0-9)]\s*(?:[+\-−*/=<>≤≥≈≠±]|→|←)\s*[A-Za-z0-9(\\]"
    r"|\d+\s*/\s*\d+"
    r"|\\[{}]"
    r"|^\([^()\n]{1,80},[^()\n]{1,80}\)$"
    r"|^[\[(][A-Za-z0-9α-ωΑ-Ω\\_{}\s+\-−*/=.,]+,[A-Za-z0-9α-ωΑ-Ω\\_{}\s+\-−*/=.,]+[\])]$"
    r"|^[A-Za-z]{1,3}\s*\([A-Za-z0-9α-ωΑ-Ω\\_{}\[\]^()+\-−*/=·∞→←≤≥≈≠±<>|&:'\s.,]+\)$"
)
_LATIN_WORD_RE = re.compile(r"[A-Za-z]+")
_NON_FORMULA_LETTER_RE = re.compile(r"[^\W\d_A-Za-zα-ωΑ-Ω]", re.UNICODE)
_FORMULA_CHARS_RE = re.compile(r"^[A-Za-z0-9α-ωΑ-Ω\\_{}\[\]^()+\-−*/=·∞→←≤≥≈≠±<>|&:'\s.,]+$")
_LATEX_ENVIRONMENT_RE = re.compile(r"\\(?:begin|end)\{[A-Za-z*]+\}")
_LATEX_TEXT_ARGUMENT_RE = re.compile(r"\\(?:text|mathrm|operatorname)\{[^{}]*\}")
_RAW_LATEX_COMMAND_RE = re.compile(
    r"\\(?:alpha|beta|gamma|delta|epsilon|varepsilon|zeta|eta|theta|iota|kappa|lambda|mu|xi|pi|rho|varrho|sigma|tau|upsilon|phi|varphi|chi|psi|omega|Gamma|Delta|Theta|Lambda|Xi|Pi|Sigma|Phi|Psi|Omega|infty|forall|exists|int|sum|prod|lim)\b"
)
_MIXED_MATH_TEXT_RE = re.compile(r"([\u3400-\u9fff，。；：、]+)")
_HTML_BLOCK_RE = re.compile(
    r"<(?P<tag>h[1-6]|p|li|blockquote)\b[^>]*>.*?</(?P=tag)>",
    re.IGNORECASE | re.DOTALL,
)
_HTML_CONTENT_RE = re.compile(
    r"</?(?:h[1-6]|p|ul|ol|li|table|thead|tbody|tfoot|tr|td|th|strong|em|blockquote|br|div|section)\b[^>]*>",
    re.IGNORECASE,
)
_DELIMITED_MATH_RE = re.compile(
    r"\\\[([\s\S]+?)\\\]|\\\((.+?)\\\)|\$\$([\s\S]+?)\$\$|\$(?!\d+\$)([^$\n]+?)\$(?!\d)"
)
_TRAILING_SENTENCE_MARKS_RE = re.compile(r"[\s.,，。；;:：]+$")
_LEADING_SENTENCE_MARKS_RE = re.compile(r"^[\s.,，。；;:：]+")
_MARKDOWN_INLINE_RE = re.compile(r"(\*\*[^*\n]+?\*\*|\*[^*\n]+?\*)")
_MARKDOWN_HEADING_RE = re.compile(r"^(#{1,3})\s+(.+)$")
_MARKDOWN_BULLET_RE = re.compile(r"^[-*]\s+(.+)$")
_MARKDOWN_ORDERED_RE = re.compile(r"^\d+[.、]\s+(.+)$")
_MARKDOWN_TABLE_SEPARATOR_RE = re.compile(r"^\s*\|?\s*:?-{3,}:?\s*(?:\|\s*:?-{3,}:?\s*)+\|?\s*$")
_FENCED_CODE_TEXT_RE = re.compile(r"^```(?P<body>[\s\S]*?)```$")
_DOCX_PAGE_UNIT_LIMIT = 34
_DOCX_TABLE_PAGE_BREAK_THRESHOLD = 0.82
_VOID_HTML_TAGS = {"area", "base", "br", "col", "embed", "hr", "img", "input", "link", "meta", "source", "track", "wbr"}
_HTML_BLOCK_TAGS = {
    "address",
    "article",
    "aside",
    "blockquote",
    "div",
    "figure",
    "footer",
    "h1",
    "h2",
    "h3",
    "h4",
    "h5",
    "h6",
    "header",
    "img",
    "li",
    "main",
    "nav",
    "ol",
    "p",
    "section",
    "table",
    "tbody",
    "td",
    "tfoot",
    "th",
    "thead",
    "tr",
    "ul",
}
_LATEX_SYMBOLS = {
    r"\to": "→",
    r"\leftarrow": "←",
    r"\rightarrow": "→",
    r"\leftrightarrow": "↔",
    r"\Leftarrow": "⇐",
    r"\Rightarrow": "⇒",
    r"\Leftrightarrow": "⇔",
    r"\Longleftarrow": "⟸",
    r"\Longrightarrow": "⟹",
    r"\Longleftrightarrow": "⟺",
    r"\infty": "∞",
    r"\cdot": "·",
    r"\times": "×",
    r"\div": "÷",
    r"\le": "≤",
    r"\ge": "≥",
    r"\leq": "≤",
    r"\geq": "≥",
    r"\approx": "≈",
    r"\ne": "≠",
    r"\neq": "≠",
    r"\pm": "±",
    r"\sim": "∼",
    r"\forall": "∀",
    r"\exists": "∃",
    r"\in": "∈",
    r"\notin": "∉",
    r"\mid": "|",
    r"\subset": "⊂",
    r"\subseteq": "⊆",
    r"\supset": "⊃",
    r"\supseteq": "⊇",
    r"\cup": "∪",
    r"\cap": "∩",
    r"\alpha": "α",
    r"\beta": "β",
    r"\gamma": "γ",
    r"\delta": "δ",
    r"\epsilon": "ε",
    r"\varepsilon": "ε",
    r"\zeta": "ζ",
    r"\eta": "η",
    r"\theta": "θ",
    r"\iota": "ι",
    r"\kappa": "κ",
    r"\lambda": "λ",
    r"\mu": "μ",
    r"\xi": "ξ",
    r"\pi": "π",
    r"\rho": "ρ",
    r"\varrho": "ρ",
    r"\sigma": "σ",
    r"\tau": "τ",
    r"\upsilon": "υ",
    r"\phi": "φ",
    r"\varphi": "φ",
    r"\chi": "χ",
    r"\psi": "ψ",
    r"\omega": "ω",
    r"\Gamma": "Γ",
    r"\Delta": "Δ",
    r"\Theta": "Θ",
    r"\Lambda": "Λ",
    r"\Xi": "Ξ",
    r"\Pi": "Π",
    r"\Sigma": "Σ",
    r"\Phi": "Φ",
    r"\Psi": "Ψ",
    r"\Omega": "Ω",
    r"\partial": "∂",
    r"\int": "∫",
    r"\sum": "∑",
    r"\prod": "∏",
    r"\dots": "…",
    r"\cdots": "⋯",
    r"\ldots": "…",
}
_LATEX_FUNCTIONS = {"sin", "cos", "tan", "ln", "log", "sqrt", "exp", "lim"}
_LATEX_TEXT_COMMANDS = {r"\text", r"\mathrm", r"\operatorname"}
_LATEX_STYLE_COMMANDS = {r"\displaystyle", r"\textstyle", r"\scriptstyle", r"\scriptscriptstyle"}
_LATEX_DELIMITER_COMMANDS = {r"\left", r"\right"}
_LATEX_SPACING_COMMANDS = {r"\quad", r"\qquad", r"\,", r"\;", r"\:", r"\!"}
_SUPERSCRIPT_CHARS = str.maketrans(
    "0123456789+-=()abcdefgijklmnoprstuvwxyz",
    "⁰¹²³⁴⁵⁶⁷⁸⁹⁺⁻⁼⁽⁾ᵃᵇᶜᵈᵉᶠᵍⁱʲᵏˡᵐⁿᵒᵖʳˢᵗᵘᵛʷˣʸᶻ",
)
_SUBSCRIPT_CHARS = str.maketrans(
    "0123456789+-=()aehijklmnoprstuvx",
    "₀₁₂₃₄₅₆₇₈₉₊₋₌₍₎ₐₑₕᵢⱼₖₗₘₙₒₚᵣₛₜᵤᵥₓ",
)


def html_to_text(content_html: str) -> str:
    without_tags = re.sub(r"</(h[1-6]|p|li|blockquote|tr)>", "\n", content_html)
    without_tags = re.sub(r"<br\s*/?>", "\n", without_tags, flags=re.IGNORECASE)
    text = re.sub(r"<[^>]+>", "", without_tags)
    text = html.unescape(text)
    lines = [line.strip() for line in text.splitlines()]
    compact_lines: list[str] = []
    for line in lines:
        if not line and (not compact_lines or not compact_lines[-1]):
            continue
        compact_lines.append(line)
    return "\n".join(compact_lines).strip()


def _markdown_text_nodes(value: str) -> list[dict[str, Any]]:
    nodes: list[dict[str, Any]] = []
    cursor = 0
    for match in _MARKDOWN_INLINE_RE.finditer(value):
        if match.start() > cursor:
            nodes.append({"type": "text", "text": value[cursor : match.start()]})
        token = match.group(0)
        if token.startswith("**") and token.endswith("**"):
            text = token[2:-2]
            mark = "bold"
        else:
            text = token[1:-1]
            mark = "italic"
        if text:
            nodes.append({"type": "text", "text": text, "marks": [{"type": mark}]})
        cursor = match.end()
    if cursor < len(value):
        nodes.append({"type": "text", "text": value[cursor:]})
    return nodes


def _inline_math_node(latex: str) -> dict[str, Any]:
    return {"type": "inlineMath", "attrs": {"latex": _normalize_latex(latex)}}


def _markdown_text_nodes_with_raw_math(value: str) -> list[dict[str, Any]]:
    nodes: list[dict[str, Any]] = []
    cursor = 0
    for match in _RAW_LATEX_COMMAND_RE.finditer(value):
        command = match.group(0)
        if not _is_likely_delimited_math(command):
            continue
        if match.start() > cursor:
            nodes.extend(_markdown_text_nodes(value[cursor : match.start()]))
        nodes.append(_inline_math_node(command))
        cursor = match.end()
    if cursor < len(value):
        nodes.extend(_markdown_text_nodes(value[cursor:]))
    return nodes


def _append_text_or_formula_nodes(nodes: list[dict[str, Any]], value: str) -> bool:
    if not value:
        return False
    stripped = value.strip()
    leading = value[: len(value) - len(value.lstrip())]
    trailing = value[len(value.rstrip()) :]
    if stripped and _is_likely_delimited_math(stripped):
        if leading:
            nodes.extend(_markdown_text_nodes(leading))
        nodes.append(_inline_math_node(stripped))
        if trailing:
            nodes.extend(_markdown_text_nodes(trailing))
        return True
    nodes.extend(_markdown_text_nodes_with_raw_math(value))
    return False


def _mixed_math_text_nodes(value: str) -> list[dict[str, Any]]:
    nodes: list[dict[str, Any]] = []
    cursor = 0
    has_formula = False
    for match in _MIXED_MATH_TEXT_RE.finditer(value):
        has_formula = _append_text_or_formula_nodes(nodes, value[cursor : match.start()]) or has_formula
        nodes.extend(_markdown_text_nodes(match.group(0)))
        cursor = match.end()
    has_formula = _append_text_or_formula_nodes(nodes, value[cursor:]) or has_formula
    return nodes if has_formula else []


def _inline_nodes(value: str) -> list[dict[str, Any]]:
    nodes: list[dict[str, Any]] = []
    cursor = 0
    for match in _DELIMITED_MATH_RE.finditer(value):
        if match.start() > cursor:
            nodes.extend(_markdown_text_nodes_with_raw_math(value[cursor : match.start()]))
        raw = match.group(0)
        latex = match.group(1) or match.group(2) or match.group(3) or match.group(4) or ""
        if latex.strip() and _is_likely_delimited_math(latex):
            nodes.append(_inline_math_node(latex))
        else:
            nodes.extend(_mixed_math_text_nodes(latex) or _markdown_text_nodes_with_raw_math(raw))
        cursor = match.end()
    if cursor < len(value):
        nodes.extend(_markdown_text_nodes_with_raw_math(value[cursor:]))
    return nodes


def _inline_html(value: str) -> str:
    parts: list[str] = []
    for node in _inline_nodes(value):
        if node.get("type") == "inlineMath":
            latex = str(node.get("attrs", {}).get("latex") or "")
            parts.append(f'<span data-type="inline-math" data-latex="{html.escape(latex, quote=True)}"></span>')
            continue
        text = html.escape(str(node.get("text") or ""))
        mark_names = {mark.get("type") for mark in node.get("marks", []) if isinstance(mark, dict)}
        if "bold" in mark_names:
            text = f"<strong>{text}</strong>"
        if "italic" in mark_names:
            text = f"<em>{text}</em>"
        parts.append(text)
    return "".join(parts)


def _paragraph_node(value: str) -> dict[str, Any]:
    content = _inline_nodes(value)
    if not content:
        return {"type": "paragraph"}
    return {"type": "paragraph", "content": content}


def _split_markdown_table_row(value: str) -> list[str]:
    stripped = value.strip().strip("|")
    return [cell.strip() for cell in stripped.split("|")]


def _is_markdown_table(lines: list[str], index: int) -> bool:
    if index + 1 >= len(lines):
        return False
    header = lines[index].strip()
    separator = lines[index + 1].strip()
    return "|" in header and bool(_MARKDOWN_TABLE_SEPARATOR_RE.match(separator))


def _display_math_block(lines: list[str], index: int) -> tuple[str, int] | None:
    line = lines[index].strip()
    delimiters = [(r"\[", r"\]"), ("$$", "$$")]
    for opener, closer in delimiters:
        if line == opener:
            formula_lines: list[str] = []
            cursor = index + 1
            while cursor < len(lines):
                current = lines[cursor].strip()
                if current == closer:
                    latex = "\n".join(formula_lines).strip()
                    return (latex, cursor + 1) if latex and _is_likely_delimited_math(latex) else None
                formula_lines.append(lines[cursor])
                cursor += 1
            return None
        if line.startswith(opener) and line.endswith(closer) and len(line) > len(opener) + len(closer):
            latex = line[len(opener) : -len(closer)].strip()
            return (latex, index + 1) if latex and _is_likely_delimited_math(latex) else None
    return None


def _table_node(rows: list[list[str]]) -> dict[str, Any]:
    table_rows: list[dict[str, Any]] = []
    for row_index, row in enumerate(rows):
        cell_type = "tableHeader" if row_index == 0 else "tableCell"
        table_rows.append(
            {
                "type": "tableRow",
                "content": [
                    {
                        "type": cell_type,
                        "content": [_paragraph_node(cell)],
                    }
                    for cell in row
                ],
            }
        )
    return {"type": "table", "content": table_rows}


def _table_html(rows: list[list[str]]) -> str:
    html_rows: list[str] = []
    for row_index, row in enumerate(rows):
        tag = "th" if row_index == 0 else "td"
        html_rows.append("<tr>" + "".join(f"<{tag}>{_inline_html(cell)}</{tag}>" for cell in row) + "</tr>")
    return "<table><tbody>" + "".join(html_rows) + "</tbody></table>"


def text_to_html(content_text: str) -> str:
    parts: list[str] = []
    lines = content_text.splitlines()
    index = 0
    while index < len(lines):
        line = lines[index].strip()
        if not line:
            index += 1
            continue

        display_math = _display_math_block(lines, index)
        if display_math:
            latex, index = display_math
            parts.append(
                f'<div data-type="block-math" data-latex="{html.escape(_normalize_latex(latex), quote=True)}"></div>'
            )
            continue

        if _is_markdown_table(lines, index):
            rows = [_split_markdown_table_row(line)]
            index += 2
            while index < len(lines) and "|" in lines[index].strip():
                rows.append(_split_markdown_table_row(lines[index]))
                index += 1
            parts.append(_table_html(rows))
            continue

        heading_match = _MARKDOWN_HEADING_RE.match(line)
        if heading_match:
            level = min(len(heading_match.group(1)), 3)
            parts.append(f"<h{level}>{_inline_html(heading_match.group(2).strip())}</h{level}>")
            index += 1
            continue

        bullet_match = _MARKDOWN_BULLET_RE.match(line)
        if bullet_match:
            items: list[str] = []
            while index < len(lines):
                item_match = _MARKDOWN_BULLET_RE.match(lines[index].strip())
                if not item_match:
                    break
                items.append(item_match.group(1).strip())
                index += 1
            parts.append("<ul>" + "".join(f"<li>{_inline_html(item)}</li>" for item in items) + "</ul>")
            continue

        ordered_match = _MARKDOWN_ORDERED_RE.match(line)
        if ordered_match:
            items = []
            while index < len(lines):
                item_match = _MARKDOWN_ORDERED_RE.match(lines[index].strip())
                if not item_match:
                    break
                items.append(item_match.group(1).strip())
                index += 1
            parts.append("<ol>" + "".join(f"<li>{_inline_html(item)}</li>" for item in items) + "</ol>")
            continue

        if line.startswith(">"):
            parts.append(f"<blockquote>{_inline_html(line.lstrip('>').strip())}</blockquote>")
        else:
            parts.append(f"<p>{_inline_html(line)}</p>")
        index += 1
    return "\n".join(parts) or "<p></p>"


def text_to_tiptap_doc(content_text: str) -> dict[str, Any]:
    nodes: list[dict[str, Any]] = []
    lines = content_text.splitlines()
    index = 0
    while index < len(lines):
        line = lines[index].strip()
        if not line:
            index += 1
            continue

        display_math = _display_math_block(lines, index)
        if display_math:
            latex, index = display_math
            nodes.append({"type": "blockMath", "attrs": {"latex": _normalize_latex(latex)}})
            continue

        if _is_markdown_table(lines, index):
            rows = [_split_markdown_table_row(line)]
            index += 2
            while index < len(lines) and "|" in lines[index].strip():
                rows.append(_split_markdown_table_row(lines[index]))
                index += 1
            nodes.append(_table_node(rows))
            continue

        heading_match = _MARKDOWN_HEADING_RE.match(line)
        if heading_match:
            level = min(len(heading_match.group(1)), 3)
            nodes.append(
                {
                    "type": "heading",
                    "attrs": {"level": level},
                    "content": _inline_nodes(heading_match.group(2).strip()),
                }
            )
            index += 1
            continue

        bullet_match = _MARKDOWN_BULLET_RE.match(line)
        if bullet_match:
            items: list[dict[str, Any]] = []
            while index < len(lines):
                item_match = _MARKDOWN_BULLET_RE.match(lines[index].strip())
                if not item_match:
                    break
                items.append({"type": "listItem", "content": [_paragraph_node(item_match.group(1).strip())]})
                index += 1
            nodes.append({"type": "bulletList", "content": items})
            continue

        ordered_match = _MARKDOWN_ORDERED_RE.match(line)
        if ordered_match:
            items = []
            while index < len(lines):
                item_match = _MARKDOWN_ORDERED_RE.match(lines[index].strip())
                if not item_match:
                    break
                items.append({"type": "listItem", "content": [_paragraph_node(item_match.group(1).strip())]})
                index += 1
            nodes.append({"type": "orderedList", "content": items})
            continue

        if line.startswith(">"):
            nodes.append({"type": "blockquote", "content": [_paragraph_node(line.lstrip(">").strip())]})
        else:
            nodes.append(_paragraph_node(line))
        index += 1
    return {"type": "doc", "content": nodes or [{"type": "paragraph"}]}


class _HTMLTreeParser(HTMLParser):
    def __init__(self) -> None:
        super().__init__(convert_charrefs=True)
        self.root: dict[str, Any] = {"tag": "root", "attrs": {}, "children": []}
        self._stack: list[dict[str, Any]] = [self.root]

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        normalized_tag = tag.lower()
        node = {
            "tag": normalized_tag,
            "attrs": {key.lower(): value or "" for key, value in attrs},
            "children": [],
        }
        self._stack[-1]["children"].append(node)
        if normalized_tag not in _VOID_HTML_TAGS:
            self._stack.append(node)

    def handle_endtag(self, tag: str) -> None:
        normalized_tag = tag.lower()
        for index in range(len(self._stack) - 1, 0, -1):
            if self._stack[index].get("tag") == normalized_tag:
                self._stack = self._stack[:index]
                return

    def handle_data(self, data: str) -> None:
        if data:
            self._stack[-1]["children"].append(data)


def _style_map(attrs: dict[str, str]) -> dict[str, str]:
    styles: dict[str, str] = {}
    for item in attrs.get("style", "").split(";"):
        if ":" not in item:
            continue
        key, value = item.split(":", 1)
        key = key.strip().lower()
        value = value.strip()
        if key and value:
            styles[key] = value
    return styles


def _block_attrs(attrs: dict[str, str]) -> dict[str, Any]:
    styles = _style_map(attrs)
    text_align = styles.get("text-align") or attrs.get("align", "")
    if text_align in {"left", "center", "right", "justify"}:
        return {"textAlign": text_align}
    return {}


def _text_style_attrs(attrs: dict[str, str]) -> dict[str, str]:
    styles = _style_map(attrs)
    text_style: dict[str, str] = {}
    if styles.get("font-size"):
        text_style["fontSize"] = styles["font-size"]
    if styles.get("font-family"):
        text_style["fontFamily"] = styles["font-family"]
    return text_style


def _with_mark(marks: list[dict[str, Any]], mark_type: str, attrs: dict[str, Any] | None = None) -> list[dict[str, Any]]:
    next_marks = [dict(mark) for mark in marks]
    if mark_type == "textStyle" and attrs:
        for mark in next_marks:
            if mark.get("type") == "textStyle":
                mark["attrs"] = {**mark.get("attrs", {}), **attrs}
                return next_marks
    if not any(mark.get("type") == mark_type for mark in next_marks):
        mark: dict[str, Any] = {"type": mark_type}
        if attrs:
            mark["attrs"] = attrs
        next_marks.append(mark)
    return next_marks


def _text_node(text: str, marks: list[dict[str, Any]]) -> dict[str, Any] | None:
    if not text:
        return None
    node: dict[str, Any] = {"type": "text", "text": text}
    if marks:
        node["marks"] = [dict(mark) for mark in marks]
    return node


def _trim_inline_content(nodes: list[dict[str, Any]]) -> list[dict[str, Any]]:
    trimmed = [dict(node) for node in nodes]
    while trimmed and trimmed[0].get("type") == "text":
        text = str(trimmed[0].get("text") or "").lstrip()
        if text:
            trimmed[0]["text"] = text
            break
        trimmed.pop(0)
    while trimmed and trimmed[-1].get("type") == "text":
        text = str(trimmed[-1].get("text") or "").rstrip()
        if text:
            trimmed[-1]["text"] = text
            break
        trimmed.pop()
    return trimmed


def _text_block_node(node_type: str, attrs: dict[str, Any], content: list[dict[str, Any]]) -> dict[str, Any]:
    node: dict[str, Any] = {"type": node_type}
    if attrs:
        node["attrs"] = attrs
    trimmed = _trim_inline_content(content)
    if trimmed:
        node["content"] = trimmed
    return node


def _html_inline_nodes(children: list[Any], marks: list[dict[str, Any]] | None = None) -> list[dict[str, Any]]:
    active_marks = marks or []
    nodes: list[dict[str, Any]] = []
    for child in children:
        if isinstance(child, str):
            text = child.replace("\xa0", " ")
            text_node = _text_node(text, active_marks)
            if text_node:
                nodes.append(text_node)
            continue
        if not isinstance(child, dict):
            continue
        tag = child.get("tag", "")
        attrs = child.get("attrs", {})
        child_children = child.get("children", [])
        node_type = (attrs.get("data-type") or "").strip()
        if node_type == "inline-math":
            latex = html.unescape((attrs.get("data-latex") or "").strip())
            if latex and _is_likely_delimited_math(latex):
                nodes.append({"type": "inlineMath", "attrs": {"latex": latex}})
            elif latex:
                text_node = _text_node(latex, active_marks)
                if text_node:
                    nodes.append(text_node)
            continue
        if tag == "br":
            nodes.append({"type": "hardBreak"})
            continue
        child_marks = active_marks
        if tag in {"strong", "b"}:
            child_marks = _with_mark(child_marks, "bold")
        elif tag in {"em", "i"}:
            child_marks = _with_mark(child_marks, "italic")
        elif tag == "u":
            child_marks = _with_mark(child_marks, "underline")
        elif tag in {"s", "strike", "del"}:
            child_marks = _with_mark(child_marks, "strike")
        elif tag == "code":
            child_marks = _with_mark(child_marks, "code")
        text_style = _text_style_attrs(attrs)
        if text_style:
            child_marks = _with_mark(child_marks, "textStyle", text_style)
        nodes.extend(_html_inline_nodes(child_children, child_marks))
    return nodes


def _has_block_children(children: list[Any]) -> bool:
    return any(isinstance(child, dict) and child.get("tag") in _HTML_BLOCK_TAGS for child in children)


def _html_table_cell_node(cell: dict[str, Any]) -> dict[str, Any]:
    cell_type = "tableHeader" if cell.get("tag") == "th" else "tableCell"
    content = _html_children_to_blocks(cell.get("children", []))
    if not content:
        content = [_text_block_node("paragraph", {}, _html_inline_nodes(cell.get("children", [])))]
    return {"type": cell_type, "content": content}


def _html_table_rows(children: list[Any]) -> list[dict[str, Any]]:
    rows: list[dict[str, Any]] = []
    for child in children:
        if not isinstance(child, dict):
            continue
        tag = child.get("tag")
        if tag in {"thead", "tbody", "tfoot"}:
            rows.extend(_html_table_rows(child.get("children", [])))
            continue
        if tag != "tr":
            continue
        cells = [
            _html_table_cell_node(cell)
            for cell in child.get("children", [])
            if isinstance(cell, dict) and cell.get("tag") in {"th", "td"}
        ]
        if cells:
            rows.append({"type": "tableRow", "content": cells})
    return rows


def _html_list_item_node(item: dict[str, Any]) -> dict[str, Any]:
    content: list[dict[str, Any]] = []
    inline_children: list[Any] = []
    for child in item.get("children", []):
        if isinstance(child, str):
            inline_children.append(child)
            continue
        if not isinstance(child, dict):
            continue
        if child.get("tag") in _HTML_BLOCK_TAGS:
            if inline_children:
                paragraph = _text_block_node("paragraph", {}, _html_inline_nodes(inline_children))
                if paragraph.get("content"):
                    content.append(paragraph)
                inline_children = []
            content.extend(_html_node_to_blocks(child))
        else:
            inline_children.append(child)
    if inline_children:
        paragraph = _text_block_node("paragraph", {}, _html_inline_nodes(inline_children))
        if paragraph.get("content"):
            content.append(paragraph)
    return {"type": "listItem", "content": content or [{"type": "paragraph"}]}


def _html_node_to_blocks(node: dict[str, Any]) -> list[dict[str, Any]]:
    tag = node.get("tag", "")
    attrs = node.get("attrs", {})
    children = node.get("children", [])
    node_type = (attrs.get("data-type") or "").strip()
    if node_type == "page-break":
        return [{"type": "pageBreak"}]
    if node_type == "block-math":
        latex = html.unescape((attrs.get("data-latex") or "").strip())
        if latex and _is_likely_delimited_math(latex):
            return [{"type": "blockMath", "attrs": {"latex": latex}}]
        if latex:
            return [{"type": "paragraph", "content": [{"type": "text", "text": latex}]}]
        return []
    if tag in {"h1", "h2", "h3", "h4", "h5", "h6"}:
        level = min(int(tag[1]), 3)
        return [
            _text_block_node(
                "heading",
                {"level": level, **_block_attrs(attrs)},
                _html_inline_nodes(children),
            )
        ]
    if tag == "p":
        return [_text_block_node("paragraph", _block_attrs(attrs), _html_inline_nodes(children))]
    if tag == "blockquote":
        content = _html_children_to_blocks(children) if _has_block_children(children) else [
            _text_block_node("paragraph", {}, _html_inline_nodes(children))
        ]
        return [{"type": "blockquote", "content": content or [{"type": "paragraph"}]}]
    if tag in {"ul", "ol"}:
        items = [
            _html_list_item_node(child)
            for child in children
            if isinstance(child, dict) and child.get("tag") == "li"
        ]
        return [{"type": "bulletList" if tag == "ul" else "orderedList", "content": items}] if items else []
    if tag == "table":
        rows = _html_table_rows(children)
        return [{"type": "table", "content": rows}] if rows else []
    if tag == "tr":
        cells = [
            _html_table_cell_node(child)
            for child in children
            if isinstance(child, dict) and child.get("tag") in {"th", "td"}
        ]
        return [{"type": "tableRow", "content": cells}] if cells else []
    if tag in {"thead", "tbody", "tfoot"}:
        return _html_children_to_blocks(children)
    if tag in {"td", "th"}:
        return [_html_table_cell_node(node)]
    if tag == "li":
        return [_html_list_item_node(node)]
    if tag == "img":
        src = (attrs.get("src") or "").strip()
        alt = (attrs.get("alt") or "").strip()
        return [{"type": "image", "attrs": {"src": src, "alt": alt}}] if src else []
    if _has_block_children(children):
        return _html_children_to_blocks(children)
    inline = _html_inline_nodes(children)
    if inline:
        return [_text_block_node("paragraph", {}, inline)]
    return []


def _html_children_to_blocks(children: list[Any]) -> list[dict[str, Any]]:
    nodes: list[dict[str, Any]] = []
    inline_children: list[Any] = []
    for child in children:
        if isinstance(child, str):
            if child.strip():
                inline_children.append(child)
            continue
        if not isinstance(child, dict):
            continue
        if child.get("tag") in _HTML_BLOCK_TAGS or (child.get("attrs", {}).get("data-type") or "").strip() in {
            "block-math",
            "page-break",
        }:
            if inline_children:
                paragraph = _text_block_node("paragraph", {}, _html_inline_nodes(inline_children))
                if paragraph.get("content"):
                    nodes.append(paragraph)
                inline_children = []
            nodes.extend(_html_node_to_blocks(child))
        else:
            inline_children.append(child)
    if inline_children:
        paragraph = _text_block_node("paragraph", {}, _html_inline_nodes(inline_children))
        if paragraph.get("content"):
            nodes.append(paragraph)
    return nodes


def html_to_tiptap_doc(content_html: str) -> dict[str, Any]:
    parser = _HTMLTreeParser()
    parser.feed(content_html)
    nodes = _html_children_to_blocks(parser.root["children"])
    return {"type": "doc", "content": nodes or [{"type": "paragraph"}]}


def _markdown_escape(value: str) -> str:
    return value.replace("\\", "\\\\").replace("|", "\\|")


def _markdown_inline(nodes: list[dict[str, Any]]) -> str:
    parts: list[str] = []
    for node in nodes:
        node_type = node.get("type")
        if node_type == "text":
            text = str(node.get("text") or "")
            mark_types = [
                mark.get("type")
                for mark in node.get("marks", [])
                if isinstance(mark, dict)
            ]
            if "code" in mark_types:
                text = f"`{text}`"
            if "bold" in mark_types:
                text = f"**{text}**"
            if "italic" in mark_types:
                text = f"*{text}*"
            parts.append(text)
        elif node_type == "hardBreak":
            parts.append("\n")
        elif node_type == "inlineMath":
            latex = str(node.get("attrs", {}).get("latex") or "").strip()
            if latex:
                parts.append(f"${latex}$")
        elif isinstance(node.get("content"), list):
            parts.append(_markdown_inline(node["content"]))
    return "".join(parts)


def _markdown_blocks(nodes: list[dict[str, Any]], *, list_depth: int = 0) -> list[str]:
    blocks: list[str] = []
    for node in nodes:
        node_type = node.get("type")
        content = node.get("content", [])
        if not isinstance(content, list):
            content = []
        if node_type == "heading":
            level = int(node.get("attrs", {}).get("level") or 1)
            blocks.append(f"{'#' * min(max(level, 1), 6)} {_markdown_inline(content).strip()}".rstrip())
        elif node_type == "paragraph":
            blocks.append(_markdown_inline(content).strip())
        elif node_type == "blockquote":
            quoted = _markdown_blocks(content)
            blocks.extend(f"> {line}" if line else ">" for line in quoted)
        elif node_type in {"bulletList", "orderedList"}:
            for index, item in enumerate(content, start=1):
                if not isinstance(item, dict):
                    continue
                item_blocks = _markdown_blocks(item.get("content", []), list_depth=list_depth + 1)
                if not item_blocks:
                    continue
                marker = "-" if node_type == "bulletList" else f"{index}."
                indent = "  " * list_depth
                blocks.append(f"{indent}{marker} {item_blocks[0]}")
                blocks.extend(f"{indent}  {line}" for line in item_blocks[1:])
        elif node_type == "table":
            rows: list[list[str]] = []
            for row in content:
                if not isinstance(row, dict) or row.get("type") != "tableRow":
                    continue
                cells: list[str] = []
                for cell in row.get("content", []):
                    if not isinstance(cell, dict):
                        continue
                    cell_text = " ".join(
                        line.strip()
                        for line in _markdown_blocks(cell.get("content", []))
                        if line.strip()
                    )
                    cells.append(_markdown_escape(cell_text))
                if cells:
                    rows.append(cells)
            if rows:
                width = max(len(row) for row in rows)
                normalized_rows = [row + [""] * (width - len(row)) for row in rows]
                blocks.append("| " + " | ".join(normalized_rows[0]) + " |")
                blocks.append("| " + " | ".join("---" for _ in range(width)) + " |")
                for row in normalized_rows[1:]:
                    blocks.append("| " + " | ".join(row) + " |")
        elif node_type == "blockMath":
            latex = str(node.get("attrs", {}).get("latex") or "").strip()
            if latex:
                blocks.append(f"$$\n{latex}\n$$")
        elif node_type == "pageBreak":
            blocks.append("---")
        elif node_type == "image":
            attrs = node.get("attrs", {})
            src = str(attrs.get("src") or "").strip()
            alt = str(attrs.get("alt") or "").strip()
            if src:
                blocks.append(f"![{alt}]({src})")
        elif content:
            blocks.extend(_markdown_blocks(content, list_depth=list_depth))
    return blocks


def html_to_markdown(content_html: str) -> str:
    parsed = html_to_tiptap_doc(content_html)
    parsed_content = parsed.get("content", [])
    if isinstance(parsed_content, list):
        markdown = "\n\n".join(line for line in _markdown_blocks(parsed_content) if line is not None).strip()
        if markdown:
            return markdown
    return html_to_text(content_html)


def looks_like_html_content(value: str) -> bool:
    return bool(_HTML_CONTENT_RE.search(value))


def document_to_markdown(document: BoardDocument) -> str:
    source_json = document.content_json if isinstance(document.content_json, dict) else {}
    content = source_json.get("content")
    if isinstance(content, list) and content:
        markdown = "\n\n".join(line for line in _markdown_blocks(content) if line is not None).strip()
        if markdown:
            return markdown
    if document.content_html.strip():
        markdown = html_to_markdown(document.content_html)
        if markdown:
            return markdown
    return document.content_text.strip()


def _sanitize_suspicious_math_node(node: dict[str, Any]) -> tuple[dict[str, Any] | None, bool]:
    node_type = node.get("type")
    if node_type == "inlineMath":
        latex = str(node.get("attrs", {}).get("latex") or "").strip()
        if latex and not _is_likely_delimited_math(latex):
            return {"type": "text", "text": latex}, True
        return dict(node), False

    if node_type == "blockMath":
        latex = str(node.get("attrs", {}).get("latex") or "").strip()
        if latex and not _is_likely_delimited_math(latex):
            return {"type": "paragraph", "content": [{"type": "text", "text": latex}]}, True
        return dict(node), False

    sanitized = dict(node)
    content = sanitized.get("content")
    changed = False
    if isinstance(content, list):
        next_content: list[dict[str, Any]] = []
        for child in content:
            if not isinstance(child, dict):
                next_content.append(child)
                continue
            next_child, child_changed = _sanitize_suspicious_math_node(child)
            changed = changed or child_changed
            if next_child is not None:
                next_content.append(next_child)
        sanitized["content"] = next_content
    return sanitized, changed


def _sanitize_suspicious_math_json(content_json: dict[str, Any]) -> tuple[dict[str, Any], bool]:
    sanitized = dict(content_json)
    content = sanitized.get("content")
    if not isinstance(content, list):
        return sanitized, False

    changed = False
    next_content: list[dict[str, Any]] = []
    for child in content:
        if not isinstance(child, dict):
            next_content.append(child)
            continue
        next_child, child_changed = _sanitize_suspicious_math_node(child)
        changed = changed or child_changed
        if next_child is not None:
            next_content.append(next_child)
    sanitized["content"] = next_content
    return sanitized, changed


def _html_attr(raw_tag: str, name: str) -> str:
    match = re.search(rf"\b{re.escape(name)}\s*=\s*(\"([^\"]*)\"|'([^']*)')", raw_tag, flags=re.IGNORECASE)
    if not match:
        return ""
    return html.unescape(match.group(2) if match.group(2) is not None else match.group(3) or "")


def _repair_suspicious_math_html(content_html: str) -> str:
    def inline_replacement(match: re.Match[str]) -> str:
        raw = match.group(0)
        latex = _html_attr(raw, "data-latex").strip()
        node_type = _html_attr(raw, "data-type").strip()
        if node_type == "inline-math" and latex and not _is_likely_delimited_math(latex):
            return html.escape(latex)
        return raw

    def block_replacement(match: re.Match[str]) -> str:
        raw = match.group(0)
        latex = _html_attr(raw, "data-latex").strip()
        node_type = _html_attr(raw, "data-type").strip()
        if node_type == "block-math" and latex and not _is_likely_delimited_math(latex):
            return f"<p>{html.escape(latex)}</p>"
        return raw

    repaired = re.sub(
        r"<span\b(?=[^>]*data-type=['\"]inline-math['\"])[^>]*>[\s\S]*?</span>",
        inline_replacement,
        content_html,
        flags=re.IGNORECASE,
    )
    repaired = re.sub(
        r"<div\b(?=[^>]*data-type=['\"]block-math['\"])[^>]*>[\s\S]*?</div>",
        block_replacement,
        repaired,
        flags=re.IGNORECASE,
    )
    return _repair_raw_math_text_html(repaired)


def _repair_raw_math_text_html(content_html: str) -> str:
    pieces = re.split(r"(<[^>]+>)", content_html)
    repaired: list[str] = []
    for piece in pieces:
        if not piece:
            continue
        if piece.startswith("<") and piece.endswith(">"):
            repaired.append(piece)
            continue
        repaired.append(_inline_html(html.unescape(piece)))
    return "".join(repaired)


def _html_has_math_nodes(content_html: str) -> bool:
    return any(
        marker in content_html
        for marker in (
            'data-type="inline-math"',
            "data-type='inline-math'",
            'data-type="block-math"',
            "data-type='block-math'",
        )
    )


def _html_has_visible_raw_math_text(content_html: str) -> bool:
    text = html_to_text(content_html)
    return "$" in text or "\\(" in text or "\\[" in text or bool(_RAW_LATEX_COMMAND_RE.search(text))


def _json_has_raw_math_text(node: dict[str, Any]) -> bool:
    if node.get("type") == "text":
        text = str(node.get("text") or "")
        return "$" in text or "\\(" in text or "\\[" in text or bool(_RAW_LATEX_COMMAND_RE.search(text))
    content = node.get("content")
    if not isinstance(content, list):
        return False
    return any(_json_has_raw_math_text(child) for child in content if isinstance(child, dict))


def build_document(
    *,
    title: str,
    content_html: str | None = None,
    content_text: str | None = None,
    content_json: dict[str, Any] | None = None,
    document_id: str | None = None,
    page_settings: DocumentPageSettings | dict[str, Any] | None = None,
) -> BoardDocument:
    normalized_html = (content_html or "").strip()
    normalized_text = (content_text or "").strip()
    repaired_html = False
    if normalized_html:
        next_html = _repair_suspicious_math_html(normalized_html)
        repaired_html = next_html != normalized_html
        normalized_html = next_html
    if not normalized_text and normalized_html:
        normalized_text = html_to_text(normalized_html)
    if not normalized_html and normalized_text:
        normalized_html = text_to_html(normalized_text)
    if not normalized_text and not normalized_html:
        normalized_html = "<p></p>"
    stale_json = bool(
        content_json
        and normalized_html
        and _html_has_math_nodes(normalized_html)
        and _json_has_raw_math_text(content_json)
    )
    rebuild_json_from_html = stale_json or (repaired_html and _html_has_math_nodes(normalized_html))
    normalized_json = content_json if content_json and not rebuild_json_from_html else (
        html_to_tiptap_doc(normalized_html) if normalized_html.strip() else text_to_tiptap_doc(normalized_text)
    )
    if isinstance(normalized_json, dict):
        normalized_json, repaired_math = _sanitize_suspicious_math_json(normalized_json)
        if repaired_math:
            normalized_html = _repair_suspicious_math_html(normalized_html)
    kwargs: dict[str, Any] = {
        "title": title,
        "content_json": normalized_json,
        "content_html": normalized_html,
        "content_text": normalized_text,
        "page_settings": page_settings or DocumentPageSettings(),
    }
    if document_id:
        kwargs["id"] = document_id
    return BoardDocument(**kwargs)


def _looks_like_markdown_document(content_text: str) -> bool:
    lines = [line.strip() for line in content_text.splitlines() if line.strip()]
    if not lines:
        return False
    for index, line in enumerate(lines):
        if _display_math_block(lines, index):
            return True
        if _MARKDOWN_HEADING_RE.match(line) or _MARKDOWN_BULLET_RE.match(line) or _MARKDOWN_ORDERED_RE.match(line):
            return True
        if _MARKDOWN_INLINE_RE.search(line):
            return True
        if _is_markdown_table(lines, index):
            return True
    return False


def upgrade_markdown_like_document(document: BoardDocument) -> BoardDocument:
    existing_document = _repair_existing_document(document)
    if not _looks_like_markdown_document(document.content_text):
        return existing_document
    upgraded = build_document(
        title=document.title,
        content_text=document.content_text,
        document_id=document.id,
        page_settings=document.page_settings,
    )
    if _would_downgrade_existing_rich_structure(existing_document, upgraded):
        return existing_document
    return upgraded


def _repair_existing_document(document: BoardDocument) -> BoardDocument:
    content_json = document.content_json if isinstance(document.content_json, dict) else {}
    sanitized_json, repaired_math = _sanitize_suspicious_math_json(content_json)
    repaired_html = _repair_suspicious_math_html(document.content_html)
    stale_json = _html_has_math_nodes(repaired_html) and _json_has_raw_math_text(sanitized_json)
    rebuild_json_from_html = stale_json or (
        repaired_html != document.content_html and _html_has_math_nodes(repaired_html)
    )
    if rebuild_json_from_html:
        sanitized_json = html_to_tiptap_doc(repaired_html)
        sanitized_json, json_repaired_math = _sanitize_suspicious_math_json(sanitized_json)
        repaired_math = repaired_math or json_repaired_math
    if not repaired_math and repaired_html == document.content_html and not stale_json:
        return document
    return BoardDocument(
        id=document.id,
        title=document.title,
        content_json=sanitized_json,
        content_html=repaired_html,
        content_text=document.content_text,
        page_settings=document.page_settings,
    )


def _would_downgrade_existing_rich_structure(current_document: BoardDocument, upgraded_document: BoardDocument) -> bool:
    current_counts = rich_structure_counts(current_document)
    current_score = rich_structure_score(current_counts)
    if current_score < 8:
        return False
    upgraded_counts = rich_structure_counts(upgraded_document)
    upgraded_score = rich_structure_score(upgraded_counts)
    if would_flatten_rich_document(current_document=current_document, new_document=upgraded_document):
        return True
    return upgraded_score < current_score


def is_document_empty(document: BoardDocument) -> bool:
    return not document.content_text.strip() and html_to_text(document.content_html) == ""


def document_changed(left: BoardDocument, right: BoardDocument) -> bool:
    return (
        left.title != right.title
        or left.content_html.strip() != right.content_html.strip()
        or left.content_text.strip() != right.content_text.strip()
        or left.page_settings.model_dump(mode="json") != right.page_settings.model_dump(mode="json")
    )


def rich_structure_counts(document: BoardDocument) -> dict[str, int]:
    counts = {
        "heading": 0,
        "bold": 0,
        "italic": 0,
        "bulletList": 0,
        "orderedList": 0,
        "listItem": 0,
        "table": 0,
        "blockquote": 0,
        "paragraph": 0,
    }

    def walk(value: Any) -> None:
        if isinstance(value, dict):
            node_type = value.get("type")
            if isinstance(node_type, str) and node_type in counts:
                counts[node_type] += 1
            marks = value.get("marks")
            if isinstance(marks, list):
                for mark in marks:
                    if isinstance(mark, dict):
                        mark_type = mark.get("type")
                        if isinstance(mark_type, str) and mark_type in counts:
                            counts[mark_type] += 1
            for child in value.values():
                walk(child)
        elif isinstance(value, list):
            for child in value:
                walk(child)

    walk(document.content_json if isinstance(document.content_json, dict) else {})
    return counts


def rich_structure_score(counts: dict[str, int]) -> int:
    return (
        counts.get("heading", 0) * 3
        + counts.get("table", 0) * 4
        + counts.get("bulletList", 0) * 2
        + counts.get("orderedList", 0) * 2
        + counts.get("listItem", 0)
        + counts.get("blockquote", 0) * 2
        + counts.get("bold", 0)
        + counts.get("italic", 0)
    )


def _compact_node_text(value: Any) -> str:
    parts: list[str] = []

    def walk(node: Any) -> None:
        if isinstance(node, dict):
            node_type = node.get("type")
            if node_type == "text":
                text = node.get("text")
                if isinstance(text, str):
                    parts.append(text)
                return
            if node_type in {"inlineMath", "blockMath"}:
                attrs = node.get("attrs")
                if isinstance(attrs, dict) and isinstance(attrs.get("latex"), str):
                    parts.append(attrs["latex"])
                return
            content = node.get("content")
            if isinstance(content, list):
                for child in content:
                    walk(child)
        elif isinstance(node, list):
            for child in node:
                walk(child)

    walk(value)
    return re.sub(r"\s+", " ", "".join(parts)).strip()


def _leading_document_blocks(document: BoardDocument, *, limit: int = 16) -> list[tuple[str, str]]:
    content_json = document.content_json if isinstance(document.content_json, dict) else {}
    content = content_json.get("content")
    if not isinstance(content, list):
        return []
    blocks: list[tuple[str, str]] = []
    for node in content:
        if not isinstance(node, dict):
            continue
        node_type = node.get("type")
        if not isinstance(node_type, str):
            continue
        text = _compact_node_text(node)
        if not text and node_type == "paragraph":
            continue
        blocks.append((node_type, text))
        if len(blocks) >= limit:
            break
    return blocks


def _would_degrade_leading_heading_hierarchy(current_document: BoardDocument, new_document: BoardDocument) -> bool:
    current_blocks = _leading_document_blocks(current_document)
    new_blocks = _leading_document_blocks(new_document)
    if not current_blocks or not new_blocks:
        return False

    current_heading_indexes = [index for index, (node_type, _) in enumerate(current_blocks[:8]) if node_type == "heading"]
    if not current_heading_indexes or current_heading_indexes[0] > 2:
        return False

    new_heading_indexes = [index for index, (node_type, _) in enumerate(new_blocks[:8]) if node_type == "heading"]
    if new_heading_indexes and new_heading_indexes[0] <= current_heading_indexes[0] + 1:
        return False

    current_heading_count = len(current_heading_indexes)
    new_heading_count = len(new_heading_indexes)
    new_leading_plain_count = sum(1 for node_type, text in new_blocks[:8] if node_type == "paragraph" and text)
    if new_heading_count == 0 and new_leading_plain_count >= 2:
        return True
    return current_heading_count >= 2 and new_heading_count * 2 < current_heading_count and new_blocks[0][0] != "heading"


def would_flatten_rich_document(
    *,
    current_document: BoardDocument,
    new_document: BoardDocument,
    operation: str | None = None,
) -> bool:
    if operation is not None and operation != "replace_document":
        return False
    if is_document_empty(current_document):
        return False

    old_counts = rich_structure_counts(current_document)
    old_score = rich_structure_score(old_counts)
    if old_score < 8:
        return False

    new_counts = rich_structure_counts(new_document)
    if _would_degrade_leading_heading_hierarchy(current_document, new_document):
        return True
    if new_counts.get("heading", 0) or new_counts.get("table", 0):
        return False

    old_primary_structure = old_counts.get("heading", 0) + old_counts.get("table", 0)
    if old_primary_structure <= 0:
        return False

    heading_hierarchy_lost = old_counts.get("heading", 0) >= 2 and new_counts.get("heading", 0) == 0
    table_structure_lost = old_counts.get("table", 0) > 0 and new_counts.get("table", 0) == 0
    if heading_hierarchy_lost or table_structure_lost:
        return True

    new_score = rich_structure_score(new_counts)
    structure_dropped = new_score <= max(2, int(old_score * 0.5))
    paragraph_heavy = new_counts.get("paragraph", 0) >= max(8, old_counts.get("paragraph", 0) // 2)
    return structure_dropped and paragraph_heavy


def append_html_section(document: BoardDocument, section_html: str) -> BoardDocument:
    next_html = "\n".join(part for part in [document.content_html.strip(), section_html.strip()] if part)
    return build_document(
        title=document.title,
        content_html=next_html,
        document_id=document.id,
        page_settings=document.page_settings,
    )


def _looks_like_html(value: str) -> bool:
    return bool(re.search(r"</?[A-Za-z][^>]*>", value))


def _replacement_html(replacement_text: str, replacement_html: str | None) -> str:
    html_candidate = (replacement_html or "").strip()
    if html_candidate and _looks_like_html(html_candidate):
        return html_candidate
    return text_to_html(replacement_text)


def _selection_match_key(value: str) -> str:
    return re.sub(r"\s+", "", html.unescape(value or ""))


def _replace_html_blocks_by_selection_text(
    *,
    content_html: str,
    selection_text: str,
    replacement_html: str,
) -> str | None:
    selected_key = _selection_match_key(selection_text)
    if not content_html.strip() or not selected_key:
        return None

    blocks: list[tuple[int, int]] = []
    source_chars: list[str] = []
    char_block_indexes: list[int] = []
    for match in _HTML_BLOCK_RE.finditer(content_html):
        block_index = len(blocks)
        blocks.append((match.start(), match.end()))
        for char in html_to_text(match.group(0)):
            if char.isspace():
                continue
            source_chars.append(char)
            char_block_indexes.append(block_index)

    source_key = "".join(source_chars)
    match_start = source_key.find(selected_key)
    if match_start < 0:
        return None

    match_end = match_start + len(selected_key) - 1
    start_block = char_block_indexes[match_start]
    end_block = char_block_indexes[match_end]
    html_start = blocks[start_block][0]
    html_end = blocks[end_block][1]
    return f"{content_html[:html_start]}{replacement_html}{content_html[html_end:]}"


def replace_selection_in_document(
    document: BoardDocument,
    *,
    selection_text: str,
    replacement_text: str,
    replacement_html: str | None = None,
) -> BoardDocument:
    selected = selection_text.strip()
    replacement = replacement_text.strip()
    if not selected:
        return document

    escaped_selection = html.escape(selected)
    block_replacement_html = _replacement_html(replacement, replacement_html)
    inline_replacement_html = html.escape(replacement).replace("\n", "<br>")
    for tag in ("p", "h1", "h2", "h3", "li", "blockquote"):
        exact_block_html = f"<{tag}>{escaped_selection}</{tag}>"
        if exact_block_html in document.content_html:
            next_html = document.content_html.replace(exact_block_html, block_replacement_html, 1)
            return build_document(
                title=document.title,
                content_html=next_html,
                document_id=document.id,
                page_settings=document.page_settings,
            )

    if escaped_selection in document.content_html:
        next_html = document.content_html.replace(escaped_selection, inline_replacement_html, 1)
        return build_document(
            title=document.title,
            content_html=next_html,
            document_id=document.id,
            page_settings=document.page_settings,
        )

    if selected in document.content_html:
        next_html = document.content_html.replace(selected, inline_replacement_html, 1)
        return build_document(
            title=document.title,
            content_html=next_html,
            document_id=document.id,
            page_settings=document.page_settings,
        )

    next_html = _replace_html_blocks_by_selection_text(
        content_html=document.content_html,
        selection_text=selected,
        replacement_html=block_replacement_html,
    )
    if next_html is not None:
        return build_document(
            title=document.title,
            content_html=next_html,
            document_id=document.id,
            page_settings=document.page_settings,
        )

    if selected in document.content_text and not document.content_html.strip():
        next_text = document.content_text.replace(selected, replacement, 1)
        return build_document(
            title=document.title,
            content_text=next_text,
            document_id=document.id,
            page_settings=document.page_settings,
        )

    next_text = f"{document.content_text.rstrip()}\n\n{replacement}".strip()
    next_html = "\n".join(
        part for part in [document.content_html.strip(), block_replacement_html.strip()] if part
    )
    return build_document(
        title=document.title,
        content_html=next_html,
        content_text=next_text,
        document_id=document.id,
        page_settings=document.page_settings,
    )


def import_docx(path: Path, *, title: str | None = None) -> BoardDocument:
    source = DocxDocument(path)
    html_parts: list[str] = []
    text_parts: list[str] = []
    inferred_title = title or path.stem

    for paragraph in source.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        style_name = (paragraph.style.name if paragraph.style else "").lower()
        escaped = html.escape(text)
        if "heading 1" in style_name or "title" in style_name:
            inferred_title = inferred_title if title else text
            html_parts.append(f"<h1>{escaped}</h1>")
        elif "heading 2" in style_name:
            html_parts.append(f"<h2>{escaped}</h2>")
        elif "heading 3" in style_name:
            html_parts.append(f"<h3>{escaped}</h3>")
        else:
            html_parts.append(f"<p>{escaped}</p>")
        text_parts.append(text)

    for table in source.tables:
        rows: list[str] = []
        for row in table.rows:
            cells = [html.escape(cell.text.strip()) for cell in row.cells]
            rows.append("<tr>" + "".join(f"<td>{cell}</td>" for cell in cells) + "</tr>")
            text_parts.append(" | ".join(cell.text.strip() for cell in row.cells))
        if rows:
            html_parts.append("<table><tbody>" + "".join(rows) + "</tbody></table>")

    return build_document(
        title=inferred_title,
        content_html="\n".join(html_parts),
        content_text="\n".join(text_parts),
    )


def _has_math_signal(value: str) -> bool:
    return bool(_MATH_SIGNAL_RE.search(value))


def _latex_validation_text(value: str) -> str:
    text = _LATEX_TEXT_ARGUMENT_RE.sub("", _LATEX_ENVIRONMENT_RE.sub("", value))
    for command in _LATEX_SPACING_COMMANDS:
        text = text.replace(command, "")
    return text


def _has_non_formula_letters(value: str) -> bool:
    without_latex_commands = re.sub(r"\\[A-Za-z]+", "", _latex_validation_text(value))
    return bool(_CJK_RE.search(without_latex_commands) or _NON_FORMULA_LETTER_RE.search(without_latex_commands))


def _latin_words_are_formula_like(value: str) -> bool:
    without_latex_commands = re.sub(r"\\[A-Za-z]+", "", _latex_validation_text(value))
    for word in _LATIN_WORD_RE.findall(without_latex_commands):
        if len(word) > 3 and word not in _LATEX_FUNCTIONS:
            return False
    return True


def _is_likely_delimited_math(value: str) -> bool:
    compact = _TRAILING_SENTENCE_MARKS_RE.sub("", _LEADING_SENTENCE_MARKS_RE.sub("", value.strip()))
    validation_text = _latex_validation_text(compact)
    if not compact or not validation_text or not _FORMULA_CHARS_RE.fullmatch(validation_text):
        return False
    if _has_non_formula_letters(compact) or not _latin_words_are_formula_like(compact):
        return False
    return _has_math_signal(compact) or bool(re.fullmatch(r"[A-Za-zα-ωΑ-Ω]", validation_text))


def _normalize_limit_subscript(value: str) -> str:
    return (
        value.replace("→", r"\to ")
        .replace("∞", r"\infty")
        .replace("+\\infty", r"+\infty")
        .replace("-\\infty", r"-\infty")
        .strip()
    )


def _normalize_latex(value: str) -> str:
    latex = _TRAILING_SENTENCE_MARKS_RE.sub("", _LEADING_SENTENCE_MARKS_RE.sub("", value.strip()))
    latex = latex.replace(r"\dfrac", r"\frac").replace(r"\tfrac", r"\frac")
    line_break_token = "\ue000"
    latex = re.sub(r"\\\\(?:\[[^\]]+\])?", line_break_token, latex)
    latex = latex.replace(r"\ ", " ")
    latex = latex.replace(r"\,", " ").replace(r"\;", " ").replace(r"\:", " ").replace(r"\!", "")
    latex = latex.replace(line_break_token, r"\\")
    replacements = [
        ("−", "-"),
        ("→", r"\to "),
        ("←", r"\leftarrow "),
        ("∞", r"\infty"),
        ("·", r"\cdot "),
        ("≤", r"\le "),
        ("≥", r"\ge "),
        ("≈", r"\approx "),
        ("≠", r"\ne "),
        ("±", r"\pm "),
    ]
    for old, new in replacements:
        latex = latex.replace(old, new)

    latex = latex.replace("，", r",\quad ")
    latex = re.sub(
        r"([A-Za-z]')\\frac\{x\}\{([A-Za-z])\}'\(([^()]+)\)",
        r"\\frac{\1(\3)}{\2'(\3)}",
        latex,
    )
    latex = re.sub(
        r"([A-Za-z])\\frac\{x\}\{([A-Za-z])\}\(([^()]+)\)",
        r"\\frac{\1(\3)}{\2(\3)}",
        latex,
    )
    latex = re.sub(
        r"(?<!\\)\blim_\{([^}]+)\}",
        lambda match: rf"\lim_{{{_normalize_limit_subscript(match.group(1))}}}",
        latex,
    )
    latex = re.sub(r"(?<!\\)\b(sin|cos|tan|ln|log|sqrt|exp)\b", r"\\\1", latex)
    latex = re.sub(r"([A-Za-z]'?\([^()]+\))\s*/\s*([A-Za-z]'?\([^()]+\))", r"\\frac{\1}{\2}", latex)
    latex = re.sub(
        r"(\\(?:sin|cos|tan|ln|log|sqrt|exp)\s+[A-Za-z0-9]+(?:\^\{?[-+\w/]+\}?)?)\s*/\s*\(([^()]+)\)",
        r"\\frac{\1}{\2}",
        latex,
    )
    latex = re.sub(
        r"(\\(?:sin|cos|tan|ln|log|sqrt|exp)\s+[A-Za-z0-9]+(?:\^\{?[-+\w/]+\}?)?)\s*/\s*([A-Za-z0-9]+(?:\^\{?[-+\w/]+\}?)?)",
        r"\\frac{\1}{\2}",
        latex,
    )
    latex = re.sub(r"\(([^()]+)\)\s*/\s*\(([^()]+)\)", r"\\frac{\1}{\2}", latex)
    latex = re.sub(r"\(([^()]+)\)\s*/\s*([A-Za-z0-9\\]+(?:\^\{?[-+\w/]+\}?)?)", r"\\frac{\1}{\2}", latex)
    latex = _normalize_top_level_slash_fractions(latex)
    return re.sub(r"\s+", " ", latex).strip()


def _plain_latex_text_argument(value: str) -> str:
    text = value
    for command in sorted(_LATEX_SPACING_COMMANDS, key=len, reverse=True):
        text = text.replace(command, " " if command != r"\!" else "")
    for command, symbol in _LATEX_SYMBOLS.items():
        text = text.replace(command, symbol)
    return re.sub(r"\s+", " ", text).strip()


def _matching_opener(value: str, end: int, opener: str, closer: str) -> int:
    depth = 0
    for index in range(end, -1, -1):
        if value[index] == closer:
            depth += 1
        elif value[index] == opener:
            depth -= 1
            if depth == 0:
                return index
    return -1


def _matching_closer(value: str, start: int, opener: str, closer: str) -> int:
    depth = 0
    for index in range(start, len(value)):
        if value[index] == opener:
            depth += 1
        elif value[index] == closer:
            depth -= 1
            if depth == 0:
                return index
    return -1


def _find_top_level_slash(value: str) -> int:
    brace_depth = 0
    paren_depth = 0
    for index, char in enumerate(value):
        if char == "{":
            brace_depth += 1
        elif char == "}":
            brace_depth = max(0, brace_depth - 1)
        elif char == "(":
            paren_depth += 1
        elif char == ")":
            paren_depth = max(0, paren_depth - 1)
        elif char == "/" and brace_depth == 0 and paren_depth == 0:
            return index
    return -1


def _latex_operand_start(value: str, slash_index: int) -> int:
    index = slash_index - 1
    while index >= 0 and value[index].isspace():
        index -= 1
    if index < 0:
        return 0

    while index >= 0:
        if value[index] in "})":
            opener, closer = ("{", "}") if value[index] == "}" else ("(", ")")
            start = _matching_opener(value, index, opener, closer)
            if start < 0:
                return index
            index = start - 1
            if index >= 0 and value[index] in {"_", "^"}:
                index -= 1
                continue
            while index >= 0 and value[index].isspace():
                index -= 1
            if index >= 0 and value[index] == "'":
                index -= 1
            while index >= 0 and (value[index].isalpha() or value[index] == "\\"):
                index -= 1
            return index + 1

        while index >= 0 and re.match(r"[A-Za-z0-9'\\.+\-∞]", value[index]):
            index -= 1
        if index >= 0 and value[index] in {"_", "^"}:
            index -= 1
            continue
        return index + 1
    return 0


def _latex_operand_end(value: str, slash_index: int) -> int:
    index = slash_index + 1
    while index < len(value) and value[index].isspace():
        index += 1
    if index < len(value) and value[index] in {"+", "-"}:
        index += 1
    if index >= len(value):
        return len(value)

    if value[index] == "\\":
        command_match = re.match(r"\\[A-Za-z]+", value[index:])
        if command_match:
            index += len(command_match.group(0))
    elif value[index] in "{(":
        closer = "}" if value[index] == "{" else ")"
        end = _matching_closer(value, index, value[index], closer)
        index = len(value) if end < 0 else end + 1
    else:
        token_match = re.match(r"[A-Za-z0-9]+(?:'\([^()]*\)|\([^()]*\)|')?", value[index:])
        if token_match:
            index += len(token_match.group(0))
        else:
            index += 1

    while index < len(value):
        if value[index].isspace():
            break
        if value[index] in {"_", "^"}:
            index += 1
            while index < len(value) and value[index].isspace():
                index += 1
            if index < len(value) and value[index] == "{":
                end = _matching_closer(value, index, "{", "}")
                index = len(value) if end < 0 else end + 1
            elif index < len(value) and value[index] == "\\":
                command_match = re.match(r"\\[A-Za-z]+", value[index:])
                index += len(command_match.group(0)) if command_match else 1
            else:
                token_match = re.match(r"[A-Za-z0-9+\-]+", value[index:])
                index += len(token_match.group(0)) if token_match else 1
            continue
        if value[index] == "/":
            nested_end = _latex_operand_end(value, index)
            index = nested_end
            continue
        break
    return index


def _strip_wrapping_parentheses(value: str) -> str:
    stripped = value.strip()
    if stripped.startswith("(") and stripped.endswith(")") and _matching_closer(stripped, 0, "(", ")") == len(stripped) - 1:
        return stripped[1:-1].strip()
    return stripped


def _normalize_top_level_slash_fractions(value: str) -> str:
    current = value
    for _ in range(24):
        slash_index = _find_top_level_slash(current)
        if slash_index < 0:
            return current
        start = _latex_operand_start(current, slash_index)
        end = _latex_operand_end(current, slash_index)
        numerator = current[start:slash_index].strip()
        denominator = current[slash_index + 1 : end].strip()
        if not numerator or not denominator:
            return current
        numerator = _normalize_top_level_slash_fractions(_strip_wrapping_parentheses(numerator))
        denominator = _normalize_top_level_slash_fractions(_strip_wrapping_parentheses(denominator))
        current = f"{current[:start]}\\frac{{{numerator}}}{{{denominator}}}{current[end:]}"
    return current


def _formula_only_latex(text: str) -> str | None:
    compact = _TRAILING_SENTENCE_MARKS_RE.sub("", text.strip())
    delimited = re.match(r"^\\\[(.+?)\\\]$", compact) or re.match(r"^\\\((.+?)\\\)$", compact)
    if delimited:
        latex = delimited.group(1)
        return _normalize_latex(latex) if _is_likely_delimited_math(latex) else None
    dollar_delimited = re.match(r"^\$\$?(.+?)\$\$?$", compact)
    if dollar_delimited:
        latex = dollar_delimited.group(1)
        return _normalize_latex(latex) if _is_likely_delimited_math(latex) else None
    if not compact or _CJK_RE.search(compact) or not _has_math_signal(compact):
        return None
    return _normalize_latex(compact)


def _math_segments(text: str) -> list[tuple[int, int, str]]:
    segments: list[tuple[int, int, str]] = []

    for match in _DELIMITED_MATH_RE.finditer(text):
        latex = match.group(1) or match.group(2) or match.group(3) or match.group(4) or ""
        if latex.strip() and _is_likely_delimited_math(latex):
            segments.append((match.start(), match.end(), _normalize_latex(latex)))

    return sorted(segments, key=lambda segment: segment[0])


def _auto_math_fragments(text: str) -> list[InlineFragment]:
    segments = _math_segments(text)
    if not segments:
        return [("text", text)] if text else []

    fragments: list[InlineFragment] = []
    cursor = 0
    for start, end, latex in segments:
        if start > cursor:
            fragments.append(("text", text[cursor:start]))
        fragments.append(("math", latex))
        cursor = end
    if cursor < len(text):
        fragments.append(("text", text[cursor:]))
    return fragments


def _fragment_text(fragments: list[InlineFragment]) -> str:
    return "".join(text for _, text in fragments)


class _DocxBlockParser(HTMLParser):
    def __init__(self) -> None:
        super().__init__()
        self.blocks: list[DocxBlock] = []
        self._tag_stack: list[str] = []
        self._buffer: list[str] = []
        self._attrs_stack: list[dict[str, Any]] = []
        self._fragments: list[InlineFragment] = []
        self._ignored_atom_depth = 0
        self._table_rows: TableRows | None = None
        self._table_row: list[list[InlineFragment]] | None = None
        self._table_cell_fragments: list[InlineFragment] | None = None
        self._table_cell_buffer: list[str] = []

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        tag = tag.lower()
        attr_map = dict(attrs)
        node_type = (attr_map.get("data-type") or "").strip()
        if self._table_rows is not None:
            self._handle_table_starttag(tag, attr_map, node_type)
            return
        if tag == "table":
            self._flush()
            self._table_rows = []
            self._table_row = None
            return
        if node_type == "block-math":
            self._flush()
            latex = html.unescape((attr_map.get("data-latex") or "").strip())
            if latex:
                self.blocks.append(("math", [("math", latex)], attr_map))
            self._ignored_atom_depth = 1
            return
        if node_type == "page-break":
            self._flush()
            self.blocks.append(("pageBreak", [], attr_map))
            self._ignored_atom_depth = 1
            return
        if node_type == "inline-math":
            self._flush_text()
            latex = html.unescape((attr_map.get("data-latex") or "").strip())
            if latex:
                self._fragments.append(("math", latex))
            self._ignored_atom_depth = 1
            return

        if tag in {"h1", "h2", "h3", "p", "li", "blockquote"}:
            self._flush()
            self._tag_stack.append(tag)
            self._attrs_stack.append(attr_map)
        elif tag == "img":
            src = (attr_map.get("src") or "").strip()
            alt = (attr_map.get("alt") or "").strip()
            self._flush()
            self.blocks.append(("img", [("text", alt)], {"src": src, "alt": alt}))
        elif tag == "br":
            self._buffer.append("\n")

    def handle_endtag(self, tag: str) -> None:
        tag = tag.lower()
        if self._ignored_atom_depth:
            self._ignored_atom_depth -= 1
            return
        if self._table_rows is not None:
            self._handle_table_endtag(tag)
            return
        if tag in {"h1", "h2", "h3", "p", "li", "blockquote"}:
            current = self._tag_stack.pop() if self._tag_stack else tag
            attrs = self._attrs_stack.pop() if self._attrs_stack else {}
            self._flush(current, attrs)

    def handle_data(self, data: str) -> None:
        if self._ignored_atom_depth:
            return
        if self._table_rows is not None:
            if self._table_cell_fragments is not None:
                self._table_cell_buffer.append(data)
            return
        self._buffer.append(data)

    def _handle_table_starttag(self, tag: str, attrs: dict[str, str | None], node_type: str) -> None:
        if tag == "tr":
            self._table_row = []
            return
        if tag in {"td", "th"}:
            if self._table_row is None:
                self._table_row = []
            self._table_cell_fragments = []
            self._table_cell_buffer = []
            return
        if tag == "br" and self._table_cell_fragments is not None:
            self._table_cell_buffer.append("\n")
            return
        if node_type in {"inline-math", "block-math"} and self._table_cell_fragments is not None:
            self._flush_table_cell_text()
            latex = html.unescape((attrs.get("data-latex") or "").strip())
            if latex:
                self._table_cell_fragments.append(("math", latex))
            self._ignored_atom_depth = 1

    def _handle_table_endtag(self, tag: str) -> None:
        if tag in {"td", "th"} and self._table_cell_fragments is not None:
            self._flush_table_cell_text()
            fragments = _trim_fragments(self._table_cell_fragments)
            if self._table_row is None:
                self._table_row = []
            self._table_row.append(fragments or [("text", "")])
            self._table_cell_fragments = None
            self._table_cell_buffer = []
            return
        if tag == "tr":
            if self._table_rows is not None and self._table_row:
                self._table_rows.append(self._table_row)
            self._table_row = None
            return
        if tag == "table":
            rows = self._table_rows or []
            if rows:
                self.blocks.append(("table", [], {"rows": rows}))
            self._table_rows = None
            self._table_row = None
            self._table_cell_fragments = None
            self._table_cell_buffer = []

    def _flush_table_cell_text(self) -> None:
        if self._table_cell_fragments is None or not self._table_cell_buffer:
            return
        text = html.unescape("".join(self._table_cell_buffer))
        self._table_cell_buffer = []
        text = re.sub(r"\s+", " ", text)
        if text:
            self._table_cell_fragments.append(("text", text))

    def _flush_text(self) -> None:
        if not self._buffer:
            return
        text = html.unescape("".join(self._buffer))
        self._buffer = []
        text = re.sub(r"\s+", " ", text)
        if text:
            self._fragments.append(("text", text))

    def _flush(self, tag: str | None = None, attrs: dict[str, Any] | None = None) -> None:
        self._flush_text()
        if not self._fragments:
            return

        fragments = _trim_fragments(self._fragments)
        self._fragments = []
        if fragments:
            self.blocks.append((tag or "p", fragments, attrs or {}))


def _trim_fragments(fragments: list[InlineFragment]) -> list[InlineFragment]:
    trimmed = list(fragments)
    while trimmed and trimmed[0][0] == "text" and not trimmed[0][1].strip():
        trimmed.pop(0)
    while trimmed and trimmed[-1][0] == "text" and not trimmed[-1][1].strip():
        trimmed.pop()
    if trimmed and trimmed[0][0] == "text":
        trimmed[0] = ("text", trimmed[0][1].lstrip())
    if trimmed and trimmed[-1][0] == "text":
        trimmed[-1] = ("text", trimmed[-1][1].rstrip())
    return trimmed


def _page_size_cm(page_size: str) -> tuple[float, float]:
    if page_size == "letter":
        return 21.59, 27.94
    if page_size == "a3":
        return 29.7, 42.0
    return 21.0, 29.7


def _margin_cm(preset: str) -> float:
    if preset == "narrow":
        return 1.27
    if preset == "wide":
        return 3.18
    return 2.54


def _apply_page_settings(target: DocxDocument, settings: DocumentPageSettings) -> None:
    section = target.sections[0]
    width_cm, height_cm = _page_size_cm(settings.page_size)
    if settings.orientation == "landscape":
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Cm(height_cm)
        section.page_height = Cm(width_cm)
    else:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = Cm(width_cm)
        section.page_height = Cm(height_cm)

    margin = Cm(_margin_cm(settings.margin_preset))
    section.top_margin = margin
    section.bottom_margin = margin
    section.left_margin = margin
    section.right_margin = margin

    cols = section._sectPr.xpath("./w:cols")
    if cols:
        cols[0].set(qn("w:num"), str(settings.columns))

    if settings.header_text:
        header_paragraph = section.header.paragraphs[0]
        header_paragraph.text = settings.header_text

    footer = section.footer
    footer.paragraphs[0].text = settings.footer_text or ""
    if settings.show_page_number:
        page_number_paragraph = footer.add_paragraph()
        page_number_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        _append_page_number_field(page_number_paragraph)


def _append_page_number_field(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(end)


def _decode_data_uri(data_uri: str) -> bytes | None:
    if not data_uri.startswith("data:") or "," not in data_uri:
        return None
    header, payload = data_uri.split(",", 1)
    if ";base64" not in header:
        return None
    try:
        import base64

        return base64.b64decode(payload)
    except Exception:
        return None


def _script_text(value: str, *, subscript: bool = False) -> str:
    parsed = _latex_inline_text(value)
    table = _SUBSCRIPT_CHARS if subscript else _SUPERSCRIPT_CHARS
    translated = parsed.translate(table)
    if subscript and re.search(r"[A-Za-zα-ωΑ-Ω]", parsed):
        for original, converted in zip(parsed, translated):
            if original.isalpha() and original == converted:
                return f"_{parsed}"
    return translated


class _InlineLatexParser:
    def __init__(self, latex: str) -> None:
        self.latex = latex
        self.index = 0

    def parse(self) -> str:
        parts: list[str] = []
        while self.index < len(self.latex):
            char = self.latex[self.index]
            if char == "}":
                break
            if char.isspace():
                if parts and parts[-1] != " ":
                    parts.append(" ")
                self.index += 1
                continue
            atom = self._parse_atom()
            parts.append(self._parse_scripts(atom))
        text = re.sub(r"\s+", " ", "".join(parts)).strip()
        text = re.sub(r"\s*([→←])\s*", r"\1", text)
        text = text.replace("± ∞", "±∞")
        return text

    def _parse_atom(self) -> str:
        char = self.latex[self.index]
        if char == "{":
            raw, self.index = _read_braced(self.latex, self.index)
            return _latex_inline_text(raw)
        if self.latex.startswith(r"\frac", self.index):
            next_index = self.index + len(r"\frac")
            numerator_raw, next_index = _read_braced(self.latex, next_index)
            denominator_raw, next_index = _read_braced(self.latex, next_index)
            self.index = next_index
            numerator = _latex_inline_text(numerator_raw)
            denominator = _latex_inline_text(denominator_raw)
            if re.search(r"\s|[+\-=*/]", numerator):
                numerator = f"({numerator})"
            if re.search(r"\s|[+\-=*/]", denominator):
                denominator = f"({denominator})"
            return f"{numerator}/{denominator}"
        if self.latex.startswith(r"\sqrt", self.index):
            next_index = self.index + len(r"\sqrt")
            radicand_raw, self.index = _read_braced(self.latex, next_index)
            return f"√({ _latex_inline_text(radicand_raw) })"
        if self.latex.startswith(r"\lim", self.index):
            self.index += len(r"\lim")
            if self.index < len(self.latex) and self.latex[self.index] == "_":
                self.index += 1
                return f"lim {_latex_inline_text(self._read_script_raw())}"
            return "lim"
        if char == "\\":
            return self._parse_command()
        match = re.match(r"[A-Za-z]+(?:')?", self.latex[self.index :])
        if match:
            token = match.group(0)
            self.index += len(token)
            return token
        match = re.match(r"\d+(?:\.\d+)?", self.latex[self.index :])
        if match:
            token = match.group(0)
            self.index += len(token)
            return token
        self.index += 1
        return char

    def _parse_command(self) -> str:
        for command in _LATEX_SPACING_COMMANDS:
            if self.latex.startswith(command, self.index):
                self.index += len(command)
                return " " if command != r"\!" else ""
        for command in _LATEX_DELIMITER_COMMANDS:
            if self.latex.startswith(command, self.index):
                self.index += len(command)
                return ""
        for command in _LATEX_STYLE_COMMANDS:
            if self.latex.startswith(command, self.index):
                self.index += len(command)
                return ""
        for command in _LATEX_TEXT_COMMANDS:
            if self.latex.startswith(command, self.index):
                self.index += len(command)
                raw, self.index = _read_braced(self.latex, self.index)
                return _plain_latex_text_argument(raw)
        if self.latex.startswith(r"\quad", self.index):
            self.index += len(r"\quad")
            return " "
        command_match = re.match(r"\\[A-Za-z]+", self.latex[self.index :])
        if not command_match:
            self.index += 1
            return ""
        command = command_match.group(0)
        self.index += len(command)
        if command in _LATEX_DELIMITER_COMMANDS:
            return ""
        if command in _LATEX_SYMBOLS:
            return _LATEX_SYMBOLS[command]
        return command[1:]

    def _parse_scripts(self, base: str) -> str:
        result = base
        while self.index < len(self.latex) and self.latex[self.index] in {"_", "^"}:
            script_kind = self.latex[self.index]
            self.index += 1
            raw = self._read_script_raw()
            result += _script_text(raw, subscript=script_kind == "_")
        return result

    def _read_script_raw(self) -> str:
        while self.index < len(self.latex) and self.latex[self.index].isspace():
            self.index += 1
        if self.index >= len(self.latex):
            return ""
        if self.latex[self.index] == "{":
            raw, self.index = _read_braced(self.latex, self.index)
            return raw
        if self.latex[self.index] == "\\":
            command_match = re.match(r"\\[A-Za-z]+", self.latex[self.index :])
            if command_match:
                raw = command_match.group(0)
                self.index += len(raw)
                return raw
        raw = self.latex[self.index]
        self.index += 1
        return raw


@lru_cache(maxsize=1024)
def _latex_inline_text(latex: str) -> str:
    return _InlineLatexParser(_normalize_latex(latex)).parse()


class _DisplayMathBox:
    def __init__(self, lines: list[str], baseline: int = 0) -> None:
        self.lines = [line.rstrip() for line in lines] or [""]
        self.width = max(1, max(len(line) for line in self.lines))
        self.lines = [line.center(self.width) for line in self.lines]
        self.baseline = min(max(0, baseline), len(self.lines) - 1)


def _combine_display_boxes(boxes: list[_DisplayMathBox], *, gap: str = " ") -> _DisplayMathBox:
    if not boxes:
        return _DisplayMathBox([""])
    baseline = max(box.baseline for box in boxes)
    descent = max(len(box.lines) - box.baseline - 1 for box in boxes)
    height = baseline + descent + 1
    combined: list[str] = []
    for row in range(height):
        pieces: list[str] = []
        for box in boxes:
            source_row = row - (baseline - box.baseline)
            if 0 <= source_row < len(box.lines):
                pieces.append(box.lines[source_row].center(box.width))
            else:
                pieces.append(" " * box.width)
        combined.append(gap.join(pieces).rstrip())
    return _DisplayMathBox(combined, baseline)


class _DisplayLatexParser:
    def __init__(self, latex: str) -> None:
        self.latex = latex
        self.index = 0

    def parse(self) -> _DisplayMathBox:
        boxes: list[_DisplayMathBox] = []
        text_buffer: list[str] = []
        while self.index < len(self.latex):
            char = self.latex[self.index]
            if char == "}":
                break
            if self.latex.startswith(r"\frac", self.index):
                self._flush_text(text_buffer, boxes)
                boxes.append(self._parse_fraction())
                boxes.append(self._parse_scripts_box())
                continue
            if self.latex.startswith(r"\lim", self.index):
                self._flush_text(text_buffer, boxes)
                boxes.append(self._parse_limit())
                continue
            if self.latex.startswith(r"\begin{cases}", self.index):
                self._flush_text(text_buffer, boxes)
                boxes.append(self._parse_cases())
                continue
            if char == "{":
                raw, self.index = _read_braced(self.latex, self.index)
                self._flush_text(text_buffer, boxes)
                boxes.append(_latex_display_box(raw))
                continue
            if char == "\\":
                text_buffer.append(self._parse_command())
                continue
            text_buffer.append(char)
            self.index += 1
        self._flush_text(text_buffer, boxes)
        return _combine_display_boxes([box for box in boxes if any(line.strip() for line in box.lines)])

    def _flush_text(self, text_buffer: list[str], boxes: list[_DisplayMathBox]) -> None:
        if not text_buffer:
            return
        text = _latex_inline_text("".join(text_buffer))
        if text:
            boxes.append(_DisplayMathBox([text]))
        text_buffer.clear()

    def _parse_fraction(self) -> _DisplayMathBox:
        next_index = self.index + len(r"\frac")
        numerator_raw, next_index = _read_braced(self.latex, next_index)
        denominator_raw, next_index = _read_braced(self.latex, next_index)
        self.index = next_index
        numerator = _latex_display_box(numerator_raw)
        denominator = _latex_display_box(denominator_raw)
        width = max(numerator.width, denominator.width) + 2
        lines = [line.center(width) for line in numerator.lines]
        lines.append("─" * width)
        lines.extend(line.center(width) for line in denominator.lines)
        return _DisplayMathBox(lines, len(numerator.lines))

    def _parse_limit(self) -> _DisplayMathBox:
        self.index += len(r"\lim")
        if self.index < len(self.latex) and self.latex[self.index] == "_":
            self.index += 1
            raw = self._read_script_raw()
            limit = _latex_inline_text(raw)
            width = max(len("lim"), len(limit))
            return _DisplayMathBox(["lim".center(width), limit.center(width)], 0)
        return _DisplayMathBox(["lim"])

    def _parse_cases(self) -> _DisplayMathBox:
        end_index = self.latex.find(r"\end{cases}", self.index)
        if end_index < 0:
            self.index += len(r"\begin{cases}")
            return _DisplayMathBox(["{"])
        raw_cases = self.latex[self.index + len(r"\begin{cases}") : end_index]
        self.index = end_index + len(r"\end{cases}")
        normalized_cases = re.sub(r"\\\\\[[^\]]+\]", r"\\\\", raw_cases)
        rows: list[str] = []
        for raw_row in re.split(r"\\\\", normalized_cases):
            row = raw_row.strip()
            if not row:
                continue
            pieces = [_latex_inline_text(piece.strip()) for piece in row.split("&") if piece.strip()]
            rows.append("  ".join(piece for piece in pieces if piece))
        if not rows:
            return _DisplayMathBox(["{ }"])
        width = max(len(row) for row in rows)
        lines = [("{ " + rows[0]).ljust(width + 2)]
        lines.extend(("  " + row).ljust(width + 2) for row in rows[1:])
        return _DisplayMathBox(lines)

    def _parse_scripts_box(self) -> _DisplayMathBox:
        if self.index >= len(self.latex) or self.latex[self.index] not in {"_", "^"}:
            return _DisplayMathBox([""])
        result = ""
        while self.index < len(self.latex) and self.latex[self.index] in {"_", "^"}:
            script_kind = self.latex[self.index]
            self.index += 1
            raw = self._read_script_raw()
            result += _script_text(raw, subscript=script_kind == "_")
        return _DisplayMathBox([result])

    def _read_script_raw(self) -> str:
        while self.index < len(self.latex) and self.latex[self.index].isspace():
            self.index += 1
        if self.index >= len(self.latex):
            return ""
        if self.latex[self.index] == "{":
            raw, self.index = _read_braced(self.latex, self.index)
            return raw
        if self.latex[self.index] == "\\":
            command_match = re.match(r"\\[A-Za-z]+", self.latex[self.index :])
            if command_match:
                raw = command_match.group(0)
                self.index += len(raw)
                return raw
        raw = self.latex[self.index]
        self.index += 1
        return raw

    def _parse_command(self) -> str:
        for command in _LATEX_SPACING_COMMANDS:
            if self.latex.startswith(command, self.index):
                self.index += len(command)
                return "  " if command in {r"\quad", r"\qquad"} else " "
        for command in _LATEX_DELIMITER_COMMANDS:
            if self.latex.startswith(command, self.index):
                self.index += len(command)
                return ""
        for command in _LATEX_STYLE_COMMANDS:
            if self.latex.startswith(command, self.index):
                self.index += len(command)
                return ""
        for command in _LATEX_TEXT_COMMANDS:
            if self.latex.startswith(command, self.index):
                self.index += len(command)
                raw, self.index = _read_braced(self.latex, self.index)
                return _plain_latex_text_argument(raw)
        command_match = re.match(r"\\[A-Za-z]+", self.latex[self.index :])
        if not command_match:
            self.index += 1
            return ""
        command = command_match.group(0)
        self.index += len(command)
        if command in _LATEX_DELIMITER_COMMANDS:
            return ""
        return _LATEX_SYMBOLS.get(command, command[1:])


@lru_cache(maxsize=1024)
def _latex_display_box(latex: str) -> _DisplayMathBox:
    return _DisplayLatexParser(_normalize_latex(latex)).parse()


def _latex_display_lines(latex: str) -> list[str]:
    normalized = _normalize_latex(latex)
    if r"\begin{cases}" in normalized:
        return [line.rstrip() for line in _latex_display_box(normalized).lines]
    if r"\\" in normalized:
        lines = [_latex_inline_text(part.strip()) for part in re.split(r"\\\\", normalized) if part.strip()]
        return lines or [""]
    text = _latex_inline_text(normalized)
    return [text] if text else [line.rstrip() for line in _latex_display_box(normalized).lines]


def _m_element(tag: str, text: str | None = None) -> OxmlElement:
    element = OxmlElement(f"m:{tag}")
    if text is not None:
        element.text = text
    return element


def _math_run(text: str) -> OxmlElement:
    run = _m_element("r")
    text_node = _m_element("t", text)
    run.append(text_node)
    return run


def _math_container(tag: str, children: list[OxmlElement]) -> OxmlElement:
    container = _m_element(tag)
    for child in children:
        container.append(child)
    return container


def _math_arg(tag: str, children: list[OxmlElement]) -> OxmlElement:
    return _math_container(tag, children or [_math_run("")])


def _math_property_value(tag: str, attr: str, value: str) -> OxmlElement:
    element = _m_element(tag)
    element.set(qn(f"m:{attr}"), value)
    return element


def _math_fraction(numerator: list[OxmlElement], denominator: list[OxmlElement]) -> OxmlElement:
    fraction = _m_element("f")
    fraction_properties = _m_element("fPr")
    fraction_properties.append(_math_property_value("type", "val", "bar"))
    fraction.append(fraction_properties)
    fraction.append(_math_arg("num", numerator))
    fraction.append(_math_arg("den", denominator))
    return fraction


def _math_script(base: list[OxmlElement], subscript: list[OxmlElement] | None, superscript: list[OxmlElement] | None) -> OxmlElement:
    if subscript is not None and superscript is not None:
        node = _m_element("sSubSup")
        node.append(_m_element("sSubSupPr"))
        node.append(_math_arg("e", base))
        node.append(_math_arg("sub", subscript))
        node.append(_math_arg("sup", superscript))
        return node
    if subscript is not None:
        node = _m_element("sSub")
        node.append(_m_element("sSubPr"))
        node.append(_math_arg("e", base))
        node.append(_math_arg("sub", subscript))
        return node
    node = _m_element("sSup")
    node.append(_m_element("sSupPr"))
    node.append(_math_arg("e", base))
    node.append(_math_arg("sup", superscript or [_math_run("")]))
    return node


def _math_limit_low(base: list[OxmlElement], limit: list[OxmlElement]) -> OxmlElement:
    node = _m_element("limLow")
    node.append(_m_element("limLowPr"))
    node.append(_math_arg("e", base))
    node.append(_math_arg("lim", limit))
    return node


def _matching_delimiter(value: str, start: int, opener: str, closer: str) -> int:
    depth = 0
    for index in range(start, len(value)):
        if value[index] == opener:
            depth += 1
        elif value[index] == closer:
            depth -= 1
            if depth == 0:
                return index
    return -1


def _read_braced(value: str, index: int) -> tuple[str, int]:
    while index < len(value) and value[index].isspace():
        index += 1
    if index >= len(value):
        return "", index
    if value[index] != "{":
        return value[index], index + 1
    end = _matching_delimiter(value, index, "{", "}")
    if end < 0:
        return value[index + 1 :], len(value)
    return value[index + 1 : end], end + 1


def _read_parenthesized(value: str, index: int) -> tuple[str, int]:
    end = _matching_delimiter(value, index, "(", ")")
    if end < 0:
        return value[index], index + 1
    return value[index : end + 1], end + 1


def _read_base(value: str, index: int) -> tuple[list[OxmlElement], int]:
    if index >= len(value):
        return [_math_run("")], index
    if value[index] == "(":
        parenthesized, next_index = _read_parenthesized(value, index)
        inner = parenthesized[1:-1] if parenthesized.startswith("(") and parenthesized.endswith(")") else ""
        if inner:
            return [_math_run("("), *_latex_to_normalized_math_children(inner), _math_run(")")], next_index
        return [_math_run(parenthesized)], next_index
    if value[index] == "{":
        braced, next_index = _read_braced(value, index)
        return _latex_to_normalized_math_children(braced), next_index
    if value[index] == "\\":
        command_match = re.match(r"\\[A-Za-z]+", value[index:])
        if command_match:
            command = command_match.group(0)
            if command in _LATEX_TEXT_COMMANDS:
                raw, next_index = _read_braced(value, index + len(command))
                return [_math_run(_plain_latex_text_argument(raw))], next_index
            if command in _LATEX_STYLE_COMMANDS or command in _LATEX_DELIMITER_COMMANDS:
                return [_math_run("")], index + len(command)
            return [_math_run(_LATEX_SYMBOLS.get(command, command[1:]))], index + len(command)
    match = re.match(r"[A-Za-z]+(?:'\([^()]*\)|\([^()]*\)|')?", value[index:])
    if match:
        token = match.group(0)
        return [_math_run(token)], index + len(token)
    return [_math_run(value[index])], index + 1


def _read_script(value: str, index: int) -> tuple[list[OxmlElement], int]:
    script, next_index = _read_braced(value, index)
    return _latex_to_normalized_math_children(script), next_index


def _apply_scripts(value: str, index: int, base: list[OxmlElement]) -> tuple[list[OxmlElement], int]:
    subscript: list[OxmlElement] | None = None
    superscript: list[OxmlElement] | None = None
    current = index
    while current < len(value) and value[current] in {"_", "^"}:
        script_kind = value[current]
        script, current = _read_script(value, current + 1)
        if script_kind == "_":
            subscript = script
        else:
            superscript = script
    if subscript is None and superscript is None:
        return base, index
    return [_math_script(base, subscript, superscript)], current


def _latex_to_math_children(latex: str) -> list[OxmlElement]:
    children: list[OxmlElement] = []
    index = 0
    while index < len(latex):
        char = latex[index]
        matched_style = next((command for command in _LATEX_STYLE_COMMANDS if latex.startswith(command, index)), "")
        if matched_style:
            index += len(matched_style)
            continue
        matched_delimiter = next((command for command in _LATEX_DELIMITER_COMMANDS if latex.startswith(command, index)), "")
        if matched_delimiter:
            index += len(matched_delimiter)
            continue
        matched_text = next((command for command in _LATEX_TEXT_COMMANDS if latex.startswith(command, index)), "")
        if matched_text:
            raw_text, index = _read_braced(latex, index + len(matched_text))
            if raw_text:
                children.append(_math_run(_plain_latex_text_argument(raw_text)))
            continue
        if latex.startswith(r"\begin{cases}", index):
            end_index = latex.find(r"\end{cases}", index)
            if end_index >= 0:
                raw_cases = latex[index + len(r"\begin{cases}") : end_index]
                children.extend(_latex_cases_children(raw_cases))
                index = end_index + len(r"\end{cases}")
                continue
        if char.isspace():
            if not children or (children[-1].tag != qn("m:r") or (children[-1].find(qn("m:t")).text or "") != " "):
                children.append(_math_run(" "))
            index += 1
            continue
        if char == "{":
            braced, index = _read_braced(latex, index)
            children.extend(_latex_to_normalized_math_children(braced))
            continue
        if latex.startswith(r"\frac", index):
            numerator_raw, after_num = _read_braced(latex, index + len(r"\frac"))
            denominator_raw, after_den = _read_braced(latex, after_num)
            fraction = _math_fraction(
                _latex_to_normalized_math_children(numerator_raw),
                _latex_to_normalized_math_children(denominator_raw),
            )
            scripted, index = _apply_scripts(latex, after_den, [fraction])
            children.extend(scripted)
            continue
        if latex.startswith(r"\lim", index):
            next_index = index + len(r"\lim")
            if next_index < len(latex) and latex[next_index] == "_":
                limit_raw, next_index = _read_braced(latex, next_index + 1)
                lim_node = _math_limit_low([_math_run("lim")], _latex_to_normalized_math_children(limit_raw))
                scripted, index = _apply_scripts(latex, next_index, [lim_node])
                children.extend(scripted)
                continue
            children.append(_math_run("lim"))
            index = next_index
            continue
        if char == "\\":
            matched_spacing = next((command for command in _LATEX_SPACING_COMMANDS if latex.startswith(command, index)), "")
            if matched_spacing:
                if matched_spacing != r"\!":
                    children.append(_math_run("    " if matched_spacing in {r"\quad", r"\qquad"} else " "))
                index += len(matched_spacing)
                continue
            command_match = re.match(r"\\[A-Za-z]+", latex[index:])
            if command_match:
                command = command_match.group(0)
                text = _LATEX_SYMBOLS.get(command, command[1:])
                node, next_index = _apply_scripts(latex, index + len(command), [_math_run(text)])
                children.extend(node)
                index = next_index
                continue
        base, next_index = _read_base(latex, index)
        scripted, index = _apply_scripts(latex, next_index, base)
        children.extend(scripted)
    return children or [_math_run(latex)]


def _latex_cases_children(raw_cases: str) -> list[OxmlElement]:
    children: list[OxmlElement] = [_math_run("{ ")]
    normalized_cases = re.sub(r"\\\\\[[^\]]+\]", r"\\\\", raw_cases)
    rows = [row.strip() for row in re.split(r"\\\\", normalized_cases) if row.strip()]
    for index, row in enumerate(rows):
        if index:
            children.append(_math_run("; "))
        row_latex = " ".join(part.strip() for part in row.split("&") if part.strip())
        children.extend(_latex_to_normalized_math_children(row_latex))
    children.append(_math_run(" }"))
    return children


def _latex_to_normalized_math_children(latex: str) -> list[OxmlElement]:
    return _latex_to_math_children(_normalize_latex(latex))


def _append_omml_math(paragraph, latex: str, *, display: bool = False) -> None:
    math = _m_element("oMath")
    for child in _latex_to_math_children(_normalize_latex(latex)):
        math.append(child)
    if display:
        math_paragraph = _m_element("oMathPara")
        math_paragraph.append(math)
        paragraph._p.append(math_paragraph)
    else:
        paragraph._p.append(math)


def _set_run_font(run, font_name: str) -> None:
    run.font.name = font_name
    r_pr = run._r.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is not None:
        r_fonts.set(qn("w:eastAsia"), font_name)
        r_fonts.set(qn("w:cs"), font_name)


def _disable_run_proofing(run) -> None:
    r_pr = run._r.get_or_add_rPr()
    if r_pr.find(qn("w:noProof")) is None:
        r_pr.append(OxmlElement("w:noProof"))


def _append_display_math_text(paragraph, latex: str) -> None:
    for index, line in enumerate(_latex_display_lines(latex)):
        if index:
            paragraph.add_run().add_break()
        run = paragraph.add_run(line)
        _set_run_font(run, "Menlo")
        _disable_run_proofing(run)
    paragraph.paragraph_format.line_spacing = 1


def _append_math(paragraph, latex: str, *, display: bool = False) -> None:
    if display:
        _append_display_math_text(paragraph, latex)
        return
    try:
        _append_omml_math(paragraph, latex, display=display)
    except Exception:
        paragraph.add_run(_latex_inline_text(latex) or latex)


def _append_fragments(
    paragraph,
    fragments: list[InlineFragment],
    *,
    auto_math: bool = True,
    display_math: bool = False,
) -> None:
    for kind, value in fragments:
        if kind == "math":
            _append_math(paragraph, value, display=display_math)
            continue
        split_fragments = _auto_math_fragments(value) if auto_math else [("text", value)]
        for split_kind, split_value in split_fragments:
            if not split_value:
                continue
            if split_kind == "math":
                _append_math(paragraph, split_value, display=display_math)
            else:
                paragraph.add_run(split_value)


def _add_fragment_paragraph(target: DocxDocument, fragments: list[InlineFragment], style: str | None = None):
    paragraph = target.add_paragraph(style=style)
    _append_fragments(paragraph, fragments)
    return paragraph


def _code_fence_body_after_opener(value: str) -> tuple[bool, list[str]]:
    stripped = value.strip()
    if not stripped.startswith("```"):
        return False, []
    body = stripped[3:].strip()
    if not body:
        return True, []
    first_token, token_separator, token_remainder = body.partition(" ")
    if token_separator and first_token.strip().lower() in {"text", "txt", "plain", "plaintext"}:
        return True, [token_remainder.strip()] if token_remainder.strip() else []
    if re.fullmatch(r"[A-Za-z0-9_-]+", body):
        return True, []
    return True, [body]


def _normalize_fenced_docx_blocks(blocks: list[DocxBlock]) -> list[DocxBlock]:
    normalized: list[DocxBlock] = []
    index = 0
    while index < len(blocks):
        tag, fragments, attrs = blocks[index]
        text = _fragment_text(fragments)
        single_block_code = _fenced_code_text(text)
        if tag == "p" and single_block_code is not None:
            normalized.append(("pre", [("text", single_block_code)], {}))
            index += 1
            continue

        is_opener, code_lines = _code_fence_body_after_opener(text)
        if tag != "p" or not is_opener:
            normalized.append((tag, fragments, attrs))
            index += 1
            continue

        index += 1
        closed = False
        while index < len(blocks):
            next_tag, next_fragments, _next_attrs = blocks[index]
            next_text = _fragment_text(next_fragments)
            stripped_next = next_text.strip()
            if next_tag == "p" and stripped_next.endswith("```"):
                before_close = stripped_next[:-3].strip()
                if before_close:
                    code_lines.append(before_close)
                closed = True
                index += 1
                break
            code_lines.append(next_text)
            index += 1

        if closed:
            normalized.append(("pre", [("text", "\n".join(line for line in code_lines if line.strip()))], {}))
        else:
            normalized.append((tag, fragments, attrs))
            for line in code_lines:
                if line.strip():
                    normalized.append(("p", [("text", line)], {}))
    return normalized


def _fenced_code_text(value: str) -> str | None:
    match = _FENCED_CODE_TEXT_RE.match(value.strip())
    if not match:
        return None
    body = match.group("body").strip()
    if not body:
        return ""
    first_line, separator, remainder = body.partition("\n")
    if separator and re.fullmatch(r"[A-Za-z0-9_-]+", first_line.strip()):
        return remainder.strip()
    first_token, token_separator, token_remainder = body.partition(" ")
    if token_separator and first_token.strip().lower() in {"text", "txt", "plain", "plaintext"}:
        return token_remainder.strip()
    return body


def _add_preformatted_paragraph(target: DocxDocument, value: str):
    paragraph = target.add_paragraph()
    lines = value.splitlines() or [value]
    for index, line in enumerate(lines):
        if index:
            paragraph.add_run().add_break()
        run = paragraph.add_run(line)
        run.font.name = "Courier New"
    return paragraph


def _add_fragment_table(target: DocxDocument, rows: TableRows):
    normalized_rows = [row for row in rows if row]
    if not normalized_rows:
        return None
    width = max(len(row) for row in normalized_rows)
    table = target.add_table(rows=len(normalized_rows), cols=width)
    table.style = "Table Grid"
    for row_index, row in enumerate(normalized_rows):
        for cell_index in range(width):
            cell = table.cell(row_index, cell_index)
            fragments = row[cell_index] if cell_index < len(row) else [("text", "")]
            paragraph = cell.paragraphs[0]
            _append_fragments(paragraph, fragments)
            paragraph.paragraph_format.keep_together = True
            if row_index < len(normalized_rows) - 1:
                paragraph.paragraph_format.keep_with_next = True
            if row_index == 0:
                for run in paragraph.runs:
                    run.bold = True
                _set_table_row_flag(table.rows[row_index], "tblHeader")
            _set_table_row_flag(table.rows[row_index], "cantSplit")
    return table


def _set_table_row_flag(row, flag_name: str) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn(f"w:{flag_name}")) is None:
        tr_pr.append(OxmlElement(f"w:{flag_name}"))


def _estimated_text_units(text: str, *, base: float = 1.0) -> float:
    compact = re.sub(r"\s+", " ", text or "").strip()
    if not compact:
        return base
    cjk_count = len(_CJK_RE.findall(compact))
    weighted_length = cjk_count + max(0, len(compact) - cjk_count) / 2
    return max(base, base + weighted_length / 34)


def _estimated_block_units(tag: str, text: str) -> float:
    if tag in {"h1"}:
        return 3.2
    if tag in {"h2"}:
        return 2.4
    if tag in {"h3"}:
        return 2.0
    if tag == "math":
        return 2.2
    if tag == "pre":
        return max(1.6, len((text or "").splitlines()) * 1.2)
    if tag in {"li"}:
        return _estimated_text_units(text, base=0.75)
    return _estimated_text_units(text)


def _estimated_table_units(rows: TableRows) -> float:
    if not rows:
        return 0
    row_units = 0.0
    for row in rows:
        cell_text = " ".join(_fragment_text(cell) for cell in row)
        row_units += _estimated_text_units(cell_text, base=1.4)
    return max(2.4, row_units + 1.2)


def _advance_page_units(current_units: float, added_units: float) -> float:
    if added_units >= _DOCX_PAGE_UNIT_LIMIT:
        return added_units % _DOCX_PAGE_UNIT_LIMIT
    next_units = current_units + added_units
    if next_units > _DOCX_PAGE_UNIT_LIMIT:
        return added_units
    return next_units


def _maybe_page_break_before_table(target: DocxDocument, rows: TableRows, current_units: float) -> float:
    table_units = _estimated_table_units(rows)
    if (
        rows
        and current_units > 0
        and table_units < _DOCX_PAGE_UNIT_LIMIT
        and current_units + table_units > _DOCX_PAGE_UNIT_LIMIT * _DOCX_TABLE_PAGE_BREAK_THRESHOLD
    ):
        target.add_page_break()
        return 0
    return current_units


def export_docx(document: BoardDocument, path: Path) -> Path:
    target = DocxDocument()
    _apply_page_settings(target, document.page_settings)
    target.add_heading(document.title, level=0)
    current_page_units = 4.0

    parser = _DocxBlockParser()
    content_html = (document.content_html or "").strip()
    content_text = (document.content_text or "").strip()
    if content_html and content_text and _html_has_visible_raw_math_text(content_html):
        content_html = text_to_html(content_text)
    elif content_html:
        content_html = _repair_suspicious_math_html(content_html)
    else:
        content_html = text_to_html(content_text)
    parser.feed(content_html)
    parser._flush()
    blocks = parser.blocks or [("p", [("text", line)], {}) for line in document.content_text.splitlines() if line.strip()]
    blocks = _normalize_fenced_docx_blocks(blocks)

    for tag, fragments, attrs in blocks:
        text = _fragment_text(fragments)
        if tag == "h1":
            paragraph = target.add_heading("", level=1)
            paragraph.paragraph_format.keep_with_next = True
            _append_fragments(paragraph, fragments)
            current_page_units = _advance_page_units(current_page_units, _estimated_block_units(tag, text))
        elif tag == "h2":
            paragraph = target.add_heading("", level=2)
            paragraph.paragraph_format.keep_with_next = True
            _append_fragments(paragraph, fragments)
            current_page_units = _advance_page_units(current_page_units, _estimated_block_units(tag, text))
        elif tag == "h3":
            paragraph = target.add_heading("", level=3)
            paragraph.paragraph_format.keep_with_next = True
            _append_fragments(paragraph, fragments)
            current_page_units = _advance_page_units(current_page_units, _estimated_block_units(tag, text))
        elif tag == "li":
            _add_fragment_paragraph(target, fragments, style="List Bullet")
            current_page_units = _advance_page_units(current_page_units, _estimated_block_units(tag, text))
        elif tag == "blockquote":
            _add_fragment_paragraph(target, fragments, style="Intense Quote")
            current_page_units = _advance_page_units(current_page_units, _estimated_block_units(tag, text))
        elif tag == "img":
            src = str(attrs.get("src") or "").strip()
            image_bytes = _decode_data_uri(src)
            if image_bytes:
                target.add_picture(io.BytesIO(image_bytes))
            elif text:
                target.add_paragraph(f"[图片] {text}")
            current_page_units = _advance_page_units(current_page_units, 10)
        elif tag == "pageBreak":
            target.add_page_break()
            current_page_units = 0
        elif tag == "table":
            rows = attrs.get("rows")
            if isinstance(rows, list):
                current_page_units = _maybe_page_break_before_table(target, rows, current_page_units)
                _add_fragment_table(target, rows)
                current_page_units = _advance_page_units(current_page_units, _estimated_table_units(rows))
        elif tag == "pre":
            _add_preformatted_paragraph(target, text)
            current_page_units = _advance_page_units(current_page_units, _estimated_block_units(tag, text))
        elif tag == "math" or (len(fragments) == 1 and fragments[0][0] == "math"):
            paragraph = target.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _append_fragments(paragraph, fragments, auto_math=False, display_math=True)
            current_page_units = _advance_page_units(current_page_units, _estimated_block_units("math", text))
        else:
            fenced_text = _fenced_code_text(text)
            if fenced_text is not None:
                _add_preformatted_paragraph(target, fenced_text)
                current_page_units = _advance_page_units(current_page_units, _estimated_block_units("pre", fenced_text))
                continue
            formula_latex = _formula_only_latex(text)
            if formula_latex:
                paragraph = target.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                _append_math(paragraph, formula_latex, display=True)
                current_page_units = _advance_page_units(current_page_units, _estimated_block_units("math", text))
            else:
                _add_fragment_paragraph(target, fragments)
                current_page_units = _advance_page_units(current_page_units, _estimated_block_units(tag, text))

    path.parent.mkdir(parents=True, exist_ok=True)
    target.save(path)
    return path
